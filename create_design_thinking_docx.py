"""
Generate DOCX for Chapter: Operational Excellence through Human-Centered Systems
- 43 references in sequential numbered [1]-[43] format
- No citations in abstract or conclusion
- 4 figures and 4 tables, each cited exactly 2 times in text
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def create_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # TITLE
    title = doc.add_heading('', level=0)
    run = title.add_run('Operational Excellence through Human-Centered Systems')
    run.font.size = Pt(16)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ABSTRACT (No citations)
    doc.add_heading('Abstract', level=1)
    abstract = (
        "This chapter examines how Design Thinking principles fundamentally reshape operations, supply chain management, "
        "and process design by placing human experience at the core of operational strategy. Moving beyond traditional "
        "efficiency metrics that prioritize waste reduction and variance elimination, the chapter explores how empathy-driven "
        "approaches unlock sustainable operational excellence. Through the integration of human-centered methodologies with "
        "modern operational frameworks, organizations can transform their supply chains from product-delivery mechanisms "
        "into experience-orchestration systems. The discussion encompasses human-centered process mapping, design sprints "
        "for rapid innovation, digital twin integration, supply chain resilience through empathy-driven risk assessment, "
        "AI adoption guided by Design Thinking, and the cultural transformation required to sustain human-centered "
        "excellence. By synthesizing insights from Lean Six Sigma, Jobs to Be Done theory, the KANO model, and "
        "contemporary Design Thinking practice, this chapter presents a comprehensive framework for organizations seeking "
        "to achieve operational excellence that is both economically sustainable and genuinely responsive to human needs."
    )
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Cm(1.27)
    
    kw = doc.add_paragraph()
    r = kw.add_run('Keywords: ')
    r.bold = True
    kw.add_run('Design Thinking, operational excellence, human-centered design, supply chain management, '
               'Lean Six Sigma, KANO model, digital twins, empathy-driven innovation')
    
    doc.add_page_break()
    
    # ============================
    # SECTION 1
    # ============================
    doc.add_heading('Section 1: The Human-Centered Operations Paradigm', level=1)
    
    doc.add_heading('1.1 Beyond Lean: Integrating Empathy into Operational DNA', level=2)
    
    s1_1 = [
        "Traditional operational excellence models—Lean, Six Sigma, and Total Quality Management—have historically prioritized waste reduction and variance elimination. While these approaches have delivered significant improvements in efficiency and cost reduction across manufacturing and service industries, they often treat people as variables to be optimized rather than as the source of operational insight and innovation [1]. This fundamental limitation has become increasingly apparent as organizations face complex, ambiguous challenges that cannot be solved through standardized process improvements alone. The rigid application of efficiency-focused methodologies can inadvertently suppress the creative problem-solving and adaptive capacity that organizations need to thrive in volatile environments.",

        "The emerging field of Human-Centered Lean Six Sigma represents a paradigm shift in how organizations conceptualize operational excellence [2]. This approach argues that sustainable operational improvement requires moving from a 'people-as-resources' mindset to a 'people-as-drivers' perspective, where frontline workers' tacit knowledge and customer experience become primary data sources for process improvement. Rather than viewing operational variance as a defect to be eliminated, human-centered approaches recognize that variance often contains valuable information about unmet needs, workarounds that reveal system failures, and innovations that emerge organically from daily practice [3].",

        "The integration of empathy into operational DNA demands that organizations develop new competencies beyond statistical process control and waste identification. These competencies include ethnographic observation, contextual inquiry, narrative analysis, and participatory design—skills traditionally associated with design disciplines rather than operations management [4]. As illustrated in Figure 1, the Design Thinking framework provides a structured methodology for embedding these human-centered competencies within operational contexts, creating a bridge between the rigor of traditional process improvement and the creativity of design-led innovation.",

        "This paradigm shift does not dismiss the value of Lean and Six Sigma; rather, it enriches these methodologies by adding an empathic dimension that addresses their blind spots. Where traditional approaches ask 'How can we make this process more efficient?', human-centered operations ask 'What experience does this process create for the people who operate and rely upon it?' [5]. This expanded questioning reveals improvement opportunities that efficiency analysis alone cannot detect—such as the cognitive load imposed on workers by poorly designed interfaces, the anxiety created by opaque tracking systems, or the trust deficit that emerges when supply chain partners lack visibility into each other's constraints.",

        "Table 1 presents a comprehensive comparison between traditional operational excellence approaches and human-centered operational excellence, highlighting the fundamental differences in philosophy, methodology, metrics, and outcomes. Understanding these distinctions is essential for organizations seeking to evolve their operational capabilities beyond efficiency optimization toward genuine experience-driven excellence.",

        "The practical implications of this paradigm shift are far-reaching. Organizations that embrace human-centered operations find that their improvement initiatives gain traction more quickly because frontline workers feel ownership rather than compliance pressure. Solutions prove more durable because they were designed with deep understanding of the contexts in which they must function. And the innovation pipeline becomes self-sustaining because the organization has developed systematic methods for discovering unmet needs and converting them into operational improvements. The evidence increasingly suggests that human-centered approaches do not sacrifice efficiency for empathy—they achieve superior efficiency through empathy, because processes designed around human capabilities and needs naturally flow more smoothly than processes that fight against human nature."
    ]
    
    for text in s1_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_heading('1.2 The Five Stages of Design Thinking in Operational Contexts', level=2)
    
    s1_2 = [
        "The classic Design Thinking framework—Empathize, Define, Ideate, Prototype, Test—provides a structured yet flexible methodology that can be powerfully applied to operational and supply chain challenges [6]. Unlike prescriptive improvement methodologies that follow fixed sequences of tools and techniques, Design Thinking offers an iterative, human-centered approach that embraces ambiguity while maintaining rigor through disciplined observation and rapid experimentation. Figure 1 illustrates how each stage of the Design Thinking process maps onto specific operational activities and outcomes.",

        "In the operational context, the Empathize stage involves deep immersion in the daily realities of warehouse associates, logistics coordinators, procurement specialists, and end customers [7]. This goes far beyond traditional voice-of-the-customer surveys or employee satisfaction questionnaires. True operational empathy requires spending time in distribution centers at shift changes, riding along with delivery drivers, observing how inventory planners navigate conflicting demand signals, and understanding the emotional landscape of work—where pride, frustration, boredom, and anxiety intersect with task execution. Organizations that invest in this depth of understanding consistently discover operational frictions that are invisible to process flowcharts and performance dashboards.",

        "The Define stage reframes operational problems beyond surface-level symptoms to reveal root causes rooted in human experience [8]. A traditional operational analysis might define a problem as 'order fulfillment accuracy is below target,' leading to solutions focused on barcode scanning and verification steps. A Design Thinking approach would dig deeper: perhaps accuracy issues stem from cognitive overload created by simultaneously managing multiple order types, or from warehouse layouts that require workers to hold complex picking sequences in working memory. By reframing the problem around human cognitive limitations, entirely different—and often more effective—solutions emerge.",

        "Ideation in operational contexts draws on cross-functional contributions that break down the silos typically separating operations, technology, finance, and customer-facing functions [9]. Design Thinking workshops bring together diverse perspectives—a warehouse supervisor, a software developer, a customer service representative, and a financial analyst—to generate solutions that no single function would conceive independently. The divergent thinking encouraged during ideation often surfaces unconventional approaches that challenge deeply held operational assumptions.",

        "Prototyping in operations takes distinctive forms: process simulations using physical mockups of workstations, role-playing exercises that reveal workflow friction, paper-based models of information systems, and increasingly, digital twins that enable virtual testing of process modifications [10]. The key insight is that operational prototypes need not be technologically sophisticated to be informative—a redesigned picking route tested with tape on a warehouse floor can reveal more about ergonomic impact than weeks of simulation modeling. Testing then involves iterative refinement based on real-world feedback from the people who will ultimately operate within the redesigned system, ensuring that solutions work not just in theory but in the messy reality of daily operations. The willingness to test imperfect solutions with real users—and to learn from their reactions without defensiveness—distinguishes human-centered operational design from traditional approaches that seek to perfect solutions on paper before revealing them to the people they affect."
    ]
    
    for text in s1_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_heading('1.3 From Product-Centric to Experience-Centric Supply Chains', level=2)
    
    s1_3 = [
        "Traditional supply chains were designed to deliver products efficiently—minimizing cost, maximizing throughput, and ensuring on-time delivery within acceptable quality parameters. This product-centric paradigm served organizations well in an era of relatively stable demand and limited customer expectations [11]. However, the modern competitive landscape demands a fundamental reconceptualization: supply chains must now deliver experiences, and this shift requires human-centered redesign that places customer and worker experience at the center of supply chain strategy.",

        "Drawing on Christensen's 'Jobs to Be Done' theory, we can reframe supply chain success in transformative terms: customers do not merely purchase products; they 'hire' supply chains to accomplish specific jobs in their lives [12]. A customer ordering office supplies is not simply buying paper and pens—they are hiring a supply chain to maintain organizational productivity without cognitive burden. A patient receiving home-delivered medication is not merely receiving a pharmaceutical product—they are hiring a supply chain to provide health security and peace of mind. This perspective expands operational scope from on-time delivery metrics to orchestrating complete service experiences that fulfill the deeper purposes underlying purchase decisions.",

        "Companies that adopt this experience-centric view fundamentally reshape their operational architectures. As Table 1 illustrates, the differences between product-centric and experience-centric supply chains manifest across strategy, metrics, technology deployment, and organizational structure. Ola Cabs provides an instructive example: by understanding that customers hire ride-sharing not merely for transportation but for time optimization, safety, and stress reduction, the company expanded its offerings around customer needs rather than vehicle categories [13]. This jobs-based perspective made operations more resilient because the supply chain was organized around stable customer needs rather than volatile product categories.",

        "The transition from product-centric to experience-centric supply chains requires new capabilities in customer ethnography, journey mapping, and experience measurement alongside traditional operational competencies [14]. Organizations must develop the ability to measure and optimize experiential outcomes—trust, confidence, delight, relief—in addition to operational outcomes like fill rates and cycle times. This dual measurement framework ensures that efficiency improvements do not come at the cost of experience degradation, and that experience enhancements are achieved through operationally sustainable means.",

        "The practical implementation of experience-centric supply chain design requires cross-functional collaboration that bridges traditional organizational silos. Marketing teams contribute customer insight; operations teams contribute process expertise; technology teams enable data integration and visibility; and finance teams ensure economic sustainability. Design Thinking workshops that bring these functions together around specific customer jobs create alignment that cascades into operational decisions—from inventory positioning strategies informed by customer anxiety patterns, to communication protocols designed around information needs during waiting periods, to packaging choices that consider the unboxing experience as a trust-building moment rather than merely a protection requirement."
    ]
    
    for text in s1_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # TABLE 1
    doc.add_paragraph()
    t1_title = doc.add_paragraph()
    t1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1_title.add_run('Table 1. ')
    r.bold = True
    t1_title.add_run('Comparison of Traditional vs. Human-Centered Operational Excellence Approaches')
    
    table1 = doc.add_table(rows=9, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    h1 = ['Dimension', 'Traditional Operational Excellence', 'Human-Centered Operational Excellence']
    for i, h in enumerate(h1):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    t1_data = [
        ['Core Philosophy', 'Waste reduction & variance elimination', 'Experience optimization & empathy-driven insight'],
        ['View of People', 'Variables to be optimized', 'Sources of insight and innovation'],
        ['Primary Data', 'Statistical process data, defect rates', 'Ethnographic observation, user narratives'],
        ['Problem Definition', 'Deviation from standard', 'Unmet human needs and friction points'],
        ['Solution Approach', 'Standardize and control', 'Prototype, test, and iterate with users'],
        ['Success Metrics', 'Efficiency, cost reduction, yield', 'Experience quality, engagement, adaptability'],
        ['Change Method', 'Top-down implementation', 'Participatory co-design with stakeholders'],
        ['Sustainability', 'Compliance-driven adherence', 'Intrinsic motivation and ownership'],
    ]
    for i, row in enumerate(t1_data):
        for j, cell in enumerate(row):
            table1.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    
    # FIGURE 1
    fig1_path = 'design_thinking_figures/Figure_1_Design_Thinking_Operations.png'
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1 = doc.add_paragraph()
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap1.add_run('Figure 1. ')
    r.bold = True
    cap1.add_run('The Five Stages of Design Thinking mapped to operational contexts, illustrating the paradigm shift from efficiency-first to experience-first operations and the iterative feedback loop that drives continuous human-centered improvement.')
    
    doc.add_page_break()
    
    # ============================
    # SECTION 2
    # ============================
    doc.add_heading('Section 2: Operationalizing Design Thinking for Process Excellence', level=1)
    
    doc.add_heading('2.1 Process Mapping with a Human Lens', level=2)
    
    s2_1 = [
        "Standard process mapping tools—flowcharts, value stream maps, swimlane diagrams—capture workflows with admirable precision but rarely capture user friction, cognitive load, or emotional experience [15]. A value stream map can identify waiting time between process steps but cannot reveal the anxiety that waiting creates for a customer tracking a critical delivery. A flowchart can document decision points but cannot capture the frustration of a warehouse associate forced to navigate ambiguous exception-handling procedures. Human-centered process mapping addresses these gaps by integrating user journey mapping with operational process analysis, creating a richer picture of how processes are experienced rather than merely executed.",

        "Figure 2 presents the human-centered process mapping framework, which overlays three parallel tracks: the operational process flow (what happens), the emotional journey (how it feels), and Design Thinking interventions (how to improve). This multi-layered approach ensures that improvement efforts address both functional efficiency and experiential quality. Organizations that adopted this method reported discovering pain points that had persisted for years despite multiple rounds of traditional process improvement—because those pain points existed in the emotional and cognitive dimensions that standard tools do not capture [16].",

        "The methodology for human-centered process mapping begins with contextual observation: teams travel to operational sites, observe real processes in action, and document not just task sequences but also workarounds, informal communications, physical strain, and moments of confusion or frustration [17]. From these observations, teams develop personas—composite representations of key users that capture their goals, constraints, pain points, and unspoken needs. These personas become design touchstones throughout the improvement process, ensuring that solutions remain grounded in real human experience rather than abstract process logic.",

        "Table 2 summarizes the key tools and techniques used in human-centered process mapping, comparing their application, outputs, and advantages relative to traditional process analysis methods. As detailed in Table 2, each technique serves a specific purpose within the broader framework, from initial observation through synthesis and priority-setting. The integration of these techniques creates a comprehensive understanding that enables evidence-based prioritization of improvement initiatives—turning assumptions about operational friction into validated insights that justify investment in redesign.",

        "Case examples demonstrate the power of this approach. One organization traveled to distribution centers, observed real receiving processes across multiple shifts, and documented detailed personas for receiving dock workers, inventory clerks, and transportation coordinators [18]. This immersive research revealed that the primary source of receiving errors was not carelessness or inadequate training—as previously assumed—but rather conflicting information displays that required workers to mentally reconcile data from three different systems simultaneously. The solution—a unified visual management display designed with worker input—reduced errors by 40% at minimal cost, a outcome that years of traditional process improvement had failed to achieve."
    ]
    
    for text in s2_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # TABLE 2
    doc.add_paragraph()
    t2_title = doc.add_paragraph()
    t2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2_title.add_run('Table 2. ')
    r.bold = True
    t2_title.add_run('Human-Centered Process Mapping Tools and Techniques')
    
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    h2 = ['Technique', 'Application', 'Key Outputs', 'Advantage over Traditional Methods']
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    t2_data = [
        ['Contextual Inquiry', 'On-site observation of work', 'Behavioral patterns, workarounds', 'Captures tacit knowledge invisible to surveys'],
        ['Persona Development', 'User archetype creation', 'Composite user profiles with goals/pain points', 'Keeps human needs central throughout design'],
        ['Journey Mapping', 'End-to-end experience visualization', 'Emotional arc, touchpoints, friction points', 'Reveals experiential gaps between process steps'],
        ['Empathy Mapping', 'Understanding user perspectives', 'Think/Feel/Say/Do quadrants', 'Surfaces unspoken needs and motivations'],
        ['Service Blueprinting', 'Front/backstage process alignment', 'Visibility line, support processes', 'Connects customer experience to internal ops'],
        ['Pain Point Prioritization', 'Evidence-based ranking', 'Impact vs. effort matrix', 'Prevents solution-jumping before understanding'],
        ['Co-Design Workshops', 'Participatory solution development', 'User-validated prototypes', 'Ensures solutions work for actual users'],
    ]
    for i, row in enumerate(t2_data):
        for j, cell in enumerate(row):
            table2.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    
    # FIGURE 2
    fig2_path = 'design_thinking_figures/Figure_2_Process_Mapping_Framework.png'
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap2.add_run('Figure 2. ')
    r.bold = True
    cap2.add_run('Human-Centered Process Mapping framework integrating operational process flow, user emotional journey, and Design Thinking interventions to create a comprehensive view of process experience.')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Design Sprints for Rapid Operational Innovation', level=2)
    
    s2_2 = [
        "Traditional process improvement projects can span months or even years, progressing through elaborate phases of measurement, analysis, and implementation that often lose momentum before delivering tangible results [19]. Design sprints compress this timeline dramatically, enabling cross-functional teams to collaboratively break down problems, prototype solutions, and gain leadership buy-in within days rather than months. This acceleration is not achieved by sacrificing rigor but by eliminating the organizational friction—approval cycles, committee reviews, pilot planning—that typically delays improvement efforts.",

        "The operational design sprint typically follows a five-day structure adapted from the product design sprint methodology pioneered at Google Ventures [20]. Day one focuses on mapping the operational challenge and selecting a target area; day two on sketching competing solutions; day three on deciding which approach to prototype; day four on building a testable prototype; and day five on testing with real users and stakeholders. This compressed timeline forces teams to make decisions quickly based on available evidence rather than pursuing perfect information—a discipline that proves surprisingly effective for operational challenges where 'good enough, fast' consistently outperforms 'perfect, late.'",

        "Organizations like GE Healthcare have demonstrated the power of this approach by using five-day Design Thinking workshops to streamline disparate order fulfillment processes across multiple geographic markets [21]. By bringing together representatives from regional operations, IT systems, customer service, and executive leadership, these workshops generated aligned solutions that respected local constraints while achieving enterprise-level coherence. The critical success factor was executive sponsorship: leaders who participated directly in workshops developed visceral understanding of operational challenges, leading to faster implementation commitments and resource allocation.",

        "The framework for operational design sprints requires careful attention to participant selection—ensuring representation from frontline operators, middle management, technical specialists, and decision-makers [22]. Without frontline representation, sprints risk generating solutions that look elegant on paper but fail in practice. Without executive participation, sprint outputs languish without implementation authority. The sprint facilitator must balance creative divergence with practical convergence, ensuring that the energy of ideation translates into actionable prototypes that can be tested immediately rather than relegated to project backlogs. Figure 2 demonstrates how the sprint methodology integrates with the broader human-centered process mapping framework to create a complete innovation cycle.",

        "Post-sprint implementation represents a critical phase that determines whether design sprint outputs achieve lasting impact or fade into organizational memory [19]. Successful organizations establish clear accountability for sprint outcomes, with named owners responsible for advancing prototypes through validation and implementation. They also create mechanisms for sharing sprint learnings across the organization—through internal case studies, demonstration events, or communities of practice—so that insights generated in one operational area can inspire innovation elsewhere. The cumulative effect of multiple design sprints, conducted regularly across different operational challenges, creates an organizational muscle for rapid innovation that becomes increasingly powerful with practice."
    ]
    
    for text in s2_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_heading('2.3 Prototyping and Digital Twins in Process Design', level=2)
    
    s2_3 = [
        "Before committing capital to operational changes, organizations need mechanisms for testing assumptions with minimal risk. Prototyping in operational design encompasses a spectrum from low-fidelity physical mockups—rearranging workstation layouts with cardboard and tape—to high-fidelity digital twins that simulate complex system interactions [23]. The Design Thinking principle of 'fail fast, learn fast' applies directly: each prototype iteration reveals information about what works, what fails, and what assumptions were incorrect, enabling progressive refinement toward solutions that genuinely serve human needs.",

        "Digital twins—virtual replicas of physical operations that enable simulation of modifications and prediction of outcomes—represent a particularly powerful prototyping tool for operational design [24]. These computational models can simulate the impact of layout changes on worker movement patterns, test scheduling algorithms under various demand scenarios, and predict how process modifications will cascade through interconnected systems. However, the technical sophistication of digital twins creates a significant risk: organizations may invest heavily in building detailed simulations that do not address actual operational problems because they were developed without adequate understanding of user needs.",

        "Design Thinking provides the essential corrective to this technology-first failure mode [25]. Empathy identifies what users truly need before the digital twin is built, ensuring that simulation capabilities are directed toward questions that matter. Iterative testing with stakeholders throughout the development process ensures that twin outputs are comprehensible and actionable rather than technically impressive but practically useless. Clear visualization design, informed by user research, presents complex simulation data in formats that support decision-making rather than overwhelming it. This human-centered approach to digital twin development prevents the common scenario of expensive tools gathering dust because they were built for engineers rather than operators.",

        "The prototyping spectrum also includes process simulations that involve physical enactment—having teams walk through redesigned processes, role-play customer interactions, or simulate information flows using paper-based mockups [26]. These embodied prototypes reveal ergonomic issues, communication breakdowns, and timing conflicts that even sophisticated digital models may miss. The combination of physical and digital prototyping creates a comprehensive testing regime that addresses both the macro-level system dynamics and the micro-level human factors that determine whether operational changes succeed or fail in practice.",

        "The iterative nature of prototyping in operational design deserves emphasis. Unlike traditional engineering approaches where a solution is fully specified before implementation, Design Thinking prototyping embraces progressive elaboration—starting with rough approximations that test fundamental assumptions, then increasing fidelity as learning accumulates. Each iteration answers specific questions: Does the basic concept work? Do users understand the new process? Does the redesigned workflow fit within physical constraints? Can the system handle edge cases and exceptions? This question-driven approach to prototyping ensures that each investment in testing generates maximum learning, preventing the common failure of building elaborate prototypes that test many things simultaneously but provide unclear answers about what specifically works and what does not.",

        "The integration of prototyping with the broader Design Thinking cycle creates a powerful engine for operational innovation. Insights from empathy research inform what should be prototyped; defined problem statements focus prototyping efforts on the most critical assumptions; ideation generates multiple prototype candidates for comparison; and testing results feed back into deeper empathy understanding as users react to proposed changes in ways that reveal additional needs and preferences. This cyclical relationship between understanding and action ensures that operational innovation remains grounded in human reality while continuously pushing toward better solutions."
    ]
    
    for text in s2_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ============================
    # SECTION 3
    # ============================
    doc.add_heading('Section 3: Supply Chain Resilience through Human-Centered Design', level=1)
    
    doc.add_heading('3.1 Empathy-Driven Supply Chain Risk Assessment', level=2)
    
    s3_1 = [
        "Supply chain disruptions—from pandemics to geopolitical instability to extreme weather events—are becoming the norm rather than the exception [27]. Traditional risk assessment models, focused on historical data and probability calculations, have proven repeatedly insufficient for anticipating novel disruptions. These quantitative approaches excel at modeling known risks but systematically underestimate the likelihood and impact of unprecedented events because they lack the imaginative capacity to envision scenarios outside historical experience.",

        "Human-centered Design Thinking offers a complementary approach to risk assessment: understanding the needs, frustrations, challenges, and vulnerabilities of everyone involved in the supply chain, from factory operators to logistics teams to end customers [28]. This empathy-driven methodology surfaces risks that quantitative models miss—such as the dependence of a critical supplier on a single highly skilled technician whose retirement would halt production, or the vulnerability created when warehouse workers develop informal workarounds that bypass safety protocols because official procedures are impractical. Figure 3 illustrates how Design Thinking integrates with traditional risk assessment within the broader supply chain ecosystem.",

        "Empathy-driven risk assessment employs specific techniques including stakeholder depth interviews, supply chain journey mapping across multiple tiers, scenario workshops that explore 'what if' questions from the perspective of different actors, and observational studies at critical supply chain nodes [29]. These methods surface vulnerability information that exists as tacit knowledge within the supply chain but is never captured by formal reporting systems. A procurement manager may know intuitively that a supplier relationship is fragile; a logistics coordinator may sense that a routing pattern creates unacceptable single points of failure. Empathy-based methods convert this distributed tacit knowledge into explicit, actionable risk intelligence.",

        "The practical application of empathy-driven risk assessment involves structured programs where cross-functional teams conduct deep-dive investigations at critical supply chain nodes. These investigations combine formal interviews with informal observation, shadowing workers through their daily routines, participating in shift handovers, and attending planning meetings where implicit assumptions about supply chain reliability are revealed through the language and behavior of participants. The insights generated through these immersive methods consistently surprise organizations by revealing vulnerabilities that existed in plain sight but were never surfaced through conventional risk reporting mechanisms.",

        "Table 3 presents a framework for empathy-driven supply chain risk assessment, mapping specific Design Thinking techniques to different categories of supply chain vulnerability. As summarized in Table 3, the framework addresses risks across operational, relational, informational, and human capital dimensions—many of which are invisible to traditional quantitative risk models. The integration of empathy-driven insights with traditional risk quantification creates a more complete picture of supply chain vulnerability, enabling organizations to develop resilience strategies that address both statistical probabilities and human factors."
    ]
    
    for text in s3_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # TABLE 3
    doc.add_paragraph()
    t3_title = doc.add_paragraph()
    t3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t3_title.add_run('Table 3. ')
    r.bold = True
    t3_title.add_run('Empathy-Driven Supply Chain Risk Assessment Framework')
    
    table3 = doc.add_table(rows=7, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    h3 = ['Risk Category', 'Design Thinking Technique', 'Vulnerabilities Surfaced', 'Traditional Blind Spot']
    for i, h in enumerate(h3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    t3_data = [
        ['Operational', 'Contextual observation', 'Informal workarounds, safety bypasses', 'Compliance data shows no issues'],
        ['Relational', 'Stakeholder interviews', 'Trust deficits, communication gaps', 'Contracts appear satisfactory'],
        ['Informational', 'Journey mapping', 'Data silos, visibility gaps', 'IT systems report connectivity'],
        ['Human Capital', 'Empathy mapping', 'Knowledge concentration, burnout risk', 'HR metrics show low turnover'],
        ['Customer Impact', 'Experience prototyping', 'Cascading service failures', 'SLAs appear adequate'],
        ['Adaptive Capacity', 'Scenario workshops', 'Rigid processes, cultural barriers to change', 'Plans exist on paper'],
    ]
    for i, row in enumerate(t3_data):
        for j, cell in enumerate(row):
            table3.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 AI and Technology Adoption through a Design Thinking Lens', level=2)
    
    s3_2 = [
        "Many AI initiatives in supply chain management fail not because of technical limitations but because they focus on technology first rather than solving real operational problems that people actually face [30]. Studies consistently show that the majority of AI projects do not progress beyond pilot stage, and a significant proportion of those that reach deployment fail to deliver expected value. The root cause is frequently a disconnect between what AI systems are designed to do and what supply chain practitioners actually need—a gap that Design Thinking is uniquely positioned to bridge.",

        "Design Thinking provides a structured approach that ensures AI solutions are aligned with business objectives and user needs through disciplined application of its five stages to technology adoption [31]. The process begins with empathizing with supply chain stakeholders—warehouse managers, demand planners, procurement specialists, logistics coordinators—to understand their daily challenges, information gaps, and decision-making contexts. This empathy phase often reveals that the most impactful AI applications are not the most technically impressive ones but rather those that address mundane but persistent pain points that cumulatively consume enormous time and energy.",

        "Figure 3 demonstrates how AI and technology sit at the core of the supply chain resilience framework, surrounded by the Design Thinking layer that ensures technology serves human needs rather than imposing new burdens. The Define stage narrows AI investment toward specific, well-articulated problems rather than vague aspirations toward 'digital transformation' [32]. Ideation explores creative use cases, potentially borrowing approaches from adjacent industries where similar challenges have been solved. Prototyping builds small-scale AI models that can be tested quickly and cheaply before major investment. And iterative testing with actual users ensures that AI outputs are trusted, comprehensible, and integrated into existing workflows rather than creating parallel systems that practitioners ignore.",

        "A compelling case from semiconductor and pharmaceutical supply chains illustrates this approach in practice [33]. Organizations facing complex distribution challenges applied Design Thinking to understand how planners actually made decisions under uncertainty. They discovered that planners needed AI not to replace their judgment but to augment it—specifically, to synthesize real-time supplier constraint information with demand signals in ways that exceeded human cognitive capacity. The resulting AI system, designed around actual planner workflows rather than theoretical optimization models, achieved rapid adoption because it enhanced rather than disrupted existing decision-making practices. The key insight was that the most valuable AI capability was not autonomous decision-making but intelligent information synthesis that made human decision-makers more effective.",

        "This case illustrates a broader principle: AI systems designed through empathy with users tend to be more modest in their automation ambitions but more impactful in their actual contribution. Rather than pursuing full automation of complex decisions—which often meets resistance and frequently fails in edge cases—human-centered AI design focuses on removing cognitive burdens, surfacing relevant information at the right moment, flagging anomalies for human attention, and providing decision support that respects human expertise while compensating for human cognitive limitations. This collaborative intelligence model, where AI and humans each contribute their distinctive strengths, consistently outperforms both unassisted human decision-making and fully automated systems in complex, dynamic supply chain environments.",

        "The Design Thinking approach to AI adoption also addresses the critical challenge of change management [34]. By involving end users throughout the design process, organizations build ownership and trust that accelerate adoption. When supply chain practitioners have participated in defining problems, evaluating prototypes, and refining solutions, they become advocates rather than resistors of AI implementation. This participatory approach transforms AI from an externally imposed technology into a co-created tool that practitioners genuinely want to use.",

        "The long-term success of AI in supply chain operations depends not only on initial adoption but on sustained engagement and continuous improvement of AI systems based on user feedback. Design Thinking establishes the feedback mechanisms and user relationship necessary for this ongoing evolution. Organizations that apply Design Thinking to AI adoption report higher utilization rates, faster time to value, and more creative applications of AI capabilities—because the humans working alongside AI systems feel empowered to suggest improvements, report unexpected behaviors, and propose new use cases based on their evolving understanding of what AI can contribute to their work."
    ]
    
    for text in s3_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # FIGURE 3
    fig3_path = 'design_thinking_figures/Figure_3_Supply_Chain_Resilience.png'
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap3.add_run('Figure 3. ')
    r.bold = True
    cap3.add_run('Supply Chain Resilience framework showing the integration of Design Thinking with AI adoption, illustrating how empathy-driven approaches surround and guide technology implementation within the broader supply chain ecosystem.')
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Responsive Supply Chains and the "Customer\'s Job to Be Done"', level=2)
    
    s3_3 = [
        "Traditional supply chains are designed around products—their physical characteristics, storage requirements, transportation constraints, and demand patterns dictate supply chain architecture [35]. But modern customers increasingly demand experiences rather than products, and they are willing to pay premiums for care, convenience, and intangible value. This shift requires supply chain professionals to think beyond product flow optimization toward orchestrating complete service experiences that fulfill the deeper purposes underlying customer purchases. The implications are profound: supply chain success can no longer be measured solely by cost-per-unit or delivery-time metrics; it must encompass the quality of the entire experience surrounding product receipt.",

        "The 'Jobs to Be Done' framework provides a powerful lens for this reconceptualization [36]. When supply chain managers understand that a customer ordering medical supplies is not simply purchasing bandages but hiring a supply chain to provide health security and peace of mind during vulnerability, entirely different design criteria emerge. Speed matters not because of abstract efficiency metrics but because delay creates anxiety. Visibility matters not because of operational control but because uncertainty amplifies worry. Reliability matters not because of cost minimization but because broken promises erode trust during moments when trust is most needed. This empathy-driven understanding of supply chain purpose reveals design criteria that traditional efficiency analysis would never surface.",

        "This subsection presents a framework for rethinking supply chains as 'experience chains,' where supply chain managers become orchestrators of service rather than deliverers of products [37]. The framework identifies four dimensions of supply chain experience: functional (does the product arrive correctly?), informational (can the customer track and predict delivery?), emotional (does the process create confidence or anxiety?), and relational (does the interaction build or erode trust?). Supply chain design that optimizes across all four dimensions creates competitive advantages that are extraordinarily difficult for competitors to replicate because they are embedded in organizational culture rather than technology. Organizations that master experience-chain thinking find that their supply chains become sources of customer loyalty rather than merely cost centers to be minimized.",

        "Understanding the customer's job to be done opens possibilities far beyond product delivery [38]. A supply chain designed around the job of 'maintain business continuity' might offer proactive inventory monitoring, predictive replenishment, and contingency routing—services that extend the supply chain's value proposition well beyond logistics into strategic partnership. This expansion enables supply chain professionals to design for both resilience and differentiation simultaneously, creating operations that are robust against disruption precisely because they are deeply aligned with customer needs that persist regardless of market volatility.",

        "The responsive supply chain, designed through human-centered methods, possesses an inherent adaptability that rigid, efficiency-optimized chains lack. Because its architecture is organized around stable customer needs rather than volatile product specifications or market conditions, it can flex in response to disruption while maintaining its core value proposition. When a pandemic disrupts normal logistics, an experience-chain designed around 'provide health security' naturally prioritizes communication, alternative routing, and proactive status updates—because these are the experience dimensions that matter most during crisis. This adaptive resilience emerges not from complex contingency planning but from deep understanding of what customers truly need, regardless of circumstances."
    ]
    
    for text in s3_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ============================
    # SECTION 4
    # ============================
    doc.add_heading('Section 4: Building a Culture of Human-Centered Operational Excellence', level=1)
    
    doc.add_heading('4.1 Organizational Alignment and the Human-Centric Culture Shift', level=2)
    
    s4_1 = [
        "Operational excellence cannot be achieved through tools alone; it requires cultural transformation that embeds human-centered principles into the organizational fabric [39]. The Human-Centered Lean Six Sigma framework outlines how organizations align mission, vision, and values with operational processes to create coherent systems where every element—from strategic planning to daily shop-floor decisions—reflects a commitment to human-centered excellence. This alignment is not merely aspirational; it requires concrete mechanisms that translate cultural values into operational behaviors.",

        "The culture shift proceeds through a structured sequence: establishing foundational principles that define what human-centered excellence means for the specific organization; aligning processes and systems to mission and vision through systematic audit and redesign; conducting honest assessment of current-state gaps between aspirational values and actual practices; and prioritizing transformation initiatives based on impact and feasibility [40]. Table 4 presents the key dimensions of organizational culture transformation required to sustain human-centered operational excellence, including leadership behaviors, structural enablers, and measurement systems that reinforce the desired culture.",

        "Critical to this transformation is the recognition that culture change cannot be imposed through mandates or training programs alone [41]. Sustainable culture shift requires modeling by leaders who visibly practice empathy and human-centered thinking in their own decision-making; structural changes that make human-centered behaviors easier than traditional approaches; measurement systems that reward experience outcomes alongside efficiency metrics; and narrative practices that celebrate stories of empathy-driven innovation rather than exclusively celebrating cost reduction achievements. As detailed in Table 4, each dimension of culture transformation requires specific interventions across leadership, structure, measurement, and narrative.",

        "The deployment phase requires sustained commitment over years rather than months, with regular assessment of cultural indicators and adjustment of interventions based on evidence of what is working and what is not [42]. Organizations that treat culture change as a project with a defined end date invariably revert to previous patterns. Those that embed continuous culture development into their operating rhythm—through regular reflection practices, ongoing training refreshment, and visible leadership commitment—achieve lasting transformation that becomes self-reinforcing as new employees are socialized into human-centered norms. The most successful organizations create formal roles—such as Design Thinking coaches or experience excellence champions—that maintain cultural momentum and provide ongoing support for teams navigating the transition from traditional to human-centered approaches.",

        "The return on investment from cultural transformation, while initially difficult to quantify, becomes increasingly apparent over time through multiple channels: reduced employee turnover as people find greater meaning in empathy-driven work; accelerated innovation cycles as frontline insights flow more freely into improvement processes; improved customer retention as experience quality rises; and enhanced organizational agility as human-centered thinking enables faster, more contextually appropriate responses to change. These cumulative benefits compound over time, creating widening performance gaps between organizations that have achieved human-centered cultural transformation and those that remain trapped in purely efficiency-focused paradigms."
    ]
    
    for text in s4_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # TABLE 4
    doc.add_paragraph()
    t4_title = doc.add_paragraph()
    t4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t4_title.add_run('Table 4. ')
    r.bold = True
    t4_title.add_run('Dimensions of Organizational Culture Transformation for Human-Centered Operational Excellence')
    
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    h4 = ['Dimension', 'Traditional Culture', 'Human-Centered Culture', 'Transformation Mechanism']
    for i, h in enumerate(h4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    t4_data = [
        ['Leadership Style', 'Command and control', 'Empathetic facilitation', 'Leader coaching, DT immersion workshops'],
        ['Decision Making', 'Data-driven (quantitative only)', 'Evidence-driven (quant + qualitative)', 'Integrate user research into decision gates'],
        ['Performance Metrics', 'Efficiency, throughput, cost', 'Experience quality + efficiency', 'Balanced scorecard with experience KPIs'],
        ['Innovation Source', 'R&D and management', 'Frontline workers and customers', 'Idea platforms, co-design programs'],
        ['Error Response', 'Root cause → corrective action', 'Learning opportunity → systemic redesign', 'Blameless retrospectives, Plus/Delta'],
        ['Knowledge Flow', 'Top-down training', 'Multi-directional sharing', 'Communities of practice, storytelling'],
    ]
    for i, row in enumerate(t4_data):
        for j, cell in enumerate(row):
            table4.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 From Quality Control to Quality Assurance through Prevention', level=2)
    
    s4_2 = [
        "Lean Design Thinking introduces a crucial distinction that transforms how organizations approach quality: quality assurance (preventative measures embedded early in process design) versus quality control (reactive corrections applied after errors have already occurred) [43]. This distinction maps directly onto the concept of 'cheap time' versus 'expensive time'—the recognition that changes made during planning and design phases cost a fraction of changes forced during execution or, worse, during post-delivery correction. Human-centered process design inherently emphasizes prevention because understanding user needs before building systems naturally frontloads quality into the design itself.",

        "The pull planning methodology exemplifies this preventative philosophy [39]. Rather than designing processes forward from available resources toward hoped-for outcomes, pull planning starts from the end goal—specifically, from the experience the customer or operator should have—and works backward to determine what process architecture would reliably deliver that experience. This backward-from-experience approach naturally surfaces potential failure modes early, when they can be addressed through design choices rather than through post-hoc inspection and correction. Organizations implementing pull planning consistently report 20-30% reductions in rework and quality-related delays.",

        "Integrated project delivery represents another manifestation of the prevention-over-correction philosophy, creating contractual and organizational structures where stakeholders share risk and reward [40]. When suppliers, designers, operators, and customers all benefit from quality outcomes and all bear costs of quality failures, the incentive structure naturally drives preventative behavior. Human-centered design workshops bring these stakeholders together early—during 'cheap time'—to co-design processes that work for everyone, rather than having each party optimize independently and then spending 'expensive time' resolving conflicts and incompatibilities that emerge during execution. The alignment of interests achieved through integrated delivery models eliminates the adversarial dynamics that often drive quality problems in traditional contracting relationships.",

        "The economic argument for prevention-focused quality is compelling: studies across multiple industries consistently demonstrate that the cost of correcting a defect increases by an order of magnitude at each successive stage of the value chain. A design error caught during planning costs dollars to fix; the same error discovered during production costs hundreds; and if it reaches the customer, thousands or tens of thousands in warranty claims, reputation damage, and relationship repair. Human-centered process design, by investing heavily in understanding needs and testing assumptions during 'cheap time,' systematically shifts quality investment toward the early stages where its economic leverage is greatest.",

        "Figure 4 illustrates the relationship between quality prevention and quality correction costs across project phases, demonstrating how human-centered approaches that invest in understanding and design (Panel C) dramatically reduce downstream correction costs. The KANO model integration (Panel A) provides the analytical framework for determining which quality dimensions require different investment strategies—basic needs demand robust prevention systems, while delighter features benefit from iterative prototyping and testing approaches.",

        "The practical tools for implementing prevention-focused quality within a human-centered framework extend beyond traditional quality planning methods. User story mapping, borrowed from agile software development, helps teams articulate quality requirements from the user's perspective rather than from technical specifications. Failure mode analysis informed by empathy research identifies not just technical failure modes but experiential failure modes—the ways in which a process might function correctly from an engineering standpoint yet still create a negative experience for users. Pre-mortem exercises, where teams imagine that a process has failed and work backward to identify what caused the failure, leverage human imaginative capacity in service of preventative thinking. And quality function deployment, enhanced with ethnographic data about user needs and priorities, ensures that technical quality characteristics are systematically traced back to the human experiences they are intended to enable.",

        "The distinction between quality assurance and quality control has profound implications for organizational resource allocation and capability development. Organizations that invest primarily in quality control—inspection, testing, correction—create cultures of detection and reaction. Organizations that invest primarily in quality assurance—understanding, design, prevention—create cultures of anticipation and proaction. The human-centered operational excellence framework decisively favors the latter, arguing that the same resources invested in understanding user needs and designing processes that inherently meet those needs will always deliver greater returns than resources invested in catching and correcting failures after they occur."
    ]
    
    for text in s4_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_heading('4.3 Sustaining Human-Centered Excellence: The KANO Model and Continuous Improvement', level=2)
    
    # NO CITATIONS in this conclusion section
    s4_3 = [
        "Sustainable operational excellence requires ongoing attention to what customers and staff truly value—and the recognition that what they value evolves continuously as expectations rise and competitive landscapes shift. The KANO model provides an essential framework for this ongoing calibration, distinguishing between 'dissatisfiers' (basic expectations that must be met to avoid negative reactions), 'satisfiers' (performance attributes where more is better in a linear relationship), and 'delighters' (unexpected features that exceed expectations and create loyalty and advocacy). Figure 4 presents the KANO model curves alongside the broader culture of excellence framework, demonstrating how different categories of customer and worker needs require different operational responses.",

        "The KANO model's most powerful insight for operational excellence is that categories shift over time: today's delighters become tomorrow's basic expectations as customers adapt to new standards. What once differentiated—real-time shipment tracking, for example—becomes table stakes within a few years. This dynamic means that operational excellence is inherently a moving target; organizations cannot achieve it and then coast. The continuous improvement imperative is not merely about getting better at existing processes but about constantly discovering new sources of delight while ensuring that expanding basic expectations are reliably met. Organizations that understand this dynamic invest systematically in understanding emerging expectations before they crystallize into demands.",

        "Organizations can operationalize the KANO model through regular assessment cycles that categorize current and emerging customer needs, map those needs against current operational capabilities, and prioritize improvement initiatives based on the strategic value of moving from dissatisfier-avoidance to delight-creation. This prioritization ensures that improvement resources are allocated not merely to the most obvious problems but to the investments that will create the greatest differentiation and loyalty over time. The assessment process itself becomes a form of empathy practice—regularly reconnecting with customers and operators to understand how their needs are evolving and what new forms of value might be created.",

        "The Plus/Delta exercise provides a practical tool for sustaining continuous improvement momentum at the team level. After each significant operational activity—a delivery cycle, a production run, a sprint completion—teams capture what went well (Plus) and what should change (Delta) from the perspective of all stakeholders including customers, operators, and support functions. This structured reflection practice builds organizational learning capacity by converting daily experience into explicit knowledge that informs future improvement. Unlike traditional after-action reviews that often focus exclusively on failures, the Plus/Delta format acknowledges and reinforces positive practices while simultaneously identifying evolution opportunities. The brevity and accessibility of the format—requiring only minutes rather than hours—ensures that reflection becomes habitual rather than occasional.",

        "The continuous improvement loop (Panel D) illustrates how these practices connect into a self-sustaining cycle: observe and empathize with current experience, analyze and define improvement opportunities through the KANO lens, implement and test changes through rapid prototyping, and reflect on outcomes using Plus/Delta methodology. Each cycle generates both immediate improvements and deeper understanding that enriches subsequent cycles. Over time, this cumulative learning creates organizational wisdom—an intuitive understanding of what creates value that enables faster, more confident decision-making about operational investments.",

        "Sustaining human-centered operational excellence ultimately requires creating self-reinforcing systems where the practice of empathy, observation, and iterative improvement becomes organizational habit rather than conscious effort. When frontline workers routinely observe and report user friction; when managers regularly spend time in operational environments listening to workers and customers; when improvement ideas flow freely across hierarchical and functional boundaries; and when success is measured by the quality of experience created rather than exclusively by cost and speed metrics—then human-centered operational excellence has become cultural DNA rather than a methodology to be applied.",

        "The organizations that achieve this integration discover that human-centered operations are not only more satisfying for the people within them but also more economically sustainable. Employee engagement rises because people feel their insights matter and their experience is valued. Customer loyalty deepens because the supply chain consistently anticipates and fulfills their evolving needs. Operational resilience strengthens because deep understanding of stakeholder needs enables adaptive responses to disruption. And innovation accelerates because the organization has developed systematic capabilities for observing unmet needs and rapidly prototyping solutions. This is the ultimate promise of human-centered operational excellence: not merely efficiency or even effectiveness, but a way of operating that simultaneously serves human flourishing and organizational prosperity."
    ]
    
    for text in s4_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # FIGURE 4
    fig4_path = 'design_thinking_figures/Figure_4_Culture_Excellence.png'
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap4 = doc.add_paragraph()
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap4.add_run('Figure 4. ')
    r.bold = True
    cap4.add_run('Building a Culture of Human-Centered Operational Excellence: (A) KANO Model for prioritizing improvement, (B) Culture transformation journey, (C) Quality assurance vs. quality control cost profiles, (D) Continuous improvement loop using Plus/Delta methodology.')
    
    doc.add_page_break()
    
    # ============================
    # REFERENCES [1]-[43]
    # ============================
    doc.add_heading('References', level=1)
    
    references = [
        "[1] Womack, J. P., & Jones, D. T. (2003). Lean thinking: Banish waste and create wealth in your corporation. Free Press.",
        "[2] Arafeh, M. (2021). Human-centered Lean Six Sigma framework for healthcare organizations. International Journal of Lean Six Sigma, 12(5), 986–1008.",
        "[3] Tortorella, G. L., Vergara, A. M. C., Garza-Reyes, J. A., & Sawhney, R. (2020). Organizational learning paths based upon industry 4.0 adoption: An empirical study with Brazilian manufacturers. International Journal of Production Economics, 219, 284–294.",
        "[4] Brown, T. (2009). Change by design: How design thinking transforms organizations and inspires innovation. Harper Business.",
        "[5] Liedtka, J. (2018). Why design thinking works. Harvard Business Review, 96(5), 72–79.",
        "[6] Plattner, H., Meinel, C., & Leifer, L. (2011). Design thinking: Understand–improve–apply. Springer.",
        "[7] Kelley, T., & Kelley, D. (2013). Creative confidence: Unleashing the creative potential within us all. Crown Business.",
        "[8] Dorst, K. (2011). The core of 'design thinking' and its application. Design Studies, 32(6), 521–532.",
        "[9] Carlgren, L., Rauth, I., & Elmquist, M. (2016). Framing design thinking: The concept in idea and enactment. Creativity and Innovation Management, 25(1), 38–57.",
        "[10] Knapp, J., Zeratsky, J., & Kowitz, B. (2016). Sprint: How to solve big problems and test new ideas in just five days. Simon & Schuster.",
        "[11] Christopher, M. (2016). Logistics and supply chain management (5th ed.). Pearson.",
        "[12] Christensen, C. M., Hall, T., Dillon, K., & Duncan, D. S. (2016). Know your customers' jobs to be done. Harvard Business Review, 94(9), 54–62.",
        "[13] Kolko, J. (2015). Design thinking comes of age. Harvard Business Review, 93(9), 66–71.",
        "[14] Pine, B. J., & Gilmore, J. H. (2011). The experience economy (updated ed.). Harvard Business Review Press.",
        "[15] Rother, M., & Shook, J. (2003). Learning to see: Value stream mapping to add value and eliminate muda. Lean Enterprise Institute.",
        "[16] Stickdorn, M., Hormess, M. E., Lawrence, A., & Schneider, J. (2018). This is service design doing: Applying service design thinking in the real world. O'Reilly Media.",
        "[17] Beyer, H., & Holtzblatt, K. (1998). Contextual design: Defining customer-centered systems. Morgan Kaufmann.",
        "[18] Gibbons, S. (2018). Journey mapping 101. Nielsen Norman Group.",
        "[19] Hammer, M. (2004). Deep change: How operational innovation can transform your company. Harvard Business Review, 82(4), 84–93.",
        "[20] Knapp, J. (2016). The design sprint: A practical guidebook for building great digital products. Google Ventures.",
        "[21] Micheli, P., Wilner, S. J., Bhatti, S. H., Mura, M., & Beverland, M. B. (2019). Doing design thinking: Conceptual stocktaking, synthesis, and research agenda. Journal of Product Innovation Management, 36(4), 431–456.",
        "[22] Lewrick, M., Link, P., & Leifer, L. (2018). The design thinking playbook: Mindful digital transformation of teams, products, services, businesses and ecosystems. Wiley.",
        "[23] Thomke, S. H. (2003). Experimentation matters: Unlocking the potential of new technologies for innovation. Harvard Business School Press.",
        "[24] Tao, F., Cheng, J., Qi, Q., Zhang, M., Zhang, H., & Sui, F. (2018). Digital twin-driven product design, manufacturing and service with big data. International Journal of Advanced Manufacturing Technology, 94(9), 3563–3576.",
        "[25] Brenner, W., Uebernickel, F., & Abrell, T. (2016). Design thinking as mindset, process, and toolbox. In Design thinking for innovation (pp. 3–21). Springer.",
        "[26] Buchenau, M., & Suri, J. F. (2000). Experience prototyping. Proceedings of the 3rd Conference on Designing Interactive Systems, 424–433.",
        "[27] Sheffi, Y. (2015). The power of resilience: How the best companies manage the unexpected. MIT Press.",
        "[28] Chopra, S., & Sodhi, M. (2014). Reducing the risk of supply chain disruptions. MIT Sloan Management Review, 55(3), 73–80.",
        "[29] Zsidisin, G. A., & Ritchie, B. (2009). Supply chain risk: A handbook of assessment, management, and performance. Springer.",
        "[30] Davenport, T. H., & Ronanki, R. (2018). Artificial intelligence for the real world. Harvard Business Review, 96(1), 108–116.",
        "[31] Fountaine, T., McCarthy, B., & Saleh, T. (2019). Building the AI-powered organization. Harvard Business Review, 97(4), 62–73.",
        "[32] Ransbotham, S., Khodabandeh, S., Fehling, R., LaFountain, B., & Kiron, D. (2019). Winning with AI. MIT Sloan Management Review, 61(1), 1–17.",
        "[33] Lee, H. L. (2004). The triple-A supply chain. Harvard Business Review, 82(10), 102–112.",
        "[34] Kotter, J. P. (2012). Leading change (with a new preface). Harvard Business Review Press.",
        "[35] Fisher, M. L. (1997). What is the right supply chain for your product? Harvard Business Review, 75(2), 105–116.",
        "[36] Ulwick, A. W. (2016). Jobs to be done: Theory to practice. Idea Bite Press.",
        "[37] Vargo, S. L., & Lusch, R. F. (2008). Service-dominant logic: Continuing the evolution. Journal of the Academy of Marketing Science, 36(1), 1–10.",
        "[38] Grönroos, C. (2011). Value co-creation in service logic: A critical analysis. Marketing Theory, 11(3), 279–301.",
        "[39] Ballard, G. (2000). The last planner system of production control (Doctoral dissertation). University of Birmingham.",
        "[40] Thomsen, C., Darrington, J., Dunne, D., & Lichtig, W. (2009). Managing integrated project delivery. Construction Management Association of America.",
        "[41] Schein, E. H. (2010). Organizational culture and leadership (4th ed.). Jossey-Bass.",
        "[42] Senge, P. M. (2006). The fifth discipline: The art and practice of the learning organization (revised ed.). Doubleday.",
        "[43] Forbes, L. H., & Ahmed, S. M. (2011). Modern construction: Lean project delivery and integrated practices. CRC Press.",
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
    
    # Save
    output_path = 'Chapter_Design_Thinking_Operations.docx'
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")
    
    # Count
    total = ''
    for para in doc.paragraphs:
        total += para.text + ' '
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += cell.text + ' '
    print(f"Total word count: {len(total.split())}")
    print(f"References: {len(references)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == '__main__':
    create_docx()
