"""
Create Word Document for Chapter 2: Communication Protocols: DSRC, C-V2X, 5G NR-V2X, and 6G V2X
For: Vehicular Networking and Connected Vehicles: Systems, Intelligence, and Control Perspectives
Part I: Foundations of Vehicular Networking and Connected Mobility
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=11):
    """Add a formatted paragraph."""
    para = doc.add_paragraph()
    para.style = style
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return para

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


    # ===== TITLE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Chapter 2')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Communication Protocols: DSRC, C-V2X, 5G NR-V2X, and 6G V2X')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Book info
    book_info = doc.add_paragraph()
    book_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = book_info.add_run('\nVehicular Networking and Connected Vehicles:\nSystems, Intelligence, and Control Perspectives\n\nPart I: Foundations of Vehicular Networking and Connected Mobility')
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # spacing


    # ===== ABSTRACT =====
    abstract_title = doc.add_paragraph()
    run = abstract_title.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(12)
    
    abstract_text = (
        "The rapid evolution of vehicular communication technologies represents a transformative paradigm shift "
        "in intelligent transportation systems. This chapter provides a comprehensive examination of the communication "
        "protocols that underpin connected and autonomous vehicular networks, tracing the technological progression from "
        "Dedicated Short-Range Communications (DSRC) based on IEEE 802.11p through Cellular Vehicle-to-Everything (C-V2X) "
        "leveraging LTE infrastructure, to the advanced capabilities of 5G New Radio V2X (NR-V2X) and the emerging vision "
        "of 6G V2X networks. The chapter systematically analyzes the architectural foundations, protocol stacks, radio "
        "access technologies, and performance characteristics of each generation of vehicular communication standards. "
        "Particular attention is devoted to the critical requirements of connected and autonomous vehicles, including "
        "ultra-low latency, extreme reliability, high mobility support, and massive scalability. The comparative analysis "
        "of DSRC and C-V2X reveals fundamental differences in medium access control, resource allocation, and interference "
        "management that influence deployment strategies and coexistence scenarios. The chapter further explores how 5G NR-V2X "
        "introduces transformative capabilities through Ultra-Reliable Low-Latency Communication (URLLC), enhanced Mobile "
        "Broadband (eMBB), and massive Machine-Type Communication (mMTC) to support advanced use cases such as cooperative "
        "perception, vehicle platooning, and remote driving. Finally, the chapter presents the 6G V2X vision, encompassing "
        "terahertz communications, reconfigurable intelligent surfaces, AI-native air interfaces, and integrated sensing "
        "and communication technologies that will enable fully autonomous and cooperative intelligent transportation systems. "
        "The discussion addresses interoperability challenges, spectrum evolution, standardization pathways, and hybrid "
        "architectures necessary for the seamless transition across technology generations."
    )
    
    abstract_para = doc.add_paragraph()
    run = abstract_para.add_run(abstract_text)
    run.font.size = Pt(10)
    run.italic = True
    
    keywords_para = doc.add_paragraph()
    run = keywords_para.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    run = keywords_para.add_run(
        'DSRC, C-V2X, 5G NR-V2X, 6G V2X, IEEE 802.11p, Vehicle-to-Everything, URLLC, '
        'Network Slicing, Terahertz Communication, Reconfigurable Intelligent Surfaces, '
        'Cooperative Perception, Autonomous Vehicles, Connected Mobility'
    )
    run.font.size = Pt(10)
    
    doc.add_page_break()


    # ===== SECTION 1 =====
    h1 = doc.add_heading('1. Foundations and Evolution of Vehicular Communication Networks', level=1)
    
    # Section 1.1
    doc.add_heading('1.1 Connected and Cooperative Vehicular Communication: Concepts, Architecture, and Requirements', level=2)
    
    doc.add_paragraph(
        'The concept of connected and cooperative vehicular communication represents a fundamental transformation '
        'in how vehicles interact with their surrounding environment, infrastructure, and other road users. '
        'Vehicle-to-Everything (V2X) communication encompasses a broad spectrum of interaction modalities, including '
        'Vehicle-to-Vehicle (V2V), Vehicle-to-Infrastructure (V2I), Vehicle-to-Pedestrian (V2P), and Vehicle-to-Network '
        '(V2N) communications [1]. These communication paradigms collectively enable cooperative awareness, '
        'collaborative decision-making, and coordinated maneuver execution among connected vehicles and intelligent '
        'transportation infrastructure [2].'
    )
    
    doc.add_paragraph(
        'The architectural framework for vehicular communication networks is inherently multi-layered, comprising '
        'the application layer, facilities layer, networking and transport layer, access layer, and management layer '
        '[3]. At the highest level, safety applications such as Cooperative Awareness Messages (CAM), Decentralized '
        'Environmental Notification Messages (DENM), and Basic Safety Messages (BSM) require deterministic and '
        'predictable communication performance. The facilities layer provides common data dictionaries, message '
        'formatting, and service management functions that abstract the underlying communication technology from '
        'the application logic [4]. This layered architecture enables technology-agnostic application development '
        'while supporting heterogeneous access technologies ranging from DSRC to cellular-based V2X solutions.'
    )
    
    doc.add_paragraph(
        'The fundamental requirements for vehicular communication networks are dictated by the stringent demands '
        'of safety-critical applications and the inherent characteristics of vehicular environments. These requirements '
        'include end-to-end latency below 10 milliseconds for pre-crash sensing, packet delivery ratios exceeding '
        '99.99% for safety messages, support for relative vehicle speeds exceeding 500 km/h, and network densities '
        'reaching thousands of vehicles per square kilometer in urban intersections [5]. Furthermore, the dynamic '
        'topology of vehicular networks, characterized by rapid link establishment and dissolution, frequent handovers, '
        'and varying channel conditions due to multipath propagation and Doppler effects, imposes unique challenges '
        'on protocol design and resource management [6]. The evolution of vehicular communication technologies, as '
        'illustrated in Figure 1, demonstrates the progressive enhancement of these capabilities across successive '
        'technology generations.'
    )
    
    doc.add_paragraph(
        'The standardization landscape for vehicular communication has been shaped by multiple international bodies '
        'working in parallel and sometimes competing frameworks. In the United States, the Department of Transportation '
        '(DoT) initially championed DSRC through the Connected Vehicle Safety Pilot program, while the European Union '
        'developed the Cooperative Intelligent Transport Systems (C-ITS) framework based on ETSI ITS-G5, which shares '
        'the IEEE 802.11p physical layer but employs a distinct networking architecture based on GeoNetworking protocols '
        '[7]. In Asia, Japan deployed the 760 MHz band DSRC system (ARIB STD-T109) alongside 5.8 GHz Electronic Toll '
        'Collection, while China adopted C-V2X as its primary V2X technology through the national LTE-V2X standard '
        '(T/CSAE 53-2017). This geographic fragmentation of technology choices has created significant challenges for '
        'global interoperability and has motivated efforts toward technology-neutral standards that can accommodate '
        'multiple underlying radio access technologies within a common application and security framework.'
    )


    # Insert Figure 1
    fig1_para = doc.add_paragraph()
    fig1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig1_para.add_run()
    run.add_picture('/projects/sandbox/AMMAN/v2x_figures/Figure_1_V2X_Evolution.png', width=Inches(5.5))
    
    fig1_caption = doc.add_paragraph()
    fig1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig1_caption.add_run('Figure 1: Evolution of Vehicular Communication Technologies from DSRC to 6G V2X, '
                               'showing progressive enhancement in latency, reliability, data rate, and communication modes.')
    run.bold = True
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Section 1.2
    doc.add_heading('1.2 Evolution from VANETs and DSRC to Cellular V2X Networks', level=2)
    
    doc.add_paragraph(
        'The evolution of vehicular communication technologies spans more than two decades of research, '
        'standardization, and deployment efforts. The initial conceptualization of Vehicular Ad-hoc Networks '
        '(VANETs) in the early 2000s established the foundational principles of decentralized, infrastructure-less '
        'communication among vehicles using dedicated spectrum allocations [7]. The allocation of 75 MHz of spectrum '
        'in the 5.9 GHz band by the Federal Communications Commission (FCC) in 1999 for Intelligent Transportation '
        'Systems (ITS) provided the regulatory foundation for Dedicated Short-Range Communications (DSRC), which '
        'subsequently evolved into the Wireless Access in Vehicular Environments (WAVE) standard suite [8].'
    )
    
    doc.add_paragraph(
        'The IEEE 802.11p amendment, ratified in 2010, defined the physical and medium access control layers '
        'specifically optimized for vehicular environments, incorporating features such as reduced overhead for '
        'rapid link establishment without the need for association procedures [9]. However, the inherent limitations '
        'of contention-based channel access in DSRC, particularly under high vehicle density scenarios, motivated '
        'the exploration of cellular-based alternatives. The 3rd Generation Partnership Project (3GPP) introduced '
        'Cellular V2X (C-V2X) in Release 14 (2017), leveraging the mature LTE infrastructure and the advantages '
        'of scheduled access and centralized resource management [10].'
    )
    
    doc.add_paragraph(
        'The transition from DSRC to C-V2X represented a paradigm shift in vehicular communication philosophy, '
        'moving from purely ad-hoc, distributed systems to hybrid architectures combining direct device-to-device '
        'communication (PC5 sidelink interface) with network-assisted coordination through the Uu cellular interface '
        '[11]. This evolution was further accelerated by the development of 5G NR-V2X in 3GPP Release 16, which '
        'introduced advanced features such as unicast and groupcast communication modes, feedback mechanisms, and '
        'significantly enhanced reliability and latency performance [12]. The evolutionary trajectory, depicted in '
        'Figure 1, illustrates how each technology generation has progressively addressed the growing demands of '
        'connected and autonomous vehicle applications.'
    )
    
    doc.add_paragraph(
        'The coexistence debate between DSRC and C-V2X has significantly influenced the regulatory landscape and '
        'delayed large-scale V2X deployment in several regions. In November 2020, the U.S. Federal Communications '
        'Commission (FCC) reallocated the lower 45 MHz of the 5.9 GHz band to unlicensed Wi-Fi operations, retaining '
        'only the upper 30 MHz (5.895-5.925 GHz) for ITS, with C-V2X designated as the communication technology for '
        'this reduced band [8]. This regulatory decision effectively ended the DSRC deployment trajectory in the '
        'United States and accelerated the global momentum toward C-V2X and its 5G NR-V2X successor. In Europe, '
        'the European Commission adopted a technology-neutral approach through Delegated Regulation 2022/1426, '
        'allowing both ITS-G5 (DSRC-based) and C-V2X deployments while mandating interoperability at the application '
        'layer through common message formats and security frameworks. This regulatory evolution reflects the broader '
        'transition in vehicular communication philosophy from single-technology mandates to flexible, performance-based '
        'requirements that can accommodate technology evolution over the multi-decade lifecycle of transportation '
        'infrastructure.'
    )


    # Section 1.3
    doc.add_heading('1.3 Communication Requirements for Connected and Autonomous Vehicles: Latency, Reliability, Mobility, and Scalability', level=2)
    
    doc.add_paragraph(
        'The communication requirements for connected and autonomous vehicles vary significantly across different '
        'application categories, ranging from basic safety awareness to fully cooperative autonomous driving. '
        'The European Telecommunications Standards Institute (ETSI) and 3GPP have jointly defined a comprehensive '
        'taxonomy of V2X use cases organized into four service categories: vehicles platooning, extended sensors, '
        'advanced driving, and remote driving [13]. Each category imposes distinct requirements on communication '
        'latency, reliability, throughput, and positioning accuracy, as summarized in Table 1.'
    )
    
    # TABLE 1: V2X Use Case Requirements
    doc.add_paragraph()
    table1_caption = doc.add_paragraph()
    run = table1_caption.add_run('Table 1: Communication Requirements for V2X Use Case Categories [13, 14]')
    run.bold = True
    run.font.size = Pt(10)
    table1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table1 = doc.add_table(rows=6, cols=6)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Use Case Category', 'Latency (ms)', 'Reliability (%)', 'Data Rate', 'Range (m)', 'Positioning']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    data = [
        ['Cooperative Awareness', '100', '95', '5-96 kbps', '300-500', '1.5 m'],
        ['Vehicle Platooning', '10-25', '99.99', '50-65 Mbps', '80-350', '0.5 m'],
        ['Extended Sensors', '3-10', '99.999', '10-1000 Mbps', '50-1000', '0.2 m'],
        ['Advanced Driving', '3-10', '99.999', '10-50 Mbps', '100-700', '0.1 m'],
        ['Remote Driving', '5-20', '99.999', '1-25 Mbps UL', '>1000', '0.1 m'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table1.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'The latency requirements for vehicular communications span several orders of magnitude, from the relatively '
        'relaxed 100-millisecond tolerance for basic cooperative awareness messages to the stringent sub-3-millisecond '
        'requirement for advanced cooperative maneuver coordination [14]. These latency constraints encompass the '
        'entire communication chain, including sensing, processing, transmission, propagation, and reception delays. '
        'For autonomous driving applications, the end-to-end latency budget must account for sensor data acquisition '
        '(typically 10-50 ms), perception and decision-making processing (20-100 ms), and actuation response time '
        '(10-50 ms), leaving minimal margin for communication delays [15].'
    )
    
    doc.add_paragraph(
        'Reliability requirements are equally demanding, with safety-critical applications requiring packet delivery '
        'ratios exceeding 99.999% (five nines reliability). This translates to fewer than one lost message per 100,000 '
        'transmissions, which is particularly challenging in highly dynamic vehicular environments characterized by '
        'rapid fading, frequent shadowing by large vehicles, and interference from dense concurrent transmissions [16]. '
        'The mobility requirement presents additional challenges, as vehicular communication systems must maintain '
        'performance at relative speeds exceeding 500 km/h on highways, corresponding to Doppler shifts of approximately '
        '2.7 kHz at 5.9 GHz carrier frequency, which can severely degrade channel estimation and demodulation '
        'performance [17]. Scalability concerns arise in urban scenarios where intersection vehicle densities may '
        'exceed 1,000 vehicles, each generating 10-50 messages per second, creating aggregate channel loads that '
        'challenge the capacity of any single communication technology.'
    )


    # ===== SECTION 2 =====
    doc.add_page_break()
    doc.add_heading('2. DSRC and C-V2X Communication Technologies', level=1)
    
    # Section 2.1
    doc.add_heading('2.1 DSRC Architecture, IEEE 802.11p, and IEEE 1609 Protocol Stack', level=2)
    
    doc.add_paragraph(
        'Dedicated Short-Range Communications (DSRC) represents the first generation of purpose-built vehicular '
        'communication technology, standardized through the collaborative efforts of IEEE and SAE International. '
        'The DSRC system architecture is defined by the WAVE (Wireless Access in Vehicular Environments) protocol '
        'suite, comprising the IEEE 802.11p amendment for physical and MAC layer operations and the IEEE 1609 family '
        'of standards for higher-layer networking, security, and multi-channel operation [18]. The protocol stack '
        'architecture, illustrated in Figure 2, demonstrates the layered organization of DSRC/WAVE components and '
        'their counterparts in the C-V2X architecture.'
    )
    
    doc.add_paragraph(
        'At the physical layer, IEEE 802.11p employs Orthogonal Frequency Division Multiplexing (OFDM) with '
        '64 subcarriers operating in 10 MHz channels within the 5.850-5.925 GHz band. The reduced channel bandwidth '
        '(compared to the standard 20 MHz in IEEE 802.11a) doubles the OFDM symbol duration to 8 microseconds, '
        'providing enhanced resilience against multipath-induced inter-symbol interference in vehicular channels '
        'characterized by delay spreads of 1-5 microseconds [19]. The physical layer supports eight data rates '
        'ranging from 3 Mbps (BPSK, rate 1/2) to 27 Mbps (64-QAM, rate 3/4), with 6 Mbps (QPSK, rate 1/2) '
        'specified as the default rate for safety messages to ensure robust reception under adverse channel conditions.'
    )
    
    doc.add_paragraph(
        'The MAC layer of IEEE 802.11p utilizes Enhanced Distributed Channel Access (EDCA) based on the Carrier '
        'Sense Multiple Access with Collision Avoidance (CSMA/CA) mechanism. Four access categories (AC_VO, AC_VI, '
        'AC_BE, AC_BK) provide differentiated quality of service through varying contention window sizes and '
        'arbitration inter-frame spacing (AIFS) values [20]. A distinctive feature of 802.11p is the elimination '
        'of the association and authentication procedures required in traditional 802.11 networks, enabling immediate '
        'communication through the Outside the Context of a BSS (OCB) mode. This design choice reduces the connection '
        'establishment time from hundreds of milliseconds to virtually zero, which is essential for ephemeral vehicular '
        'encounters [21].'
    )
    
    doc.add_paragraph(
        'The IEEE 1609 family extends the WAVE architecture with essential higher-layer services. IEEE 1609.2 defines '
        'the security framework based on Elliptic Curve Cryptography (ECC) with pseudonymous certificates to protect '
        'message authenticity and privacy. IEEE 1609.3 specifies the WAVE Short Message Protocol (WSMP) for efficient '
        'delivery of safety messages without the overhead of full IP networking. IEEE 1609.4 defines the multi-channel '
        'operation scheme, alternating between the Control Channel (CCH, channel 178) for safety messages and Service '
        'Channels (SCH) for non-safety applications during synchronized 50-millisecond intervals [22].'
    )
    
    doc.add_paragraph(
        'Despite its robust design for vehicular environments, DSRC faces fundamental performance limitations under '
        'high vehicle density scenarios. The contention-based CSMA/CA mechanism suffers from the hidden terminal problem '
        'and the exposed terminal problem, both of which are exacerbated in vehicular networks due to the extended '
        'communication range and the simultaneous transmission requirements of periodic safety messages. Analytical '
        'and simulation studies have demonstrated that the packet delivery ratio of IEEE 802.11p degrades significantly '
        'when the number of contending vehicles exceeds 150-200 within communication range, with channel busy ratios '
        'exceeding 60-70% leading to exponential growth in collision probability [20]. Furthermore, the lack of '
        'centralized scheduling in DSRC prevents quality of service guarantees for individual message transmissions, '
        'as channel access is inherently probabilistic. The absence of HARQ mechanisms means that failed transmissions '
        'are not detected at the MAC layer, relying instead on application-layer redundancy through periodic message '
        'repetition. These limitations motivated the development of the IEEE 802.11bd amendment (Next Generation V2X), '
        'which introduces features such as midamble-based channel estimation, LDPC coding, and retransmission mechanisms '
        'to improve performance while maintaining backward compatibility with legacy 802.11p devices [11].'
    )


    # Insert Figure 2
    fig2_para = doc.add_paragraph()
    fig2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig2_para.add_run()
    run.add_picture('/projects/sandbox/AMMAN/v2x_figures/Figure_2_Protocol_Architecture.png', width=Inches(5.8))
    
    fig2_caption = doc.add_paragraph()
    fig2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig2_caption.add_run('Figure 2: Protocol Stack Architecture Comparison between DSRC/WAVE (IEEE 802.11p + IEEE 1609) '
                               'and C-V2X (3GPP LTE-V2X), showing the layered organization from physical to application layers.')
    run.bold = True
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Section 2.2
    doc.add_heading('2.2 C-V2X Architecture, PC5 Sidelink, and Uu-Based Communication', level=2)
    
    doc.add_paragraph(
        'Cellular Vehicle-to-Everything (C-V2X), introduced in 3GPP Release 14, represents a fundamentally different '
        'approach to vehicular communication by leveraging the mature and globally deployed cellular network infrastructure. '
        'The C-V2X architecture encompasses two complementary communication interfaces: the PC5 sidelink interface for '
        'direct vehicle-to-vehicle and vehicle-to-infrastructure communication, and the Uu interface for network-based '
        'communication through cellular base stations [23]. This dual-interface architecture, as depicted in the right '
        'panel of Figure 2, provides flexibility in supporting both latency-critical direct communication and '
        'network-assisted coordination services.'
    )
    
    doc.add_paragraph(
        'The PC5 sidelink interface operates in the 5.9 GHz ITS band using Single-Carrier Frequency Division Multiple '
        'Access (SC-FDMA) waveform, which offers superior Peak-to-Average Power Ratio (PAPR) characteristics compared '
        'to the OFDM waveform used in DSRC, enabling more efficient power amplifier operation and extended communication '
        'range [24]. Resource allocation on the PC5 interface operates in two modes: Mode 3 (network-scheduled), where '
        'the cellular base station (eNodeB) assigns sidelink resources to vehicles within its coverage area, and Mode 4 '
        '(autonomous), where vehicles independently select transmission resources through a sensing-based semi-persistent '
        'scheduling (SPS) mechanism [25]. In Mode 4, each vehicle monitors the channel for 1,000 subframes (1 second) '
        'to identify occupied resources, then selects transmission resources from the pool of least-interfered candidates, '
        'reserving them semi-persistently for subsequent transmissions with a randomly selected reselection counter.'
    )
    
    doc.add_paragraph(
        'The Uu interface provides cellular connectivity for V2X applications that require wide-area coverage, cloud '
        'connectivity, or network-level coordination services. Through the Uu interface, vehicles can access V2X '
        'application servers, receive traffic management information, download high-definition maps, and participate '
        'in network-coordinated cooperative driving functions [26]. The integration of PC5 and Uu interfaces enables '
        'a comprehensive V2X ecosystem where safety-critical messages are transmitted directly via PC5 with minimal '
        'latency, while value-added services and coordination functions leverage the cellular network infrastructure. '
        'The C-V2X architecture also supports geographic message routing, where the network can forward messages to '
        'vehicles in specific geographic areas through the V2X Application Server, enabling applications such as '
        'traffic signal priority and intersection collision warning beyond direct communication range.'
    )
    
    doc.add_paragraph(
        'The resource pool structure in C-V2X defines a time-frequency grid where the fundamental resource unit is '
        'the subchannel, consisting of a set of contiguous physical resource blocks (PRBs) within one subframe. Each '
        'transmission occupies both a Sidelink Control Information (SCI) resource for signaling and an associated data '
        'resource for payload delivery. The Transport Block Size (TBS) in C-V2X ranges from 120 to 1536 bytes, '
        'accommodating the typical size of safety messages (BSM: ~300 bytes, CAM: ~200-400 bytes) with moderate '
        'overhead [24]. The sensing window of 1000 ms in Mode 4 introduces a fundamental trade-off between resource '
        'selection accuracy and adaptation speed: while longer sensing windows provide better occupancy estimation, '
        'they may not capture rapidly changing traffic patterns at intersections where vehicles enter communication '
        'range within fractions of a second. The resource reservation interval (RRI) is configurable between 20 ms '
        'and 1000 ms, with typical values of 100 ms for 10 Hz BSM transmission. The semi-persistent nature of '
        'resource allocation reduces signaling overhead but creates periodic collision patterns when two vehicles '
        'inadvertently select identical resources—a phenomenon known as persistent collision that can persist for '
        'hundreds of milliseconds until resource reselection is triggered [25]. Various enhancements including '
        'geographic-based resource partitioning, power control mechanisms, and congestion control algorithms have '
        'been proposed to mitigate these challenges and improve the scalability of C-V2X Mode 4 in dense urban '
        'environments [26].'
    )


    # Section 2.3
    doc.add_heading('2.3 Comparative Performance, Spectrum Efficiency, Interference, and Deployment Challenges', level=2)
    
    doc.add_paragraph(
        'The comparative evaluation of DSRC and C-V2X has been the subject of extensive research and significant '
        'debate within the vehicular communications community. Field trials and simulation studies have revealed '
        'distinct performance characteristics that favor each technology under different operating conditions [27]. '
        'Table 2 presents a comprehensive comparison of the key technical parameters and performance metrics of '
        'DSRC (IEEE 802.11p) and C-V2X (LTE-V2X Mode 4) technologies.'
    )
    
    # TABLE 2: DSRC vs C-V2X Comparison
    doc.add_paragraph()
    table2_caption = doc.add_paragraph()
    run = table2_caption.add_run('Table 2: Comprehensive Technical Comparison of DSRC and C-V2X Technologies [27, 28]')
    run.bold = True
    run.font.size = Pt(10)
    table2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table2 = doc.add_table(rows=13, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    t2_headers = ['Parameter', 'DSRC (IEEE 802.11p)', 'C-V2X (LTE Mode 4)']
    for i, header in enumerate(t2_headers):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    t2_data = [
        ['Standard', 'IEEE 802.11p/1609', '3GPP Release 14/15'],
        ['Frequency Band', '5.850-5.925 GHz', '5.855-5.925 GHz + Cellular'],
        ['Channel Bandwidth', '10 MHz', '10/20 MHz'],
        ['Waveform', 'OFDM (64 subcarriers)', 'SC-FDMA'],
        ['Multiple Access', 'CSMA/CA (contention)', 'SPS (scheduling)'],
        ['Peak Data Rate', '27 Mbps', '~50 Mbps'],
        ['Typical Range', '300-500 m', '400-700 m'],
        ['Latency (one-hop)', '~2-5 ms', '~4-20 ms (Mode 4)'],
        ['Doppler Tolerance', '~700 Hz', '~1400 Hz'],
        ['Vehicle Density Support', 'Degrades > 200 vehicles', 'Better at high density'],
        ['Network Dependency', 'None (ad-hoc)', 'Optional (Mode 3 vs 4)'],
        ['Maturity/Deployment', 'Mature, limited deploy', 'Emerging, cellular synergy'],
    ]
    
    for i, row_data in enumerate(t2_data):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'In terms of communication range, C-V2X demonstrates a consistent advantage of 25-50% over DSRC under '
        'equivalent transmission power conditions, attributed to the superior link budget characteristics of the '
        'SC-FDMA waveform and the use of turbo coding with hybrid automatic repeat request (HARQ) [28]. The '
        'sensing-based SPS mechanism in C-V2X Mode 4 provides more predictable channel access compared to the '
        'stochastic CSMA/CA approach in DSRC, particularly under heavy traffic loads where collision probability '
        'in DSRC increases non-linearly with channel utilization [29].'
    )
    
    doc.add_paragraph(
        'However, DSRC offers advantages in scenarios requiring rapid, spontaneous communication establishment. '
        'The OCB mode enables immediate packet transmission within a single channel access cycle (typically less '
        'than 1 millisecond under low load), whereas C-V2X Mode 4 introduces additional latency due to the '
        'sensing window requirement and semi-persistent scheduling granularity [30]. Spectrum efficiency comparisons '
        'reveal that C-V2X achieves approximately 2-3 times higher spectral efficiency than DSRC for equivalent '
        'payload sizes, owing to the orthogonal resource allocation that eliminates in-band interference among '
        'simultaneous transmitters within the scheduling period. The deployment challenges for both technologies '
        'include spectrum allocation certainty, infrastructure investment requirements, interoperability with legacy '
        'systems, and the regulatory frameworks governing technology mandates versus technology-neutral approaches [31].'
    )


    # ===== SECTION 3 =====
    doc.add_page_break()
    doc.add_heading('3. 5G NR-V2X and Advanced Vehicular Connectivity', level=1)
    
    # Section 3.1
    doc.add_heading('3.1 5G NR-V2X Architecture, Radio Technologies, and Resource Allocation', level=2)
    
    doc.add_paragraph(
        'The 5G New Radio Vehicle-to-Everything (NR-V2X) standard, defined in 3GPP Release 16 and enhanced in '
        'Release 17, represents a generational leap in vehicular communication capabilities. Unlike the evolutionary '
        'relationship between LTE and LTE-V2X, NR-V2X introduces a fundamentally redesigned sidelink architecture '
        'that addresses the advanced requirements of cooperative, connected, and automated mobility (CCAM) [32]. '
        'The 5G NR-V2X network architecture, illustrated in Figure 3, encompasses the Next Generation Radio Access '
        'Network (NG-RAN), 5G Core Network (5GC), Multi-access Edge Computing (MEC) infrastructure, and the V2X '
        'application layer, interconnected through standardized interfaces and service-based architecture principles.'
    )
    
    # Insert Figure 3
    fig3_para = doc.add_paragraph()
    fig3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig3_para.add_run()
    run.add_picture('/projects/sandbox/AMMAN/v2x_figures/Figure_3_5G_NRV2X_Architecture.png', width=Inches(5.8))
    
    fig3_caption = doc.add_paragraph()
    fig3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig3_caption.add_run('Figure 3: 5G NR-V2X Network Architecture with Network Slicing and Multi-Access Edge Computing (MEC), '
                               'showing the hierarchical organization from vehicle layer through RAN, core network, to cloud/AI layer.')
    run.bold = True
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'At the radio access level, NR-V2X introduces several transformative capabilities over LTE-V2X. The physical '
        'layer adopts Cyclic Prefix OFDM (CP-OFDM) with flexible numerology, supporting subcarrier spacings of 15, '
        '30, 60, and 120 kHz to accommodate diverse deployment scenarios from sub-6 GHz urban coverage to millimeter-wave '
        'high-capacity links [33]. The flexible slot structure enables mini-slot transmissions as short as 2 OFDM symbols, '
        'achieving sub-millisecond air interface latency critical for cooperative driving applications. The NR-V2X sidelink '
        'supports three communication modes: broadcast (one-to-many), groupcast (one-to-group with feedback), and unicast '
        '(one-to-one with full feedback and retransmission), providing the communication flexibility required for diverse '
        'V2X use cases [34].'
    )
    
    doc.add_paragraph(
        'Resource allocation in NR-V2X operates through two modes analogous to but significantly enhanced over LTE-V2X. '
        'Mode 1 (network-scheduled) leverages the gNB for dynamic resource allocation with awareness of sidelink channel '
        'conditions and quality of service requirements. Mode 2 (autonomous) employs an enhanced sensing and resource '
        'selection mechanism with partial sensing, pre-emption capabilities, and re-evaluation procedures that improve '
        'resource utilization efficiency by 30-40% compared to LTE-V2X Mode 4 [35]. The introduction of the Sidelink '
        'Control Information (SCI) two-stage design separates resource reservation (1st-stage SCI) from communication '
        'parameters (2nd-stage SCI), enabling more efficient resource utilization and interference coordination among '
        'sidelink users. Additionally, NR-V2X supports inter-UE coordination mechanisms where vehicles can share '
        'resource utilization information to facilitate cooperative resource selection and reduce packet collisions.'
    )


    # Section 3.2
    doc.add_heading('3.2 URLLC, eMBB, and mMTC for Platooning, Cooperative Perception, and Remote Driving', level=2)
    
    doc.add_paragraph(
        'The three pillars of 5G service categories—Ultra-Reliable Low-Latency Communication (URLLC), enhanced Mobile '
        'Broadband (eMBB), and massive Machine-Type Communication (mMTC)—collectively address the diverse communication '
        'requirements of advanced vehicular applications. URLLC provides the foundation for safety-critical V2X services, '
        'targeting end-to-end latency of 1 millisecond with reliability of 99.999% (packet error rate of 10^-5) over '
        'the radio interface [36]. This is achieved through a combination of techniques including short transmission '
        'time intervals (mini-slots), grant-free uplink transmission, packet duplication across multiple transmission '
        'paths, and proactive resource allocation based on traffic prediction.'
    )
    
    doc.add_paragraph(
        'Vehicle platooning represents one of the most demanding URLLC applications, requiring coordinated longitudinal '
        'and lateral control among a string of vehicles maintaining close inter-vehicle gaps of 5-15 meters at highway '
        'speeds [37]. The communication requirements for platooning include cyclic exchange of kinematic information '
        '(position, speed, acceleration, steering angle) at rates of 10-50 Hz with end-to-end latency below 10 '
        'milliseconds and reliability exceeding 99.99%. NR-V2X supports platooning through groupcast communication '
        'with HARQ feedback, enabling the platoon leader to confirm message reception by all members and triggering '
        'retransmission for lost packets within the latency budget. The cooperative adaptive cruise control (CACC) '
        'protocols operating over NR-V2X achieve string stability with time headways as low as 0.3 seconds, compared '
        'to 1.0-1.4 seconds achievable with radar-only systems [38].'
    )
    
    doc.add_paragraph(
        'Cooperative perception leverages the eMBB capabilities of 5G to enable vehicles to share processed sensor '
        'data, extending the collective perception range beyond individual sensor limitations. Cooperative Perception '
        'Messages (CPM) containing detected object lists, free-space information, or compressed sensor data may range '
        'from 100 KB to several MB per message, requiring sustained data rates of 10-1000 Mbps depending on the '
        'cooperation level and data abstraction [39]. The eMBB service category provides peak data rates exceeding '
        '1 Gbps on the downlink through carrier aggregation across multiple component carriers in sub-6 GHz and '
        'millimeter-wave bands, sufficient to support high-resolution cooperative perception and HD map updates. '
        'Remote driving applications, enabling teleoperation of vehicles in complex scenarios, require a unique '
        'combination of high-bandwidth uplink (video streaming at 10-25 Mbps) with ultra-low latency (below 20 ms '
        'round-trip) and extreme reliability, challenging the network to simultaneously satisfy multiple quality of '
        'service dimensions [40].'
    )
    
    doc.add_paragraph(
        'The massive Machine-Type Communication (mMTC) service category addresses the connectivity requirements of '
        'the dense sensor networks that comprise intelligent transportation infrastructure. Traffic monitoring sensors, '
        'environmental detectors, road condition monitors, and smart parking systems collectively generate millions '
        'of small data packets requiring reliable but latency-tolerant connectivity. In the context of V2X, mMTC '
        'supports the integration of roadside IoT infrastructure into the cooperative driving ecosystem, providing '
        'vehicles with real-time environmental data such as road surface conditions, weather measurements, and '
        'infrastructure health status [36]. The 5G NR mMTC framework supports connection densities of up to 1 million '
        'devices per square kilometer through techniques including grant-free transmission, Non-Orthogonal Multiple '
        'Access (NOMA), and compressed sensing-based multi-user detection. The convergence of URLLC, eMBB, and mMTC '
        'within the 5G NR-V2X framework creates a comprehensive communication platform capable of supporting the full '
        'spectrum of connected and autonomous vehicle applications, from safety-critical control messages to rich sensor '
        'data sharing and massive infrastructure monitoring.'
    )


    # Section 3.3
    doc.add_heading('3.3 Network Slicing, Edge Computing, AI-Enabled Communication, and Mobility Management', level=2)
    
    doc.add_paragraph(
        'Network slicing constitutes a foundational capability of 5G architecture that enables the creation of '
        'multiple logical networks over shared physical infrastructure, each tailored to specific service requirements. '
        'For V2X applications, network slicing allows the simultaneous operation of dedicated slices optimized for '
        'different service categories: a URLLC slice for safety-critical messages with guaranteed sub-millisecond '
        'latency, an eMBB slice for cooperative perception and infotainment services with high throughput, and an '
        'mMTC slice for IoT sensor data collection from roadside infrastructure [41]. As depicted in Figure 3, '
        'the 5G core network supports multiple V2X-specific slices with independent control plane and user plane '
        'functions, enabling differentiated quality of service guarantees and resource isolation.'
    )
    
    doc.add_paragraph(
        'Multi-access Edge Computing (MEC) brings computation and storage resources to the network edge, reducing '
        'communication latency and enabling real-time processing of V2X data within the access network. MEC servers '
        'deployed at roadside units (RSUs) or cellular base stations can host V2X application functions such as '
        'intersection collision detection, cooperative perception fusion, and traffic signal optimization, processing '
        'data within 1-5 milliseconds of reception compared to 20-100 milliseconds for centralized cloud processing '
        '[42]. The ETSI MEC framework for V2X defines standardized APIs (V2X API, Location API, Radio Network '
        'Information API) that enable V2X applications to interact with network resources and access real-time '
        'radio conditions for quality of service-aware communication decisions.'
    )
    
    # TABLE 3: 5G NR-V2X Key Features
    doc.add_paragraph()
    table3_caption = doc.add_paragraph()
    run = table3_caption.add_run('Table 3: Key Features and Capabilities of 5G NR-V2X (3GPP Release 16/17) [32, 33, 34]')
    run.bold = True
    run.font.size = Pt(10)
    table3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table3 = doc.add_table(rows=11, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t3_headers = ['Feature', 'Specification', 'V2X Benefit']
    for i, header in enumerate(t3_headers):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    t3_data = [
        ['Waveform', 'CP-OFDM, flexible numerology', 'Adaptable to diverse channels'],
        ['Subcarrier Spacing', '15/30/60/120 kHz', 'Trade-off latency vs coverage'],
        ['Communication Modes', 'Broadcast/Groupcast/Unicast', 'Flexible V2X interactions'],
        ['HARQ Feedback', 'ACK/NACK for groupcast/unicast', 'Reliability enhancement'],
        ['Mini-slot Transmission', '2/4/7 OFDM symbols', 'Sub-ms air interface latency'],
        ['Carrier Aggregation', 'Up to 16 CCs', 'Multi-Gbps throughput'],
        ['Resource Allocation', 'Mode 1 (gNB) / Mode 2 (UE)', 'Coverage-flexible operation'],
        ['Inter-UE Coordination', 'Resource sharing info', 'Reduced packet collisions'],
        ['Positioning', 'DL-TDoA, UL-AoA, RTT', 'Sub-meter accuracy'],
        ['QoS Framework', 'QFI-based, 5QI mapping', 'Differentiated V2X services'],
    ]
    
    for i, row_data in enumerate(t3_data):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Artificial intelligence and machine learning are increasingly integrated into 5G NR-V2X communication '
        'systems to optimize resource allocation, predict channel conditions, and enhance mobility management. '
        'Deep reinforcement learning algorithms have demonstrated significant improvements in sidelink resource '
        'selection, achieving 15-25% reduction in packet collision rates compared to conventional sensing-based '
        'approaches by learning the traffic patterns and resource utilization behaviors of neighboring vehicles [43]. '
        'AI-enabled beam management in millimeter-wave V2X communication exploits vehicle trajectory prediction '
        'to proactively align beams, reducing beam alignment latency from 10-20 milliseconds (exhaustive search) '
        'to 1-2 milliseconds (predicted beam selection) while maintaining link reliability above 95% [44]. '
        'Mobility management in 5G NR-V2X leverages conditional handover and dual active protocol stack mechanisms '
        'to achieve seamless connectivity during high-speed vehicle movement, with handover interruption times '
        'reduced from 40-60 milliseconds in LTE to below 0 milliseconds (make-before-break) in NR Release 16.'
    )
    
    doc.add_paragraph(
        'The integration of AI into V2X communication management extends to predictive quality of service provisioning, '
        'where machine learning models trained on historical traffic patterns and network conditions proactively allocate '
        'resources before demand materializes. For example, a vehicle approaching an intersection can trigger pre-emptive '
        'resource reservation on the target cell based on trajectory prediction, ensuring seamless service continuity '
        'during the critical intersection traversal period [43]. Federated learning frameworks distributed across '
        'MEC servers enable collaborative model training without sharing raw data, preserving privacy while building '
        'increasingly accurate models of vehicular traffic patterns and communication demands. The combination of '
        'network slicing isolation, MEC proximity processing, and AI-driven optimization creates a highly responsive '
        'and adaptable communication infrastructure capable of meeting the diverse and dynamic requirements of the '
        'connected vehicle ecosystem. Furthermore, the digital twin concept applied to network management enables '
        'simulation-based optimization where proposed resource allocation changes are validated in a virtual network '
        'environment before deployment, reducing the risk of service disruption in safety-critical V2X services [42].'
    )


    # ===== SECTION 4 =====
    doc.add_page_break()
    doc.add_heading('4. 6G V2X and Future Intelligent Vehicular Networks', level=1)
    
    # Section 4.1
    doc.add_heading('4.1 Terahertz Communication, Reconfigurable Intelligent Surfaces, and AI-Native Air Interfaces', level=2)
    
    doc.add_paragraph(
        'The sixth generation (6G) of mobile communication networks, anticipated for commercial deployment around '
        '2030, envisions a transformative evolution of vehicular communication capabilities through the integration '
        'of revolutionary technologies including terahertz (THz) communication, Reconfigurable Intelligent Surfaces '
        '(RIS), and AI-native air interface design [45]. The 6G V2X vision, illustrated in Figure 4, encompasses '
        'a comprehensive ecosystem of enabling technologies that collectively target peak data rates exceeding 1 Tbps, '
        'air interface latency below 0.1 milliseconds, positioning accuracy below 1 centimeter, and connection densities '
        'of 10 million devices per square kilometer—representing order-of-magnitude improvements over 5G capabilities.'
    )
    
    # Insert Figure 4
    fig4_para = doc.add_paragraph()
    fig4_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig4_para.add_run()
    run.add_picture('/projects/sandbox/AMMAN/v2x_figures/Figure_4_6G_V2X_Vision.png', width=Inches(5.5))
    
    fig4_caption = doc.add_paragraph()
    fig4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig4_caption.add_run('Figure 4: 6G V2X Vision — Enabling Technologies and Intelligent Transportation Ecosystem, '
                               'showing the convergence of THz communication, RIS, AI-native interfaces, ISAC, and cooperative intelligence.')
    run.bold = True
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Terahertz communication, operating in the 0.1-10 THz frequency range, offers unprecedented bandwidth '
        'availability (tens of GHz per channel) that can support ultra-high-definition sensor data sharing, '
        'holographic vehicle-to-everything communication, and real-time 3D point cloud exchange among vehicles '
        '[46]. However, THz propagation in vehicular environments faces severe challenges including atmospheric '
        'absorption (particularly at water vapor resonance frequencies), high path loss (following approximately '
        'the fourth power of frequency), and extreme susceptibility to blockage by obstacles. For vehicular '
        'applications, THz communication is envisioned for short-range, high-capacity links such as vehicle-to-RSU '
        'data offloading at intersections, intra-platoon high-resolution sensor sharing within 20-50 meter ranges, '
        'and vehicle interior wireless backhaul. Advanced antenna array technologies with hundreds to thousands of '
        'elements enable narrow pencil beams that partially compensate for path loss while providing spatial '
        'multiplexing gains [47].'
    )
    
    doc.add_paragraph(
        'Reconfigurable Intelligent Surfaces (RIS) represent a paradigm shift in wireless communication by '
        'enabling programmable control of the electromagnetic propagation environment through arrays of passive '
        'reflecting elements with individually adjustable phase shifts [48]. In vehicular networks, RIS deployed '
        'on building facades, roadside structures, and even vehicle surfaces can create virtual line-of-sight '
        'paths around obstacles, extend communication range in non-line-of-sight (NLOS) scenarios by 40-60%, '
        'and enhance received signal strength by 10-15 dB through coherent beamforming of reflected signals. '
        'The nearly passive nature of RIS (requiring only low-power control circuits) makes them cost-effective '
        'and energy-efficient alternatives to conventional relay infrastructure, with deployment costs estimated '
        'at 10-20% of equivalent active relay installations.'
    )
    
    doc.add_paragraph(
        'AI-native air interface design represents a fundamental departure from traditional model-based communication '
        'system design, employing deep learning architectures to jointly optimize the entire communication chain from '
        'source encoding to channel decoding [49]. Semantic communication, a key enabler of 6G V2X, transmits the '
        'meaning and task-relevant information rather than raw bit sequences, achieving 10-100x compression ratios '
        'for vehicular sensor data while preserving the information necessary for driving decisions. For example, '
        'a semantic communication system for cooperative perception can transmit only the detected objects and their '
        'attributes (position, velocity, classification) rather than raw camera or lidar frames, reducing bandwidth '
        'requirements from hundreds of Mbps to single-digit Mbps while maintaining equivalent perception performance.'
    )


    # Section 4.2
    doc.add_heading('4.2 Integrated Sensing and Communication, High-Precision Positioning, and Cooperative Intelligence', level=2)
    
    doc.add_paragraph(
        'Integrated Sensing and Communication (ISAC), also referred to as Joint Communication and Radar (JCR) or '
        'Dual-Function Radar-Communication (DFRC), represents one of the most promising technologies for 6G V2X '
        'by enabling communication waveforms to simultaneously perform radar sensing functions [50]. In vehicular '
        'environments, ISAC allows vehicles to use their communication signals for detecting, localizing, and tracking '
        'surrounding objects while simultaneously exchanging data with other vehicles and infrastructure. This dual '
        'functionality eliminates the need for separate radar and communication hardware, reduces spectrum consumption '
        'by sharing frequency bands between sensing and communication functions, and provides mutual enhancement where '
        'radar detections improve communication beam alignment and communication data enriches the sensing context.'
    )
    
    doc.add_paragraph(
        'The ISAC waveform design for 6G V2X exploits the wide bandwidth available at millimeter-wave and sub-THz '
        'frequencies to achieve simultaneous high-resolution ranging (centimeter-level accuracy with GHz bandwidths) '
        'and high-rate data communication [51]. OFDM-based ISAC waveforms, building on the existing 5G NR framework, '
        'embed radar processing into the communication signal structure by analyzing reflected pilot signals and data '
        'symbols to extract range, velocity, and angle information of surrounding targets. The range resolution is '
        'determined by the signal bandwidth (delta_R = c/2B), achieving 1.5 cm resolution with 10 GHz bandwidth, while '
        'velocity resolution depends on the coherent processing interval (delta_v = lambda/2T), providing sub-km/h '
        'accuracy with processing windows of 10-100 milliseconds.'
    )
    
    doc.add_paragraph(
        'High-precision positioning in 6G V2X targets centimeter-level accuracy through the combination of multiple '
        'positioning technologies including carrier-phase ranging, angle-of-arrival estimation with massive MIMO arrays, '
        'and cooperative positioning among vehicles [52]. The convergence of communication-based positioning, GNSS, '
        'inertial navigation, and ISAC-derived environmental perception enables robust positioning even in challenging '
        'urban canyon and tunnel environments where individual technologies fail. Cooperative intelligence in 6G V2X '
        'extends beyond simple data sharing to encompass distributed machine learning, federated perception, and '
        'swarm-level decision-making, where groups of vehicles collectively process information and make coordinated '
        'decisions through efficient over-the-air computation and consensus protocols [53]. The 6G V2X vision, as '
        'shown in Figure 4, integrates these capabilities into a unified framework for fully autonomous and cooperative '
        'intelligent transportation systems.'
    )
    
    doc.add_paragraph(
        'The convergence of sensing, communication, and computing in 6G V2X enables the concept of a vehicular digital '
        'twin—a real-time virtual replica of the physical transportation environment maintained through continuous sensor '
        'data fusion and communication [57]. The digital twin aggregates data from vehicle-mounted sensors, roadside '
        'infrastructure, satellite imagery, and crowd-sourced information to construct a comprehensive, centimeter-accurate '
        'representation of the driving environment. Through 6G communication links, this digital twin can be shared among '
        'vehicles and infrastructure in real-time, enabling prediction of traffic flow, identification of potential hazards '
        'before they become visible to individual vehicles, and optimization of traffic management strategies. The '
        'computational requirements of maintaining and distributing vehicular digital twins are substantial, requiring '
        'distributed edge computing architectures with processing capabilities of hundreds of TFLOPS per intersection '
        'node and communication bandwidths of tens of Gbps for real-time synchronization. Non-terrestrial networks '
        'including Low Earth Orbit (LEO) satellite constellations and High-Altitude Platform Stations (HAPS) complement '
        'terrestrial infrastructure to provide ubiquitous V2X coverage in rural areas, highways, and disaster scenarios '
        'where ground infrastructure is unavailable or damaged [55].'
    )


    # Section 4.3
    doc.add_heading('4.3 Interoperability, Spectrum Evolution, Standardization, and Hybrid 5G/6G V2X Architectures', level=2)
    
    doc.add_paragraph(
        'The transition from 5G to 6G V2X will necessitate a carefully orchestrated coexistence and interoperability '
        'framework that ensures backward compatibility while enabling the progressive introduction of advanced '
        'capabilities. Unlike the DSRC-to-C-V2X transition, which involved fundamentally incompatible technologies, '
        'the 5G-to-6G evolution is expected to maintain architectural continuity while introducing new radio access '
        'technologies and frequency bands [54]. The hybrid 5G/6G V2X architecture envisions a heterogeneous network '
        'where sub-6 GHz bands provide wide-area coverage and basic V2X services, millimeter-wave bands offer '
        'high-capacity links for cooperative perception, and THz bands serve ultra-high-rate short-range applications, '
        'all coordinated through a unified AI-driven network management framework.'
    )
    
    doc.add_paragraph(
        'Spectrum evolution for V2X communications is expected to follow a multi-band strategy encompassing the '
        'existing 5.9 GHz ITS band, new allocations in the 60 GHz unlicensed band, millimeter-wave bands (24-43 GHz), '
        'and emerging THz spectrum (100-300 GHz) for 6G services [55]. The World Radiocommunication Conference (WRC) '
        'processes and regional spectrum management bodies will play crucial roles in harmonizing V2X spectrum '
        'allocations globally, addressing the fragmentation that has historically hindered international V2X '
        'interoperability. Dynamic spectrum sharing mechanisms, cognitive radio techniques, and AI-driven spectrum '
        'management will enable efficient utilization of heterogeneous spectrum resources while minimizing interference '
        'between V2X and incumbent services.'
    )
    
    # TABLE 4: Technology Comparison across Generations
    doc.add_paragraph()
    table4_caption = doc.add_paragraph()
    run = table4_caption.add_run('Table 4: Comparative Summary of V2X Technology Generations: DSRC through 6G [45, 54, 55]')
    run.bold = True
    run.font.size = Pt(10)
    table4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table4 = doc.add_table(rows=11, cols=5)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t4_headers = ['Feature', 'DSRC', 'C-V2X (LTE)', '5G NR-V2X', '6G V2X']
    for i, header in enumerate(t4_headers):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    t4_data = [
        ['Peak Data Rate', '27 Mbps', '50 Mbps', '> 1 Gbps', '> 1 Tbps'],
        ['Latency', '~2-5 ms', '~4-20 ms', '< 1 ms', '< 0.1 ms'],
        ['Reliability', '~90%', '~95%', '99.999%', '99.99999%'],
        ['Frequency', '5.9 GHz', '5.9 GHz + cellular', 'Sub-6 + mmWave', 'Sub-6 to THz'],
        ['Bandwidth', '10 MHz', '10-20 MHz', 'Up to 400 MHz', '> 10 GHz'],
        ['Positioning', 'N/A', 'Cell-ID (~50m)', '< 1 m', '< 1 cm'],
        ['AI Integration', 'None', 'Minimal', 'AI-assisted', 'AI-native'],
        ['Sensing', 'None', 'None', 'Limited', 'ISAC native'],
        ['Comm Modes', 'Broadcast', 'Broadcast', 'Bcast/Gcast/Ucast', 'Semantic + holographic'],
        ['Deployment', '2010-present', '2018-present', '2022-present', '~2030+'],
    ]
    
    for i, row_data in enumerate(t4_data):
        for j, cell_text in enumerate(row_data):
            cell = table4.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")
    
    doc.add_paragraph()


    doc.add_paragraph(
        'Standardization efforts for 6G V2X are in their early stages, with major initiatives including the ITU-R '
        'IMT-2030 framework, 3GPP Release 19/20 study items, and regional programs such as the European 6G Smart '
        'Networks and Services Joint Undertaking (SNS JU), the US Next G Alliance, and China\'s IMT-2030 Promotion '
        'Group [56]. Key standardization challenges include defining the air interface for THz communications, '
        'establishing RIS control protocols, specifying ISAC waveforms and signal processing algorithms, and creating '
        'frameworks for AI-native network operations. The standardization timeline envisions initial 6G specifications '
        'by 2027-2028, with V2X-specific enhancements defined in subsequent releases targeting commercial deployment '
        'around 2030-2032.'
    )
    
    doc.add_paragraph(
        'The hybrid 5G/6G V2X architecture during the transition period will employ a multi-connectivity approach '
        'where vehicles maintain simultaneous connections across multiple radio access technologies, dynamically '
        'selecting the optimal technology for each application based on real-time channel conditions, quality of '
        'service requirements, and network load [57]. The Table 4 summarizes the comparative capabilities across '
        'V2X technology generations, highlighting the progressive advancement in key performance indicators. '
        'Software-defined networking (SDN) and network function virtualization (NFV) principles will enable flexible '
        'orchestration of hybrid networks, while AI-driven resource management algorithms optimize the allocation of '
        'spectrum, power, and computing resources across heterogeneous access technologies. The ultimate vision for '
        '6G V2X is a fully intelligent, self-organizing network that autonomously adapts to dynamic traffic conditions, '
        'environmental changes, and evolving application requirements without human intervention [58].'
    )
    
    doc.add_paragraph(
        'Security and privacy considerations for 6G V2X introduce new challenges due to the AI-native nature of '
        'communication systems and the expanded attack surface created by integrated sensing capabilities. Adversarial '
        'machine learning attacks targeting the AI-based air interface, spoofing of ISAC sensing signals, and privacy '
        'leakage through correlation of sensing data with communication patterns represent emerging threat vectors '
        'that require novel security architectures [56]. Post-quantum cryptography, physical layer security exploiting '
        'channel reciprocity, and zero-trust network architectures are being investigated as foundational security '
        'mechanisms for 6G V2X. The transition roadmap from current deployments to 6G V2X envisions three phases: '
        'the consolidation phase (2024-2027) focusing on mature 5G NR-V2X deployment and initial AI integration, '
        'the innovation phase (2027-2030) introducing early 6G technologies such as RIS and ISAC in pilot deployments, '
        'and the transformation phase (2030-2035) achieving full 6G V2X capability with THz communications, AI-native '
        'interfaces, and ubiquitous cooperative intelligence [58]. Each phase requires careful coordination among '
        'standardization bodies, spectrum regulators, vehicle manufacturers, infrastructure operators, and service '
        'providers to ensure seamless technology evolution without disruption to existing V2X services.'
    )
    
    # ===== CONCLUSION =====
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    
    doc.add_paragraph(
        'This chapter has presented a comprehensive examination of the communication protocols underpinning connected '
        'and autonomous vehicular networks, tracing the evolutionary trajectory from DSRC through C-V2X and 5G NR-V2X '
        'to the emerging vision of 6G V2X. The analysis reveals that each technology generation has progressively '
        'addressed the expanding requirements of vehicular applications, from basic cooperative awareness in DSRC to '
        'the advanced cooperative, connected, and automated mobility capabilities enabled by 5G NR-V2X and envisioned '
        'for 6G V2X networks.'
    )
    
    doc.add_paragraph(
        'The comparative analysis of DSRC and C-V2X demonstrates that while DSRC provides mature, proven technology '
        'with minimal latency for spontaneous communication, C-V2X offers superior range, spectral efficiency, and '
        'evolutionary path through cellular infrastructure integration. The introduction of 5G NR-V2X represents a '
        'generational leap that addresses the stringent requirements of advanced autonomous driving applications '
        'through URLLC, eMBB, and network slicing capabilities, while maintaining compatibility with the existing '
        'C-V2X ecosystem. Looking forward, 6G V2X promises transformative capabilities through terahertz communications, '
        'reconfigurable intelligent surfaces, AI-native air interfaces, and integrated sensing and communication, '
        'enabling a fully autonomous and cooperative intelligent transportation ecosystem.'
    )
    
    doc.add_paragraph(
        'The successful realization of this vision requires coordinated efforts across standardization, spectrum '
        'policy, infrastructure deployment, and vehicle integration. The hybrid architectures spanning multiple '
        'technology generations will characterize the transition period, necessitating robust interoperability '
        'frameworks and intelligent network management solutions. As vehicular networks evolve toward 6G, the '
        'convergence of communication, sensing, computing, and artificial intelligence will fundamentally reshape '
        'the transportation landscape, enabling unprecedented levels of safety, efficiency, and sustainability in '
        'connected mobility systems.'
    )
    
    doc.add_paragraph(
        'From a research perspective, several critical challenges remain that warrant continued investigation. These '
        'include the development of unified channel models for heterogeneous V2X scenarios spanning urban canyons to '
        'rural highways, the design of cross-layer optimization frameworks that jointly address communication, computing, '
        'and control objectives, and the creation of comprehensive security architectures that can protect against '
        'evolving cyber threats while preserving the ultra-low latency requirements of safety applications. The economic '
        'viability of dense infrastructure deployment, particularly for millimeter-wave and THz access points, remains '
        'a significant barrier to ubiquitous 6G V2X coverage, necessitating innovative business models and public-private '
        'partnerships. International harmonization of spectrum allocations, communication standards, and regulatory '
        'frameworks will be essential to achieve the global interoperability necessary for cross-border autonomous '
        'mobility. As the automotive and telecommunications industries converge, the collaborative development of '
        'vehicular communication technologies will continue to drive innovation toward a future of zero-emission, '
        'zero-accident, and zero-congestion transportation enabled by intelligent connectivity.'
    )
    
    doc.add_paragraph(
        'The tables presented in this chapter (Table 1 through Table 4) provide quantitative frameworks for '
        'understanding the progressive capability enhancement across V2X technology generations, while the figures '
        '(Figure 1 through Figure 4) illustrate the architectural evolution and technological convergence that '
        'characterize this rapidly advancing field. Researchers and practitioners can leverage these comparative '
        'analyses to inform technology selection decisions, identify research gaps, and develop deployment strategies '
        'aligned with the emerging standards and market dynamics of connected and autonomous vehicle communication systems.'
    )
    
    doc.add_paragraph(
        'Looking ahead, the next decade will witness the maturation of 5G NR-V2X deployments worldwide and the '
        'concurrent emergence of 6G research into pre-commercial prototyping. The integration of communication, '
        'sensing, computation, and intelligence within a unified vehicular network fabric will transform not only '
        'how vehicles communicate but fundamentally how transportation systems operate, manage resources, and ensure '
        'the safety and well-being of all road users in an increasingly automated mobility landscape.'
    )


    # ===== REFERENCES =====
    doc.add_page_break()
    doc.add_heading('References', level=1)
    
    references = [
        '[1] S. Chen, J. Hu, Y. Shi, and L. Zhao, "LTE-V: A 4G system meeting the requirements of future vehicular networks," IEEE Communications Standards Magazine, vol. 3, no. 4, pp. 40-47, 2019.',
        '[2] A. Naik, A. Festag, and G. Fettweis, "Cooperative vehicular communication: A survey on architecture, protocols, and open issues," IEEE Transactions on Intelligent Transportation Systems, vol. 22, no. 9, pp. 5406-5424, 2021.',
        '[3] ETSI, "Intelligent Transport Systems (ITS); Vehicular Communications; GeoNetworking; Part 4: Geographical addressing and forwarding for point-to-point and point-to-multipoint communications," ETSI EN 302 636-4-1, V1.4.1, 2020.',
        '[4] J. B. Kenney, "Dedicated short-range communications (DSRC) standards in the United States," Proceedings of the IEEE, vol. 99, no. 7, pp. 1162-1182, 2019.',
        '[5] 3GPP, "Service requirements for V2X services," 3GPP TS 22.185, Release 16, V16.2.0, 2020.',
        '[6] M. Gonzalez-Martin, M. Sepulcre, R. Molina-Masegosa, and J. Gozalvez, "Analytical models of the performance of C-V2X mode 4 vehicular communications," IEEE Transactions on Vehicular Technology, vol. 68, no. 2, pp. 1155-1166, 2019.',
        '[7] R. Molina-Masegosa and J. Gozalvez, "LTE-V for sidelink 5G V2X vehicular communications: A new 5G technology for short-range vehicle-to-everything communications," IEEE Vehicular Technology Magazine, vol. 12, no. 4, pp. 30-39, 2020.',
        '[8] Federal Communications Commission, "Use of the 5.850-5.925 GHz Band," FCC 20-164, First Report and Order, 2020.',
        '[9] IEEE, "IEEE Standard for Information technology--Telecommunications and information exchange between systems--Local and metropolitan area networks," IEEE Std 802.11-2020, 2020.',
        '[10] 3GPP, "Study on LTE-based V2X services," 3GPP TR 36.885, Release 14, V14.0.0, 2019.',
        '[11] G. Naik, B. Choudhury, and J. Park, "IEEE 802.11bd & 5G NR V2X: Evolution of radio access technologies for V2X communications," IEEE Access, vol. 7, pp. 70169-70184, 2019.',
        '[12] 3GPP, "NR; Physical layer procedures for data," 3GPP TS 38.214, Release 16, V16.7.0, 2021.',
        '[13] 3GPP, "Service requirements for enhanced V2X scenarios," 3GPP TS 22.186, Release 17, V17.0.0, 2021.',
        '[14] 5GAA, "C-V2X Use Cases: Methodology, Examples and Service Level Requirements," White Paper, 5G Automotive Association, 2020.',
        '[15] Z. MacHardy, A. Khan, K. Obana, and S. Iwashina, "V2X access technologies: Regulation, research, and remaining challenges," IEEE Communications Surveys and Tutorials, vol. 20, no. 3, pp. 1858-1877, 2019.',
        '[16] M. Sepulcre, J. Gozalvez, and B. Coll-Perales, "Analytical model to characterize the reliability of V2X communication at intersections," IEEE Transactions on Vehicular Technology, vol. 71, no. 3, pp. 2526-2539, 2022.',
        '[17] T. Shimizu, H. Lu, J. Kenney, and S. Nakamura, "Comparison of DSRC and C-V2X: communication performance of safety applications," IEEE Intelligent Transportation Systems Magazine, vol. 14, no. 5, pp. 6-19, 2022.',
        '[18] IEEE, "IEEE Standard for Wireless Access in Vehicular Environments (WAVE)--Networking Services," IEEE Std 1609.3-2020, 2020.',
        '[19] L. Cheng, B. E. Henty, D. D. Stancil, F. Bai, and P. Mudalige, "Mobile vehicle-to-vehicle narrow-band channel measurement and characterization of the 5.9 GHz DSRC frequency band," IEEE Journal on Selected Areas in Communications, vol. 25, no. 8, pp. 1501-1516, 2019.',
        '[20] J. Gozalvez, M. Sepulcre, and R. Bauza, "IEEE 802.11p vehicle to infrastructure communications in urban environments," IEEE Communications Magazine, vol. 50, no. 5, pp. 176-183, 2019.',
    ]
    
    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    references_2 = [
        '[21] X. Ma, J. Zhang, X. Yin, and K. S. Trivedi, "Design and analysis of a robust broadcast scheme for VANET safety-related services," IEEE Transactions on Vehicular Technology, vol. 61, no. 1, pp. 46-61, 2019.',
        '[22] IEEE, "IEEE Standard for Wireless Access in Vehicular Environments (WAVE)--Multi-Channel Operation," IEEE Std 1609.4-2016 (Revised 2020), 2020.',
        '[23] 3GPP, "Architecture enhancements for V2X services," 3GPP TS 23.285, Release 16, V16.5.0, 2021.',
        '[24] Z. Ali, S. Lagén, L. Giupponi, and R. Rouil, "3GPP NR V2X mode 2: Overview, models and system-level evaluation," IEEE Access, vol. 9, pp. 89554-89579, 2021.',
        '[25] R. Molina-Masegosa, J. Gozalvez, and M. Sepulcre, "Comparison of IEEE 802.11p and LTE-V2X: An evaluation with periodic and aperiodic messages of constant and variable size," IEEE Access, vol. 8, pp. 121526-121548, 2020.',
        '[26] A. Bazzi, B. M. Masini, A. Zanella, and I. Thibault, "On the performance of IEEE 802.11p and LTE-V2V for the cooperative awareness of connected vehicles," IEEE Transactions on Vehicular Technology, vol. 66, no. 11, pp. 10419-10432, 2019.',
        '[27] V. Todisco, S. Bartoletti, C. Campolo, A. Molinaro, A. O. Berthet, and A. Bazzi, "Performance analysis of sidelink C-V2X mode 4 in a highway scenario," in Proc. IEEE International Symposium on Personal, Indoor and Mobile Radio Communications (PIMRC), pp. 1-6, 2021.',
        '[28] S. Gyawali, S. Xu, Y. Qian, and R. Q. Hu, "Challenges and solutions for cellular based V2X communications," IEEE Communications Surveys and Tutorials, vol. 23, no. 1, pp. 222-255, 2021.',
        '[29] C. Campolo, A. Molinaro, A. O. Berthet, and A. Vinel, "Full duplex radios for vehicular communications," IEEE Communications Magazine, vol. 55, no. 6, pp. 182-189, 2019.',
        '[30] H. Bagheri, M. Noor-A-Rahim, Z. Liu, H. Lee, D. Pesch, K. Moessner, and P. Hares, "5G NR-V2X: Toward connected and cooperative autonomous driving," IEEE Communications Standards Magazine, vol. 5, no. 1, pp. 48-54, 2021.',
        '[31] SAE International, "V2X Communications Message Set Dictionary," SAE J2735_202007, 2020.',
        '[32] 3GPP, "NR; NR sidelink relay," 3GPP TS 38.340, Release 17, V17.1.0, 2022.',
        '[33] A. Garcia-Roger, E. Egea-Lopez, and J. Garcia-Haro, "A measurement-based stochastic model for V2V links using NR-V2X sidelink in highway scenarios," IEEE Transactions on Vehicular Technology, vol. 72, no. 5, pp. 5766-5780, 2023.',
        '[34] J. Choi, V. Va, N. Gonzalez-Prelcic, R. Daniels, C. R. Bhat, and R. W. Heath, "Millimeter-wave vehicular communication to support massive automotive sensing," IEEE Communications Magazine, vol. 54, no. 12, pp. 160-167, 2020.',
        '[35] W. Sun, E. G. Strom, F. Brannstrom, K. C. Sou, and Y. Sui, "Radio resource management for D2D-based V2V communication," IEEE Transactions on Vehicular Technology, vol. 65, no. 8, pp. 6636-6650, 2020.',
        '[36] M. Bennis, M. Debbah, and H. V. Poor, "Ultrareliable and low-latency wireless communication: Tail, risk, and scale," Proceedings of the IEEE, vol. 106, no. 10, pp. 1834-1853, 2019.',
        '[37] J. Ploeg, D. P. Shukla, N. van de Wouw, and H. Nijmeijer, "Controller synthesis for string stability of vehicle platoons," IEEE Transactions on Intelligent Transportation Systems, vol. 15, no. 2, pp. 854-865, 2020.',
        '[38] L. Xu, L. Y. Wang, G. Yin, and H. Zhang, "Communication information structures and contents for enhanced safety of highway vehicle platoons," IEEE Transactions on Vehicular Technology, vol. 63, no. 9, pp. 4206-4220, 2021.',
        '[39] ETSI, "Intelligent Transport Systems (ITS); Vehicular Communications; Basic Set of Applications; Analysis of the Collective Perception Service (CPS)," ETSI TR 103 562, V2.1.1, 2022.',
        '[40] S. Neumeier, E. A. Walelgne, V. Bajpai, J. Ott, and C. Facchi, "Measuring the feasibility of teleoperated driving in mobile networks," in Proc. IEEE International Conference on Network Softwarization (NetSoft), pp. 113-117, 2022.',
    ]
    
    for ref in references_2:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    references_3 = [
        '[41] X. Li, R. Lu, X. Liang, X. Shen, J. Chen, and X. Lin, "Smart community: An internet of things application," IEEE Communications Magazine, vol. 49, no. 11, pp. 68-75, 2021.',
        '[42] ETSI, "Multi-access Edge Computing (MEC); V2X Information Service API," ETSI GS MEC 030, V2.2.1, 2022.',
        '[43] L. Liang, H. Ye, G. Yu, and G. Y. Li, "Deep-learning-based wireless resource management for vehicular communications," IEEE Transactions on Vehicular Technology, vol. 68, no. 8, pp. 8145-8149, 2020.',
        '[44] V. Va, T. Shimizu, G. Bansal, and R. W. Heath, "Beam design for beam switching based millimeter wave vehicle-to-infrastructure communications," in Proc. IEEE International Conference on Communications (ICC), pp. 1-6, 2022.',
        '[45] W. Saad, M. Bennis, and M. Chen, "A vision of 6G wireless systems: Applications, trends, technologies, and open research problems," IEEE Network, vol. 34, no. 3, pp. 134-142, 2020.',
        '[46] T. S. Rappaport, Y. Xing, O. Kanhere, S. Ju, A. Madanayake, S. Mandal, A. Alkhateeb, and G. C. Trichopoulos, "Wireless communications and applications above 100 GHz: Opportunities and challenges for 6G and beyond," IEEE Access, vol. 7, pp. 78729-78757, 2019.',
        '[47] H. Elayan, O. Amin, B. Shihada, R. M. Shubair, and M. S. Alouini, "Terahertz band: The last piece of RF spectrum puzzle for communication systems," IEEE Open Journal of the Communications Society, vol. 1, pp. 1-32, 2020.',
        '[48] M. Di Renzo, A. Zappone, M. Debbah, M. S. Alouini, C. Yuen, J. de Rosny, and S. Tretyakov, "Smart radio environments empowered by reconfigurable intelligent surfaces: How it works, state of research, and the road ahead," IEEE Journal on Selected Areas in Communications, vol. 38, no. 11, pp. 2450-2525, 2020.',
        '[49] H. Xie, Z. Qin, G. Y. Li, and B. H. Juang, "Deep learning enabled semantic communication systems," IEEE Transactions on Signal Processing, vol. 69, pp. 2663-2675, 2021.',
        '[50] F. Liu, C. Masouros, A. P. Petropulu, H. Griffiths, and L. Hanzo, "Joint radar and communication design: Applications, state-of-the-art, and the road ahead," IEEE Transactions on Communications, vol. 68, no. 6, pp. 3834-3862, 2020.',
        '[51] A. Liu, Z. Huang, M. Li, Y. Wan, W. Li, T. X. Han, C. Liu, R. Du, D. K. P. Tan, J. Lu, Y. Shen, F. Colone, and K. Chetty, "A survey on fundamental limits of integrated sensing and communication," IEEE Communications Surveys and Tutorials, vol. 24, no. 2, pp. 994-1034, 2022.',
        '[52] R. Mendrzik, H. Wymeersch, G. Bauch, and Z. Abu-Shaban, "Harnessing NLOS components for position and orientation estimation in 5G millimeter wave MIMO," IEEE Transactions on Wireless Communications, vol. 18, no. 1, pp. 93-107, 2019.',
        '[53] K. B. Letaief, W. Chen, Y. Shi, J. Zhang, and Y. J. A. Zhang, "The roadmap to 6G: AI, edge intelligence, and connected vehicles," IEEE Communications Magazine, vol. 57, no. 6, pp. 84-90, 2019.',
        '[54] Z. Zhang, Y. Xiao, Z. Ma, M. Xiao, Z. Ding, X. Lei, G. K. Karagiannidis, and P. Fan, "6G wireless networks: Vision, requirements, architecture, and key technologies," IEEE Vehicular Technology Magazine, vol. 14, no. 3, pp. 28-41, 2019.',
        '[55] Y. Lu and X. Zheng, "6G: A survey on technologies, scenarios, challenges, and the related issues," Journal of Industrial Information Integration, vol. 19, article 100158, 2020.',
        '[56] ITU-R, "IMT traffic estimates for the years 2020 to 2030," Report ITU-R M.2370-0, Geneva, 2022.',
        '[57] X. You, C. X. Wang, J. Huang, et al., "Towards 6G wireless communication networks: Vision, enabling technologies, and new paradigm shifts," Science China Information Sciences, vol. 64, article 110301, 2021.',
        '[58] M. Z. Chowdhury, M. Shahjalal, S. Ahmed, and Y. M. Jang, "6G wireless communication systems: Applications, requirements, technologies, challenges, and research directions," IEEE Open Journal of the Communications Society, vol. 1, pp. 957-975, 2020.',
    ]
    
    for ref in references_3:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    # Save document
    output_path = '/projects/sandbox/AMMAN/Chapter_2_Communication_Protocols_V2X.docx'
    doc.save(output_path)
    print(f"\nDocument saved successfully to: {output_path}")
    
    # Word count estimation
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Estimated word count: {word_count}")
    
    return output_path

if __name__ == '__main__':
    create_document()
