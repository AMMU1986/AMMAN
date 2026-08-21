"""
Generate DOCX for Chapter: Advancing Education 5.0 through AI Readiness and Acceptance
- 43 references in sequential numbered [1]-[43] format
- No citations in abstract
- 4 figures and 4 tables, each cited exactly 2 times in body text
- Target ~8300 words
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def create_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # TITLE
    title = doc.add_heading('', level=0)
    run = title.add_run('Advancing Education 5.0 through AI Readiness and Acceptance: Insights from the Academic Community')
    run.font.size = Pt(16)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ABSTRACT (No citations)
    doc.add_heading('Abstract', level=1)
    abstract = (
        "The emergence of Education 5.0 represents a paradigm shift from technology-centered digital transformation "
        "toward human-centered, intelligent, and socially responsible higher education. Artificial intelligence has "
        "become central to this transformation through its capacity to personalize learning, automate administrative "
        "processes, support research, and improve institutional decision-making. However, successful AI integration "
        "requires more than technological infrastructure; it demands comprehensive institutional readiness encompassing "
        "human competencies, organizational capacity, ethical governance, cultural openness, and stakeholder acceptance. "
        "This chapter presents a data-driven AI readiness assessment framework organized into four interconnected "
        "layers: inputs, readiness dimensions, acceptance indicators, and transformation outcomes. Drawing on insights "
        "from faculty and student perspectives, the chapter examines how perceived usefulness, ease of use, trust, "
        "self-efficacy, and institutional support shape AI acceptance within academic communities. Strategic pathways "
        "for building AI-ready universities are proposed, including continuous assessment cycles, human-centered "
        "governance, workforce development, and institutional benchmarking. The chapter concludes with future "
        "directions emphasizing predictive readiness models, cross-cultural comparative studies, and longitudinal "
        "research connecting AI readiness to sustained educational transformation outcomes."
    )
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Cm(1.27)

    kw = doc.add_paragraph()
    r = kw.add_run('Keywords: ')
    r.bold = True
    kw.add_run('Education 5.0, artificial intelligence readiness, technology acceptance, higher education, '
               'AI literacy, institutional transformation, human-centered education, digital governance')

    doc.add_page_break()

    # ============================
    # SECTION 1
    # ============================
    doc.add_heading('1. Foundations of AI Readiness in Higher Education 5.0', level=1)

    doc.add_heading('1.1 Education 5.0 and the Emergence of AI-Enabled Universities', level=2)

    s1_1 = [
        "Higher Education 5.0 represents a fundamental transition from technology-centered digital transformation toward human-centered, intelligent, adaptive, and socially responsible education [1]. While earlier stages of educational development emphasized digitization, online learning, automation, and data-driven management, Education 5.0 places greater emphasis on collaboration between humans and intelligent technologies. This paradigm recognizes that technology should serve human flourishing rather than merely optimize institutional efficiency, demanding educational systems that develop creativity, critical thinking, ethical reasoning, and adaptive expertise alongside technical competencies [2].",

        "Artificial intelligence has become one of the most influential technologies supporting this transformation because of its ability to analyze large datasets, personalize learning experiences, automate administrative processes, support research discovery, and improve institutional decision-making [3]. Generative AI systems including large language models have further accelerated this transformation by offering capabilities in content creation, tutoring, assessment support, and research assistance that were previously unimaginable. However, the rapid proliferation of AI tools has created urgent questions about institutional preparedness, appropriate governance, academic integrity, and the fundamental purpose of higher education in an age of artificial intelligence [4].",

        "AI readiness refers to the capacity of a higher education institution and its academic community to adopt, implement, govern, and continuously improve AI-enabled systems [5]. Readiness extends far beyond technological infrastructure to encompass faculty competencies, student awareness, organizational culture, leadership commitment, data governance, ethical preparedness, financial resources, and acceptance of AI-based technologies. An institution may possess advanced computing infrastructure but remain poorly prepared for AI adoption if its faculty lack appropriate competencies or if users distrust algorithmic systems. Conversely, an institution with modest infrastructure but strong human capital and supportive culture may achieve more effective AI integration than a technologically advanced but culturally resistant counterpart.",

        "Education 5.0 therefore requires an integrated perspective in which technology and people develop together [6]. AI should augment rather than replace educators, researchers, administrators, and students. In this context, AI readiness becomes a strategic institutional capability rather than a purely technical characteristic. Institutions that develop comprehensive AI readiness are better positioned to harness AI's benefits while managing its risks, ensuring that technological adoption serves educational mission rather than merely following technological trends. The strategic nature of AI readiness means that it requires deliberate cultivation through sustained investment, leadership attention, and organizational learning rather than emerging spontaneously from technology procurement alone.",

        "The urgency of developing institutional AI readiness has intensified dramatically since the public release of generative AI systems in late 2022 [4]. Students and faculty began using AI tools regardless of institutional policies or preparedness, creating de facto AI integration that outpaced governance development. Institutions that had invested in readiness foundations—including AI literacy programs, ethical frameworks, and flexible governance structures—navigated this disruption more successfully than those caught without preparation. This experience demonstrates that AI readiness is not merely aspirational but practically necessary for institutions seeking to maintain academic quality, integrity, and relevance in rapidly evolving technological environments. Figure 1 presents the multi-layered AI readiness assessment framework that structures this comprehensive approach to institutional preparedness.",
    ]

    for text in s1_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('1.2 Dimensions of Artificial Intelligence Readiness', level=2)

    s1_2 = [
        "A comprehensive AI readiness assessment should examine multiple dimensions simultaneously, recognizing that institutional preparedness is multifaceted and that strength in one dimension cannot compensate for weakness in another [7]. The first dimension is technological readiness, encompassing network infrastructure, cloud computing services, computing capacity, learning management systems, data platforms, cybersecurity protocols, and system interoperability. Technological readiness provides the foundation upon which AI applications operate, but it represents only one component of comprehensive institutional preparedness.",

        "The second dimension is human readiness [8]. Faculty members and academic staff require knowledge of AI concepts, data literacy, prompt engineering capabilities, AI-assisted teaching methods, and critical evaluation of AI-generated outputs. Students similarly require AI literacy to use intelligent tools responsibly and effectively. Human readiness assessment examines not only current competency levels but also learning capacity, motivation for professional development, and attitudes toward AI-mediated change. Table 1 presents a detailed breakdown of AI readiness dimensions with their associated indicators, assessment methods, and benchmarking criteria.",

        "The third dimension is organizational readiness, which concerns institutional leadership, strategic planning, policy frameworks, funding mechanisms, technical support infrastructure, and change-management capacity [9]. Strong leadership is particularly important because AI adoption can alter teaching practices, assessment methods, research workflows, and administrative responsibilities in ways that challenge established roles and power structures. Organizations with clear AI strategies, dedicated budgets, and change-management expertise navigate these transitions more successfully than those that approach AI adoption in an ad hoc manner.",

        "The fourth dimension is ethical and governance readiness [10]. Higher education institutions must address privacy protection, algorithmic bias, transparency requirements, accountability mechanisms, intellectual property considerations, academic integrity policies, and responsible data use protocols. These dimensions are particularly critical when AI systems process student records, research data, examination information, or personal information. Institutions lacking robust ethical frameworks risk both legal liability and erosion of stakeholder trust that can undermine AI adoption even when technical implementation succeeds.",

        "The fifth dimension is cultural readiness, reflecting institutional openness toward innovation, experimentation, and constructive failure [11]. A culture that encourages professional development, values curiosity, tolerates responsible experimentation, and rewards innovation is substantially more likely to support successful AI transformation than one characterized by risk aversion, rigid hierarchies, and resistance to change. Figure 2 illustrates these five interconnected dimensions and their relationship to overall institutional AI readiness, demonstrating that comprehensive preparedness requires balanced development across all dimensions rather than excellence in any single area.",

        "As presented in Table 1, each readiness dimension can be assessed through specific indicators using quantitative and qualitative methods. The multidimensional nature of AI readiness means that institutions may exhibit varying maturity levels across dimensions—for example, possessing advanced technological infrastructure (Level 4) while demonstrating limited ethical governance (Level 2). Such imbalances identify specific intervention priorities and prevent the common error of assuming that technological investment alone constitutes adequate preparation for AI-enabled transformation.",
    ]

    for text in s1_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 1
    doc.add_paragraph()
    t1_title = doc.add_paragraph()
    t1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1_title.add_run('Table 1. ')
    r.bold = True
    t1_title.add_run('AI Readiness Dimensions, Indicators, and Assessment Methods for Higher Education Institutions')

    table1 = doc.add_table(rows=6, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    h1 = ['Readiness Dimension', 'Key Indicators', 'Assessment Methods', 'Benchmark Criteria']
    for i, h in enumerate(h1):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t1_data = [
        ['Technological', 'Network bandwidth, cloud services, LMS capability, cybersecurity, data platforms', 'Infrastructure audit, system performance metrics, interoperability testing', 'ISO 27001, EDUCAUSE benchmarks, uptime >99.5%'],
        ['Human', 'AI literacy scores, data competencies, prompt engineering, critical AI evaluation', 'Competency assessments, self-report surveys, performance tasks', 'UNESCO AI competency framework, DIGCOMP 2.2'],
        ['Organizational', 'AI strategy existence, budget allocation, leadership commitment, support structures', 'Document analysis, leadership interviews, resource mapping', 'Dedicated AI budget >2%, named AI leadership'],
        ['Ethical & Governance', 'Privacy policies, bias auditing, transparency protocols, integrity guidelines', 'Policy review, compliance audit, stakeholder consultation', 'GDPR compliance, UNESCO AI ethics, institutional IRB'],
        ['Cultural', 'Innovation climate, experimentation tolerance, collaboration norms, change receptivity', 'Climate surveys, innovation indices, focus groups', 'Innovation culture score >3.5/5, training participation >60%'],
    ]
    for i, row in enumerate(t1_data):
        for j, cell in enumerate(row):
            table1.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    doc.add_heading('1.3 The Role of Academic Community Acceptance', level=2)

    s1_3 = [
        "Technology adoption depends not only on institutional readiness but also on individual acceptance by the academic community members who must ultimately use, interact with, and integrate AI systems into their professional practice [12]. Faculty members may perceive AI as an opportunity for improving teaching effectiveness and research productivity, while others may express concerns about job displacement, reduced academic autonomy, misinformation propagation, plagiarism facilitation, or excessive dependence on automated systems. These individual-level perceptions significantly influence whether institutional AI investments translate into actual transformation or remain underutilized despite adequate infrastructure.",

        "Students may similarly demonstrate different levels of acceptance depending on their AI literacy, previous technology experience, perceived usefulness for academic tasks, ease of use, and concerns about fairness, privacy, and academic assessment implications [13]. A student who understands how AI generates outputs and recognizes its limitations may engage productively with AI tools, while a student who lacks this understanding may either uncritically accept AI outputs or refuse engagement entirely. Both responses represent suboptimal outcomes that targeted educational interventions can address.",

        "Academic community acceptance can be examined through established technology acceptance constructs including perceived usefulness, perceived ease of use, trust, perceived risks, self-efficacy, social influence, facilitating conditions, behavioral intention, and actual use [14]. These variables—drawn from the Technology Acceptance Model (TAM), Unified Theory of Acceptance and Use of Technology (UTAUT), and trust-based frameworks—provide validated instruments for assessing whether AI systems are likely to be adopted effectively after implementation. Figure 3 presents the academic community AI acceptance model that integrates these constructs with education-specific factors.",

        "Consequently, an AI readiness framework should combine objective institutional indicators with subjective perceptions of academic stakeholders [15]. This integration ensures that readiness assessments capture both the organizational capacity to deploy AI systems and the human willingness to engage with them productively. Institutions that invest heavily in infrastructure while neglecting acceptance-building activities frequently discover that sophisticated AI tools remain unused because the academic community was neither consulted during design nor prepared for adoption. The gap between institutional capability and stakeholder acceptance represents one of the most common failure modes in educational technology implementation, and AI readiness frameworks that bridge this gap provide significantly more accurate predictions of transformation success than either capability or acceptance assessment alone.",

        "The temporal dynamics of acceptance deserve attention. Initial acceptance measured at the point of AI introduction may differ substantially from sustained engagement measured months or years later. Early enthusiasm may fade as novelty effects diminish and implementation challenges accumulate. Conversely, initial skepticism may gradually transform into acceptance as evidence of AI benefits accumulates and users develop competence through experience. Readiness frameworks should therefore incorporate longitudinal assessment capabilities that track acceptance trajectories rather than relying on single-point-in-time measurements that capture only a snapshot of evolving attitudes.",
    ]

    for text in s1_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # FIGURE 1
    fig1_path = 'education5_figures/Figure_1_AI_Readiness_Framework.png'
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1 = doc.add_paragraph()
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap1.add_run('Figure 1. ')
    r.bold = True
    cap1.add_run('Multi-layered AI Readiness Assessment Framework for Higher Education 5.0, showing the four interconnected layers from institutional data inputs through readiness dimensions and acceptance indicators to transformation outcomes.')

    # FIGURE 2
    doc.add_paragraph()
    fig2_path = 'education5_figures/Figure_2_Dimensions_Readiness.png'
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap2.add_run('Figure 2. ')
    r.bold = True
    cap2.add_run('Five interconnected dimensions of AI Readiness in Higher Education: Technological, Human, Organizational, Ethical & Governance, and Cultural readiness, with their key components and mutual relationships.')

    doc.add_page_break()

    # ============================
    # SECTION 2
    # ============================
    doc.add_heading('2. Data-Driven AI Readiness Assessment Framework', level=1)

    doc.add_heading('2.1 Conceptual Structure of the Assessment Framework', level=2)

    s2_1 = [
        "A data-driven AI readiness framework for Higher Education 5.0 can be organized into four interconnected layers: inputs, readiness dimensions, acceptance indicators, and transformation outcomes [16]. This layered architecture ensures systematic progression from data collection through analysis to actionable insight, while maintaining clear connections between institutional capabilities, stakeholder behaviors, and educational outcomes. As previously illustrated in Figure 1, each layer feeds into subsequent layers while receiving feedback from outcome assessment that informs continuous improvement. The framework's data-driven nature distinguishes it from purely normative readiness models by grounding assessment in empirical evidence rather than aspirational targets alone.",

        "The input layer consists of institutional and stakeholder data that collectively represent the raw material from which readiness assessment is constructed [17]. Institutional data may include infrastructure availability and performance metrics, AI-related capital and operational investments, training program participation rates and completion statistics, digital platform adoption levels and usage intensity, research capacity indicators including AI-related publications and grants, and governance mechanism documentation including policies, committees, and audit records. Stakeholder data may include faculty and student surveys measuring attitudes and perceptions, interviews providing contextual depth, focus groups surfacing collective concerns, usage analytics tracking actual AI engagement behaviors, competency assessments measuring AI literacy levels, and longitudinal perception tracking identifying trend directions.",

        "The readiness layer evaluates technological, human, organizational, ethical, governance, and cultural capabilities through the dimensional framework described previously [18]. Each dimension is represented through measurable indicators that can be normalized to common scales and aggregated into composite scores using weighted combinations. The weighting structure should reflect institutional priorities and context—a research-intensive university may weight human research competencies more heavily, while a teaching-focused institution may emphasize pedagogical AI literacy. The acceptance layer measures whether academic stakeholders are willing and prepared to use AI technologies, connecting institutional capability with actual human behavior through validated acceptance constructs drawn from technology adoption research.",

        "The final transformation layer evaluates outcomes including improved teaching quality as measured by student learning gains, personalized learning effectiveness assessed through adaptive pathway completion rates, research productivity gains quantified through publications and grants, administrative efficiency improvements tracked through process cycle times, student engagement enhancement measured through participation analytics, institutional innovation capacity reflected in new program development, and competitive positioning assessed through rankings and reputation metrics [18]. This outcome layer provides the ultimate validation of whether readiness investments translate into genuine educational improvement.",

        "A simplified readiness score can be expressed as AIR = Sum(wi multiplied by Ri), where AIR represents the overall AI readiness index, Ri represents the normalized score of readiness dimension i, and wi represents the weight assigned to that dimension [19]. The weights can be established through expert judgment using Delphi methods, analytic hierarchy processes that structure pairwise comparisons, statistical methods such as principal component analysis that identify empirical importance, or machine-learning-based feature importance extraction that determines which readiness dimensions most strongly predict transformation outcomes. Weight determination should reflect institutional priorities—for example, an institution prioritizing ethical AI may assign greater weight to governance readiness, while one emphasizing immediate capability development may weight human readiness more heavily. Regular weight recalibration ensures that the readiness index remains aligned with evolving institutional strategy.",
    ]

    for text in s2_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('2.2 Data Collection and Analytical Methods', level=2)

    s2_2 = [
        "The framework should combine quantitative and qualitative evidence to capture both measurable indicators and contextual understanding [20]. Surveys can measure attitudes, acceptance levels, AI literacy scores, perceived usefulness, trust in AI systems, and specific concerns. Institutional records provide objective information regarding infrastructure capacity, training completion rates, AI tool adoption, budget allocation, and research output metrics. Qualitative methods including interviews, focus groups, and observational studies provide depth and context that quantitative indicators alone cannot capture.",

        "Likert-scale questionnaires can assess stakeholder perceptions across acceptance dimensions [21]. For example, faculty may be asked to rate statements concerning confidence in using generative AI for teaching preparation, adequacy of institutional support for AI adoption, severity of ethical concerns regarding AI-generated content, and expectations regarding AI's long-term impact on academic practice. Students can similarly evaluate AI tool accessibility, perceived usefulness for learning, reliability of AI-generated information, and concerns about AI's effects on assessment fairness. Table 2 presents the key constructs, measurement items, and psychometric properties for the AI acceptance assessment instrument.",

        "Statistical techniques provide rigorous analytical approaches for identifying relationships among readiness and acceptance variables [22]. Descriptive statistics characterize the current state; correlation analysis identifies associations between variables; exploratory factor analysis discovers underlying constructs; confirmatory factor analysis validates theoretical measurement models; regression analysis identifies predictors of acceptance; and structural equation modelling tests complex relationships among multiple variables simultaneously. These techniques collectively enable evidence-based conclusions about which readiness factors most strongly influence acceptance and which acceptance factors most reliably predict actual AI use.",

        "Machine-learning approaches complement traditional statistics by identifying non-linear patterns and enabling predictive classification [23]. Clustering algorithms such as k-means, hierarchical clustering, or DBSCAN can identify distinct groups within the academic community—for example, AI-ready enthusiasts, developing adopters, cautious observers, and AI-resistant skeptics. Each cluster may require different intervention strategies: enthusiasts benefit from advanced training and leadership roles; developing adopters need structured support and mentoring; cautious observers require evidence and gradual exposure; and resistant skeptics need addressed concerns and demonstrated value before productive engagement becomes likely.",

        "Predictive models including decision trees, random forests, and neural networks can estimate the probability that specific faculty members or student groups will adopt particular AI applications based on their competency profiles, attitudinal scores, institutional support levels, and previous technology experience [24]. These predictions enable targeted intervention design that allocates limited resources toward the populations and barriers where impact will be greatest, rather than implementing uniform approaches that may be optimal for no one. For instance, a model might identify that faculty in social sciences with moderate AI literacy but low institutional support represent a high-potential group where modest support investment would yield substantial adoption gains—insights that generic programs would miss entirely. Table 2 details the specific measurement constructs and their validated assessment approaches that underpin these analytical methods.",

        "The analytical methodology should also accommodate the dynamic nature of AI readiness. Cross-sectional surveys provide snapshots, but panel designs that track the same individuals over time reveal how attitudes, competencies, and behaviors evolve in response to interventions, experience, and changing institutional contexts. Time-series analysis of usage data can identify adoption trajectories—accelerating, plateauing, or declining engagement—enabling proactive responses to emerging patterns before they solidify into permanent behaviors. The combination of survey-based attitudinal data with behavioral usage data provides particularly powerful insights, revealing gaps between stated intentions and actual behaviors that identify where additional support or barrier removal is needed.",
    ]

    for text in s2_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 2
    doc.add_paragraph()
    t2_title = doc.add_paragraph()
    t2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2_title.add_run('Table 2. ')
    r.bold = True
    t2_title.add_run('AI Acceptance Assessment Constructs, Measurement Approaches, and Psychometric Targets')

    table2 = doc.add_table(rows=9, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    h2 = ['Construct', 'Definition', 'Measurement Approach', 'Target Reliability']
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t2_data = [
        ['Perceived Usefulness', 'Belief that AI improves professional performance', '4-5 Likert items adapted from TAM', "Cronbach's alpha > 0.80"],
        ['Perceived Ease of Use', 'Belief that AI use requires minimal effort', '4-5 Likert items adapted from TAM', "Cronbach's alpha > 0.80"],
        ['Trust in AI', 'Confidence in AI reliability, accuracy, fairness', '5-6 items from trust frameworks', "Cronbach's alpha > 0.85"],
        ['Self-Efficacy', 'Confidence in personal ability to use AI effectively', '4 items adapted from CSE scale', "Cronbach's alpha > 0.80"],
        ['Perceived Risk', 'Concerns about privacy, bias, job displacement', '5-6 items covering risk dimensions', "Cronbach's alpha > 0.75"],
        ['Social Influence', 'Peer and leadership encouragement of AI use', '3-4 items from UTAUT', "Cronbach's alpha > 0.75"],
        ['Behavioral Intention', 'Plan to use AI in professional practice', '3 items measuring future use intent', "Cronbach's alpha > 0.85"],
        ['Actual Use', 'Frequency and depth of current AI engagement', 'Usage logs, self-reported frequency', 'Test-retest r > 0.70'],
    ]
    for i, row in enumerate(t2_data):
        for j, cell in enumerate(row):
            table2.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    doc.add_heading('2.3 Readiness Levels and Institutional Benchmarking', level=2)

    s2_3 = [
        "The assessment framework can classify institutions into five readiness levels that represent progressive stages of AI integration maturity [25]. Level 1 (Initial) characterizes institutions with limited AI infrastructure, minimal awareness among stakeholders, absence of dedicated policies, and nascent competency development. At this level, AI engagement is typically limited to individual faculty experiments without institutional coordination or support. Level 2 (Emerging) describes institutions where initial AI experimentation occurs but institutional support remains fragmented, ad hoc, and dependent on individual champions rather than systematic organizational commitment. Training opportunities may exist but participation is low and disconnected from strategic objectives.",

        "Level 3 (Developing) represents institutions with established infrastructure, formal training programs, emerging governance mechanisms, and growing acceptance among early-majority adopters [26]. At this level, AI policies exist and are communicated, dedicated support staff assist faculty with AI integration, and student AI literacy receives curricular attention. Level 4 (Advanced) characterizes institutions where AI is meaningfully integrated into teaching, research, and administration with robust governance, widespread competent use, and systematic outcome measurement. Faculty routinely incorporate AI into pedagogical practice, assessment design reflects AI-augmented environments, and research productivity benefits from AI tools. Level 5 (Transformative) describes institutions where AI operates as part of a mature, human-centered Education 5.0 ecosystem with continuous adaptation, predictive governance, and systemic integration across all institutional functions. At this highest level, AI readiness is no longer a distinct initiative but is embedded within institutional culture and normal operations.",

        "Benchmarking allows institutions to compare their performance across dimensions rather than relying solely on a single overall score [27]. For example, an institution may demonstrate Level 4 technological readiness but Level 2 ethical readiness—an imbalance that identifies a specific area requiring urgent intervention before technological capabilities outpace governance capacity. Cross-institutional benchmarking additionally enables identification of best practices from peer institutions, sector-wide tracking of AI readiness development over time, and early warning of emerging challenges that leading institutions encounter before they affect the broader sector. Figure 4 presents the strategic pathway from initial readiness assessment through the transformation process to Education 5.0 outcomes, including the five readiness levels as a maturity continuum.",

        "The benchmarking process should be iterative rather than episodic [28]. Because AI technologies, stakeholder expectations, regulatory requirements, and institutional contexts change rapidly, a readiness assessment conducted once provides only a temporal snapshot that quickly becomes outdated. Continuous monitoring through institutional dashboards—incorporating real-time infrastructure metrics, periodic stakeholder surveys, training completion tracking, and AI usage analytics—enables responsive adaptation rather than reactive crisis management when readiness gaps emerge. Dashboard design should balance comprehensiveness with usability, presenting key indicators in accessible formats that support decision-making without overwhelming institutional leaders with data complexity.",

        "The readiness level classification also enables strategic resource allocation by identifying which investments will produce the greatest maturity advancement given current institutional position. A Level 1 institution benefits most from foundational infrastructure and awareness-building investments. A Level 3 institution benefits more from governance formalization and systematic faculty development. A Level 4 institution requires investments in advanced analytics, predictive systems, and continuous improvement mechanisms to progress toward transformative maturity. As illustrated in Figure 2, the interconnected nature of readiness dimensions means that advancement requires balanced development—an institution cannot achieve Level 5 transformative maturity through excellence in technology alone if ethical governance remains at Level 2. This differentiated investment logic prevents the common error of applying generic AI strategies regardless of institutional context and current capability.",
    ]

    for text in s2_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================
    # SECTION 3
    # ============================
    doc.add_heading('3. Insights from the Academic Community', level=1)

    doc.add_heading('3.1 Faculty Perspectives on AI Adoption', level=2)

    s3_1 = [
        "Faculty members represent a critical component of AI-enabled educational transformation because their acceptance directly affects the implementation of AI-supported teaching, assessment, research, and academic administration [29]. Unlike students who primarily engage with AI as consumers of educational services, faculty occupy dual roles as both users and designers of AI-enhanced learning experiences. Their decisions about whether and how to integrate AI into curricula, assessments, research methodologies, and supervisory practices collectively determine the institution's actual AI transformation regardless of strategic intentions declared at leadership levels.",

        "Faculty acceptance is generally influenced by perceived usefulness, ease of use, professional relevance, institutional support quality, training accessibility, and trust in AI system reliability [30]. AI can reduce repetitive administrative work, assist in developing learning materials, support individualized feedback generation, facilitate systematic literature analysis, enable data-driven pedagogical decisions, and support personalized instructional strategies. Faculty who perceive these benefits as genuinely useful for their specific disciplinary context are substantially more likely to invest the time required to develop AI competencies and integrate AI tools into established practices.",

        "However, legitimate concerns remain regarding inaccurate AI outputs, academic integrity challenges, authorship ambiguity, student overdependence on AI assistance, privacy risks when processing student data, and the potential erosion of professional judgment through algorithmic delegation [31]. Faculty in humanities and social sciences may additionally express concerns about AI's capacity to handle nuanced interpretation, cultural context, and ethical reasoning that characterize their disciplines. These disciplinary differences suggest that faculty development programs must be adapted to specific academic contexts rather than delivered through generic workshops that fail to address discipline-specific concerns and applications.",

        "Figure 3 presents the academic community AI acceptance model showing how external factors including institutional support and training interact with core acceptance constructs to influence behavioral intention and actual use. Effective faculty development should include AI fundamentals, responsible generative AI use, AI-assisted pedagogy, critical evaluation of AI outputs, assessment redesign for AI-augmented environments, data protection compliance, and discipline-specific application exploration [32]. Faculty should also participate actively in institutional AI policy development because participation increases ownership, trust, and compliance while ensuring that policies reflect practical classroom realities rather than abstract theoretical ideals. Development programs should be sustained rather than episodic, recognizing that AI capabilities evolve rapidly and that competencies developed today may require updating within months as new tools and approaches emerge.",

        "The format and delivery of faculty development programs significantly affects their uptake and impact. Programs that integrate into existing workloads—through embedded coaching during course development, just-in-time support when challenges arise, and peer mentoring relationships that provide ongoing guidance—achieve substantially higher engagement than standalone workshop events that faculty attend once and subsequently forget. Recognition and reward mechanisms—including consideration of AI-enhanced teaching in promotion criteria, teaching innovation awards, and reduced workload during intensive development periods—signal institutional commitment while reducing the personal cost of investing in new capabilities [33].",

        "Table 3 summarizes common faculty concerns regarding AI adoption alongside evidence-based institutional responses that address these concerns constructively. As presented in Table 3, each concern category has corresponding intervention strategies that can transform resistance into productive engagement when implemented with genuine attention to faculty perspectives and adequate resource allocation. The key insight is that faculty resistance typically reflects legitimate professional concerns rather than irrational technophobia, and institutions that dismiss resistance rather than addressing its underlying causes invariably fail to achieve meaningful adoption [33].",
    ]

    for text in s3_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 3
    doc.add_paragraph()
    t3_title = doc.add_paragraph()
    t3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t3_title.add_run('Table 3. ')
    r.bold = True
    t3_title.add_run('Faculty Concerns Regarding AI Adoption and Evidence-Based Institutional Responses')

    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    h3 = ['Faculty Concern', 'Root Cause', 'Institutional Response Strategy']
    for i, h in enumerate(h3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t3_data = [
        ['AI output inaccuracy', 'Lack of AI literacy; over-reliance on outputs', 'Critical evaluation training; verification protocols; discipline-specific prompting workshops'],
        ['Academic integrity threats', 'Unclear policies; assessment design vulnerability', 'Revised assessment strategies; clear AI-use policies; authentic assessment redesign'],
        ['Job displacement fears', 'Uncertainty about AI role boundaries', 'Augmentation framing; new role creation; participatory policy development'],
        ['Loss of academic autonomy', 'Top-down technology mandates', 'Faculty-led adoption; choice in tools; pedagogical freedom preservation'],
        ['Privacy and data concerns', 'Unclear data governance; vendor opacity', 'Transparent data policies; institutional AI platforms; vendor accountability contracts'],
        ['Workload for AI integration', 'Insufficient support and time allocation', 'Reduced teaching loads during transition; instructional design support; peer mentoring'],
    ]
    for i, row in enumerate(t3_data):
        for j, cell in enumerate(row):
            table3.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    # FIGURE 3
    fig3_path = 'education5_figures/Figure_3_Acceptance_Model.png'
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap3.add_run('Figure 3. ')
    r.bold = True
    cap3.add_run('Academic Community AI Acceptance Model showing relationships among external factors, core acceptance constructs (perceived usefulness, ease of use, trust), behavioral intention, and actual AI use, with moderating variables.')

    doc.add_paragraph()

    doc.add_heading('3.2 Student Perspectives on AI-Enhanced Learning', level=2)

    s3_2 = [
        "Students are both users and beneficiaries of Education 5.0 technologies, occupying a unique position as the primary population whose learning outcomes AI systems are designed to improve [34]. AI can provide personalized learning pathways adapted to individual pace and prior knowledge, intelligent tutoring that identifies and addresses specific misconceptions, automated formative feedback that enables rapid iteration, adaptive content presentation that matches cognitive load capacity, language assistance for international students, and career guidance informed by labour market analytics. The breadth of potential AI applications for students spans the entire educational experience from admission through graduation and into alumni engagement.",

        "Student acceptance depends strongly on usability, accessibility, perceived learning benefits, trust in AI reliability, and existing digital competence [35]. Students with stronger AI literacy may be more comfortable experimenting with intelligent tools, recognizing their outputs as starting points for refinement rather than finished products. Students with limited digital skills may require additional scaffolding including guided tutorials, structured prompting frameworks, and explicit instruction in critical evaluation of AI-generated content before they can engage productively with AI tools. Importantly, student acceptance is not uniform across demographic groups: socioeconomic background, prior educational experience, disciplinary context, and cultural factors all moderate acceptance patterns in ways that institutions must understand to ensure equitable AI adoption.",

        "At the same time, students may worry about surveillance through learning analytics, privacy implications of AI processing their academic data, biased algorithms that might disadvantage certain groups, inaccurate information that could mislead their learning, and the consequences of AI use for academic assessment fairness [36]. These concerns are not irrational and should not be dismissed as resistance to progress. Learning analytics systems that track student behavior in granular detail can feel intrusive. Assessment algorithms that make high-stakes decisions based on patterns invisible to students can feel opaque and unjust. Institutions must therefore establish transparent rules explaining when and how AI may be used in assignments, examinations, research, and collaborative activities. These rules should be developed collaboratively with student representatives to ensure they are perceived as fair and comprehensible rather than arbitrary or punitive.",

        "Rather than simply restricting AI use—an approach that is both practically unenforceable and educationally counterproductive—universities should develop AI literacy as a core academic competency [37]. Students should learn how to formulate effective prompts that elicit useful AI outputs, verify AI-generated information against authoritative sources, identify hallucinations and systematic biases in AI outputs, protect sensitive personal and research information when using AI platforms, cite AI-assisted work appropriately and transparently, and maintain independent reasoning capacity that AI augments rather than replaces. This competency-building approach prepares students for professional environments where AI collaboration is standard while preserving the critical thinking capacities that higher education fundamentally exists to develop [38].",

        "The pedagogical integration of AI literacy requires coordination across curricula rather than isolation within single courses. Introductory courses can establish foundational understanding of AI capabilities and limitations. Intermediate courses can develop discipline-specific AI application skills. Advanced courses can engage students with ethical, societal, and professional implications of AI in their fields. This progressive development ensures that AI literacy grows alongside disciplinary expertise, creating graduates who can leverage AI within their professional contexts rather than possessing generic AI skills disconnected from domain knowledge.",
    ]

    for text in s3_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('3.3 Institutional Leadership, Governance, and Change Management', level=2)

    s3_3 = [
        "Leadership determines whether AI readiness becomes an institutional strategy with dedicated resources and accountability or remains a collection of isolated experiments dependent on individual enthusiasm [39]. University leaders should establish clear AI visions aligned with educational missions and institutional values, communicating not merely that AI will be adopted but articulating why AI adoption serves educational purposes and how it will be implemented in ways that respect academic values including autonomy, integrity, inclusivity, and excellence.",

        "Governance should encompass comprehensive policies covering data privacy, cybersecurity, intellectual property, academic integrity, responsible AI use, procurement standards, algorithmic transparency, and accountability mechanisms [40]. AI governance committees can coordinate academic, technical, legal, administrative, and student perspectives, ensuring that policies reflect the full range of stakeholder interests rather than privileging any single perspective. Effective governance is enabling rather than restrictive: it creates clarity about boundaries within which innovation can flourish freely, reducing the uncertainty that often paralyzes adoption more effectively than prohibition.",

        "Change management is equally important because AI adoption can create significant uncertainty by modifying established roles, workflows, assessment practices, and professional identities [41]. Institutions should communicate the purpose of AI transformation clearly and consistently, emphasizing augmentation rather than replacement of human expertise. Change management should address emotional responses—anxiety, excitement, frustration, curiosity—alongside practical questions about new procedures and expectations. Faculty and staff who feel heard, supported, and genuinely involved in shaping change are dramatically more likely to engage constructively than those who perceive transformation as something done to them rather than with them.",

        "A successful implementation strategy can follow an iterative cycle: assess current readiness, prepare infrastructure and people, pilot AI applications in controlled contexts, evaluate outcomes against defined criteria, scale successful pilots across the institution, and monitor continuously for emerging challenges and opportunities [42]. As depicted in Figure 4, this cycle operates continuously rather than terminating after initial implementation, because AI technologies, user expectations, regulatory requirements, and institutional contexts evolve rapidly and demand ongoing adaptive response. The strategic pathway from initial readiness assessment through transformation stages provides a structured approach that balances urgency with deliberation.",

        "The monitoring phase deserves particular emphasis because it distinguishes sustainable AI transformation from temporary adoption enthusiasm. Monitoring should track not only technical metrics—system usage, performance indicators, error rates—but also human metrics including faculty satisfaction, student experience quality, workload impacts, and emergent ethical concerns. Early identification of adoption fatigue, growing resistance, or unintended consequences enables corrective intervention before minor issues become entrenched problems. Monitoring should also celebrate successes and surface innovations, creating positive feedback loops that sustain transformation momentum across the institution.",
    ]

    for text in s3_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================
    # SECTION 4
    # ============================
    doc.add_heading('4. Strategic Pathways toward Education 5.0 Transformation', level=1)

    doc.add_heading('4.1 Building an AI-Ready Academic Ecosystem', level=2)

    s4_1 = [
        "An AI-ready university requires coordinated investment across infrastructure, people, policies, and partnerships that collectively create an ecosystem supportive of sustainable AI integration [16]. Institutions should develop centralized AI platforms where appropriate—providing standardized, privacy-compliant, and institutionally governed access to AI capabilities—while maintaining interoperability with existing learning-management systems, student-information systems, research-management platforms, and administrative systems. This integration ensures that AI enhances rather than fragments the institutional digital ecosystem.",

        "AI laboratories and innovation centers can provide controlled environments for experimentation where faculty and students explore AI applications without risk to production systems or student outcomes [29]. These spaces serve multiple functions: they enable low-stakes experimentation that builds confidence; they generate evidence about what works in specific institutional contexts; they create visible demonstrations of AI potential that inspire broader adoption; and they provide training venues where structured learning complements self-directed exploration. Faculty communities of practice facilitate peer learning, enabling early adopters to share experiences, challenges, solutions, and strategies with colleagues at earlier adoption stages. These communities are particularly valuable because faculty often trust peer recommendations more than institutional mandates, and practical insights from disciplinary colleagues carry greater weight than generic training materials. Interdisciplinary research groups investigating AI applications across engineering, medicine, business, humanities, social sciences, and other disciplines ensure that AI readiness develops across the full breadth of institutional activity rather than concentrating in technically-oriented faculties.",

        "Investment should extend substantially beyond technology [7]. Continuous AI literacy programs for faculty and staff, instructional-design support for AI-enhanced curriculum development, cybersecurity awareness training, ethical-reasoning workshops, and discipline-specific AI application exploration are essential for sustainable transformation. Table 4 presents the key components of an AI-ready academic ecosystem across infrastructure, human development, governance, and partnership dimensions, providing a comprehensive checklist for institutional planning.",

        "Partnership strategies should encompass technology vendors (ensuring institutional values are embedded in procurement), peer institutions (enabling benchmarking and shared learning), industry partners (connecting academic AI development with workforce requirements), and regulatory bodies (anticipating compliance requirements before they become mandatory) [43]. These partnerships create an institutional network that supports AI readiness development while managing the complexity and cost that no single institution can bear alone. Consortium approaches—where multiple institutions collectively negotiate vendor agreements, develop shared training resources, establish common ethical frameworks, and conduct collaborative research—can dramatically reduce per-institution costs while accelerating capability development through shared learning and mutual accountability.",

        "The temporal dimension of ecosystem development deserves attention. Building an AI-ready academic ecosystem is not an event but a process that unfolds over years. Early investments focus on foundational infrastructure and awareness-building. Intermediate phases develop systematic capabilities and governance. Mature phases achieve integrated operation and continuous improvement. Institutions that expect immediate transformation from initial investments inevitably experience frustration and may prematurely abandon promising initiatives before they mature sufficiently to deliver value. Strategic patience combined with measurable interim milestones enables sustained commitment without sacrificing accountability for progress.",
    ]

    for text in s4_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 4
    doc.add_paragraph()
    t4_title = doc.add_paragraph()
    t4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t4_title.add_run('Table 4. ')
    r.bold = True
    t4_title.add_run('Key Components of an AI-Ready Academic Ecosystem')

    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    h4 = ['Ecosystem Dimension', 'Key Components', 'Implementation Priority', 'Success Indicators']
    for i, h in enumerate(h4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t4_data = [
        ['Infrastructure', 'Centralized AI platform, cloud computing, secure data lakes, LMS integration', 'High (foundational)', 'Platform uptime >99%, API integration complete, data governance active'],
        ['Human Development', 'AI literacy programs, communities of practice, instructional design support, mentoring', 'High (ongoing)', 'Faculty training >80%, student AI literacy score >3.5/5, peer networks active'],
        ['Governance & Policy', 'AI ethics board, usage policies, integrity guidelines, procurement standards', 'High (immediate)', 'Policies published, compliance >90%, incident response tested'],
        ['Partnerships', 'Vendor agreements, peer institution networks, industry advisory boards, regulatory engagement', 'Medium (strategic)', 'Active partnerships >5, joint research projects, advisory board meetings quarterly'],
    ]
    for i, row in enumerate(t4_data):
        for j, cell in enumerate(row):
            table4.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    doc.add_heading('4.2 From AI Adoption to Human-Centered Transformation', level=2)

    s4_2 = [
        "The ultimate objective of Education 5.0 is not maximum AI deployment but improved human and institutional outcomes [1]. AI should support educators in making better pedagogical decisions informed by learning analytics, enable students to learn more effectively through personalized pathways, assist researchers in discovering new knowledge through AI-augmented literature analysis and hypothesis generation, and help administrators deliver more responsive and equitable services. The distinction between AI deployment and AI-enabled transformation is crucial: institutions can deploy numerous AI tools while achieving minimal transformation if those tools are not integrated into workflows, accepted by users, or aligned with educational objectives.",

        "Human-centered transformation requires maintaining meaningful human oversight over AI-mediated decisions [6]. High-impact decisions involving student assessment, academic progression, admission, resource allocation, or employment should not depend exclusively on automated systems regardless of their statistical accuracy. Humans must retain authority to override algorithmic recommendations when contextual factors—personal circumstances, cultural considerations, developmental trajectories—are relevant but invisible to AI systems operating on structured data alone. This principle preserves both ethical accountability and the educational relationship that constitutes higher education's fundamental value proposition.",

        "Universities should consequently evaluate AI systems according to both technical performance and social outcomes [30]. Technical indicators may include prediction accuracy, system reliability, processing speed, and integration quality. Social indicators should include learning gains, student satisfaction, faculty workload impacts, accessibility improvements, equity effects across demographic groups, research productivity contributions, trust levels, and ethical compliance. As summarized in Table 4, success indicators span both technical functionality and human outcomes, reflecting the dual nature of Education 5.0's value proposition.",

        "The measurement of AI transformation outcomes should follow established educational evaluation methodologies adapted for AI contexts [19]. Pre-post comparisons of learning outcomes, controlled experiments comparing AI-enhanced and traditional approaches, longitudinal tracking of student progression, and qualitative investigation of experience quality collectively provide the evidence base for determining whether AI adoption genuinely advances educational quality or merely adds technological complexity without proportionate benefit. Mixed-methods evaluation designs that combine quantitative outcome measurement with qualitative process understanding provide the richest insights for institutional decision-making.",

        "Equity considerations must be central to transformation evaluation. AI systems may inadvertently advantage students with stronger digital backgrounds while disadvantaging those with limited technology access or different learning preferences. Assessment of AI transformation outcomes should therefore disaggregate results by relevant demographic categories—socioeconomic status, prior achievement, first-generation status, disability, language background—to identify whether AI integration creates or exacerbates educational inequities. Institutions committed to inclusive excellence must ensure that AI transformation serves all students equitably rather than amplifying existing advantages for the already-privileged.",
    ]

    for text in s4_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('4.3 Future Directions for Data-Driven AI Readiness', level=2)

    # NO CITATIONS in this concluding section
    s4_3 = [
        "Future AI readiness assessment should move toward continuous and predictive models that provide dynamic institutional intelligence rather than static periodic snapshots. Instead of conducting comprehensive readiness assessments once every few years, universities can develop institutional dashboards that continuously monitor infrastructure performance, AI usage patterns, training participation rates, stakeholder perception trends, governance compliance indicators, and transformation outcome metrics. These dashboards enable real-time identification of emerging readiness gaps and proactive intervention before minor issues escalate into significant barriers.",

        "Artificial intelligence can itself support readiness assessment through automated pattern recognition within institutional data. Machine learning algorithms can identify correlations between training investments and subsequent adoption rates, predict which departments or individuals are likely to encounter adoption difficulties, and recommend targeted interventions based on profiles of successful transformation in comparable contexts. However, predictive analytics applied to institutional readiness should remain transparent, subject to human oversight, and designed to support rather than replace human judgment about institutional strategy and priorities.",

        "Future research should investigate cross-cultural differences in AI acceptance, recognizing that cultural values regarding technology, authority, privacy, and innovation significantly influence how academic communities respond to AI integration. Discipline-specific readiness research is needed because the relevance, risks, and appropriate applications of AI vary substantially across fields ranging from medicine and engineering to philosophy and fine arts. Generative AI adoption specifically requires investigation because its conversational interface, content creation capabilities, and assessment implications distinguish it from previous AI applications in ways that existing acceptance models may not fully capture.",

        "The ethical dimensions of AI readiness assessment itself warrant critical examination. Readiness frameworks implicitly assume that AI adoption is desirable and that institutions should progress toward greater AI integration. However, legitimate arguments exist for deliberate restraint in certain contexts—for example, where AI integration may compromise disciplinary practices that depend on unmediated human engagement, or where surveillance concerns outweigh efficiency benefits. Future frameworks should accommodate principled non-adoption as a legitimate outcome of readiness assessment rather than treating resistance as uniformly problematic.",

        "Longitudinal studies will be particularly valuable for determining whether improvements in measured readiness actually translate into sustained educational transformation rather than temporary adoption enthusiasm that fades once novelty effects dissipate. Research connecting readiness indicators with long-term outcomes—student learning gains measured over multiple semesters, graduate employability tracked through career outcomes, research impact assessed through citation patterns and societal influence, institutional reputation evaluated through stakeholder perception—would provide the evidence base for determining which readiness investments deliver genuine return and which represent costly but unproductive institutional signaling. Such research requires patience and sustained funding commitment that short-term project cycles rarely support, making this an area where national research councils and international bodies have important roles to play.",

        "The relationship between AI readiness and workforce transformation in higher education deserves dedicated investigation. As AI assumes greater roles in content delivery, assessment feedback, administrative processing, research literature analysis, and even grant writing support, the competencies required of academic professionals will evolve significantly. Understanding how readiness development relates to professional identity evolution, career satisfaction, role redefinition, and workforce planning will be essential for institutions navigating the human dimensions of AI transformation responsibly. Research in this area should examine both the opportunities—new roles in AI governance, learning design, and human-AI collaboration—and the risks—deskilling, cognitive atrophy, and professional identity threat—that accompany AI integration into academic work.",

        "Overall, the transition toward Higher Education 5.0 requires a balanced approach in which technological capability, human competence, organizational capacity, ethical governance, and stakeholder acceptance develop together in coordinated fashion. A data-driven AI readiness framework provides universities with a structured mechanism for identifying gaps, prioritizing investments, evaluating progress, and developing responsible AI strategies aligned with educational mission. The academic community should remain at the center of this transformation, ensuring that AI becomes an enabler of creativity, inclusion, personalization, research excellence, and human development rather than simply another layer of technological automation imposed upon educational institutions without genuine engagement with the people whose work and learning it is intended to enhance.",
    ]

    for text in s4_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # FIGURE 4
    fig4_path = 'education5_figures/Figure_4_Transformation_Roadmap.png'
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap4 = doc.add_paragraph()
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap4.add_run('Figure 4. ')
    r.bold = True
    cap4.add_run('Strategic pathway from AI Readiness assessment to Education 5.0 Transformation, showing the six-phase implementation cycle (Assess-Prepare-Pilot-Evaluate-Scale-Monitor) and five institutional readiness levels.')

    doc.add_page_break()

    # ============================
    # REFERENCES [1]-[43]
    # ============================
    doc.add_heading('References', level=1)

    references = [
        "[1] Carayannis, E. G., & Morawska-Jancelewicz, J. (2022). The futures of Europe: Society 5.0 and Industry 5.0 as driving forces of future universities. Journal of the Knowledge Economy, 13(4), 3445–3471.",
        "[2] Mhlanga, D. (2023). Open AI in Education, the responsible and ethical use of ChatGPT towards lifelong learning. Education and Information Technologies, 28(8), 7541–7568.",
        "[3] Zawacki-Richter, O., Marín, V. I., Bond, M., & Gouverneur, F. (2019). Systematic review of research on artificial intelligence applications in higher education. International Journal of Educational Technology in Higher Education, 16(1), 1–27.",
        "[4] Bearman, M., Ryan, J., & Ajjawi, R. (2023). Discourses of artificial intelligence in higher education: A critical literature review. Higher Education, 86(2), 369–385.",
        "[5] Jöhnk, J., Weißert, M., & Wyrtki, K. (2021). Ready or not, AI comes—An interview study of organizational AI readiness factors. Business & Information Systems Engineering, 63(1), 5–20.",
        "[6] Xu, X., Lu, Y., Vogel-Heuser, B., & Wang, L. (2021). Industry 4.0 and Industry 5.0—Inception, conception and perception. Journal of Manufacturing Systems, 61, 530–535.",
        "[7] Alsheibani, S., Cheung, Y., & Messom, C. (2018). Artificial intelligence adoption: AI-readiness at firm-level. Proceedings of the Pacific Asia Conference on Information Systems, 37, 1–12.",
        "[8] Long, D., & Magerko, B. (2020). What is AI literacy? Competencies and design considerations. Proceedings of the 2020 CHI Conference on Human Factors in Computing Systems, 1–16.",
        "[9] Pumplun, L., Taber, C., & Buxmann, P. (2021). Beyond plug and pray: How organizations can get more out of artificial intelligence. Business & Information Systems Engineering, 63(2), 159–174.",
        "[10] Jobin, A., Ienca, M., & Vayena, E. (2019). The global landscape of AI ethics guidelines. Nature Machine Intelligence, 1(9), 389–399.",
        "[11] Schein, E. H. (2010). Organizational culture and leadership (4th ed.). Jossey-Bass.",
        "[12] Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340.",
        "[13] Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425–478.",
        "[14] Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology: Extending the unified theory of acceptance and use of technology. MIS Quarterly, 36(1), 157–178.",
        "[15] Agrawal, A., Gans, J., & Goldfarb, A. (2018). Prediction machines: The simple economics of artificial intelligence. Harvard Business Review Press.",
        "[16] Crompton, H., & Burke, D. (2023). Artificial intelligence in higher education: The state of the field. International Journal of Educational Technology in Higher Education, 20(1), 1–22.",
        "[17] Dwivedi, Y. K., Kshetri, N., Hughes, L., Slade, E. L., Jeyaraj, A., Kar, A. K., & Wright, R. (2023). Opinion paper: So what if ChatGPT wrote it? Multidisciplinary perspectives on opportunities, challenges and implications of generative conversational AI for research, practice and policy. International Journal of Information Management, 71, 102642.",
        "[18] Mishra, P., & Koehler, M. J. (2006). Technological pedagogical content knowledge: A framework for teacher knowledge. Teachers College Record, 108(6), 1017–1054.",
        "[19] Saaty, T. L. (2008). Decision making with the analytic hierarchy process. International Journal of Services Sciences, 1(1), 83–98.",
        "[20] Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). SAGE Publications.",
        "[21] Likert, R. (1932). A technique for the measurement of attitudes. Archives of Psychology, 22(140), 1–55.",
        "[22] Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate data analysis (8th ed.). Cengage Learning.",
        "[23] Hastie, T., Tibshirani, R., & Friedman, J. (2009). The elements of statistical learning: Data mining, inference, and prediction (2nd ed.). Springer.",
        "[24] Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
        "[25] UNESCO. (2021). AI and education: Guidance for policy-makers. UNESCO Publishing.",
        "[26] JISC. (2022). National centre for AI in tertiary education: AI readiness framework. JISC.",
        "[27] EDUCAUSE. (2023). Horizon Report: Teaching and Learning Edition. EDUCAUSE.",
        "[28] Kaplan, R. S., & Norton, D. P. (1996). The balanced scorecard: Translating strategy into action. Harvard Business School Press.",
        "[29] Moorhouse, B. L. (2024). Beginning and first-year teachers' use of generative AI for lesson planning. British Journal of Educational Technology, 55(3), 1150–1167.",
        "[30] Chan, C. K. Y., & Hu, W. (2023). Students' voices on generative AI: Perceptions, benefits, and challenges in higher education. International Journal of Educational Technology in Higher Education, 20(1), 43.",
        "[31] Cotton, D. R. E., Cotton, P. A., & Shipway, J. R. (2023). Chatting and cheating: Ensuring academic integrity in the era of ChatGPT. Innovations in Education and Teaching International, 61(2), 228–239.",
        "[32] Ng, D. T. K., Leung, J. K. L., Chu, S. K. W., & Qiao, M. S. (2021). Conceptualizing AI literacy: An exploratory review. Computers and Education: Artificial Intelligence, 2, 100041.",
        "[33] Selwyn, N. (2022). The future of AI and education: Some cautionary notes. European Journal of Education, 57(4), 620–631.",
        "[34] Holmes, W., Bialik, M., & Fadel, C. (2019). Artificial intelligence in education: Promises and implications for teaching and learning. Center for Curriculum Redesign.",
        "[35] Sánchez-Prieto, J. C., Olmos-Migueláñez, S., & García-Peñalvo, F. J. (2017). MLearning and pre-service teachers: An assessment of the behavioral intention using an expanded TAM model. Computers in Human Behavior, 72, 644–654.",
        "[36] Prinsloo, P., & Slade, S. (2016). Student vulnerability, agency, and learning analytics: An exploration. Journal of Learning Analytics, 3(1), 159–182.",
        "[37] Ng, D. T. K., Leung, J. K. L., Su, J., Ng, R. C. W., & Chu, S. K. W. (2023). Teachers' AI digital competencies and twenty-first century skills in the post-pandemic world. Educational Technology Research and Development, 71(1), 137–161.",
        "[38] Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, 102274.",
        "[39] Kotter, J. P. (2012). Leading change (with a new preface). Harvard Business Review Press.",
        "[40] Floridi, L., Cowls, J., Beltrametti, M., Chatila, R., Chazerand, P., Dignum, V., & Vayena, E. (2018). AI4People—An ethical framework for a good AI society. Minds and Machines, 28(4), 689–707.",
        "[41] Fullan, M. (2015). The new meaning of educational change (5th ed.). Teachers College Press.",
        "[42] Rogers, E. M. (2003). Diffusion of innovations (5th ed.). Free Press.",
        "[43] Williamson, B., & Eynon, R. (2020). Historical threads, missing links, and future directions in AI in education. Learning, Media and Technology, 45(3), 223–235.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)

    # Save
    output_path = 'Chapter_Education5_AI_Readiness.docx'
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

    # Verification
    total = ''
    for para in doc.paragraphs:
        total += para.text + ' '
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += cell.text + ' '
    print(f"Total word count: {len(total.split())}")
    print(f"References: {len(references)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == '__main__':
    create_docx()
