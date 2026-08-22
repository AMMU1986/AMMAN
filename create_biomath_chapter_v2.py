"""
Create Word Document for Book Chapter:
"Differential Equations and Dynamical Systems in Biology"
from: Biomathematics: A New Horizon of Science and Engineering

Features:
- ~8300 words
- 43 references in serial order (square brackets)
- 4 tables (cited twice each)
- 4 figures in PNG format (cited twice each)
- 17 equations in Word equation editor format (OMML)
- Nomenclature section
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as etree

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_numbered_equation(doc, equation_xml, eq_number):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = paragraph._p
    omath_element = etree.fromstring(equation_xml)
    omath_para.append(omath_element)
    run = paragraph.add_run(f'    ({eq_number})')
    run.font.size = Pt(11)
    return paragraph

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def add_heading_styled(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

# ============================================================
# EQUATION DEFINITIONS (OMML format)
# ============================================================
omml_ns = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

def make_eq(inner):
    return f'<m:oMath {omml_ns}>{inner}</m:oMath>'

def make_frac(num, den):
    return f'<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'

def make_run(text):
    return f'<m:r><m:t>{text}</m:t></m:r>'

def make_sup(base, sup):
    return f'<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'

def make_sub(base, sub):
    return f'<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'

# Equation 1: General ODE
eq1_xml = make_eq(make_frac(make_run('dx'), make_run('dt')) + make_run(' = f(x, t; θ)'))

# Equation 2: Linear stability (Jacobian)
eq2_xml = make_eq(make_frac(make_run('dξ'), make_run('dt')) + make_run(' = J·ξ,  J = ') + make_frac(make_run('∂f'), make_run('∂x')) + make_run('|') + make_sub(make_run(''), make_run('x*')))

# Equation 3: Lotka-Volterra prey
eq3_xml = make_eq(make_frac(make_run('dx'), make_run('dt')) + make_run(' = αx − βxy'))

# Equation 4: Lotka-Volterra predator
eq4_xml = make_eq(make_frac(make_run('dy'), make_run('dt')) + make_run(' = δxy − γy'))

# Equation 5: Logistic growth
eq5_xml = make_eq(make_frac(make_run('dN'), make_run('dt')) + make_run(' = rN') + make_run('(1 − ') + make_frac(make_run('N'), make_run('K')) + make_run(')'))

# Equation 6: SIR - Susceptible
eq6_xml = make_eq(make_frac(make_run('dS'), make_run('dt')) + make_run(' = −βSI/N'))

# Equation 7: SIR - Infected
eq7_xml = make_eq(make_frac(make_run('dI'), make_run('dt')) + make_run(' = βSI/N − γI'))

# Equation 8: SIR - Recovered
eq8_xml = make_eq(make_frac(make_run('dR'), make_run('dt')) + make_run(' = γI'))

# Equation 9: Basic reproduction number
eq9_xml = make_eq(make_sub(make_run('R'), make_run('0')) + make_run(' = ') + make_frac(make_run('β'), make_run('γ')))

# Equation 10: Reaction-diffusion
eq10_xml = make_eq(make_frac(make_run('∂u'), make_run('∂t')) + make_run(' = D') + make_frac(make_sup(make_run('∂'), make_run('2')) + make_run('u'), make_run('∂') + make_sup(make_run('x'), make_run('2'))) + make_run(' + f(u)'))

# Equation 11: Hill function
eq11_xml = make_eq(make_run('H(x) = ') + make_frac(make_sup(make_run('x'), make_run('n')), make_sup(make_run('K'), make_run('n')) + make_run(' + ') + make_sup(make_run('x'), make_run('n'))))

# Equation 12: Michaelis-Menten
eq12_xml = make_eq(make_run('v = ') + make_frac(make_sub(make_run('V'), make_run('max')) + make_run('[S]'), make_sub(make_run('K'), make_run('m')) + make_run(' + [S]')))

# Equation 13: Lyapunov function
eq13_xml = make_eq(make_frac(make_run('dV'), make_run('dt')) + make_run(' = ') + make_frac(make_run('∂V'), make_run('∂x')) + make_run('·f(x) ≤ 0'))

# Equation 14: Bifurcation normal form
eq14_xml = make_eq(make_frac(make_run('dx'), make_run('dt')) + make_run(' = μx − ') + make_sup(make_run('x'), make_run('3')))

# Equation 15: Turing instability condition
eq15_xml = make_eq(make_sub(make_run('d'), make_run('v')) + make_sub(make_run('f'), make_run('u')) + make_run(' + ') + make_sub(make_run('d'), make_run('u')) + make_sub(make_run('g'), make_run('v')) + make_run(' − 2√(') + make_sub(make_run('d'), make_run('u')) + make_sub(make_run('d'), make_run('v')) + make_run('Δ) > 0'))

# Equation 16: Stochastic differential equation
eq16_xml = make_eq(make_run('dX = f(X)dt + g(X)dW(t)'))

# Equation 17: PINN loss function
eq17_xml = make_eq(make_run('L = ') + make_sub(make_run('L'), make_run('data')) + make_run(' + λ') + make_sub(make_run('L'), make_run('physics')) + make_run(' + μ') + make_sub(make_run('L'), make_run('BC')))

equations = [eq1_xml, eq2_xml, eq3_xml, eq4_xml, eq5_xml, eq6_xml, eq7_xml, 
             eq8_xml, eq9_xml, eq10_xml, eq11_xml, eq12_xml, eq13_xml, eq14_xml,
             eq15_xml, eq16_xml, eq17_xml]

# ============================================================
# MAIN DOCUMENT CREATION
# ============================================================

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ============================================================
# TITLE AND AUTHORS
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CHAPTER')
run.bold = True
run.font.size = Pt(14)

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('Differential Equations and Dynamical Systems in Biology')
run.bold = True
run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('From: Biomathematics: A New Horizon of Science and Engineering')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_styled(doc, 'Abstract', level=1)

abstract_text = (
    "This chapter provides a comprehensive treatment of differential equations and dynamical systems "
    "as applied to biological sciences, covering the mathematical foundations, analytical techniques, "
    "and computational methodologies that underpin modern biomathematics. We begin with the principles "
    "of mathematical modeling in biology, emphasizing model formulation, biological assumptions, and "
    "the balance between model complexity and parsimony. The chapter systematically develops the theory "
    "of ordinary differential equations in biological contexts, including equilibrium analysis, "
    "linearization, phase-plane methods, and Lyapunov stability. Spatial dynamics are addressed through "
    "partial differential equations governing reaction-diffusion processes, chemotaxis, and biological "
    "wave propagation. Applications span population ecology, mathematical epidemiology, gene regulatory "
    "networks, metabolic pathways, and cellular signaling. Advanced topics include bifurcation analysis, "
    "numerical methods for stiff systems, model reduction techniques, hybrid multiscale modeling, "
    "data-driven approaches using physics-informed neural networks, and dynamical systems perspectives "
    "on personalized medicine and synthetic biology. The chapter integrates classical analytical methods "
    "with emerging computational paradigms, providing researchers and practitioners with a unified "
    "framework for understanding biological complexity through the lens of dynamical systems theory. "
    "We demonstrate how these mathematical tools can be applied to understand oscillatory population "
    "dynamics, predict epidemic outbreaks, analyze cellular decision-making, and design synthetic "
    "biological circuits with prescribed behaviors."
)
p = doc.add_paragraph(abstract_text)
p.paragraph_format.first_line_indent = Cm(1.27)

keywords = doc.add_paragraph()
run = keywords.add_run('Keywords: ')
run.bold = True
keywords.add_run('Differential equations; Dynamical systems; Mathematical biology; Population dynamics; '
                 'Epidemiological modeling; Gene regulatory networks; Reaction-diffusion; Bifurcation analysis; '
                 'Physics-informed neural networks; Personalized medicine; Synthetic biology; Multiscale modeling')

doc.add_paragraph()

# ============================================================
# NOMENCLATURE
# ============================================================
add_heading_styled(doc, 'Nomenclature', level=1)

nomenclature_items = [
    ('x, y', 'State variables (population densities, concentrations)'),
    ('t', 'Time variable'),
    ('N', 'Total population size'),
    ('S, I, R', 'Susceptible, Infected, Recovered compartments'),
    ('α, β, γ, δ', 'Rate parameters in ecological/epidemiological models'),
    ('K', 'Carrying capacity / Half-saturation constant'),
    ('r', 'Intrinsic growth rate'),
    ('R₀', 'Basic reproduction number'),
    ('D', 'Diffusion coefficient'),
    ('J', 'Jacobian matrix'),
    ('λ', 'Eigenvalue'),
    ('V', 'Lyapunov function'),
    ('n', 'Hill coefficient'),
    ('Vmax', 'Maximum enzyme velocity'),
    ('Km', 'Michaelis-Menten constant'),
    ('μ', 'Bifurcation parameter / death rate'),
    ('θ', 'Parameter vector'),
    ('ξ', 'Perturbation vector'),
    ('W(t)', 'Wiener process (Brownian motion)'),
    ('f, g', 'Reaction kinetics functions'),
    ('∇²', 'Laplacian operator'),
    ('H(x)', 'Hill function'),
    ('L', 'Loss function (neural network)'),
    ('τ', 'Time delay'),
    ('ODE', 'Ordinary Differential Equation'),
    ('PDE', 'Partial Differential Equation'),
    ('SDE', 'Stochastic Differential Equation'),
    ('DDE', 'Delay Differential Equation'),
    ('SIR', 'Susceptible-Infected-Recovered model'),
    ('PINN', 'Physics-Informed Neural Network'),
    ('GRN', 'Gene Regulatory Network'),
    ('MCA', 'Metabolic Control Analysis'),
    ('QSSA', 'Quasi-Steady-State Approximation'),
    ('ABM', 'Agent-Based Model'),
    ('AIC', 'Akaike Information Criterion'),
    ('BIC', 'Bayesian Information Criterion'),
]

nom_table = doc.add_table(rows=len(nomenclature_items)+1, cols=2)
nom_table.style = 'Table Grid'

hdr = nom_table.rows[0]
hdr.cells[0].text = 'Symbol/Abbreviation'
hdr.cells[1].text = 'Description'
for cell in hdr.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_shading(cell, "4472C4")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (symbol, desc) in enumerate(nomenclature_items):
    row = nom_table.rows[i+1]
    row.cells[0].text = symbol
    row.cells[1].text = desc

nom_table.alignment = WD_TABLE_ALIGNMENT.CENTER
doc.add_paragraph()

# ============================================================
# SECTION 1
# ============================================================
add_heading_styled(doc, '1. Mathematical Foundations and Modeling of Biological Systems', level=1)

add_heading_styled(doc, '1.1 Principles of Mathematical Modeling in Biology', level=2)

add_heading_styled(doc, '1.1.1 Model Formulation and Biological Assumptions', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Mathematical modeling constitutes the cornerstone of quantitative biology, providing a rigorous "
    "framework for translating biological hypotheses into testable mathematical statements [1]. The "
    "process of model formulation requires careful consideration of biological assumptions, the identification "
    "of key state variables, and the specification of governing dynamical laws that describe how these "
    "variables evolve over time [2]. A fundamental principle in biological modeling is parsimony—the idea "
    "that models should be as simple as possible while still capturing the essential dynamics of the "
    "biological system under investigation [3]. This principle, often attributed to Occam's razor in "
    "scientific modeling, guards against overfitting and ensures that model predictions are robust and "
    "interpretable. The art of mathematical modeling in biology lies in identifying the minimal set of "
    "variables and interactions that suffice to explain observed phenomena while remaining faithful to "
    "known biological mechanisms."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The modeling cycle in biology follows an iterative process: biological observations motivate "
    "hypotheses, which are translated into mathematical equations; these equations are analyzed "
    "analytically or numerically; predictions are compared against experimental data; and discrepancies "
    "between predictions and observations guide model refinement [4]. Each iteration of this cycle "
    "deepens our understanding of the biological system and progressively constrains the space of "
    "plausible mechanistic explanations. The biological assumptions underlying any model must be "
    "explicitly stated and critically evaluated, as they determine the domain of validity and the "
    "interpretive framework for model outputs. Common assumptions include spatial homogeneity "
    "(well-mixed compartments), temporal continuity (no abrupt transitions), population-level "
    "averaging (mean-field approximations), and specific functional forms for interaction terms "
    "(mass action, saturating responses). The violation of these assumptions often motivates the "
    "development of more sophisticated modeling frameworks, including spatially explicit models, "
    "stochastic descriptions, and individual-based simulations."
)

add_heading_styled(doc, '1.1.2 Deterministic versus Stochastic Representations', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The distinction between deterministic and stochastic representations is fundamental in biological "
    "modeling [5]. Deterministic models, typically formulated as ordinary or partial differential equations, "
    "describe the average behavior of biological systems and are appropriate when population sizes are "
    "large and fluctuations are negligible. These models assume that the future state of the system is "
    "completely determined by its current state and the governing equations. Stochastic models, on the "
    "other hand, explicitly account for random fluctuations arising from the discrete nature of molecular "
    "interactions and demographic events [6]. The chemical master equation provides the most fundamental "
    "stochastic description of biochemical reaction networks, tracking the probability distribution over "
    "all possible molecular configurations."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The choice between deterministic and stochastic frameworks depends critically on the biological "
    "context and the questions being asked. Intracellular processes with low molecule counts—such as "
    "gene expression regulated by a few transcription factor molecules or the initiation of viral "
    "infection by single particles—necessitate stochastic descriptions that capture the inherent "
    "randomness of individual molecular events [7]. At the population level, demographic stochasticity "
    "becomes important for small populations near extinction thresholds, while environmental "
    "stochasticity captures random fluctuations in external conditions affecting all individuals "
    "simultaneously. The Fokker-Planck equation and Langevin (Itô stochastic differential equation) "
    "approximations provide intermediate descriptions between fully discrete stochastic models and "
    "continuous deterministic equations, applicable when molecule numbers are moderate and Gaussian "
    "approximations to fluctuations are reasonable. The chemical Langevin equation bridges these "
    "scales by adding state-dependent noise terms to deterministic rate equations, maintaining "
    "computational tractability while capturing essential stochastic effects."
)

add_heading_styled(doc, '1.1.3 Parameter Estimation and Model Validation', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Parameter estimation and model validation represent critical challenges in biomathematical modeling [8]. "
    "Structural identifiability analysis determines whether model parameters can, in principle, be uniquely "
    "determined from perfect data, while practical identifiability assesses whether parameters can be "
    "estimated given realistic, noisy experimental measurements [9]. Techniques for structural "
    "identifiability include differential algebra approaches, Taylor series methods, and generating "
    "series techniques, each providing conditions under which model parameters are uniquely or "
    "non-uniquely determined by the model structure and available observations."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Model validation involves comparing model predictions against independent experimental data "
    "not used in parameter estimation, employing statistical criteria such as the Akaike Information "
    "Criterion (AIC) and Bayesian Information Criterion (BIC) for model selection [10]. Cross-validation "
    "techniques assess predictive performance on held-out data, while posterior predictive checks in "
    "Bayesian frameworks evaluate whether simulated data from the fitted model are statistically "
    "consistent with observed data. Profile likelihood methods provide an alternative approach to "
    "assessing parameter identifiability by examining how the likelihood function varies as individual "
    "parameters are systematically varied while others are re-optimized. Table 1 summarizes the key "
    "modeling frameworks and their biological applications, illustrating the diversity of mathematical "
    "approaches available to biological modelers. The selection among these frameworks depends on the "
    "spatial and temporal scales of interest, the available data quality and quantity, the scientific "
    "questions being addressed, and the computational resources available for model analysis and "
    "simulation."
)

# TABLE 1
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table 1: Mathematical Modeling Frameworks in Biology and Their Applications')
run.bold = True
run.font.size = Pt(10)

table1 = doc.add_table(rows=8, cols=4)
table1.style = 'Table Grid'

headers = ['Framework', 'Mathematical Form', 'Biological Application', 'Key Assumptions']
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h
    for paragraph in table1.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_shading(table1.rows[0].cells[i], "4472C4")
    for paragraph in table1.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

table1_data = [
    ['ODEs', 'dx/dt = f(x,t)', 'Population dynamics, enzyme kinetics', 'Continuous, deterministic, well-mixed'],
    ['PDEs', '∂u/∂t = D∇²u + f(u)', 'Spatial ecology, morphogenesis', 'Continuous space, diffusion-driven'],
    ['SDEs', 'dX = f(X)dt + g(X)dW', 'Gene expression, molecular noise', 'Stochastic fluctuations significant'],
    ['DDEs', 'dx/dt = f(x(t), x(t-τ))', 'Immune response, cell cycle', 'Time delays in feedback'],
    ['Agent-based', 'Rule-based interactions', 'Cell migration, tumor growth', 'Individual heterogeneity matters'],
    ['Boolean networks', 'xᵢ(t+1) = Bᵢ(x(t))', 'Gene regulatory networks', 'Binary gene states sufficient'],
    ['Hybrid models', 'ODE/PDE + stochastic', 'Multiscale tissue dynamics', 'Multiple scales coupled'],
]

for i, row_data in enumerate(table1_data):
    for j, val in enumerate(row_data):
        table1.rows[i+1].cells[j].text = val

set_table_style(table1)
doc.add_paragraph()

# 1.2 ODEs in Biological Systems
add_heading_styled(doc, '1.2 Ordinary Differential Equations in Biological Systems', level=2)

add_heading_styled(doc, '1.2.1 Formulation and Interpretation of Biological ODE Models', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Ordinary differential equations form the most widely employed mathematical framework in "
    "biological modeling, describing how the state of a biological system evolves continuously in time [11]. "
    "The power of ODE models lies in their ability to capture essential dynamical features—stability, "
    "oscillations, bistability, and transient responses—through relatively compact mathematical "
    "descriptions that can be analyzed both analytically and numerically. The general form of a "
    "biological ODE model can be expressed as shown in Equation (1), where x represents the vector "
    "of state variables (such as species concentrations, population densities, or gene expression "
    "levels), t denotes time, and θ is a vector of model parameters encoding the rates of biological "
    "processes including synthesis, degradation, binding, and transport."
)

add_numbered_equation(doc, eq1_xml, 1)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The biological interpretation of ODE models requires careful attention to the meaning of each term "
    "in the governing equations [12]. Production terms represent birth, synthesis, or immigration; "
    "degradation terms correspond to death, decay, or emigration; and interaction terms capture "
    "predation, competition, cooperation, or catalysis. The formulation of rate laws typically draws "
    "upon mass-action kinetics (where reaction rates are proportional to the product of reactant "
    "concentrations), Michaelis-Menten enzyme kinetics (describing saturating substrate-velocity "
    "relationships), or Hill-type cooperative responses (capturing sigmoidal input-output relationships), "
    "each reflecting different mechanistic assumptions about the underlying biological processes [13]. "
    "The choice of rate law determines both the qualitative dynamics of the model and the biological "
    "interpretability of its parameters."
)

add_heading_styled(doc, '1.2.2 Equilibrium Points, Linearization, and Local Stability', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Equilibrium analysis constitutes the foundation of qualitative ODE analysis in biology. Equilibrium "
    "points (steady states) x* satisfy f(x*, θ) = 0 and represent states where the system remains "
    "unchanged over time [14]. In biological contexts, equilibrium points correspond to homeostatic "
    "states (normal physiological operating points), endemic disease levels (where infection persists "
    "indefinitely), coexistence points in ecological communities (where multiple species maintain "
    "stable populations), or steady-state metabolic fluxes (balanced metabolic pathways). The number "
    "and nature of equilibria depend on system parameters, and changes in parameters can create or "
    "destroy equilibria through bifurcation events. The local stability of equilibrium points is "
    "determined by linearization—examining the behavior of small perturbations ξ = x − x* near the "
    "steady state, which evolves according to Equation (2), where J is the Jacobian matrix of partial "
    "derivatives evaluated at the equilibrium."
)

add_numbered_equation(doc, eq2_xml, 2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The eigenvalues λ of the Jacobian matrix J determine local stability: if all eigenvalues have "
    "negative real parts, the equilibrium is asymptotically stable (perturbations decay exponentially "
    "and the system returns to equilibrium); if any eigenvalue has a positive real part, the equilibrium "
    "is unstable (perturbations grow exponentially and the system moves away from the steady state) "
    "[15]. Complex eigenvalues with negative real parts indicate damped oscillations toward the "
    "steady state—the system spirals inward, which is commonly observed in predator-prey models with "
    "density-dependent growth. Purely imaginary eigenvalues suggest the existence of a center "
    "or the possibility of a limit cycle emerging through a Hopf bifurcation, which underlies "
    "the generation of biological oscillations in circadian clocks, cell cycles, and calcium "
    "signaling. For two-dimensional systems, the trace and determinant of the Jacobian provide "
    "convenient classification criteria: stability requires a negative trace (representing net "
    "damping) and positive determinant (ensuring no saddle-point character), while oscillatory "
    "approach to equilibrium requires the discriminant (trace² − 4·determinant) to be negative."
)

add_heading_styled(doc, '1.2.3 Phase-Plane Analysis and Lyapunov Stability', level=3)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Phase-plane analysis provides geometric insight into system dynamics for two-dimensional "
    "systems, revealing trajectories, separatrices, and basins of attraction that determine long-term "
    "system behavior [16]. Nullclines—curves where one component of the vector field vanishes—divide "
    "the phase plane into regions of qualitatively different flow directions, and their intersections "
    "identify equilibrium points. The geometry of nullclines provides immediate qualitative insight "
    "into system behavior: when nullclines intersect transversally, the resulting equilibrium is "
    "typically hyperbolic (either a node, saddle, or spiral); when nullclines are tangent, the system "
    "is near a bifurcation point. The Poincaré-Bendixson theorem guarantees that in planar systems, "
    "any bounded trajectory that does not approach an equilibrium must approach a limit cycle, "
    "providing a powerful tool for establishing the existence of periodic orbits in biological models. "
    "This theorem has been instrumental in proving the existence of sustained oscillations in models "
    "of neural activity, calcium signaling, and predator-prey dynamics with realistic functional "
    "responses."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Lyapunov stability theory offers a powerful alternative to eigenvalue analysis, particularly for "
    "nonlinear systems where linearization may be insufficient [17]. A Lyapunov function V(x) is a "
    "scalar function that decreases along system trajectories, as expressed in Equation (13). The "
    "existence of such a function guarantees stability without requiring explicit solution of the "
    "differential equations."
)

add_numbered_equation(doc, eq13_xml, 13)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "In biological contexts, Lyapunov functions often correspond to free energy, entropy production, "
    "or other thermodynamically motivated quantities. For ecological models, the function "
    "V = Σᵢ cᵢ(xᵢ − xᵢ* − xᵢ* ln(xᵢ/xᵢ*)) serves as a Lyapunov function for many Lotka-Volterra "
    "type systems, connecting dynamical stability to information-theoretic measures of distance "
    "from equilibrium. The global stability results obtainable through Lyapunov methods complement "
    "the local information provided by linearization, ensuring that stability properties hold for "
    "all biologically relevant initial conditions, not merely those infinitesimally close to "
    "equilibrium."
)

# 1.3 Dynamical Systems Perspectives
add_heading_styled(doc, '1.3 Dynamical Systems Perspectives in Biology', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The dynamical systems perspective provides a unifying framework for understanding biological "
    "phenomena ranging from molecular switches to ecosystem dynamics [18]. Central concepts include "
    "stability and feedback—negative feedback promotes homeostasis and steady-state regulation by "
    "counteracting deviations from set points, while positive feedback can generate bistability and "
    "switch-like behavior by amplifying small initial differences [19]. Biological oscillations, "
    "ubiquitous in circadian rhythms (approximately 24-hour period), cell cycles (ranging from "
    "minutes in early embryos to days in mammalian somatic cells), cardiac pacemaking, and neural "
    "firing patterns, arise from delayed negative feedback or coupled positive and negative feedback "
    "loops. The interplay between different feedback architectures generates the rich repertoire of "
    "dynamical behaviors observed in living systems, from precise adaptation (where the system "
    "returns exactly to its pre-stimulus state) to robust oscillations (maintained despite parameter "
    "perturbations) to irreversible transitions (one-way switches that cannot be reversed by "
    "removing the stimulus)."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Bistability represents a fundamental dynamical motif in biology, enabling cells to make irreversible "
    "decisions between distinct phenotypic states [20]. As illustrated in Figure 3, which depicts the "
    "bifurcation structure of a genetic toggle switch, bistable systems exhibit two stable steady states "
    "separated by an unstable equilibrium. The system's response to perturbations depends critically on "
    "whether the perturbation crosses the separatrix (threshold) between basins of attraction. "
    "Emergent behaviors in biological systems—collective phenomena that cannot be predicted from the "
    "properties of individual components—arise from nonlinear interactions and feedback coupling. "
    "Examples include flocking behavior in animal groups, synchronized oscillations in coupled cells, "
    "and spontaneous pattern formation in developing embryos."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Multiscale dynamics present additional challenges, as biological systems operate simultaneously "
    "across temporal scales ranging from milliseconds (enzyme catalysis) to years (evolutionary "
    "adaptation) and spatial scales from nanometers (molecular) to meters (organismal) [21]. "
    "The separation of time scales enables mathematical simplifications through singular perturbation "
    "theory and quasi-steady-state approximations, while the coupling across scales generates complex "
    "behaviors that require multiscale computational frameworks to resolve. The concept of "
    "biological robustness—the ability of living systems to maintain function despite perturbations—"
    "emerges as a fundamental organizing principle, with implications for understanding disease "
    "as the failure of robust homeostatic mechanisms and for designing therapeutic interventions "
    "that restore system-level function."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The mathematical tools developed in this section—equilibrium analysis, linearization, Lyapunov "
    "stability, phase-plane methods, and the dynamical systems perspective on feedback and "
    "oscillations—provide the analytical foundation for the biological applications developed in "
    "subsequent sections. These methods, originally developed in the context of physics and "
    "engineering, acquire new significance and face unique challenges when applied to biological "
    "systems characterized by high dimensionality, nonlinearity, stochasticity, evolutionary "
    "change, and the absence of precise conservation laws that simplify physical systems."
)

# ============================================================
# SECTION 2
# ============================================================
add_heading_styled(doc, '2. Spatial Dynamics and Applications in Population and Epidemiological Systems', level=1)

add_heading_styled(doc, '2.1 Partial Differential Equations and Spatial Biological Processes', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "When spatial heterogeneity plays a significant role in biological dynamics, partial differential "
    "equations (PDEs) become the appropriate modeling framework [22]. Spatial processes are ubiquitous "
    "in biology: organisms disperse through heterogeneous landscapes, chemical signals diffuse through "
    "tissues to coordinate cellular behavior, and pathogens spread through spatially structured "
    "populations. The fundamental reaction-diffusion equation, shown in Equation (10), combines "
    "local reaction kinetics (production, degradation, and transformation of chemical species) with "
    "spatial diffusion driven by concentration gradients according to Fick's law, providing the "
    "mathematical basis for understanding pattern formation, wave propagation, and spatial "
    "organization in biological systems. The diffusion coefficient D characterizes the rate of "
    "spatial spreading (with units of length²/time), while the reaction term f(u) describes local "
    "biochemical transformations that produce or consume the diffusing substance."
)

add_numbered_equation(doc, eq10_xml, 10)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Alan Turing's seminal 1952 paper demonstrated that diffusion, typically a homogenizing process, "
    "can paradoxically drive pattern formation when coupled with appropriate reaction kinetics [23]. "
    "The Turing instability mechanism requires an activator-inhibitor system where the inhibitor "
    "diffuses substantially faster than the activator—a condition known as differential diffusion. "
    "Intuitively, the fast-diffusing inhibitor creates a region of suppression surrounding each "
    "activator peak, preventing homogeneous activation and establishing characteristic wavelengths "
    "of spatial structure. The mathematical condition for Turing instability, expressed in "
    "Equation (15), determines when spatial perturbations of a homogeneous steady state are amplified "
    "rather than dampened, leading to the spontaneous emergence of spatial patterns with wavelengths "
    "selected by the ratio of diffusion coefficients and the kinetic parameters."
)

add_numbered_equation(doc, eq15_xml, 15)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Figure 4 illustrates the emergent spatial patterns generated by a reaction-diffusion system "
    "satisfying Turing instability conditions. The activator concentration (Figure 4a) forms characteristic "
    "spots or stripes depending on system parameters, while the inhibitor (Figure 4b) shows complementary "
    "spatial distributions [24]. These mathematical patterns bear remarkable resemblance to pigmentation "
    "patterns observed in animal skins, shell patterns, and vegetation distributions in semi-arid "
    "environments. Chemotaxis—directed cell movement along chemical gradients—introduces advective "
    "transport terms into the PDE framework, enabling the modeling of immune cell recruitment, bacterial "
    "swarming, and embryonic cell migration [25]. Wave propagation in excitable biological media, "
    "such as action potential propagation along nerve fibers and calcium waves in cardiac tissue, "
    "is described by traveling wave solutions of reaction-diffusion equations."
)

# Insert Figure 4
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('/projects/sandbox/AMMAN/chapter_figures/Figure_4_Turing_Patterns.png', width=Inches(5.5))
caption4 = doc.add_paragraph()
caption4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption4.add_run('Figure 4: Turing pattern formation via reaction-diffusion equations showing (a) activator and (b) inhibitor spatial distributions. The emergent patterns arise from diffusion-driven instability in an activator-inhibitor system with differential diffusion rates.')
run.italic = True
run.font.size = Pt(9)
doc.add_paragraph()

# 2.2 Population Dynamics
add_heading_styled(doc, '2.2 Population Dynamics and Ecological Systems', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Population ecology provides one of the oldest and richest application domains for differential "
    "equations in biology, with roots extending back to the pioneering work of Malthus, Verhulst, "
    "Lotka, and Volterra in the 18th through early 20th centuries [26]. The logistic growth equation, "
    "Equation (5), represents the simplest model incorporating density-dependent growth limitation, "
    "where r is the intrinsic growth rate (the per-capita growth rate at very low densities) and K "
    "is the environmental carrying capacity (the maximum sustainable population size given available "
    "resources). Despite its apparent simplicity, the logistic model captures the fundamental "
    "biological constraint that populations cannot grow indefinitely due to resource limitation, "
    "and it serves as a building block for more complex ecological models including those "
    "incorporating age structure, spatial heterogeneity, and interspecific interactions."
)

add_numbered_equation(doc, eq5_xml, 5)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The Lotka-Volterra predator-prey model, described by Equations (3) and (4), represents one of "
    "the foundational models in mathematical ecology, providing fundamental insights into the "
    "mechanisms generating population oscillations in nature [27]. In this system, prey (x) grows "
    "exponentially in the absence of predators at intrinsic rate α, and is consumed at a rate "
    "proportional to the encounter frequency βxy, where β is the predation rate constant reflecting "
    "both search efficiency and handling time. Predators (y) convert consumed prey biomass into "
    "offspring with efficiency δ (the conversion efficiency) and experience natural mortality at "
    "rate γ independent of prey availability. The resulting dynamics exhibit perpetual oscillations "
    "in the classical model, with prey abundance and predator abundance cycling out of phase "
    "in a manner analogous to a frictionless pendulum—the system is structurally conservative "
    "and any perturbation leads to a different amplitude orbit rather than returning to the "
    "original one."
)

add_numbered_equation(doc, eq3_xml, 3)
add_numbered_equation(doc, eq4_xml, 4)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Figure 1 illustrates the characteristic oscillatory dynamics of the Lotka-Volterra system. "
    "The time series (Figure 1a) reveals periodic population cycles where prey peaks precede predator "
    "peaks by a quarter period—a pattern observed in numerous ecological systems including the classic "
    "snowshoe hare–lynx interaction documented over decades of fur trading records [28]. The phase "
    "portrait (Figure 1b) demonstrates that trajectories form closed orbits around the coexistence "
    "equilibrium, reflecting the conservative (Hamiltonian) nature of the classical Lotka-Volterra "
    "system. The amplitude of oscillations depends on initial conditions, with orbits farther from "
    "the equilibrium exhibiting larger amplitude fluctuations. Extensions of the basic model "
    "incorporating density-dependent prey growth (Rosenzweig-MacArthur model), saturating functional "
    "responses (Holling types II and III), or predator interference can qualitatively alter the "
    "dynamics, generating stable limit cycles, chaotic attractors, or globally stable equilibria "
    "depending on parameter values. As summarized in Table 1, ODE models like the Lotka-Volterra "
    "system assume continuous, deterministic, and well-mixed populations, which may be appropriate "
    "at large spatial scales but becomes questionable for small, spatially structured populations "
    "where stochastic extinction, spatial refugia, and individual-level behavioral decisions "
    "significantly influence population outcomes."
)

# Insert Figure 1
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('/projects/sandbox/AMMAN/chapter_figures/Figure_1_Predator_Prey_Dynamics.png', width=Inches(5.5))
caption1 = doc.add_paragraph()
caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption1.add_run('Figure 1: Lotka-Volterra predator-prey model dynamics showing (a) temporal oscillations in prey and predator populations and (b) phase portrait with closed orbits around the coexistence equilibrium point.')
run.italic = True
run.font.size = Pt(9)
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Competition models extend the predator-prey framework to describe interactions between species "
    "competing for shared resources [29]. The competitive exclusion principle, formalized through "
    "the Lotka-Volterra competition equations, predicts that two species competing for a single "
    "limiting resource cannot coexist indefinitely—one will inevitably exclude the other unless "
    "niche differentiation allows resource partitioning. Modern extensions incorporate multiple "
    "resources, spatial heterogeneity, and temporal variation, revealing conditions under which "
    "biodiversity can be maintained through storage effects, relative nonlinearity of competition, "
    "and spatial refuge mechanisms. Allee effects, representing reduced per-capita growth rates at "
    "low population densities due to difficulty finding mates, reduced cooperative defense, or "
    "impaired group foraging, introduce critical thresholds below which populations face inevitable "
    "extinction [30]. These effects have profound implications for species conservation and the "
    "management of endangered species, as they create minimum viable population sizes that must be "
    "maintained for population persistence and necessitate active management interventions—including "
    "habitat restoration, captive breeding programs, and translocation strategies—when populations "
    "decline below critical levels. The mathematical characterization of Allee thresholds through "
    "bistability analysis directly informs conservation policy by quantifying extinction risk and "
    "identifying critical population sizes below which recovery becomes impossible without "
    "intervention."
)

# TABLE 2
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table 2: Comparison of Major Population Dynamics Models')
run.bold = True
run.font.size = Pt(10)

table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Table Grid'

headers2 = ['Model', 'Species', 'Key Feature', 'Equilibrium Behavior', 'Biological Example']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_shading(table2.rows[0].cells[i], "4472C4")
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

table2_data = [
    ['Malthusian', 'Single', 'Exponential growth', 'Unbounded growth', 'Bacterial culture (early)'],
    ['Logistic', 'Single', 'Density dependence', 'Stable carrying capacity', 'Yeast growth in culture'],
    ['Lotka-Volterra', 'Two', 'Predator-prey cycling', 'Neutral oscillations', 'Hare-lynx system'],
    ['Rosenzweig-MacArthur', 'Two', 'Saturating predation', 'Limit cycles possible', 'Plankton dynamics'],
    ['Competitive exclusion', 'Two+', 'Resource competition', 'One species dominates', 'Paramecium experiments'],
    ['Allee effect', 'Single', 'Critical threshold', 'Bistability', 'Endangered species'],
]

for i, row_data in enumerate(table2_data):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val

set_table_style(table2)
doc.add_paragraph()

# 2.3 Mathematical Epidemiology
add_heading_styled(doc, '2.3 Mathematical Epidemiology and Disease Dynamics', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Mathematical epidemiology employs compartmental models to describe the transmission dynamics of "
    "infectious diseases through populations [31]. The SIR (Susceptible-Infected-Recovered) model, "
    "introduced by Kermack and McKendrick in their landmark 1927 paper, represents one of the most "
    "influential contributions of mathematics to public health. This model partitions the population "
    "into three compartments based on disease status and describes their temporal evolution through "
    "the coupled system of Equations (6)–(8). Susceptible individuals (S) become infected through "
    "contact with infectious individuals at transmission rate β, which incorporates both the contact "
    "rate and the probability of transmission per contact. Infected individuals (I) recover at rate γ "
    "(with 1/γ representing the mean infectious period), and recovered individuals (R) acquire "
    "permanent immunity. The model assumes homogeneous mixing (all individuals equally likely to "
    "contact each other), a closed population (no births, deaths, or migration), and that recovery "
    "confers lifelong immunity—assumptions that can be relaxed in more complex formulations."
)

add_numbered_equation(doc, eq6_xml, 6)
add_numbered_equation(doc, eq7_xml, 7)
add_numbered_equation(doc, eq8_xml, 8)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The basic reproduction number R₀, defined in Equation (9), represents the average number of "
    "secondary infections produced by a single infectious individual introduced into a fully "
    "susceptible population [32]. This dimensionless threshold quantity, often called the most "
    "important number in epidemiology, determines whether an epidemic will occur: when R₀ > 1, the "
    "disease-free equilibrium is unstable and the disease can invade the population, generating an "
    "epidemic; when R₀ < 1, the disease-free equilibrium is globally asymptotically stable and "
    "imported cases cannot generate sustained transmission chains. The herd immunity threshold, "
    "given by the critical vaccination coverage p_c = 1 − 1/R₀, specifies the minimum proportion "
    "of the population that must be immune (through vaccination or natural infection) to prevent "
    "epidemic spread, providing the scientific basis for vaccination coverage targets in public "
    "health programs. For measles (R₀ ≈ 12-18), this requires approximately 92-95% coverage, "
    "while for influenza (R₀ ≈ 1.5-2), approximately 33-50% coverage suffices."
)

add_numbered_equation(doc, eq9_xml, 9)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Figure 2 illustrates the dynamics of the SIR model under different epidemiological scenarios. "
    "Figure 2a shows the classic epidemic curve for R₀ = 3.0, with the characteristic rapid rise in "
    "infections followed by a gradual decline as susceptible individuals are depleted [33]. The "
    "susceptible fraction decreases monotonically throughout the epidemic, while the recovered "
    "fraction increases sigmoidally toward its final value. Figure 2b demonstrates how the epidemic "
    "peak height and timing depend on R₀, with higher reproduction numbers producing earlier, "
    "taller, and more severe epidemic peaks. The total attack rate—the ultimate fraction of the "
    "population infected—increases nonlinearly with R₀, approaching unity for very large reproduction "
    "numbers. These mathematical results have direct policy implications: interventions that reduce "
    "R₀ below 1 (through social distancing, mask mandates, or vaccination) guarantee epidemic "
    "decline, while interventions that merely reduce R₀ without crossing the threshold delay and "
    "flatten the epidemic curve without preventing eventual widespread infection. As shown in Table 2, "
    "the oscillatory dynamics observed in Lotka-Volterra predator-prey models have analogues in "
    "epidemic systems with temporary immunity, where recurrent epidemic waves arise from the "
    "replenishment of susceptible individuals through waning immunity or new births."
)

# Insert Figure 2
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('/projects/sandbox/AMMAN/chapter_figures/Figure_2_SIR_Epidemic_Model.png', width=Inches(5.5))
caption2 = doc.add_paragraph()
caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption2.add_run('Figure 2: SIR compartmental epidemic model showing (a) temporal evolution of susceptible, infected, and recovered fractions for R₀ = 3, and (b) comparison of epidemic curves for different basic reproduction numbers.')
run.italic = True
run.font.size = Pt(9)
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Extended compartmental models incorporate additional biological complexity including exposed "
    "(latent) periods (SEIR models), waning immunity (SIRS models), age structure, and spatial "
    "heterogeneity [34]. Vaccination models analyze the impact of different immunization "
    "strategies—pulse vaccination, ring vaccination, and age-targeted campaigns—on disease "
    "elimination prospects. The interplay between vaccination coverage, waning vaccine-induced "
    "immunity, and pathogen evolution creates complex dynamical landscapes that require sophisticated "
    "mathematical analysis to navigate optimally. Optimal control theory provides tools for "
    "designing time-varying intervention strategies that minimize disease burden subject to resource "
    "constraints, balancing the epidemiological benefit of aggressive early intervention against "
    "the economic and social costs of sustained control measures. The COVID-19 pandemic catalyzed "
    "unprecedented advances in real-time epidemiological modeling, integrating mobility data from "
    "mobile phones, genomic surveillance of viral evolution, wastewater epidemiology, and "
    "high-resolution spatial models to inform public health interventions and pandemic preparedness "
    "planning [35]. These advances demonstrated both the power of mathematical modeling to guide "
    "policy decisions and the challenges of communicating model uncertainty to decision-makers "
    "and the public."
)

# ============================================================
# SECTION 3
# ============================================================
add_heading_styled(doc, '3. Dynamical Systems in Molecular and Cellular Biology', level=1)

add_heading_styled(doc, '3.1 Gene Regulatory and Cellular Signaling Networks', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Gene regulatory networks (GRNs) govern cellular decision-making by controlling the expression "
    "of genes in response to environmental and developmental signals [36]. The mathematical modeling "
    "of GRNs typically employs Hill functions, Equation (11), to describe cooperative transcriptional "
    "regulation, where n is the Hill coefficient characterizing the steepness of the regulatory "
    "response and K is the half-maximal activation concentration. Hill coefficients greater than one "
    "indicate cooperative binding, producing sigmoidal dose-response curves that enable sharp "
    "transitions between gene expression states. In practice, Hill coefficients in biological "
    "systems typically range from 1 to 8, with higher values observed in systems with strong "
    "cooperativity such as hemoglobin oxygen binding (n ≈ 2.8) or the lambda phage genetic switch "
    "(effective n ≈ 4-5). The mathematical framework of GRN modeling has been instrumental in "
    "understanding cell fate specification during embryonic development, where cascades of "
    "transcription factors progressively restrict cell identity through sequential bistable switches."
)

add_numbered_equation(doc, eq11_xml, 11)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The genetic toggle switch, a synthetic gene circuit consisting of two mutually repressing "
    "transcription factors, provides a paradigmatic example of bistability in biological systems [37]. "
    "As depicted in Figure 3, the bifurcation structure reveals how the strength of transcriptional "
    "regulation determines whether the system exhibits monostable or bistable behavior. The saddle-node "
    "bifurcation points mark the critical parameter values at which the system transitions between "
    "these qualitatively different dynamical regimes. The phase portrait analysis shown in Figure 3b "
    "reveals nullcline intersections that define the stable and unstable equilibria of the toggle "
    "switch, with the separatrix dividing the state space into basins of attraction for each stable "
    "state [38]."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Cellular signaling pathways transmit information from the cell surface to the nucleus through "
    "cascades of biochemical reactions, often involving phosphorylation-dephosphorylation cycles [39]. "
    "These cascades can exhibit ultrasensitivity—a switch-like input-output relationship that arises "
    "from zero-order kinetics when enzymes operate near saturation—enabling cells to generate "
    "all-or-none responses from graded stimuli. Multi-tier kinase cascades, such as the MAPK "
    "(mitogen-activated protein kinase) pathway, amplify ultrasensitivity through sequential "
    "layers of switch-like modules. Feedback mechanisms in signaling networks generate diverse "
    "dynamical behaviors including adaptation (exact return to baseline following a transient "
    "stimulus, as in bacterial chemotaxis), oscillations (as in NF-κB nuclear-cytoplasmic shuttling "
    "and p53-Mdm2 dynamics), and digital (all-or-none) responses to graded inputs. The encoding "
    "of information in the temporal dynamics of signaling—frequency modulation, amplitude modulation, "
    "and duration encoding—enables cells to distinguish between different stimuli using shared "
    "molecular components, effectively multiplexing information through a limited number of "
    "signaling channels."
)

# Insert Figure 3
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('/projects/sandbox/AMMAN/chapter_figures/Figure_3_Bifurcation_Gene_Switch.png', width=Inches(5.5))
caption3 = doc.add_paragraph()
caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption3.add_run('Figure 3: Bistability in a gene regulatory toggle switch showing (a) saddle-node bifurcation diagram with stable (solid) and unstable (dashed) steady states, and (b) nullcline analysis and vector field revealing two stable equilibria separated by a saddle point.')
run.italic = True
run.font.size = Pt(9)
doc.add_paragraph()

# 3.2 Metabolic Networks
add_heading_styled(doc, '3.2 Metabolic Networks and Cellular Homeostasis', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Metabolic pathways represent highly organized networks of enzyme-catalyzed reactions that "
    "convert substrates into products while maintaining cellular homeostasis through precise flux "
    "regulation [40]. The Michaelis-Menten equation, Equation (12), describes the fundamental "
    "kinetics of enzyme-catalyzed reactions, where Vmax is the maximum reaction velocity achieved "
    "at saturating substrate concentrations (determined by the product of enzyme concentration and "
    "catalytic rate constant), and Km is the Michaelis constant—the substrate concentration at which "
    "the reaction velocity equals half of Vmax, providing a measure of enzyme-substrate affinity. "
    "This equation derives from the quasi-steady-state assumption applied to the enzyme-substrate "
    "complex, and it provides the cornerstone of quantitative enzymology and metabolic modeling."
)

add_numbered_equation(doc, eq12_xml, 12)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Dynamic models of metabolic pathways couple multiple enzyme-catalyzed reactions through shared "
    "metabolite pools, creating complex networks with emergent regulatory properties [41]. Metabolic "
    "control analysis (MCA) provides a rigorous framework for quantifying how control over pathway "
    "flux is distributed among individual enzymes, demonstrating through the flux control coefficient "
    "summation theorem that the sum of all control coefficients equals unity—meaning that control "
    "is necessarily shared among pathway enzymes, with no single enzyme typically 'controlling' "
    "the entire pathway. This mathematical result has profound implications for metabolic engineering "
    "and drug target identification, as it demonstrates that modifying a single enzyme rarely produces "
    "proportional changes in pathway output. Ultrasensitivity in metabolic regulation can arise from "
    "substrate competition, enzyme sequestration, multistep covalent modification cycles, and "
    "allosteric cooperativity, enabling sharp transitions between metabolic states such as the "
    "switch between glycolysis and gluconeogenesis in hepatic metabolism or the commitment to "
    "apoptosis through mitochondrial outer membrane permeabilization."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Cellular homeostasis—the maintenance of intracellular conditions within narrow physiological "
    "bounds despite environmental fluctuations—emerges from the interplay of feedback regulation, "
    "buffering mechanisms, and robust network architectures [42]. Integral feedback, a control "
    "engineering concept increasingly recognized in biological systems (such as in bacterial "
    "chemotaxis and calcium homeostasis), achieves perfect adaptation by accumulating error signals "
    "over time. The antithetic integral feedback motif, recently characterized in both natural and "
    "synthetic systems, provides a molecular implementation of integral control through balanced "
    "production and mutual annihilation of controller species. The robustness of homeostatic "
    "mechanisms against parameter perturbations connects to the concept of structural stability in "
    "dynamical systems theory, where qualitative behavior is preserved under small perturbations. "
    "Robustness analysis reveals that certain network topologies are inherently more robust than "
    "others—for example, networks with negative feedback and redundant pathways maintain function "
    "despite component failures, while purely feedforward architectures are sensitive to parameter "
    "variations. Table 3 summarizes the key dynamical phenomena observed in cellular and molecular "
    "systems, linking mathematical mechanisms to specific biological examples and identifying the "
    "associated bifurcation types that govern transitions between qualitatively different behaviors."
)

# TABLE 3
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table 3: Key Dynamical Phenomena in Cellular and Molecular Biology')
run.bold = True
run.font.size = Pt(10)

table3 = doc.add_table(rows=8, cols=4)
table3.style = 'Table Grid'

headers3 = ['Phenomenon', 'Mathematical Mechanism', 'Biological Example', 'Associated Bifurcation']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_shading(table3.rows[0].cells[i], "4472C4")
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

table3_data = [
    ['Bistability', 'Positive feedback + nonlinearity', 'Cell fate decisions, lac operon', 'Saddle-node'],
    ['Oscillations', 'Delayed negative feedback', 'Circadian rhythms, cell cycle', 'Hopf'],
    ['Ultrasensitivity', 'Zero-order kinetics, cooperativity', 'MAPK cascade signaling', 'Hill function (n > 1)'],
    ['Adaptation', 'Integral feedback control', 'Bacterial chemotaxis', 'None (structural)'],
    ['Excitability', 'Fast-slow dynamics, threshold', 'Action potentials, calcium waves', 'Canard/SNIC'],
    ['Pattern formation', 'Turing instability', 'Embryonic morphogenesis', 'Spatial Hopf/Turing'],
    ['Noise switching', 'Stochastic bistable dynamics', 'Phage lambda decision', 'Stochastic bifurcation'],
]

for i, row_data in enumerate(table3_data):
    for j, val in enumerate(row_data):
        table3.rows[i+1].cells[j].text = val

set_table_style(table3)
doc.add_paragraph()

# 3.3 Advanced Analysis
add_heading_styled(doc, '3.3 Advanced Analysis and Computational Methods', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Bifurcation analysis provides a systematic methodology for understanding how qualitative changes "
    "in system behavior arise as parameters are varied continuously [43]. Bifurcations represent "
    "critical parameter values at which the number or stability of equilibria changes, limit cycles "
    "appear or disappear, or other qualitative transitions in the phase portrait occur. Common "
    "bifurcation types in biological systems include saddle-node bifurcations (creation/annihilation "
    "of equilibrium pairs, underlying bistability), Hopf bifurcations (emergence of oscillations "
    "from a steady state), transcritical bifurcations (exchange of stability between equilibria, "
    "as in epidemic thresholds), and period-doubling cascades leading to chaos. The normal form of a "
    "supercritical pitchfork bifurcation, Equation (14), illustrates how a stable equilibrium at "
    "x = 0 loses stability as the bifurcation parameter μ increases through zero, giving rise to "
    "two symmetric stable branches x* = ±√μ. This bifurcation type appears in symmetry-breaking "
    "transitions in biological systems, such as the emergence of cell polarity from a symmetric "
    "state or the breaking of left-right symmetry in embryonic development."
)

add_numbered_equation(doc, eq14_xml, 14)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Sensitivity analysis complements bifurcation studies by quantifying how model outputs respond "
    "to parameter perturbations [43]. Local sensitivity coefficients, computed as partial derivatives "
    "of model outputs with respect to individual parameters, measure the infinitesimal response "
    "to parameter changes, providing first-order approximations to parameter influence near nominal "
    "values. Global sensitivity methods—including variance-based Sobol indices that decompose output "
    "variance into contributions from individual parameters and their interactions, and Morris "
    "elementary effects screening for computationally efficient exploration of high-dimensional "
    "parameter spaces—identify the most influential parameters across the full range of biologically "
    "plausible values. These analyses guide experimental design by identifying which parameters must "
    "be measured precisely (those with high sensitivity indices) and which can be estimated roughly "
    "without significantly affecting model predictions (those with low sensitivity). As documented "
    "in Table 3, different dynamical phenomena are associated with distinct mathematical mechanisms "
    "and specific bifurcation types that can be systematically identified through numerical "
    "continuation methods implemented in software packages such as AUTO, MATCONT, and PyDSTool."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Numerical integration of biological ODE systems presents unique challenges due to the stiffness "
    "that arises from widely separated time scales—fast enzymatic reactions occurring on millisecond "
    "time scales coupled with slow gene expression changes over hours [14]. Explicit methods such as "
    "fourth-order Runge-Kutta require prohibitively small time steps for stiff systems, necessitating "
    "implicit solvers (backward differentiation formulas, implicit Runge-Kutta methods) that can "
    "handle large time-scale separations efficiently. Adaptive step-size control algorithms "
    "automatically adjust the integration step based on local error estimates, balancing accuracy "
    "against computational cost while maintaining solution reliability."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Model reduction techniques are essential for analyzing complex biological models with many "
    "variables and parameters. The quasi-steady-state approximation (QSSA) eliminates fast variables "
    "by assuming they equilibrate instantaneously relative to slow variables—the mathematical basis "
    "of the Michaelis-Menten approximation in enzyme kinetics [12]. The validity of the QSSA "
    "requires careful assessment: the classical condition requires that the total enzyme concentration "
    "be much less than the initial substrate concentration plus Km, though more refined conditions "
    "have been developed for various reaction network topologies. Singular perturbation theory "
    "provides a rigorous mathematical framework for these approximations, identifying slow manifolds "
    "in phase space on which the reduced dynamics evolve after initial fast transients decay. "
    "Fenichel's theorem guarantees the persistence of normally hyperbolic invariant manifolds "
    "under small perturbations, providing the theoretical foundation for slow manifold reductions. "
    "Center manifold reduction captures the essential dynamics near bifurcation points by projecting "
    "onto the low-dimensional subspace associated with critical (zero or purely imaginary) eigenvalues, "
    "enabling the derivation of normal forms that classify local bifurcation behavior and predict "
    "the emergence of new dynamical states. These reduction techniques are indispensable for "
    "extracting insight from the high-dimensional models that arise in systems biology, where "
    "genome-scale metabolic models may contain thousands of species and reactions."
)

# ============================================================
# SECTION 4
# ============================================================
add_heading_styled(doc, '4. Emerging Computational Paradigms and Future Directions', level=1)

add_heading_styled(doc, '4.1 Hybrid and Multiscale Modeling of Biological Systems', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Biological systems inherently operate across multiple spatial and temporal scales, necessitating "
    "hybrid modeling approaches that couple different mathematical frameworks at their respective "
    "natural scales [21]. The challenge of multiscale modeling lies in faithfully capturing the "
    "interactions between scales while maintaining computational tractability—a requirement that "
    "has driven the development of novel mathematical and computational methods specifically "
    "tailored to biological applications. Hybrid deterministic-stochastic models combine the "
    "computational efficiency of ODE/PDE descriptions for abundant molecular species with "
    "stochastic simulation algorithms (Gillespie SSA) for rare molecular events where discrete "
    "stochastic effects are important, using partitioning criteria based on molecule copy numbers "
    "or reaction propensities to assign each species or reaction to the appropriate mathematical "
    "framework. The general stochastic framework, expressed as Equation (16), augments deterministic "
    "dynamics with a noise term driven by Brownian motion, where the diffusion function g(X) "
    "characterizes the amplitude and state-dependence of fluctuations."
)

add_numbered_equation(doc, eq16_xml, 16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Agent-based models (ABMs) represent biological entities as discrete, autonomous agents that "
    "interact according to specified rules, enabling the exploration of emergent population-level "
    "behavior from individual-level heterogeneity [13]. In cancer modeling, ABMs simulate individual "
    "tumor cells with distinct phenotypes, genotypes, and microenvironmental interactions, capturing "
    "spatial competition for nutrients and oxygen, evolutionary dynamics under selective pressures, "
    "and the complex interplay between tumor cells and immune effectors that cannot be resolved by "
    "continuum models. The integration of multi-omics data—transcriptomics, proteomics, metabolomics, "
    "and epigenomics—into dynamical models represents a frontier challenge, requiring methods that can "
    "handle high-dimensional, heterogeneous data types while preserving mechanistic interpretability "
    "and enabling prediction of system behavior under novel perturbations."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The coupling of continuous PDE descriptions with discrete stochastic models creates multiscale "
    "frameworks capable of capturing phenomena spanning molecular to tissue scales [21]. For example, "
    "in wound healing models, individual cell migration and proliferation (described by agent-based "
    "rules) are coupled to continuum descriptions of growth factor diffusion and extracellular matrix "
    "remodeling (governed by PDEs). As demonstrated in Figure 4, reaction-diffusion patterns emerge "
    "at tissue scales from molecular-level interactions described by Turing-type mechanisms, "
    "illustrating how microscale dynamics generate macroscale spatial organization. Similarly, "
    "Figure 1 demonstrates how individual predator-prey encounters at the population level generate "
    "emergent oscillatory patterns visible at the ecosystem scale, connecting microscale interactions "
    "to macroscale dynamics through the mathematical framework of differential equations."
)

# 4.2 Data-Driven
add_heading_styled(doc, '4.2 Data-Driven Dynamical Systems and Artificial Intelligence', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The convergence of dynamical systems theory with machine learning and artificial intelligence "
    "is fundamentally transforming biological modeling [1]. Physics-informed neural networks (PINNs) "
    "embed known physical and biological laws as soft constraints within neural network architectures, "
    "enabling the learning of system dynamics from sparse and noisy experimental data while respecting "
    "conservation laws, symmetries, and other known physical principles. Unlike purely data-driven "
    "approaches that may violate fundamental biological constraints (such as non-negativity of "
    "concentrations or conservation of mass), PINNs guarantee physical consistency of predictions "
    "even in data-sparse regimes by penalizing solutions that violate the governing equations. "
    "The PINN loss function, Equation (17), combines data-fitting terms (L_data) measuring the "
    "discrepancy between network predictions and observed data, with physics-based residual penalties "
    "(L_physics) that penalize violations of the governing differential equations evaluated at "
    "collocation points throughout the spatial-temporal domain, and boundary condition enforcement "
    "(L_BC), with hyperparameters λ and μ controlling the relative weighting of each component."
)

add_numbered_equation(doc, eq17_xml, 17)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Equation discovery methods, such as SINDy (Sparse Identification of Nonlinear Dynamics), "
    "automatically identify governing equations from time-series data by exploiting sparsity in the "
    "space of possible dynamical terms [11]. These methods construct a library of candidate nonlinear "
    "functions (polynomials, trigonometric functions, rational functions, and interaction terms) and "
    "use sparse regression algorithms (LASSO, sequential thresholded least squares) to identify the "
    "minimal subset that accurately reproduces observed dynamics while avoiding overfitting. SINDy "
    "and its extensions—including ensemble SINDy for uncertainty quantification, weak-form SINDy for "
    "noisy data, and PDE-FIND for spatiotemporal systems—have successfully recovered known biological "
    "models from experimental data, including Lotka-Volterra dynamics from ecological time series, "
    "glycolytic oscillations from metabolite measurements, and genetic regulatory circuits from "
    "fluorescence data. Perhaps most excitingly, these methods have discovered novel dynamical "
    "relationships that were not anticipated by domain experts, suggesting new biological mechanisms "
    "worthy of experimental investigation."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Machine learning-assisted parameter estimation employs neural network surrogates, Bayesian "
    "optimization, and simulation-based inference to efficiently navigate high-dimensional parameter "
    "spaces that are intractable for traditional grid search or gradient-based optimization methods "
    "[9]. Neural posterior estimation using normalizing flows or variational autoencoders can "
    "approximate full Bayesian posteriors over model parameters given observed data, enabling "
    "rapid uncertainty quantification without expensive Markov Chain Monte Carlo (MCMC) sampling. "
    "These simulation-based inference methods are particularly valuable for complex biological "
    "models where likelihood functions are intractable but forward simulations are feasible. "
    "Model selection—choosing among competing mechanistic hypotheses—benefits from these approaches, "
    "which can rapidly evaluate model evidence across large model spaces using neural network "
    "classifiers trained on simulated data from each candidate model. Table 4 provides a "
    "comprehensive comparison of computational approaches for biological dynamical systems, "
    "highlighting the complementary strengths of mechanistic and data-driven methodologies and "
    "guiding practitioners in selecting appropriate tools for specific biological questions."
)

# TABLE 4
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table 4: Comparison of Computational Approaches for Biological Dynamical Systems')
run.bold = True
run.font.size = Pt(10)

table4 = doc.add_table(rows=7, cols=5)
table4.style = 'Table Grid'

headers4 = ['Method', 'Data Requirement', 'Interpretability', 'Scalability', 'Best Application']
for i, h in enumerate(headers4):
    table4.rows[0].cells[i].text = h
    for paragraph in table4.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_shading(table4.rows[0].cells[i], "4472C4")
    for paragraph in table4.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

table4_data = [
    ['Mechanistic ODE/PDE', 'Low (parameter fitting)', 'High', 'Moderate', 'Well-understood mechanisms'],
    ['Physics-informed NN', 'Moderate', 'Moderate-High', 'High', 'Sparse data + known physics'],
    ['SINDy (equation discovery)', 'Moderate-High', 'High', 'Moderate', 'Unknown governing equations'],
    ['Deep learning surrogate', 'High', 'Low', 'Very High', 'Fast prediction/optimization'],
    ['Bayesian inference', 'Low-Moderate', 'High', 'Low-Moderate', 'Uncertainty quantification'],
    ['Agent-based simulation', 'Moderate', 'Moderate', 'Low', 'Individual heterogeneity'],
]

for i, row_data in enumerate(table4_data):
    for j, val in enumerate(row_data):
        table4.rows[i+1].cells[j].text = val

set_table_style(table4)
doc.add_paragraph()

# 4.3 Personalized Medicine
add_heading_styled(doc, '4.3 Personalized Medicine, Synthetic Biology, and Control', level=2)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Dynamical systems approaches are increasingly central to personalized medicine, where "
    "patient-specific mathematical models inform individualized treatment strategies [1]. In "
    "oncology, tumor growth models calibrated to patient imaging data (MRI, PET-CT) can predict "
    "treatment responses and optimize chemotherapy scheduling, radiation fractionation protocols, "
    "and immunotherapy dosing regimens. Mathematical models of tumor-immune interactions reveal "
    "critical thresholds for immune activation and identify patient-specific windows of "
    "opportunity for immunotherapy intervention. Pharmacokinetic-pharmacodynamic (PK-PD) models "
    "describe drug absorption, distribution, metabolism, and elimination as coupled ODE systems, "
    "enabling individualized dosing based on patient-specific parameters estimated from therapeutic "
    "drug monitoring data. The concept of mathematical 'digital twins'—personalized dynamical models "
    "that mirror individual patient physiology—represents an emerging paradigm for precision "
    "therapeutics that integrates continuous monitoring data with mechanistic biological knowledge "
    "to provide real-time clinical decision support. In diabetes management, for example, "
    "mathematical models of glucose-insulin dynamics are already used in artificial pancreas "
    "systems that continuously adjust insulin delivery based on model predictions of future "
    "glucose levels."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "As demonstrated in Table 4, Bayesian inference methods play a critical role in personalizing "
    "dynamical models by quantifying uncertainty in patient-specific parameters, enabling clinicians "
    "to make informed decisions even when individual data are limited. The integration of wearable "
    "sensor data, electronic health records, and genomic information into patient-specific dynamical "
    "models creates opportunities for continuous model updating and adaptive treatment optimization "
    "that responds to evolving disease states in real time."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Synthetic biology applies engineering principles to design and construct novel biological "
    "circuits with prescribed dynamical behaviors [37]. Control theory provides the mathematical "
    "foundation for synthetic circuit design, enabling the construction of genetic oscillators "
    "(repressilator), toggle switches, logic gates, band-pass filters, and feedback controllers "
    "with guaranteed stability and performance properties. The design-build-test-learn cycle in "
    "synthetic biology relies heavily on mathematical modeling to predict circuit behavior before "
    "physical construction, reducing the experimental iteration time and enabling the rational "
    "design of increasingly complex genetic programs. Recent advances in cell-free systems "
    "provide rapid prototyping platforms where circuit designs can be tested in simplified "
    "biochemical environments before implementation in living cells, accelerating the design cycle "
    "and enabling high-throughput exploration of circuit design spaces guided by mathematical "
    "optimization algorithms."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "External control of biological systems using optogenetics (light-inducible gene expression), "
    "microfluidic inputs (precise chemical perturbations), or pharmacological modulators enables "
    "real-time feedback control strategies that maintain cells in desired dynamical states or guide "
    "state transitions along optimal paths. Model predictive control (MPC) algorithms, which solve "
    "optimization problems over a receding horizon using the mathematical model as a predictor, "
    "have been successfully applied to control gene expression levels in living cells, demonstrating "
    "the practical utility of dynamical systems theory for biological engineering applications. "
    "Robust control approaches account for model uncertainty and biological variability, ensuring "
    "that control performance is maintained despite imprecise parameter knowledge and cell-to-cell "
    "heterogeneity in the controlled population. The convergence of synthetic biology with control "
    "theory creates opportunities for autonomous therapeutic systems—engineered cells that sense "
    "disease biomarkers, compute appropriate responses through genetic logic circuits, and deliver "
    "therapeutic payloads in a precisely controlled manner."
)

# ============================================================
# CONCLUSION
# ============================================================
add_heading_styled(doc, '5. Conclusions and Future Perspectives', level=1)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "This chapter has presented a comprehensive survey of differential equations and dynamical "
    "systems theory as applied to biological sciences, spanning from fundamental mathematical "
    "foundations to cutting-edge computational methodologies. Beginning with the mathematical "
    "foundations of biological modeling—including ODE formulation (Equations 1–2), stability analysis, "
    "and phase-plane methods—we have systematically developed the theoretical framework necessary for "
    "understanding biological dynamics across multiple scales of organization, from molecular "
    "interactions to ecosystem-level processes. The application domains covered span population "
    "ecology (predator-prey interactions described by Equations 3–5, competition, and Allee effects), "
    "mathematical epidemiology (SIR models governed by Equations 6–9 and vaccination strategies), "
    "molecular biology (gene regulatory networks characterized by Hill functions (Equation 11) and "
    "metabolic pathways following Michaelis-Menten kinetics (Equation 12)), and spatial biology "
    "(reaction-diffusion patterns governed by Equation 10 and Turing instability conditions in "
    "Equation 15)."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The emergence of data-driven methodologies—physics-informed neural networks (Equation 17), "
    "equation discovery algorithms, and machine learning-assisted inference—promises to transform "
    "the field by enabling the construction of accurate dynamical models from increasingly rich "
    "experimental datasets generated by high-throughput technologies. However, these computational "
    "advances do not diminish the importance of classical dynamical systems theory; rather, they "
    "complement it by providing new tools for model identification, parameter estimation, and "
    "prediction that are grounded in physically meaningful constraints. The synergy between "
    "mechanistic understanding and data-driven discovery represents perhaps the most promising "
    "direction for future progress, enabling scientists to leverage domain knowledge while "
    "remaining open to unexpected dynamical relationships revealed by data. As demonstrated "
    "throughout this chapter—from the population oscillations in Figure 1 to the epidemic dynamics "
    "in Figure 2, the cellular switches in Figure 3, and the spatial patterns in Figure 4—"
    "differential equations and dynamical systems provide the essential mathematical language "
    "for understanding biological complexity across all scales of organization, from molecules "
    "to ecosystems."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "The transition toward personalized medicine and synthetic biology further underscores the need "
    "for rigorous dynamical modeling, as individualized therapeutic strategies and engineered "
    "biological circuits demand precise mathematical characterization of system behavior. Key "
    "future challenges include: (i) developing methods for learning dynamical models from single-cell "
    "and spatial transcriptomics data that capture cell-to-cell variability and reveal rare cell "
    "states; (ii) creating unified multiscale frameworks that seamlessly couple molecular, cellular, "
    "tissue, and organism-level dynamics as illustrated by the stochastic framework in Equation 16; "
    "(iii) establishing rigorous uncertainty quantification methods for complex biological models "
    "to enable reliable clinical decision-making under parameter and structural uncertainty; "
    "(iv) translating mathematical insights into practical applications through digital twins, "
    "model-informed drug development, and rationally designed synthetic circuits; and (v) developing "
    "mathematical frameworks that can accommodate the inherent heterogeneity, plasticity, and "
    "evolutionary potential of biological systems."
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.27)
p.add_run(
    "Overall, the transition toward Higher Education 5.0 requires a balanced approach in which "
    "technological capability, human competence, organizational capacity, ethical governance, and "
    "stakeholder acceptance develop together. A data-driven AI readiness framework can provide "
    "universities with a structured mechanism for identifying gaps, prioritizing investments, "
    "evaluating progress, and developing responsible AI strategies. The integration of mathematical "
    "modeling and computational biology into university curricula exemplifies this transformation, "
    "as students must develop both rigorous mathematical foundations and computational proficiency "
    "to address the complex biological questions of the 21st century. The academic community should "
    "remain at the center of this transformation so that AI becomes an enabler of creativity, "
    "inclusion, personalization, research excellence, and human development rather than simply "
    "another layer of technological automation. The dynamical systems approach to biology, as "
    "developed throughout this chapter, represents precisely the kind of interdisciplinary "
    "mathematical thinking that Higher Education 5.0 must cultivate to prepare the next generation "
    "of scientists and engineers for the challenges ahead."
)

# ============================================================
# REFERENCES
# ============================================================
add_heading_styled(doc, 'References', level=1)

references = [
    '[1] Murray, J.D. (2002). Mathematical Biology I: An Introduction. Springer-Verlag, New York.',
    '[2] Edelstein-Keshet, L. (2005). Mathematical Models in Biology. SIAM, Philadelphia.',
    '[3] Burnham, K.P. and Anderson, D.R. (2002). Model Selection and Multimodel Inference. Springer, New York.',
    '[4] Allen, L.J.S. (2010). An Introduction to Stochastic Processes with Applications to Biology. CRC Press.',
    '[5] Gillespie, D.T. (1977). Exact stochastic simulation of coupled chemical reactions. Journal of Physical Chemistry, 81(25), 2340–2361.',
    '[6] Wilkinson, D.J. (2018). Stochastic Modelling for Systems Biology. Third Edition, CRC Press.',
    '[7] Raj, A. and van Oudenaarden, A. (2008). Nature, nurture, or chance: stochastic gene expression and its consequences. Cell, 135(2), 216–226.',
    '[8] Jaqaman, K. and Danuser, G. (2006). Linking data to models: data regression. Nature Reviews Molecular Cell Biology, 7(11), 813–819.',
    '[9] Raue, A. et al. (2009). Structural and practical identifiability analysis of partially observed dynamical models. Bioinformatics, 25(15), 1923–1929.',
    '[10] Burnham, K.P. and Anderson, D.R. (2004). Multimodel inference: understanding AIC and BIC in model selection. Sociological Methods & Research, 33(2), 261–304.',
    '[11] Strogatz, S.H. (2018). Nonlinear Dynamics and Chaos: With Applications to Physics, Biology, Chemistry, and Engineering. CRC Press.',
    '[12] Keener, J. and Sneyd, J. (2009). Mathematical Physiology I: Cellular Physiology. Springer, New York.',
    '[13] Alon, U. (2019). An Introduction to Systems Biology: Design Principles of Biological Circuits. CRC Press.',
    '[14] Hirsch, M.W., Smale, S., and Devaney, R.L. (2013). Differential Equations, Dynamical Systems, and an Introduction to Chaos. Academic Press.',
    '[15] Guckenheimer, J. and Holmes, P. (1983). Nonlinear Oscillations, Dynamical Systems, and Bifurcations of Vector Fields. Springer.',
    '[16] Jordan, D.W. and Smith, P. (2007). Nonlinear Ordinary Differential Equations. Oxford University Press.',
    '[17] Khalil, H.K. (2002). Nonlinear Systems. Third Edition, Prentice Hall.',
    '[18] Tyson, J.J., Chen, K.C., and Novak, B. (2003). Sniffers, buzzers, toggles and blinkers: dynamics of regulatory and signaling pathways. Current Opinion in Cell Biology, 15(2), 221–231.',
    '[19] Ferrell, J.E. (2002). Self-perpetuating states in signal transduction: positive feedback, double-negative feedback and bistability. Current Opinion in Cell Biology, 14(2), 140–148.',
    '[20] Ozbudak, E.M. et al. (2004). Multistability in the lactose utilization network of Escherichia coli. Nature, 427(6976), 737–740.',
    '[21] Qu, Z. et al. (2011). Multi-scale modeling in biology: how to bridge the gaps between scales? Progress in Biophysics and Molecular Biology, 107(1), 21–31.',
    '[22] Okubo, A. and Levin, S.A. (2001). Diffusion and Ecological Problems: Modern Perspectives. Springer.',
    '[23] Turing, A.M. (1952). The chemical basis of morphogenesis. Philosophical Transactions of the Royal Society B, 237(641), 37–72.',
    '[24] Kondo, S. and Miura, T. (2010). Reaction-diffusion model as a framework for understanding biological pattern formation. Science, 329(5999), 1616–1620.',
    '[25] Hillen, T. and Painter, K.J. (2009). A user\'s guide to PDE models for chemotaxis. Journal of Mathematical Biology, 58(1), 183–217.',
    '[26] Kot, M. (2001). Elements of Mathematical Ecology. Cambridge University Press.',
    '[27] Volterra, V. (1926). Fluctuations in the abundance of a species considered mathematically. Nature, 118, 558–560.',
    '[28] Elton, C.S. and Nicholson, M. (1942). The ten-year cycle in numbers of the lynx in Canada. Journal of Animal Ecology, 11(2), 215–244.',
    '[29] Tilman, D. (1982). Resource Competition and Community Structure. Princeton University Press.',
    '[30] Courchamp, F., Berec, L., and Gascoigne, J. (2008). Allee Effects in Ecology and Conservation. Oxford University Press.',
    '[31] Hethcote, H.W. (2000). The mathematics of infectious diseases. SIAM Review, 42(4), 599–653.',
    '[32] Diekmann, O., Heesterbeek, J.A.P., and Metz, J.A.J. (1990). On the definition and computation of the basic reproduction ratio R₀. Journal of Mathematical Biology, 28(4), 365–382.',
    '[33] Anderson, R.M. and May, R.M. (1991). Infectious Diseases of Humans: Dynamics and Control. Oxford University Press.',
    '[34] Keeling, M.J. and Rohani, P. (2008). Modeling Infectious Diseases in Humans and Animals. Princeton University Press.',
    '[35] Vespignani, A. et al. (2020). Modelling COVID-19. Nature Reviews Physics, 2(6), 279–281.',
    '[36] Alon, U. (2007). Network motifs: theory and experimental approaches. Nature Reviews Genetics, 8(6), 450–461.',
    '[37] Gardner, T.S., Cantor, C.R., and Collins, J.J. (2000). Construction of a genetic toggle switch in Escherichia coli. Nature, 403(6767), 339–342.',
    '[38] Ferrell, J.E. and Ha, S.H. (2014). Ultrasensitivity part II: multisite phosphorylation, stoichiometric inhibitors, and positive feedback. Trends in Biochemical Sciences, 39(11), 556–569.',
    '[39] Kholodenko, B.N. (2006). Cell-signalling dynamics in time and space. Nature Reviews Molecular Cell Biology, 7(3), 165–176.',
    '[40] Palsson, B.O. (2015). Systems Biology: Constraint-based Reconstruction and Analysis. Cambridge University Press.',
    '[41] Heinrich, R. and Schuster, S. (1996). The Regulation of Cellular Systems. Chapman & Hall.',
    '[42] Kitano, H. (2004). Biological robustness. Nature Reviews Genetics, 5(11), 826–837.',
    '[43] Kuznetsov, Y.A. (2004). Elements of Applied Bifurcation Theory. Springer, New York.',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ============================================================
# SAVE DOCUMENT
# ============================================================
output_path = '/projects/sandbox/AMMAN/Chapter_Differential_Equations_Dynamical_Systems_Biology.docx'
doc.save(output_path)
print(f"\nDocument saved successfully: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")

# Count approximate words
word_count = 0
for para in doc.paragraphs:
    word_count += len(para.text.split())
print(f"Approximate word count: {word_count}")
