"""
Generate Word Document: Chapter - Internet of Things (IoT) & Smart Cities
Complete with 43 references, 4 tables, 4 figures (~8300 words)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Page setup
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_figure(image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = cap.add_run(caption)
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.name = 'Times New Roman'



# ============ TITLE ============
title = doc.add_heading('Internet of Things (IoT) & Smart Cities', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

# ============ ABSTRACT ============
add_heading_styled('Abstract', level=1)

abstract_text = (
    "The Internet of Things (IoT) represents a transformative paradigm that is fundamentally reshaping urban environments "
    "into intelligent, responsive, and sustainable ecosystems. This chapter provides a comprehensive examination of IoT architecture "
    "and its integration within smart city frameworks, encompassing the full technology stack from perception layers through edge, "
    "fog, and cloud computing paradigms. The discussion extends to advanced communication technologies including 5G/6G networks, "
    "Low-Power Wide-Area Networks (LPWAN), mesh topologies, and Time-Sensitive Networking (TSN) protocols that form the backbone "
    "of urban connectivity. The chapter explores artificial intelligence, machine learning, and digital twin technologies as "
    "enabling tools for intelligent urban management. Practical applications across smart transportation, energy systems, water "
    "distribution, waste management, and environmental monitoring are examined in detail. Furthermore, the chapter addresses "
    "critical challenges in governance, cybersecurity, privacy, interoperability, and citizen-centric service delivery. "
    "Strategic pathways for future development including AI-driven autonomous city management, blockchain-based trust frameworks, "
    "and resilient human-centric urban models are presented. The analysis synthesizes current research findings and industry "
    "developments to provide a holistic perspective on how IoT technologies can deliver competitive advantage while ensuring "
    "sustainability, equity, and resilience in urban environments."
)
add_para(abstract_text)

add_para("Keywords: Internet of Things, Smart Cities, Edge Computing, Digital Twins, 5G/6G, Urban Infrastructure, "
         "Artificial Intelligence, Cybersecurity, Sustainable Development, Autonomous Systems", italic=True)



# ============ SECTION 1 ============
add_heading_styled('1. IoT Architecture and Digital Transformation of Smart Cities', level=1)

sec_1_intro = (
    "The digital transformation of urban environments through Internet of Things technologies represents "
    "one of the most significant sociotechnical shifts of the twenty-first century. As global urbanization "
    "continues to accelerate, with projections indicating that 68% of the world population will reside in "
    "cities by 2050, the imperative for intelligent urban management systems becomes increasingly urgent. "
    "IoT-enabled smart cities leverage interconnected sensors, advanced computing architectures, and "
    "artificial intelligence to optimize resource utilization, enhance service delivery, improve "
    "environmental sustainability, and elevate quality of life for urban inhabitants. This section "
    "examines the foundational architectural elements that enable smart city functionality, from the "
    "physical sensing layer through distributed computing paradigms to the communication technologies "
    "that bind these elements into cohesive systems."
)
add_para(sec_1_intro)

add_heading_styled('1.1 IoT Architecture: Sensors, Connectivity, Edge-Fog-Cloud Computing', level=2)

sec_1_1_p1 = (
    "The Internet of Things (IoT) architecture for smart cities represents a sophisticated multi-layered framework "
    "designed to seamlessly integrate billions of heterogeneous devices into cohesive urban intelligence systems [1]. "
    "At the foundation lies the perception layer, comprising diverse sensors, actuators, and embedded systems that "
    "continuously monitor urban parameters including temperature, humidity, air quality, traffic flow, noise levels, "
    "and structural integrity [2]. Modern IoT sensors have evolved significantly, with micro-electromechanical systems "
    "(MEMS) achieving unprecedented miniaturization while maintaining accuracy levels exceeding 99.5% for critical "
    "measurements [3]. The proliferation of low-cost sensor nodes, now available at sub-dollar price points, has "
    "enabled massive-scale deployments with cities like Barcelona deploying over 19,500 sensors across metropolitan "
    "infrastructure [4]."
)
add_para(sec_1_1_p1)

sec_1_1_p2 = (
    "The edge computing layer represents the first tier of computational intelligence, positioned at or near data "
    "generation sources to enable real-time processing with latencies below 10 milliseconds [5]. Edge nodes perform "
    "critical functions including data filtering, anomaly detection, and local decision-making without requiring "
    "round-trip communication to centralized servers. This architectural approach reduces backbone network traffic "
    "by 60-75% while enabling time-critical applications such as autonomous vehicle coordination and emergency "
    "response systems [6]. The fog computing layer extends this distributed intelligence by providing intermediate "
    "processing capabilities at regional aggregation points, typically co-located with telecommunications "
    "infrastructure [7]. Fog nodes facilitate data fusion from multiple edge sources, perform pattern recognition "
    "across geographic clusters, and manage load balancing to optimize resource utilization. The cloud computing "
    "layer provides virtually unlimited computational resources for big data analytics, artificial intelligence "
    "model training, and long-term data archival [8]. This hierarchical architecture, illustrated in Figure 1, "
    "enables smart cities to process an estimated 2.5 quintillion bytes of urban data daily while maintaining "
    "service quality across diverse application requirements."
)
add_para(sec_1_1_p2)



# Insert Figure 1
add_figure('/projects/sandbox/AMMAN/iot_figures/Figure_1_IoT_Architecture.png',
           'Figure 1: Multi-Layer IoT Architecture for Smart City Infrastructure')

sec_1_1_p3 = (
    "The convergence of edge, fog, and cloud computing within IoT architectures has given rise to the concept of "
    "computational continuum, where processing tasks are dynamically allocated based on latency requirements, "
    "data sensitivity, and available resources [9]. Modern smart city platforms employ container orchestration "
    "frameworks such as Kubernetes adapted for IoT workloads, enabling seamless migration of computational tasks "
    "across architectural layers. Research indicates that optimized task offloading algorithms can reduce overall "
    "system energy consumption by 40-55% compared to cloud-only architectures while maintaining quality-of-service "
    "guarantees [10]. The integration of hardware security modules at each architectural layer ensures end-to-end "
    "data integrity, with trusted execution environments providing isolated processing for sensitive urban data "
    "including citizen biometrics and financial transactions."
)
add_para(sec_1_1_p3)

sec_1_1_p4 = (
    "The scalability of IoT architectures is fundamentally dependent on efficient resource management across "
    "the computing continuum. Microservices architectures decompose monolithic smart city applications into "
    "independently deployable components, each optimized for specific operational requirements. Service mesh "
    "technologies provide observability, traffic management, and security policies across distributed IoT "
    "microservices without requiring application-level modifications. The adoption of serverless computing "
    "paradigms at the edge layer enables event-driven processing of IoT data streams with automatic scaling "
    "from zero to millions of concurrent invocations, eliminating the need for pre-provisioned infrastructure. "
    "Digital data management strategies employ tiered storage architectures where hot data resides at the edge "
    "for immediate access, warm data is maintained in fog-layer databases for recent analytics, and cold data "
    "is archived in cost-efficient cloud storage for historical analysis and compliance requirements. This "
    "architectural sophistication enables smart cities to manage the exponential growth in IoT data volumes, "
    "projected to reach 79.4 zettabytes globally by 2025, while controlling infrastructure costs and "
    "maintaining responsive service delivery to urban stakeholders."
)
add_para(sec_1_1_p4)

sec_1_1_p5 = (
    "Emerging paradigms in IoT architecture include intent-based networking, where administrators specify "
    "desired outcomes rather than explicit configurations, and the network autonomously adapts to achieve "
    "stated objectives. Software-defined IoT architectures decouple the control plane from the data plane, "
    "enabling centralized policy management across heterogeneous device populations while maintaining "
    "distributed data processing for performance optimization. The concept of computing-aware networking "
    "integrates real-time awareness of computational resource availability into routing decisions, ensuring "
    "that IoT data flows are directed to processing nodes with available capacity. These architectural "
    "innovations collectively enable smart city IoT deployments to achieve the scalability, reliability, "
    "and performance characteristics required for mission-critical urban applications while maintaining "
    "the flexibility to accommodate rapidly evolving technology landscapes and expanding application requirements."
)
add_para(sec_1_1_p5)

add_heading_styled('1.2 5G/6G, LPWAN, Mesh Networks, and Time-Sensitive Networking', level=2)

sec_1_2_p1 = (
    "Communication infrastructure forms the nervous system of smart city ecosystems, with diverse network "
    "technologies serving distinct application requirements across the urban landscape [11]. Fifth-generation "
    "(5G) cellular networks provide three fundamental service categories: enhanced Mobile Broadband (eMBB) "
    "delivering peak speeds of 20 Gbps, Ultra-Reliable Low-Latency Communications (URLLC) achieving latencies "
    "below 1 millisecond with 99.999% reliability, and massive Machine-Type Communications (mMTC) supporting "
    "up to one million connected devices per square kilometer [12]. These capabilities enable previously "
    "impossible applications including real-time holographic urban planning, remote surgical procedures in "
    "smart hospitals, and vehicle-to-everything (V2X) communications for autonomous transportation systems. "
    "The emerging sixth-generation (6G) networks, anticipated for commercial deployment by 2030, promise "
    "terahertz-band communications with data rates exceeding 1 Tbps, sub-microsecond latency, and integrated "
    "sensing-communication capabilities [13]."
)
add_para(sec_1_2_p1)



sec_1_2_p2 = (
    "Low-Power Wide-Area Networks (LPWAN) address the requirements of massive-scale sensor deployments where "
    "devices must operate for 10-15 years on a single battery while transmitting data over distances of 5-15 "
    "kilometers [14]. LoRaWAN technology has emerged as the dominant LPWAN protocol for smart city applications, "
    "with global deployments exceeding 170 countries and supporting diverse use cases from smart parking to "
    "environmental monitoring. Narrowband IoT (NB-IoT) leverages existing cellular infrastructure to provide "
    "carrier-grade reliability with deep indoor penetration, achieving coverage gains of 20 dB over conventional "
    "LTE [15]. The communication network topology for smart cities, as depicted in Figure 2, illustrates the "
    "complementary nature of these technologies in creating comprehensive urban connectivity."
)
add_para(sec_1_2_p2)

# Insert Figure 2
add_figure('/projects/sandbox/AMMAN/iot_figures/Figure_2_Communication_Networks.png',
           'Figure 2: Communication Technologies and Network Topology for Smart City IoT Ecosystems')

sec_1_2_p3 = (
    "Mesh networking protocols including Zigbee, Thread, and Wi-Fi 6E enable local device-to-device "
    "communication with self-healing capabilities essential for resilient urban infrastructure [16]. These "
    "networks dynamically reconfigure routing paths in response to node failures, maintaining connectivity "
    "even under adverse conditions. Time-Sensitive Networking (TSN) standards, defined by IEEE 802.1 working "
    "groups, provide deterministic communication guarantees required for industrial IoT applications within "
    "smart city contexts [17]. TSN enables bounded latency, zero congestion loss, and precise time "
    "synchronization across heterogeneous network segments, making it essential for applications such as "
    "synchronized traffic signal control and distributed energy resource coordination. Table 1 presents "
    "a comparative analysis of key communication technologies deployed in smart city environments."
)
add_para(sec_1_2_p3)

sec_1_2_p4 = (
    "The heterogeneous nature of smart city communication requirements necessitates sophisticated network "
    "orchestration frameworks that dynamically allocate connectivity resources based on application needs. "
    "Network slicing technology, a cornerstone of 5G architecture, enables the creation of logically "
    "isolated virtual networks on shared physical infrastructure, each tailored for specific smart city "
    "service categories. A critical communications slice may guarantee sub-millisecond latency for "
    "autonomous vehicle coordination, while a massive IoT slice optimizes for connection density supporting "
    "millions of environmental sensors per square kilometer. Non-terrestrial networks (NTN) integrating "
    "Low Earth Orbit (LEO) satellite constellations with terrestrial IoT infrastructure provide ubiquitous "
    "coverage including rural peri-urban areas and disaster-affected zones where ground-based infrastructure "
    "is unavailable or damaged. The convergence of communication and sensing capabilities in future networks "
    "enables passive environmental monitoring through signal analysis, where communication waveforms "
    "simultaneously carry data and sense the physical environment, creating a pervasive sensing fabric "
    "that augments dedicated IoT sensor deployments without additional hardware investment."
)
add_para(sec_1_2_p4)



# TABLE 1
doc.add_paragraph()
table1_caption = doc.add_paragraph()
table1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tc1 = table1_caption.add_run('Table 1: Comparative Analysis of Smart City Communication Technologies')
run_tc1.bold = True
run_tc1.font.size = Pt(10)

table1 = doc.add_table(rows=7, cols=6)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['Technology', 'Range', 'Data Rate', 'Latency', 'Power', 'Key Application']
data1 = [
    ['5G NR', '1-10 km', '20 Gbps', '<1 ms', 'High', 'Autonomous vehicles, AR/VR'],
    ['LoRaWAN', '5-15 km', '50 kbps', '1-2 s', 'Ultra-low', 'Environmental sensors'],
    ['NB-IoT', '10-35 km', '250 kbps', '1.5-10 s', 'Very low', 'Smart metering'],
    ['Zigbee', '10-100 m', '250 kbps', '15 ms', 'Low', 'Building automation'],
    ['Wi-Fi 6E', '50-100 m', '9.6 Gbps', '<5 ms', 'Medium', 'High-density venues'],
    ['TSN/Ethernet', '100 m', '10 Gbps', '<10 us', 'Medium', 'Industrial control']
]

for i, header in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for row_idx, row_data in enumerate(data1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table1.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()



add_heading_styled('1.3 AI, Machine Learning, and Digital Twins for Intelligent Urban Systems', level=2)

sec_1_3_p1 = (
    "Artificial intelligence and machine learning constitute the cognitive engine of smart city ecosystems, "
    "transforming raw sensor data into actionable urban intelligence [18]. Deep learning architectures, "
    "particularly convolutional neural networks (CNNs) and transformer models, have demonstrated exceptional "
    "performance in urban analytics tasks including traffic prediction with 94% accuracy, air quality "
    "forecasting achieving mean absolute errors below 5 micrograms per cubic meter, and anomaly detection "
    "in critical infrastructure with false positive rates under 0.1% [19]. Federated learning paradigms "
    "enable collaborative model training across distributed urban IoT nodes without centralizing sensitive "
    "data, addressing privacy concerns while maintaining model accuracy within 2-3% of centralized "
    "approaches [20]. Reinforcement learning algorithms optimize real-time urban operations including "
    "traffic signal timing, energy distribution, and emergency resource allocation, with studies "
    "demonstrating 25-35% improvements in system efficiency compared to traditional rule-based "
    "approaches [21]."
)
add_para(sec_1_3_p1)

sec_1_3_p2 = (
    "Digital twin technology creates high-fidelity virtual replicas of physical urban infrastructure, "
    "enabling simulation, prediction, and optimization without disrupting real-world operations [22]. "
    "City-scale digital twins integrate data from thousands of IoT sensors to maintain real-time "
    "representations of traffic networks, utility systems, building energy performance, and environmental "
    "conditions. Singapore's Virtual Singapore project exemplifies this approach, creating a semantically "
    "enriched 3D model of the entire city-state that supports urban planning, emergency response simulation, "
    "and policy evaluation [23]. The convergence of digital twin technology with generative AI enables "
    "scenario planning capabilities where urban planners can simulate the impacts of infrastructure "
    "modifications, climate events, or policy changes before implementation. Research indicates that "
    "digital twin-enabled predictive maintenance reduces infrastructure failure rates by 45-60% while "
    "extending asset lifespans by 20-30% [24]. The integration of these technologies within the IoT "
    "architecture creates a self-learning urban ecosystem capable of continuous optimization, as "
    "illustrated in the applications framework shown in Figure 3."
)
add_para(sec_1_3_p2)

sec_1_3_p3 = (
    "The application of explainable AI (XAI) in smart city contexts addresses the critical challenge of "
    "transparency in algorithmic urban governance. Decision support systems that recommend policy actions "
    "based on IoT data must provide interpretable justifications that city administrators and elected "
    "officials can understand, evaluate, and communicate to constituents. Attention mechanisms in "
    "transformer architectures enable identification of which IoT data streams most significantly "
    "influence predictions, creating audit trails for automated decisions. Graph neural networks (GNNs) "
    "have emerged as particularly effective for modeling the networked structure of urban systems, "
    "capturing spatial dependencies between infrastructure components and propagating information across "
    "topological relationships. Transfer learning approaches enable models trained on data-rich cities "
    "to be adapted for resource-constrained municipalities, democratizing access to advanced AI capabilities "
    "regardless of local data availability or technical expertise. The convergence of these AI approaches "
    "with comprehensive IoT sensing creates urban intelligence systems that not only detect current "
    "conditions but anticipate future states with sufficient lead time for proactive intervention, "
    "fundamentally shifting urban management from reactive to predictive paradigms."
)
add_para(sec_1_3_p3)



# ============ SECTION 2 ============
add_heading_styled('2. IoT Applications for Intelligent and Sustainable Urban Infrastructure', level=1)

sec_2_intro = (
    "The translation of IoT architectural capabilities into practical urban applications represents the "
    "critical value creation stage of smart city development. While the preceding section established the "
    "technological foundations, this section examines how these capabilities are deployed across key urban "
    "domains to deliver tangible improvements in efficiency, sustainability, and citizen well-being. The "
    "applications discussed span transportation systems that move people and goods, energy and utility "
    "networks that sustain urban metabolism, and environmental monitoring systems that safeguard public "
    "health. Each domain demonstrates unique IoT integration patterns while contributing to the holistic "
    "optimization of urban systems through data sharing and cross-domain intelligence. The economic "
    "justification for IoT investment in urban infrastructure rests on demonstrable returns including "
    "reduced operational costs, improved asset utilization, decreased environmental externalities, and "
    "enhanced citizen satisfaction with public services. Global smart city market projections indicate "
    "compound annual growth rates exceeding 25% through 2030, driven by urbanization pressures, "
    "climate imperatives, and technological maturation that collectively create compelling business "
    "cases for IoT-enabled urban transformation across diverse economic and geographic contexts."
)
add_para(sec_2_intro)

add_heading_styled('2.1 Smart Transportation, Connected Vehicles, and Intelligent Traffic Management', level=2)

sec_2_1_p1 = (
    "Smart transportation represents the most visible and economically significant application domain for "
    "IoT in urban environments, with global investment exceeding $130 billion annually [25]. Intelligent "
    "Traffic Management Systems (ITMS) leverage networks of inductive loop detectors, video analytics "
    "cameras, radar sensors, and connected vehicle data to optimize traffic flow across metropolitan road "
    "networks. Adaptive signal control technologies such as SCOOT and SCATS dynamically adjust signal "
    "timing based on real-time demand, achieving traffic delay reductions of 20-40% compared to fixed-time "
    "plans [26]. The integration of vehicle-to-infrastructure (V2I) communication enables proactive "
    "traffic management, with connected vehicles receiving speed advisories and route recommendations "
    "that collectively reduce congestion by 15-25% during peak periods."
)
add_para(sec_2_1_p1)

sec_2_1_p2 = (
    "Connected and autonomous vehicles (CAVs) represent the frontier of smart transportation, with IoT "
    "sensor fusion combining LiDAR, radar, cameras, and V2X communications to achieve Society of Automotive "
    "Engineers (SAE) Level 4-5 automation [27]. Urban deployments of autonomous shuttle services in cities "
    "including Helsinki, Singapore, and Dubai have demonstrated the viability of driverless public "
    "transportation in controlled environments. Mobility-as-a-Service (MaaS) platforms integrate multimodal "
    "transportation options including public transit, ride-sharing, bike-sharing, and micro-mobility through "
    "unified IoT-enabled booking and payment systems [28]. Smart parking solutions employing magnetometer "
    "sensors and computer vision reduce parking search time by 43%, decreasing associated emissions by an "
    "estimated 30% in dense urban cores. The comprehensive application of IoT across transportation "
    "domains is visualized in Figure 3, demonstrating the interconnected nature of smart city services."
)
add_para(sec_2_1_p2)

sec_2_1_p3 = (
    "Electric vehicle (EV) charging infrastructure managed through IoT platforms represents a critical "
    "intersection of smart transportation and smart energy systems. Intelligent charging stations equipped "
    "with bidirectional power electronics enable vehicle-to-grid (V2G) services where parked EVs serve "
    "as distributed energy storage assets, providing grid balancing services during peak demand periods. "
    "IoT-based fleet management systems optimize routing, scheduling, and maintenance for public transit "
    "and commercial vehicle fleets, with telematics data enabling predictive maintenance that reduces "
    "vehicle downtime by 30-45%. Drone-based delivery systems, operating as extensions of smart "
    "transportation networks, leverage IoT infrastructure for airspace management, autonomous navigation, "
    "and integration with ground-based logistics systems. The evolution toward cooperative intelligent "
    "transportation systems (C-ITS) enables real-time information exchange between vehicles, infrastructure, "
    "and vulnerable road users, with studies indicating potential reductions of 25-40% in traffic fatalities "
    "through cooperative collision avoidance and intersection management systems."
)
add_para(sec_2_1_p3)

# Insert Figure 3
add_figure('/projects/sandbox/AMMAN/iot_figures/Figure_3_Applications_Framework.png',
           'Figure 3: IoT Applications Framework for Sustainable Urban Infrastructure')



add_heading_styled('2.2 Smart Energy, Grids, Water Distribution, and Waste Management', level=2)

sec_2_2_p1 = (
    "IoT-enabled smart energy systems are revolutionizing urban utility management through intelligent "
    "monitoring, automated control, and predictive optimization of generation, distribution, and consumption "
    "assets [29]. Advanced Metering Infrastructure (AMI) deployments provide bidirectional communication "
    "between utilities and consumers at 15-minute intervals, enabling dynamic pricing signals, demand "
    "response programs, and granular consumption analytics. Smart grid technologies integrate distributed "
    "energy resources including rooftop solar installations, battery storage systems, and electric vehicle "
    "charging infrastructure into cohesive microgrids capable of autonomous operation during grid "
    "disruptions [30]. IoT sensors deployed across distribution networks enable real-time fault detection "
    "and isolation, reducing outage durations by 50-70% through automated switching and self-healing "
    "capabilities. Table 2 presents the key performance metrics achieved by IoT deployments across "
    "urban utility sectors."
)
add_para(sec_2_2_p1)

# TABLE 2
doc.add_paragraph()
table2_caption = doc.add_paragraph()
table2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tc2 = table2_caption.add_run('Table 2: IoT Performance Metrics Across Urban Utility Sectors')
run_tc2.bold = True
run_tc2.font.size = Pt(10)

table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Utility Sector', 'IoT Technology', 'Key Metric', 'Improvement', 'ROI Period']
data2 = [
    ['Electricity', 'Smart Meters/AMI', 'Peak Demand Reduction', '15-25%', '3-5 years'],
    ['Water Supply', 'Acoustic Sensors', 'Leak Detection Rate', '85-95%', '2-4 years'],
    ['Gas Distribution', 'IoT Monitors', 'Leak Response Time', '70% faster', '3-6 years'],
    ['Waste Collection', 'Fill-level Sensors', 'Collection Efficiency', '30-40%', '1-3 years'],
    ['Street Lighting', 'Smart Controls', 'Energy Savings', '50-70%', '2-4 years'],
    ['District Heating', 'Flow Sensors', 'Distribution Loss', '20-35% reduction', '4-7 years']
]

for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for row_idx, row_data in enumerate(data2):
    for col_idx, cell_text in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()



sec_2_2_p2 = (
    "Smart water management systems employ IoT sensor networks to monitor pressure, flow, quality, and "
    "acoustic signatures across distribution infrastructure, enabling early detection of leaks that "
    "traditionally account for 25-35% of treated water losses in aging urban networks [31]. Real-time "
    "water quality monitoring using multi-parameter probes provides continuous assessment of turbidity, "
    "pH, chlorine residual, and microbial indicators, enabling rapid response to contamination events. "
    "Smart waste management solutions integrate ultrasonic fill-level sensors in collection containers "
    "with route optimization algorithms, reducing collection vehicle mileage by 30-40% and associated "
    "emissions proportionally [32]. Pneumatic waste collection systems in new urban developments "
    "eliminate surface-level waste containers entirely, with IoT-controlled vacuum networks transporting "
    "waste underground to central processing facilities."
)
add_para(sec_2_2_p2)

sec_2_2_p3 = (
    "Building energy management systems (BEMS) represent a significant application of IoT for urban "
    "sustainability, given that buildings account for approximately 40% of total urban energy consumption "
    "and 33% of greenhouse gas emissions. IoT-enabled BEMS integrate sensors for occupancy detection, "
    "thermal comfort monitoring, lighting level assessment, and equipment performance tracking with "
    "AI-driven control algorithms that continuously optimize heating, ventilation, air conditioning, "
    "and lighting systems. Deep reinforcement learning approaches to building control have demonstrated "
    "energy savings of 20-35% while maintaining or improving occupant comfort compared to conventional "
    "rule-based systems. District-level energy optimization coordinates individual building systems "
    "with shared resources including combined heat and power plants, thermal storage, and shared "
    "renewable generation assets. The integration of weather forecasting, occupancy prediction, and "
    "energy market pricing into IoT-enabled building management creates anticipatory systems that "
    "pre-condition spaces during low-cost periods and participate in demand response programs during "
    "grid stress events, transforming buildings from passive energy consumers into active grid resources "
    "that contribute to overall urban energy resilience and efficiency."
)
add_para(sec_2_2_p3)

add_heading_styled('2.3 Environmental Monitoring, Pollution Control, and Urban Public Health', level=2)

sec_2_3_p1 = (
    "IoT-based environmental monitoring networks provide unprecedented spatial and temporal resolution "
    "for tracking urban environmental quality, enabling evidence-based interventions to protect public "
    "health [33]. Dense networks of low-cost air quality sensors measuring particulate matter (PM2.5, "
    "PM10), nitrogen dioxide, ozone, and volatile organic compounds create hyperlocal pollution maps "
    "that reveal micro-environmental variations invisible to traditional sparse monitoring stations. "
    "Machine learning algorithms calibrate and fuse data from heterogeneous sensor types to achieve "
    "reference-grade accuracy while maintaining the spatial coverage advantages of distributed networks [34]. "
    "Cities implementing comprehensive IoT air quality monitoring have achieved 15-25% reductions in "
    "population exposure to harmful pollutants through targeted interventions including dynamic traffic "
    "restrictions, industrial emission alerts, and citizen notification systems."
)
add_para(sec_2_3_p1)

sec_2_3_p2 = (
    "Urban noise mapping through IoT acoustic sensor networks characterizes the spatiotemporal distribution "
    "of environmental noise, informing urban planning decisions and enabling enforcement of noise regulations [35]. "
    "Integration of environmental monitoring with public health surveillance systems enables rapid identification "
    "of exposure-disease relationships, supporting predictive health analytics. The COVID-19 pandemic accelerated "
    "deployment of IoT-enabled public health infrastructure including thermal screening systems, occupancy "
    "monitoring for social distancing enforcement, and indoor air quality systems with real-time ventilation "
    "optimization [36]. Wearable IoT devices and smartphone-based sensing platforms enable citizen science "
    "approaches to environmental monitoring, democratizing data collection while enhancing spatial coverage "
    "beyond what fixed infrastructure alone can achieve."
)
add_para(sec_2_3_p2)

sec_2_3_p3 = (
    "The integration of satellite remote sensing with ground-based IoT networks creates multi-scale environmental "
    "monitoring systems that combine the spatial coverage of orbital platforms with the temporal resolution and "
    "measurement specificity of in-situ sensors. Urban heat island monitoring through networks of temperature "
    "sensors, combined with satellite thermal imagery and computational fluid dynamics models, enables targeted "
    "interventions including reflective surface installations, strategic tree planting, and water feature "
    "deployment in heat-vulnerable neighborhoods. Biodiversity monitoring using acoustic sensors, camera traps, "
    "and environmental DNA sampling stations tracks urban ecosystem health, informing conservation strategies "
    "within metropolitan environments. The application of IoT to urban agriculture, including vertical farming "
    "facilities with comprehensive environmental control systems, contributes to food security while reducing "
    "transportation-related emissions associated with conventional supply chains. These environmental monitoring "
    "capabilities, when integrated with public health databases and epidemiological models, enable predictive "
    "health impact assessments that quantify the benefits of environmental interventions in terms of reduced "
    "hospitalizations, improved life expectancy, and economic productivity gains, providing robust cost-benefit "
    "justifications for continued investment in smart city environmental infrastructure."
)
add_para(sec_2_3_p3)



# ============ SECTION 3 ============
add_heading_styled('3. Smart Governance, Security, and Citizen-Centric Services', level=1)

sec_3_intro = (
    "The deployment of IoT infrastructure across urban environments creates both unprecedented opportunities "
    "for improved governance and significant challenges in security, privacy, and equitable access. Smart "
    "governance leverages IoT-generated intelligence to enhance public administration, improve emergency "
    "response capabilities, and enable evidence-based policymaking. However, the same data collection "
    "capabilities that enable intelligent governance also create surveillance risks, privacy vulnerabilities, "
    "and potential for algorithmic discrimination that must be addressed through robust technical safeguards "
    "and governance frameworks. This section examines the dual nature of IoT-enabled governance, exploring "
    "both its transformative potential and the essential protections required to maintain democratic values "
    "and citizen trust in an increasingly data-rich urban environment. The balance between security and "
    "freedom, between efficiency and privacy, represents perhaps the defining challenge of smart city "
    "development, requiring thoughtful engagement from technologists, policymakers, civil society "
    "organizations, and citizens themselves to establish norms and boundaries that reflect societal values."
)
add_para(sec_3_intro)

add_heading_styled('3.1 IoT-Enabled Governance, Public Safety, and Emergency Response', level=2)

sec_3_1_p1 = (
    "IoT technologies are transforming urban governance by enabling data-driven decision-making, improving "
    "service delivery efficiency, and enhancing transparency in public administration [37]. Smart city "
    "platforms aggregate data from diverse IoT sources to provide city administrators with comprehensive "
    "operational dashboards that visualize real-time urban performance across multiple dimensions including "
    "mobility, energy consumption, environmental quality, and public safety. These integrated command centers, "
    "exemplified by Rio de Janeiro's Centro de Operacoes and Dubai's Pulse Platform, enable coordinated "
    "multi-agency responses to urban events ranging from traffic incidents to natural disasters [38]. "
    "Predictive analytics applied to historical IoT data enables proactive governance approaches, with "
    "machine learning models forecasting infrastructure maintenance needs, crime hotspots, and demand "
    "patterns for public services with lead times of days to weeks."
)
add_para(sec_3_1_p1)

sec_3_1_p2 = (
    "Public safety applications of IoT encompass intelligent surveillance systems, gunshot detection "
    "networks, flood early warning systems, and structural health monitoring of critical infrastructure [39]. "
    "Computer vision algorithms processing feeds from city-wide camera networks enable automated detection "
    "of security threats, traffic violations, and emergency situations, with response time improvements of "
    "40-60% compared to human-only monitoring. IoT-enabled emergency response systems integrate caller "
    "location data, connected building fire alarms, vehicle telematics, and hospital capacity information "
    "to optimize dispatching and resource allocation. Table 3 summarizes key smart governance applications "
    "and their measured impacts across leading smart city implementations. The interconnected nature of these "
    "systems, as shown in Figure 1, demonstrates how the layered IoT architecture supports diverse "
    "governance applications through a unified technology platform."
)
add_para(sec_3_1_p2)

sec_3_1_p3 = (
    "Disaster resilience through IoT infrastructure represents a critical public safety application, with "
    "sensor networks providing early warning capabilities for floods, earthquakes, landslides, and severe "
    "weather events. Seismic sensor arrays deployed across urban areas enable earthquake early warning "
    "systems that provide seconds to minutes of advance notice, triggering automated protective actions "
    "including elevator positioning, gas supply shutdown, and bridge closure. Flood monitoring systems "
    "combine river level sensors, rainfall gauges, drainage network monitors, and hydrological models "
    "to provide spatially explicit inundation forecasts with lead times of hours to days. IoT-enabled "
    "building structural health monitoring systems continuously assess the integrity of critical "
    "infrastructure including bridges, tunnels, and high-rise buildings through vibration analysis, "
    "strain measurement, and tilt detection, enabling condition-based maintenance that prevents "
    "catastrophic failures. The integration of these diverse safety systems through unified IoT "
    "platforms enables cascading event management, where the detection of an initial trigger such as "
    "an earthquake automatically activates secondary monitoring systems, initiates evacuation protocols, "
    "and pre-positions emergency response resources based on predicted impact patterns."
)
add_para(sec_3_1_p3)

# TABLE 3
doc.add_paragraph()
table3_caption = doc.add_paragraph()
table3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tc3 = table3_caption.add_run('Table 3: Smart Governance Applications and Measured Impacts')
run_tc3.bold = True
run_tc3.font.size = Pt(10)

table3 = doc.add_table(rows=7, cols=5)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Application Domain', 'IoT Infrastructure', 'City Example', 'Key Impact', 'Scale']
data3 = [
    ['Traffic Management', 'Sensors + AI cameras', 'Singapore', '25% congestion reduction', 'City-wide'],
    ['Crime Prevention', 'Surveillance + Analytics', 'Chicago', '30% faster response', '700+ sensors'],
    ['Emergency Response', 'Integrated IoT platform', 'Rio de Janeiro', '40% response improvement', 'Metro area'],
    ['Infrastructure Mgmt', 'Structural sensors', 'Amsterdam', '45% maintenance savings', '1200+ bridges'],
    ['Energy Governance', 'Smart grid IoT', 'Copenhagen', '42% carbon reduction', 'City-wide'],
    ['Citizen Services', 'Multi-channel IoT', 'Seoul', '35% satisfaction increase', '10M citizens']
]

for i, header in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for row_idx, row_data in enumerate(data3):
    for col_idx, cell_text in enumerate(row_data):
        cell = table3.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()



add_heading_styled('3.2 Data Security, Privacy, Interoperability, and Scalable IoT Ecosystems', level=2)

sec_3_2_p1 = (
    "The massive scale and heterogeneity of smart city IoT deployments create unprecedented cybersecurity "
    "challenges, with attack surfaces expanding exponentially as billions of devices connect to urban "
    "networks [39]. IoT devices, often constrained in computational resources, present vulnerabilities "
    "that adversaries exploit through botnet recruitment, man-in-the-middle attacks, and firmware "
    "manipulation. The Mirai botnet attack of 2016, which compromised over 600,000 IoT devices, "
    "demonstrated the catastrophic potential of unsecured IoT infrastructure [40]. Modern security "
    "frameworks for smart city IoT employ defense-in-depth strategies incorporating device identity "
    "management through Public Key Infrastructure (PKI), network segmentation using software-defined "
    "networking, encrypted communications via lightweight cryptographic protocols (e.g., DTLS, OSCORE), "
    "and behavioral analytics for anomaly detection. Zero-trust architectures eliminate implicit trust "
    "between network segments, requiring continuous verification of device identity and authorization "
    "for every data transaction [41]."
)
add_para(sec_3_2_p1)

sec_3_2_p2 = (
    "Privacy preservation in IoT-enabled smart cities requires balancing the analytical value of urban "
    "data against citizens' fundamental rights to privacy and data protection. Differential privacy "
    "mechanisms inject calibrated noise into aggregated datasets to prevent re-identification of "
    "individual behavior patterns while maintaining statistical utility for urban analytics [41]. "
    "Homomorphic encryption enables computation on encrypted data, allowing cloud-based analytics "
    "without exposing raw sensor readings to service providers. The European General Data Protection "
    "Regulation (GDPR) and similar frameworks globally impose strict requirements on IoT data "
    "processing, necessitating privacy-by-design approaches in smart city system architectures. "
    "Interoperability challenges arise from the fragmented IoT ecosystem comprising hundreds of "
    "competing protocols, data formats, and vendor-specific platforms. Standards-based approaches "
    "including FIWARE NGSI-LD, oneM2M, and W3C Web of Things provide semantic interoperability "
    "layers that enable cross-domain data exchange and application portability across smart city "
    "platforms [42]."
)
add_para(sec_3_2_p2)

sec_3_2_p3 = (
    "Scalability of IoT ecosystems demands architectural approaches that accommodate exponential growth "
    "in device populations without proportional increases in management complexity or infrastructure costs. "
    "Autonomous device management platforms employing AI-driven operations (AIOps) automate device "
    "provisioning, configuration, monitoring, and lifecycle management across millions of heterogeneous "
    "IoT endpoints. Digital twin-based device management creates virtual representations of physical "
    "devices, enabling remote diagnostics, over-the-air firmware updates, and predictive failure detection "
    "without physical access to deployed equipment. The economic sustainability of large-scale IoT "
    "deployments requires careful consideration of total cost of ownership including device procurement, "
    "connectivity subscriptions, data storage, computational resources, and ongoing maintenance. "
    "Shared IoT infrastructure models, where multiple city departments and third-party service providers "
    "utilize common sensor networks and communication infrastructure, reduce per-application costs while "
    "maximizing data reuse. The establishment of urban IoT testbeds and sandboxes enables controlled "
    "experimentation with new technologies and applications before full-scale deployment, reducing risks "
    "associated with technology adoption while building institutional capacity for IoT system management."
)
add_para(sec_3_2_p3)

add_heading_styled('3.3 Digital Inclusion, Citizen Participation, and Responsible Data Governance', level=2)

sec_3_3_p1 = (
    "The success of smart city initiatives ultimately depends on inclusive design that ensures equitable "
    "access to IoT-enabled services across all demographic groups, avoiding the creation of digital "
    "divides that exacerbate existing social inequalities [37]. Digital inclusion strategies must address "
    "barriers including affordability of connected devices, digital literacy requirements for service "
    "access, language and cultural accessibility, and physical accessibility for persons with disabilities. "
    "Participatory design methodologies involving citizens in the co-creation of IoT services ensure "
    "that technology deployments align with community needs and values rather than imposed technocratic "
    "visions. Living lab approaches, as exemplified by Barcelona's Superblocks program and Amsterdam's "
    "Smart Citizens Lab, demonstrate how citizen engagement can drive innovation while building social "
    "acceptance of IoT technologies [38]."
)
add_para(sec_3_3_p1)

sec_3_3_p2 = (
    "Responsible data governance frameworks establish transparent rules for the collection, processing, "
    "sharing, and retention of urban IoT data, ensuring accountability and public trust [39]. Data "
    "trusts and cooperatives represent emerging governance models where citizens collectively control "
    "the use of their data, receiving tangible benefits from its exploitation while maintaining "
    "veto rights over unacceptable uses. Open data initiatives make non-sensitive urban IoT datasets "
    "publicly available, enabling civic innovation, academic research, and commercial service development. "
    "The integration of communication technologies illustrated in Figure 2 must be accompanied by "
    "robust governance frameworks that ensure these powerful connectivity capabilities serve public "
    "interests while protecting individual rights. Ethical AI frameworks guide the development and "
    "deployment of algorithmic decision-making systems within smart city contexts, requiring "
    "explainability, fairness auditing, and human oversight of consequential automated decisions [40]."
)
add_para(sec_3_3_p2)

sec_3_3_p3 = (
    "The concept of data sovereignty extends responsible governance to the municipal level, where cities "
    "assert control over urban data assets generated within their jurisdiction regardless of the "
    "nationality of technology providers. Municipal data platforms aggregate and curate IoT data from "
    "public and private sources, creating authoritative urban data repositories that support evidence-based "
    "policymaking. Data impact assessments, analogous to environmental impact assessments, evaluate the "
    "potential consequences of new IoT deployments on privacy, equity, and social cohesion before "
    "approval. Algorithmic accountability mechanisms require regular auditing of automated decision "
    "systems for bias, discrimination, and error patterns, with results published transparently and "
    "remediation mandated for identified issues. These governance innovations collectively ensure that "
    "the enormous power of IoT-enabled urban intelligence serves democratic values and public interest, "
    "building the social license necessary for continued technology adoption while preventing the "
    "emergence of techno-authoritarian urban governance models that prioritize efficiency over rights."
)
add_para(sec_3_3_p3)



# ============ SECTION 4 ============
add_heading_styled('4. Future Directions and Strategic Pathways for IoT-Enabled Smart Cities', level=1)

sec_4_intro = (
    "The trajectory of IoT-enabled smart city development points toward increasingly sophisticated, "
    "autonomous, and integrated urban systems that fundamentally redefine the relationship between "
    "citizens, technology, and urban infrastructure. Emerging technologies including artificial general "
    "intelligence, quantum computing, advanced robotics, and brain-computer interfaces will expand the "
    "capabilities of urban IoT systems far beyond current implementations. Simultaneously, evolving "
    "societal expectations around sustainability, equity, transparency, and democratic participation "
    "will shape how these technologies are deployed and governed. This section explores the strategic "
    "pathways available to cities as they navigate toward increasingly intelligent urban futures, "
    "examining the technological enablers, governance innovations, and design philosophies that will "
    "determine whether smart cities fulfill their promise of sustainable human flourishing."
)
add_para(sec_4_intro)

add_heading_styled('4.1 AI-Driven Autonomous Cities, Edge Intelligence, and Predictive Urban Management', level=2)

sec_4_1_p1 = (
    "The evolution toward AI-driven autonomous cities represents the next frontier in smart city development, "
    "where urban systems achieve self-optimization capabilities through advanced machine intelligence "
    "operating at the network edge [42]. Edge AI processors, incorporating neural processing units (NPUs) "
    "with performance exceeding 100 TOPS (Tera Operations Per Second), enable complex inference tasks "
    "including real-time video analytics, natural language processing, and predictive modeling to execute "
    "locally without cloud dependencies. This architectural evolution, illustrated in the strategic roadmap "
    "presented in Figure 4, transforms IoT infrastructure from passive data collection systems into active "
    "intelligent agents capable of autonomous decision-making within defined operational boundaries. "
    "Neuromorphic computing architectures, inspired by biological neural systems, promise orders-of-magnitude "
    "improvements in energy efficiency for edge AI workloads, enabling always-on intelligence in battery-powered "
    "IoT devices [42]."
)
add_para(sec_4_1_p1)

# Insert Figure 4
add_figure('/projects/sandbox/AMMAN/iot_figures/Figure_4_Strategic_Roadmap.png',
           'Figure 4: Strategic Roadmap for AI-Driven Autonomous Smart Cities')

sec_4_1_p2 = (
    "Predictive urban management leverages digital twins enhanced with generative AI capabilities to "
    "simulate and optimize city operations across multiple time horizons [22]. Short-term predictions "
    "(minutes to hours) enable real-time traffic routing, energy load balancing, and emergency response "
    "optimization. Medium-term forecasts (days to weeks) support infrastructure maintenance scheduling, "
    "event planning, and resource procurement. Long-term projections (months to years) inform capital "
    "investment decisions, urban expansion planning, and climate adaptation strategies. The convergence "
    "of large language models with urban IoT data enables natural language interfaces for city management, "
    "allowing administrators to query complex urban datasets and receive synthesized insights through "
    "conversational interaction [24]. Swarm intelligence algorithms coordinate fleets of autonomous "
    "urban robots for tasks including last-mile delivery, infrastructure inspection, and environmental "
    "cleanup, operating as extensions of the IoT sensor network with actuation capabilities."
)
add_para(sec_4_1_p2)

sec_4_1_p3 = (
    "The transition toward autonomous city operations introduces novel challenges in system reliability, "
    "fail-safe mechanisms, and human oversight architectures. Safety-critical autonomous systems require "
    "formal verification of decision algorithms, with mathematical proofs ensuring that automated actions "
    "remain within defined safety envelopes under all possible input conditions. Graceful degradation "
    "architectures ensure that autonomous systems revert to safe states when confronted with unprecedented "
    "situations, maintaining essential services while alerting human operators to intervene. The concept "
    "of shared autonomy enables graduated automation levels, where IoT systems handle routine decisions "
    "independently while escalating complex or sensitive situations to human decision-makers through "
    "intuitive interfaces that present relevant context and recommended actions. Quantum computing "
    "integration with smart city IoT infrastructure promises transformative capabilities in optimization "
    "problems that are computationally intractable for classical systems, including large-scale traffic "
    "network optimization, portfolio optimization for distributed energy resources, and cryptographic "
    "protocols for securing IoT communications against future quantum attacks."
)
add_para(sec_4_1_p3)



add_heading_styled('4.2 Blockchain, Trust, and Transparent Digital Urban Ecosystems', level=2)

sec_4_2_p1 = (
    "Blockchain and distributed ledger technologies (DLT) address fundamental trust challenges in smart "
    "city IoT ecosystems by providing immutable, transparent, and decentralized record-keeping for urban "
    "transactions and data provenance [40]. Smart contracts automate multi-party agreements in areas "
    "including peer-to-peer energy trading, automated insurance claims processing, and supply chain "
    "verification for municipal procurement. Blockchain-based identity management systems provide "
    "self-sovereign digital identities that enable citizens to control access to their personal data "
    "across smart city services without relying on centralized identity providers [41]. The strategic "
    "roadmap shown in Figure 4 positions blockchain integration as a critical enabler for the "
    "transition from intelligent to autonomous city operations."
)
add_para(sec_4_2_p1)

sec_4_2_p2 = (
    "Decentralized IoT data marketplaces built on blockchain infrastructure enable transparent and fair "
    "exchange of urban data assets between public agencies, private companies, and citizens [40]. These "
    "platforms employ token economics to incentivize data sharing while ensuring equitable value "
    "distribution among contributors. Blockchain-based voting and governance systems enable secure "
    "digital democracy, allowing citizens to participate in urban planning decisions through "
    "tamper-proof digital referenda. The convergence of blockchain with IoT creates opportunities for "
    "decentralized autonomous organizations (DAOs) in urban governance, where community-owned "
    "infrastructure operates according to collectively agreed smart contract rules [41]. However, "
    "challenges remain in blockchain scalability, energy consumption of consensus mechanisms, and "
    "regulatory frameworks governing decentralized urban systems. Table 4 presents a comparative "
    "analysis of blockchain platforms evaluated for smart city IoT applications."
)
add_para(sec_4_2_p2)

sec_4_2_p3 = (
    "The implementation of trusted execution environments (TEEs) in conjunction with blockchain technology "
    "addresses the oracle problem in smart city IoT applications, ensuring that real-world sensor data "
    "feeding into smart contracts is authentic and untampered. Verifiable credentials issued through "
    "blockchain-based identity systems enable privacy-preserving authentication for smart city services, "
    "where citizens prove eligibility for services without revealing unnecessary personal information. "
    "Supply chain transparency enabled by IoT-blockchain integration provides citizens with verifiable "
    "information about the provenance of municipal procurement, from construction materials to food "
    "served in public institutions. Carbon credit tokenization and automated verification through IoT "
    "environmental sensors create transparent markets for emissions reductions, incentivizing private "
    "sector investment in urban sustainability initiatives. The maturation of layer-2 scaling solutions "
    "and energy-efficient consensus mechanisms addresses earlier concerns about blockchain sustainability, "
    "making decentralized approaches viable for the transaction volumes characteristic of city-scale "
    "IoT deployments processing millions of events daily."
)
add_para(sec_4_2_p3)

sec_4_2_p4 = (
    "Interoperability between blockchain platforms and legacy municipal IT systems presents implementation "
    "challenges that require middleware solutions and standardized application programming interfaces. "
    "Regulatory uncertainty surrounding cryptocurrency, decentralized finance, and tokenized assets "
    "creates risks for municipalities exploring blockchain-based urban services, necessitating close "
    "engagement with regulatory bodies during pilot deployments. The governance of blockchain protocols "
    "themselves introduces meta-governance challenges, where decisions about protocol upgrades, parameter "
    "adjustments, and dispute resolution mechanisms must align with democratic accountability principles. "
    "Despite these challenges, the trajectory toward blockchain integration in smart city infrastructure "
    "appears well-established, with major technology companies, standards organizations, and international "
    "development agencies actively supporting research and pilot deployments across diverse urban contexts "
    "from advanced economies to rapidly urbanizing developing nations."
)
add_para(sec_4_2_p4)

# TABLE 4
doc.add_paragraph()
table4_caption = doc.add_paragraph()
table4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tc4 = table4_caption.add_run('Table 4: Blockchain Platforms for Smart City IoT Applications')
run_tc4.bold = True
run_tc4.font.size = Pt(10)

table4 = doc.add_table(rows=7, cols=5)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

headers4 = ['Platform', 'Consensus', 'TPS', 'IoT Suitability', 'Smart City Use Case']
data4 = [
    ['Ethereum 2.0', 'Proof of Stake', '100,000+', 'Medium', 'DeFi, Digital Identity'],
    ['Hyperledger Fabric', 'PBFT', '20,000+', 'High', 'Supply Chain, Governance'],
    ['IOTA Tangle', 'DAG/Coordinator', '1,000+', 'Very High', 'Microtransactions, Data Market'],
    ['Algorand', 'Pure PoS', '46,000+', 'High', 'Carbon Credits, Payments'],
    ['Polkadot', 'NPoS', '1,500+', 'Medium-High', 'Cross-chain Interop'],
    ['Hedera Hashgraph', 'Asynchronous BFT', '10,000+', 'High', 'Audit Trails, IoT Data']
]

for i, header in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for row_idx, row_data in enumerate(data4):
    for col_idx, cell_text in enumerate(row_data):
        cell = table4.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()



add_heading_styled('4.3 Resilient, Sustainable, and Human-Centric Smart City Models for Competitive Advantage', level=2)

sec_4_3_p1 = (
    "The ultimate vision for IoT-enabled smart cities transcends technological optimization to encompass "
    "resilient, sustainable, and human-centric urban models that deliver enduring competitive advantage "
    "for cities and their inhabitants [41]. Urban resilience frameworks leverage IoT infrastructure to "
    "enhance city capacity for absorbing shocks, adapting to changing conditions, and transforming in "
    "response to systemic challenges including climate change, pandemic events, and economic disruptions. "
    "IoT-enabled climate adaptation strategies include real-time flood monitoring and automated response "
    "systems, urban heat island mitigation through smart green infrastructure management, and predictive "
    "models for extreme weather impacts on critical infrastructure [42]. Circular economy principles "
    "integrated with IoT systems enable material flow tracking, waste stream optimization, and industrial "
    "symbiosis facilitation, where one organization's waste becomes another's input through automated "
    "marketplace matching."
)
add_para(sec_4_3_p1)

sec_4_3_p2 = (
    "Sustainability metrics embedded within IoT platforms provide real-time tracking of progress toward "
    "Sustainable Development Goals (SDGs) at the urban level, enabling evidence-based sustainability "
    "governance [41]. Carbon accounting systems leveraging IoT data from energy meters, transportation "
    "sensors, and industrial monitors provide granular emissions inventories that support carbon trading "
    "and offset verification. Nature-based solutions enhanced with IoT monitoring, including smart "
    "urban forests, green roofs with automated irrigation, and constructed wetlands with water quality "
    "sensors, demonstrate the integration of ecological and technological approaches to urban "
    "sustainability [43]. The human-centric design philosophy ensures that technology serves as an "
    "enabler of human flourishing rather than a mechanism of surveillance and control, with IoT systems "
    "designed to enhance personal agency, community connection, and quality of life [43]."
)
add_para(sec_4_3_p2)

sec_4_3_p2b = (
    "The concept of regenerative cities extends sustainability beyond minimizing harm to actively restoring "
    "ecological systems through IoT-optimized interventions. Smart green infrastructure networks equipped "
    "with soil moisture sensors, species identification cameras, and microclimate monitors create data-driven "
    "approaches to urban ecology that maximize biodiversity outcomes per unit of green space investment. "
    "Circular economy platforms leveraging IoT tracking of material flows, product lifecycles, and waste "
    "streams enable industrial symbiosis at urban scales, where algorithmic matching connects waste "
    "generators with potential consumers of secondary materials in real-time markets. Social sustainability "
    "metrics including community cohesion indicators, access equity measurements, and well-being indices "
    "ensure that technological optimization does not come at the cost of social fabric, with IoT systems "
    "designed to strengthen rather than replace human connections and community institutions."
)
add_para(sec_4_3_p2b)

sec_4_3_p2c = (
    "The economic models underpinning smart city investments are evolving from traditional capital expenditure "
    "frameworks to platform-based approaches that generate recurring value through data monetization, "
    "efficiency gains, and innovation ecosystem development. Public-private partnerships structured around "
    "outcome-based contracts align incentives between technology providers and municipal governments, "
    "ensuring that IoT deployments deliver measurable public value rather than merely technological "
    "sophistication. The concept of city-as-a-platform enables municipalities to provide foundational "
    "IoT infrastructure upon which private sector innovators build value-added services, creating economic "
    "multiplier effects that justify public investment in connectivity and data infrastructure. International "
    "benchmarking and knowledge sharing through networks such as the Open and Agile Smart Cities initiative "
    "accelerate learning and reduce implementation risks by enabling cities to adapt proven approaches "
    "from peer municipalities while avoiding documented pitfalls and failure modes."
)
add_para(sec_4_3_p2c)

sec_4_3_p2d = (
    "Workforce development and human capital strategies are essential complements to technological investment "
    "in smart cities. The transition to IoT-enabled urban management creates demand for new skills in data "
    "science, systems integration, cybersecurity, and human-computer interaction design, requiring "
    "educational institutions and workforce development programs to evolve accordingly. Citizen digital "
    "literacy programs ensure that residents can effectively utilize IoT-enabled services and participate "
    "meaningfully in data governance decisions. The economic displacement risks associated with automation "
    "of urban services must be addressed through proactive transition support, reskilling programs, and "
    "inclusive economic development strategies that ensure smart city benefits are broadly shared rather "
    "than concentrated among technology-skilled populations."
)
add_para(sec_4_3_p2d)

sec_4_3_p3 = (
    "Competitive advantage for smart cities derives not merely from technology deployment but from the "
    "creation of innovation ecosystems that attract talent, investment, and entrepreneurship [43]. "
    "IoT infrastructure serves as a platform upon which diverse stakeholders including startups, research "
    "institutions, and civic organizations build novel services and solutions. Open innovation approaches, "
    "supported by standardized IoT APIs and open data platforms, lower barriers to entry for urban "
    "technology development. Cities that successfully balance technological advancement with social "
    "inclusion, environmental sustainability, and democratic governance position themselves for "
    "long-term prosperity in an increasingly urbanized and digitized world [43]. The strategic "
    "integration of all technological layers discussed throughout this chapter, from the foundational "
    "IoT architecture presented in Figure 1 through the communication networks of Figure 2, the "
    "applications framework of Figure 3, and the evolutionary roadmap of Figure 4, creates a "
    "comprehensive blueprint for cities pursuing digital transformation that serves both economic "
    "competitiveness and human well-being."
)
add_para(sec_4_3_p3)



# ============ CONCLUSION ============
add_heading_styled('5. Conclusion', level=1)

conclusion_p1 = (
    "This chapter has presented a comprehensive analysis of the Internet of Things and its transformative "
    "role in creating intelligent, sustainable, and human-centric smart cities. The multi-layered IoT "
    "architecture, encompassing perception, edge, fog, cloud, and application layers, provides the "
    "technological foundation upon which diverse urban services are constructed and integrated [1, 2]. "
    "The critical importance of communication infrastructure, from LPWAN technologies serving massive "
    "sensor deployments to 5G/6G networks enabling ultra-reliable low-latency applications, has been "
    "demonstrated as the connective tissue that unifies disparate urban systems into coherent digital "
    "ecosystems [11-17]. The integration of artificial intelligence, machine learning, and digital twin "
    "technologies transforms raw IoT data streams into actionable urban intelligence, enabling predictive "
    "and eventually autonomous management of city operations across transportation, energy, water, waste, "
    "and environmental domains [18-24]."
)
add_para(conclusion_p1)

conclusion_p2 = (
    "The practical applications examined across smart transportation, energy systems, water distribution, "
    "waste management, and environmental monitoring demonstrate that IoT technologies have progressed "
    "beyond experimental deployments to deliver measurable benefits at urban scale [25-36]. Connected "
    "and autonomous vehicles, intelligent traffic management, smart grids, and precision environmental "
    "monitoring represent mature applications delivering documented improvements in efficiency, "
    "sustainability, and quality of life. However, the realization of these benefits depends critically "
    "on addressing challenges in cybersecurity, privacy preservation, interoperability, and digital "
    "inclusion that can undermine public trust and limit adoption if not proactively managed [37-42]. "
    "The governance frameworks, security architectures, and citizen engagement strategies discussed "
    "provide essential complements to technological capabilities, ensuring that smart city development "
    "serves democratic values and equitable outcomes."
)
add_para(conclusion_p2)

conclusion_p3 = (
    "Looking forward, the convergence of edge intelligence, blockchain-based trust systems, and "
    "human-centric design principles points toward a future of increasingly autonomous yet accountable "
    "urban management [42, 43]. The strategic roadmap from connected to cognitive cities provides a "
    "phased approach that enables cities to build capabilities incrementally while managing risks and "
    "maintaining citizen engagement throughout the transformation journey. Success in this endeavor "
    "requires not only technological excellence but also institutional innovation, cross-sector "
    "collaboration, and unwavering commitment to sustainability, equity, and democratic governance. "
    "Cities that achieve this balance will secure enduring competitive advantage in attracting talent, "
    "investment, and innovation while delivering superior quality of life for their inhabitants. The "
    "IoT-enabled smart city thus represents not merely a technological achievement but a comprehensive "
    "reimagining of urban civilization for the challenges and opportunities of the twenty-first century."
)
add_para(conclusion_p3)

conclusion_p4 = (
    "The research and implementation landscape continues to evolve rapidly, with new technological "
    "capabilities, governance models, and application domains emerging continuously. The frameworks "
    "and architectures described in this chapter provide foundational understanding, but practitioners "
    "must remain attentive to emerging developments in quantum computing, artificial general intelligence, "
    "advanced materials, and biotechnology that may fundamentally alter smart city possibilities within "
    "the coming decade. Equally important is ongoing attention to social dynamics, political contexts, "
    "and cultural factors that shape technology adoption and determine whether smart city investments "
    "translate into genuine improvements in human welfare and urban sustainability. The interdisciplinary "
    "nature of smart city development demands collaboration across engineering, computer science, urban "
    "planning, social science, economics, and public administration to realize the full potential of "
    "IoT-enabled urban transformation while safeguarding the values and rights that define thriving "
    "democratic societies. Ultimately, the measure of smart city success lies not in the sophistication "
    "of deployed technology but in the demonstrable improvement of outcomes for urban residents across "
    "dimensions of health, prosperity, environmental quality, social connection, safety, and personal "
    "fulfillment, ensuring that the enormous investments in urban IoT infrastructure generate lasting "
    "returns for all segments of the urban population."
)
add_para(conclusion_p4)

# ============ REFERENCES ============
add_heading_styled('References', level=1)

references_all = [
    "[1] Zanella, A., Bui, N., Castellani, A., Vangelista, L., & Zorzi, M. (2014). Internet of Things for smart cities. IEEE Internet of Things Journal, 1(1), 22-32.",
    "[2] Gubbi, J., Buyya, R., Marusic, S., & Palaniswami, M. (2013). Internet of Things (IoT): A vision, architectural elements, and future directions. Future Generation Computer Systems, 29(7), 1645-1660.",
    "[3] Al-Fuqaha, A., Guizani, M., Mohammadi, M., Aledhari, M., & Ayyash, M. (2015). Internet of Things: A survey on enabling technologies, protocols, and applications. IEEE Communications Surveys & Tutorials, 17(4), 2347-2376.",
    "[4] Cisco Systems. (2023). Barcelona Smart City: IoT deployment case study. Cisco Smart City Solutions White Paper.",
    "[5] Shi, W., Cao, J., Zhang, Q., Li, Y., & Xu, L. (2016). Edge computing: Vision and challenges. IEEE Internet of Things Journal, 3(5), 637-646.",
    "[6] Satyanarayanan, M. (2017). The emergence of edge computing. Computer, 50(1), 30-39.",
    "[7] Bonomi, F., Milito, R., Zhu, J., & Addepalli, S. (2012). Fog computing and its role in the internet of things. Proceedings of the First MCC Workshop on Mobile Cloud Computing, 13-16.",
    "[8] Dastjerdi, A.V., & Buyya, R. (2016). Fog computing: Helping the Internet of Things realize its potential. Computer, 49(8), 112-116.",
    "[9] Moreschini, S., Pecorelli, F., Li, X., & Nikkola, V. (2022). Cloud continuum: The definition. IEEE Access, 10, 131876-131886.",
    "[10] Mao, Y., You, C., Zhang, J., Huang, K., & Letaief, K.B. (2017). A survey on mobile edge computing. IEEE Communications Surveys & Tutorials, 19(4), 2322-2358.",
    "[11] Lin, J., Yu, W., Zhang, N., Yang, X., Zhang, H., & Zhao, W. (2017). A survey on internet of things: Architecture, enabling technologies, security and privacy. IEEE Internet of Things Journal, 4(5), 1125-1142.",
    "[12] Agiwal, M., Roy, A., & Saxena, N. (2016). Next generation 5G wireless networks: A comprehensive survey. IEEE Communications Surveys & Tutorials, 18(3), 1617-1655.",
    "[13] Saad, W., Bennis, M., & Chen, M. (2020). A vision of 6G wireless systems: Applications, trends, technologies, and open research problems. IEEE Network, 34(3), 134-142.",
    "[14] Raza, U., Kulkarni, P., & Sooriyabandara, M. (2017). Low power wide area networks: An overview. IEEE Communications Surveys & Tutorials, 19(2), 855-873.",
    "[15] Wang, Y.P.E., Lin, X., Adhikary, A., Grovlen, A., Sui, Y., & Blankenship, Y. (2017). A primer on 3GPP narrowband internet of things. IEEE Communications Magazine, 55(3), 117-123.",
    "[16] Darroudi, S.M., & Gomez, C. (2017). Bluetooth low energy mesh networks: A survey. Sensors, 17(7), 1467.",
    "[17] Finn, N., Thubert, P., Voit, E., & Bellagamba, E. (2019). Deterministic networking architecture. IETF RFC 8655.",
    "[18] Mohammadi, M., Al-Fuqaha, A., Sorour, S., & Guizani, M. (2018). Deep learning for IoT big data and streaming analytics. IEEE Communications Surveys & Tutorials, 20(4), 2923-2960.",
    "[19] Zhang, J., Zheng, Y., & Qi, D. (2017). Deep spatio-temporal residual networks for citywide crowd flows prediction. Proceedings of AAAI Conference on Artificial Intelligence, 31(1), 1655-1661.",
    "[20] McMahan, B., Moore, E., Ramage, D., Hampson, S., & Arcas, B.A. (2017). Communication-efficient learning of deep networks from decentralized data. Proceedings of AISTATS, 54, 1273-1282.",
    "[21] Wei, H., Zheng, G., Yao, H., & Li, Z. (2018). IntelliLight: A reinforcement learning approach for intelligent traffic light control. Proceedings of ACM SIGKDD, 2496-2505.",
    "[22] Batty, M. (2018). Digital twins. Environment and Planning B: Urban Analytics and City Science, 45(5), 817-820.",
    "[23] National Research Foundation Singapore. (2024). Virtual Singapore: A 3D city model platform for knowledge sharing and collaboration. NRF Technical Report.",
    "[24] Fuller, A., Fan, Z., Day, C., & Barlow, C. (2020). Digital twin: Enabling technologies, challenges and open research. IEEE Access, 8, 108952-108971.",
    "[25] McKinsey Global Institute. (2023). Smart cities: Digital solutions for a more livable future. McKinsey & Company Report.",
    "[26] Guo, J., Huang, W., & Williams, B.M. (2014). Adaptive Kalman filter approach for stochastic short-term traffic flow rate prediction. Transportation Research Part C, 43, 50-64.",
    "[27] Badue, C., Guidolini, R., Carneiro, R.V., et al. (2021). Self-driving cars: A survey. Expert Systems with Applications, 165, 113816.",
    "[28] Jittrapirom, P., Caiati, V., Feneri, A.M., et al. (2017). Mobility as a service: A critical review of definitions, assessments of schemes, and key challenges. Urban Planning, 2(2), 13-25.",
    "[29] Fang, X., Misra, S., Xue, G., & Yang, D. (2012). Smart grid - The new and improved power grid: A survey. IEEE Communications Surveys & Tutorials, 14(4), 944-980.",
    "[30] Tuballa, M.L., & Abundo, M.L. (2016). A review of the development of smart grid technologies. Renewable and Sustainable Energy Reviews, 59, 710-725.",
    "[31] Puust, R., Kapelan, Z., Savic, D.A., & Koppel, T. (2010). A review of methods for leakage management in pipe networks. Urban Water Journal, 7(1), 25-45.",
    "[32] Gutierrez, J.M., Jensen, M., Henius, M., & Riaz, T. (2015). Smart waste collection system based on location intelligence. Procedia Computer Science, 61, 120-127.",
    "[33] Kumar, P., Morawska, L., Martani, C., et al. (2015). The rise of low-cost sensing for managing air pollution in cities. Environment International, 75, 199-205.",
    "[34] Concas, F., Mineraud, J., Lagerspetz, E., et al. (2021). Low-cost outdoor air quality monitoring and sensor calibration: A survey and critical analysis. ACM Computing Surveys, 54(3), 1-44.",
    "[35] Mydlarz, C., Sharma, M., Lockerman, Y., Steers, B., Silva, C., & Bello, J.P. (2019). The life of a New York City noise sensor network. Sensors, 19(6), 1415.",
    "[36] Nasajpour, M., Pouriyeh, S., Parizi, R.M., Dorodchi, M., Valero, M., & Arabnia, H.R. (2020). Internet of Things for current COVID-19 and future pandemics: An exploratory study. Journal of Healthcare Informatics Research, 4(4), 325-364.",
    "[37] Meijer, A., & Bolivar, M.P.R. (2016). Governing the smart city: A review of the literature on smart urban governance. International Review of Administrative Sciences, 82(2), 392-408.",
    "[38] Kitchin, R. (2014). The real-time city? Big data and smart urbanism. GeoJournal, 79(1), 1-14.",
    "[39] Frustaci, M., Pace, P., Aloi, G., & Fortino, G. (2018). Evaluating critical security issues of the IoT world. IEEE Internet of Things Journal, 5(4), 2483-2495.",
    "[40] Novo, O. (2018). Blockchain meets IoT: An architecture for scalable access management in IoT. IEEE Internet of Things Journal, 5(2), 1184-1195.",
    "[41] Bibri, S.E., & Krogstie, J. (2017). Smart sustainable cities of the future: An extensive interdisciplinary literature review. Sustainable Cities and Society, 31, 183-212.",
    "[42] Deng, S., Zhao, H., Fang, W., Yin, J., Dustdar, S., & Zomaya, A.Y. (2020). Edge intelligence: The confluence of edge computing and artificial intelligence. IEEE Internet of Things Journal, 7(8), 7457-7469.",
    "[43] Yigitcanlar, T., Kamruzzaman, M., Foth, M., et al. (2019). Can cities become smart without being sustainable? A systematic review of the literature. Sustainable Cities and Society, 45, 348-365.",
]

for ref in references_all:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'



# Save the document
output_path = '/projects/sandbox/AMMAN/Chapter_IoT_Smart_Cities.docx'
doc.save(output_path)
print(f"Document saved successfully: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
