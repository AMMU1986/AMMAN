"""
Create DOCX file for the Design Thinking Integration chapter.
Includes formatted text, tables, and embedded figures.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
import re

doc = Document()

# Page setup
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.size = Pt(10)
    cap_run.italic = True
    cap_run.font.name = 'Times New Roman'
    doc.add_paragraph()


def add_table(doc, headers, rows, caption):
    # Caption above table
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.bold = True
    cap_run.font.size = Pt(10)
    cap_run.font.name = 'Times New Roman'
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
    
    doc.add_paragraph()


def parse_md_and_build_docx(md_path, doc):
    """Parse the markdown file and build the DOCX document."""
    with open(md_path, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    table_caption = ""
    
    while i < len(lines):
        line = lines[i]
        
        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# ') and not line.startswith('## '):
            # Main title
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            add_heading(doc, text, level=1)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, level=2)
            i += 1
            continue
        
        # Handle figure placeholders
        if line.startswith('**[Figure'):
            # Extract figure number
            fig_match = re.search(r'Figure (\d+)', line)
            if fig_match:
                fig_num = fig_match.group(1)
                fig_files = {
                    '1': 'chapter_figures/Figure_1_DT_Analytical_Integration.png',
                    '2': 'chapter_figures/Figure_2_DT_Systems_Integration.png',
                    '3': 'chapter_figures/Figure_3_DT_Creative_Complementarity.png',
                    '4': 'chapter_figures/Figure_4_AI_Augmented_Framework.png',
                }
                fig_captions = {
                    '1': 'Figure 1: Integrated Framework of Design Thinking and Analytical Decision-Making in Strategic Innovation',
                    '2': 'Figure 2: Integration of Design Thinking and Systems Thinking for Strategic Innovation',
                    '3': 'Figure 3: Complementarity of Design Thinking and Creative Thinking in the Innovation Process',
                    '4': 'Figure 4: AI-Augmented Multimethod Framework for Future Strategic Thinking',
                }
                if fig_num in fig_files:
                    fig_path = os.path.join(os.path.dirname(md_path), fig_files[fig_num])
                    add_figure(doc, fig_path, fig_captions[fig_num])
            i += 1
            continue
        
        # Handle table caption lines (bold with Table N:)
        if line.startswith('**Table') and '**' in line[2:]:
            table_caption = line.replace('**', '').strip()
            i += 1
            continue
        
        # Handle table start
        if line.strip().startswith('|') and not in_table:
            in_table = True
            # Parse header
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_headers = cells
            table_rows = []
            i += 1
            # Skip separator line
            if i < len(lines) and '---' in lines[i]:
                i += 1
            continue
        
        # Handle table rows
        if in_table and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        
        # End of table
        if in_table and not line.strip().startswith('|'):
            in_table = False
            if table_headers and table_rows:
                add_table(doc, table_headers, table_rows, table_caption)
            table_headers = []
            table_rows = []
            table_caption = ""
            # Don't increment i, process current line
            continue
        
        # Handle bold keyword line
        if line.startswith('**Keywords:'):
            text = line.replace('**', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
            i += 1
            continue
        
        # Handle bold subheadings (like **New Product...**)
        if line.startswith('**') and line.endswith('**') and len(line) > 4:
            text = line[2:-2].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip() and not line.startswith('#') and not line.startswith('|'):
            add_para(doc, line.strip())
            i += 1
            continue
        
        # Empty lines
        i += 1
    
    # Handle any remaining table
    if in_table and table_headers and table_rows:
        add_table(doc, table_headers, table_rows, table_caption)


# Build the document
md_file = '/projects/sandbox/AMMAN/Chapter_Design_Thinking_Integration.md'
parse_md_and_build_docx(md_file, doc)

# Save
output_path = '/projects/sandbox/AMMAN/Chapter_Design_Thinking_Integration.docx'
doc.save(output_path)
print(f"DOCX saved to: {output_path}")

