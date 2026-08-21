"""
Create Chapter 12 Word Document: Industrial Translation, Scale-Up, Regulatory Aspects,
Artificial Intelligence Integration, Market Potential, and Future Perspectives
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, 
                           font_size=None, alignment=None, space_after=None, space_before=None):
    """Add a formatted paragraph to the document."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
    return para

def add_body_text(doc, text, space_after=6):
    """Add body text with proper formatting."""
    para = doc.add_paragraph(style='Normal')
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.5
    # Handle italic text within body (for species names)
    parts = text.split('*')
    for i, part in enumerate(parts):
        run = para.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:  # Odd indices are italic
            run.italic = True
    return para

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # =========================================================================
    # TITLE
    # =========================================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Chapter 12')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle_para.add_run(
        'Industrial Translation, Scale-Up, Regulatory Aspects, Artificial Intelligence '
        'Integration, Market Potential, and Future Perspectives'
    )
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # =========================================================================
    # ABSTRACT
    # =========================================================================
    abstract_heading = doc.add_paragraph()
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    abstract_run = abstract_heading.add_run('Abstract')
    abstract_run.bold = True
    abstract_run.font.size = Pt(13)
    
    abstract_text = (
        "Microbial-derived bio-nanomaterials represent a paradigm shift in sustainable nanotechnology, "
        "offering green, cost-effective, and scalable alternatives to conventional chemical and physical "
        "synthesis methods. This chapter provides a comprehensive examination of the multidimensional "
        "challenges and opportunities associated with translating laboratory-scale microbial nanomaterial "
        "synthesis to industrial production. The discussion encompasses scale-up strategies, manufacturing "
        "technologies, quality-by-design approaches, and economic considerations essential for commercial "
        "viability. Regulatory frameworks governing biomedical and environmental applications are critically "
        "analyzed, alongside toxicological assessments, biosafety protocols, and environmental risk evaluation. "
        "The transformative role of artificial intelligence in optimizing microbial synthesis parameters, "
        "predicting nanomaterial properties, enabling real-time process control, and accelerating materials "
        "discovery is thoroughly explored. Market potential, commercialization pathways, intellectual property "
        "landscapes, and investment opportunities are evaluated across biomedical and environmental sectors. "
        "Finally, future perspectives addressing next-generation bio-nanomaterials, convergence of biotechnology "
        "with nanotechnology and artificial intelligence, and a comprehensive research roadmap are presented "
        "to guide the sustainable development of this rapidly evolving field."
    )
    
    abstract_para = doc.add_paragraph()
    abstract_para.paragraph_format.space_after = Pt(12)
    abstract_para.paragraph_format.line_spacing = 1.5
    run = abstract_para.add_run(abstract_text)
    run.font.size = Pt(11)
    run.italic = True
    
    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.space_after = Pt(18)
    kw_run1 = kw_para.add_run('Keywords: ')
    kw_run1.bold = True
    kw_run1.font.size = Pt(11)
    kw_run2 = kw_para.add_run(
        'Microbial bio-nanomaterials; Industrial scale-up; Regulatory frameworks; '
        'Artificial intelligence; Machine learning; Quality-by-design; Commercialization; '
        'Green nanotechnology; Sustainable manufacturing; Future perspectives'
    )
    kw_run2.font.size = Pt(11)
    
    # =========================================================================
    # SECTION I
    # =========================================================================
    section_heading = doc.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(18)
    section_heading.paragraph_format.space_after = Pt(12)
    section_run = section_heading.add_run(
        'Section I: Industrial Translation and Scale-Up of Microbial-Derived Bio-Nanomaterials'
    )
    section_run.bold = True
    section_run.font.size = Pt(13)
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- 12.1 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.1 From Laboratory Synthesis to Industrial Production')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc, 
        "The transition of microbial-derived bio-nanomaterials from laboratory proof-of-concept "
        "demonstrations to reproducible, large-scale industrial manufacturing constitutes one of the "
        "most significant challenges facing contemporary green nanotechnology [1]. While numerous studies "
        "have successfully demonstrated the biosynthesis of metallic, metal oxide, and semiconductor "
        "nanoparticles using bacteria, fungi, algae, and yeast at bench scale, the fundamental requirements "
        "for industrial translation demand systematic approaches to process standardization, reproducibility, "
        "and economic viability [2]. The laboratory-to-industry gap is characterized by multiple technical, "
        "economic, and regulatory barriers that must be addressed through integrated strategies combining "
        "microbiology, process engineering, and materials science [3]."
    )
    
    add_body_text(doc,
        "The fundamental challenge of industrial translation lies in maintaining the precise control "
        "over nanoparticle characteristics that is achievable in small-scale laboratory settings while "
        "operating at volumetric scales orders of magnitude larger. Laboratory synthesis typically "
        "involves carefully controlled flask experiments with volumes ranging from milliliters to a few "
        "liters, where parameters such as mixing uniformity, temperature homogeneity, and precursor "
        "distribution are inherently well-controlled. At industrial scales exceeding thousands of liters, "
        "maintaining equivalent parameter uniformity requires sophisticated engineering solutions "
        "addressing heat and mass transfer limitations, mixing dead zones, and gradients in dissolved "
        "oxygen and nutrient concentrations that inevitably arise in large-volume systems."
    )
    
    add_body_text(doc,
        "Furthermore, the economic landscape of industrial biosynthesis must account for the total "
        "cost of goods including upstream processing costs for media preparation and sterilization, "
        "the fermentation process itself, and extensive downstream processing required to isolate "
        "and purify nanoparticles from complex biological matrices. Supply chain considerations for "
        "metal precursors, culture media components, and specialty chemicals must ensure reliable "
        "sourcing at industrial quantities while maintaining consistent quality specifications. "
        "Workforce requirements span microbiology, biochemical engineering, analytical chemistry, "
        "and quality assurance disciplines, necessitating interdisciplinary teams with specialized "
        "training in both biological processes and nanomaterial characterization."
    )
    
    add_body_text(doc,
        "The selection of appropriate microbial strains represents the foundational decision in scaling "
        "biosynthesis processes. Industrial strains must demonstrate robust growth characteristics, high "
        "nanoparticle yields, tolerance to process perturbations, and genetic stability over extended "
        "cultivation periods [4]. Organisms such as *Bacillus subtilis*, *Pseudomonas aeruginosa*, "
        "*Aspergillus niger*, *Saccharomyces cerevisiae*, and *Chlorella vulgaris* have emerged as "
        "promising candidates due to their well-characterized metabolic pathways, established fermentation "
        "protocols, and generally recognized as safe (GRAS) status [5]. The substrate selection, including "
        "metal precursor concentration, carbon and nitrogen sources, and trace elements, must be optimized "
        "for both nanoparticle quality and production economics [6]."
    )
    
    add_body_text(doc,
        "Process standardization requires the development of detailed standard operating procedures "
        "encompassing inoculum preparation, fermentation conditions, metal salt addition timing and "
        "concentration, and harvest protocols [7]. Batch-to-batch consistency remains a critical challenge, "
        "as biological systems inherently exhibit variability in enzyme expression levels, metabolic activity, "
        "and reducing capacity. Statistical process control methods, including control charts, capability "
        "indices, and design of experiments approaches, provide systematic frameworks for monitoring and "
        "maintaining product consistency [8]. The implementation of process analytical technology (PAT) "
        "enables real-time monitoring of critical process parameters including pH, dissolved oxygen, "
        "temperature, optical density, and nanoparticle formation kinetics [9]."
    )
    
    add_body_text(doc,
        "Seed train development represents another critical aspect of industrial translation, requiring "
        "optimization of inoculum propagation stages to ensure consistent cell density, viability, and "
        "metabolic state at the time of production-scale inoculation. The number of seed stages, "
        "transfer volumes, and timing of transfers must be validated to prevent genetic drift, maintain "
        "production strain characteristics, and ensure reproducible performance across manufacturing "
        "campaigns. Master and working cell bank systems, following pharmaceutical industry best "
        "practices, provide characterized starting materials with documented identity, purity, and "
        "stability profiles essential for regulatory compliance and manufacturing consistency."
    )
    
    # TABLE 1
    add_formatted_paragraph(doc, 
        'Table 1: Comparison of Microbial Strains for Industrial Bio-Nanomaterial Production',
        bold=True, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    
    table1 = doc.add_table(rows=7, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Table 1 headers
    headers1 = ['Microbial Strain', 'Nanoparticle Type', 'Size Range (nm)', 'Key Advantages', 'Scalability Rating']
    for i, header in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Table 1 data
    table1_data = [
        ['Bacillus subtilis', 'Ag, Au, ZnO', '5-50', 'GRAS status, robust growth, high yield', 'High'],
        ['Pseudomonas aeruginosa', 'Ag, Au, Pd', '10-100', 'Versatile, multiple metals, fast synthesis', 'Medium-High'],
        ['Aspergillus niger', 'Ag, Au, Fe₃O₄', '20-80', 'High enzyme production, scalable fermentation', 'High'],
        ['Saccharomyces cerevisiae', 'Au, Ag, CdS', '2-30', 'Food-grade, well-established bioprocess', 'Very High'],
        ['Chlorella vulgaris', 'Ag, Au, TiO₂', '5-40', 'Photosynthetic, CO₂ utilization, sustainable', 'Medium'],
        ['Lactobacillus sp.', 'Ag, Au, Se', '10-60', 'Probiotic, biocompatible, safe handling', 'High'],
    ]
    
    for i, row_data in enumerate(table1_data):
        for j, cell_text in enumerate(row_data):
            cell = table1.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "D6E4F0")
    
    add_formatted_paragraph(doc, '', space_after=12)
    
    # --- 12.2 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.2 Scale-Up Strategies and Manufacturing Technologies')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Bioreactor design and selection constitute critical determinants of successful scale-up for "
        "microbial nanomaterial production [10]. Stirred-tank bioreactors remain the most widely employed "
        "configuration due to their well-understood mixing characteristics, scalability, and established "
        "design correlations. However, the unique requirements of nanomaterial biosynthesis, including "
        "controlled metal ion delivery, maintenance of optimal redox conditions, and prevention of "
        "nanoparticle aggregation, necessitate modifications to conventional bioreactor designs [11]. "
        "Alternative configurations including airlift reactors, membrane bioreactors, packed-bed systems, "
        "and microfluidic platforms offer specific advantages for particular microbial systems and "
        "nanoparticle types [12]."
    )
    
    add_body_text(doc,
        "The selection of appropriate bioreactor configuration depends on multiple factors including "
        "the microbial system employed, oxygen requirements, shear sensitivity, nanoparticle formation "
        "mechanism (intracellular versus extracellular), and target production scale. For extracellular "
        "biosynthesis, where nanoparticles form in the culture supernatant through interaction of "
        "secreted enzymes and metabolites with metal precursors, stirred-tank and airlift configurations "
        "provide adequate mixing while enabling straightforward product recovery. Intracellular "
        "biosynthesis processes require additional cell disruption and extraction steps, favoring "
        "reactor configurations that maximize biomass production and facilitate efficient harvest. "
        "Immobilized cell systems using biofilm reactors or entrapped cultures offer advantages of "
        "continuous operation and simplified product separation for certain applications."
    )
    
    add_body_text(doc,
        "Process intensification strategies aim to maximize volumetric productivity while minimizing "
        "resource consumption and waste generation. Continuous production approaches, including chemostat "
        "cultivation and perfusion systems, offer advantages over batch processes including steady-state "
        "operation, consistent product quality, reduced downtime, and improved space-time yields [13]. "
        "Fed-batch strategies with programmed nutrient and metal precursor feeding profiles enable precise "
        "control over microbial growth phases and nanoparticle nucleation and growth kinetics [14]. The "
        "integration of in-situ product removal techniques prevents product inhibition and enables "
        "continuous harvesting of nanoparticles from the culture medium."
    )
    
    add_body_text(doc,
        "Scale-up correlations derived from dimensional analysis and similarity principles guide "
        "the transition from small-scale to large-scale bioreactors while maintaining equivalent "
        "process performance. Key dimensionless groups including Reynolds number (characterizing "
        "flow regime), power number (relating impeller power to fluid properties), and Damköhler "
        "number (comparing reaction rate to transport rate) provide frameworks for maintaining "
        "process similarity across scales. However, perfect geometric and dynamic similarity is "
        "rarely achievable in practice, requiring empirical validation of scale-up correlations "
        "through intermediate pilot-scale studies. Computational fluid dynamics (CFD) modeling "
        "provides additional insights into mixing patterns, shear stress distributions, and "
        "mass transfer characteristics at different scales, enabling rational selection of "
        "operating parameters for large-scale bioreactors prior to expensive pilot trials."
    )
    
    add_body_text(doc,
        "Optimization of microbial growth conditions and nanoparticle synthesis parameters requires "
        "systematic approaches including response surface methodology, Taguchi methods, and evolutionary "
        "optimization algorithms [15]. Critical parameters encompassing temperature, pH, agitation speed, "
        "aeration rate, metal precursor concentration, incubation time, and biomass-to-substrate ratio "
        "must be simultaneously optimized to achieve desired nanoparticle characteristics including size "
        "distribution, morphology, crystallinity, and surface properties [16]. As illustrated in Figure 1, "
        "the scale-up process involves multiple stages from flask-scale optimization through pilot-plant "
        "validation to full industrial production."
    )
    
    # Insert Figure 1
    add_formatted_paragraph(doc, '', space_after=6)
    fig1_path = '/projects/sandbox/AMMAN/chapter12_figures/Figure_1_Scaleup_Pathway.png'
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig1_caption = doc.add_paragraph()
    fig1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig1_caption.paragraph_format.space_after = Pt(12)
    cap_run = fig1_caption.add_run(
        'Figure 1: Schematic representation of the scale-up pathway from laboratory to industrial '
        'production of microbial-derived bio-nanomaterials, showing critical process parameters, '
        'quality attributes, and economic considerations at each stage.'
    )
    cap_run.bold = True
    cap_run.font.size = Pt(9)
    
    add_body_text(doc,
        "Downstream processing encompasses the critical operations of nanoparticle separation, "
        "purification, stabilization, and formulation [17]. Centrifugation, ultrafiltration, dialysis, "
        "and chromatographic methods are employed individually or in combination to separate nanoparticles "
        "from residual biomass, culture medium components, and unreacted precursors. Stabilization "
        "strategies including capping agent addition, lyophilization, spray drying, and encapsulation "
        "ensure long-term storage stability and preservation of functional properties [18]. The "
        "formulation stage adapts purified nanomaterials for specific end-use applications, incorporating "
        "considerations of dispersibility, biocompatibility, and application-specific performance requirements."
    )
    
    add_body_text(doc,
        "The challenge of downstream processing is further complicated by the need to preserve the "
        "biological capping agents that naturally form on microbially-synthesized nanoparticles. These "
        "biogenic surface coatings, composed of proteins, polysaccharides, and other biomolecules, "
        "often contribute significantly to nanoparticle stability, biocompatibility, and functional "
        "properties. Overly aggressive purification procedures that strip these biological layers "
        "may compromise the unique advantages of biosynthesized nanomaterials. Therefore, purification "
        "strategies must balance the removal of unwanted contaminants with preservation of beneficial "
        "surface functionalities, requiring careful optimization for each specific nanoparticle system "
        "and intended application."
    )
    
    # --- 12.3 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.3 Quality-by-Design and Economic Considerations')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "The Quality-by-Design (QbD) paradigm, originally developed for pharmaceutical manufacturing, "
        "provides a systematic framework for ensuring consistent quality in bio-nanomaterial production [19]. "
        "This approach requires identification of critical quality attributes (CQAs) including nanoparticle "
        "size, polydispersity index, zeta potential, crystalline structure, surface chemistry, and biological "
        "activity. Corresponding critical process parameters (CPPs) that influence these quality attributes "
        "must be identified through risk assessment and experimental investigation [20]. The establishment "
        "of a design space defining acceptable ranges for all CPPs ensures that operation within this "
        "multidimensional space consistently produces nanomaterials meeting predefined quality specifications."
    )
    
    add_body_text(doc,
        "Risk assessment tools including Ishikawa diagrams, failure mode and effects analysis (FMEA), "
        "and fault tree analysis provide systematic methodologies for identifying potential sources of "
        "variability and their impact on product quality. These tools enable prioritization of critical "
        "parameters requiring tight control versus those with minimal impact on final product attributes. "
        "Process validation studies, typically conducted at pilot scale, confirm that the identified "
        "design space boundaries consistently produce acceptable product when operated under anticipated "
        "commercial manufacturing conditions. Continued process verification during commercial "
        "manufacturing ensures ongoing state of control through statistical monitoring of quality "
        "attributes and process parameters."
    )
    
    add_body_text(doc,
        "Economic analysis of microbial nanomaterial production must comprehensively evaluate capital "
        "expenditure, operating costs, raw material inputs, energy consumption, labor requirements, and "
        "waste treatment costs [21]. Techno-economic assessments comparing biosynthesis with conventional "
        "chemical reduction and physical methods consistently demonstrate the potential for cost advantages, "
        "particularly when utilizing waste substrates, operating under ambient conditions, and eliminating "
        "hazardous reagents [22]. However, current limitations including lower volumetric productivity, "
        "longer processing times, and more complex purification requirements partially offset these "
        "advantages and represent targets for process optimization."
    )
    
    add_body_text(doc,
        "Sensitivity analysis of techno-economic models reveals that nanoparticle yield concentration "
        "and downstream processing efficiency represent the most influential cost drivers for microbial "
        "biosynthesis. Improvements in these parameters through strain engineering, process optimization, "
        "and advanced purification technologies would most rapidly improve economic competitiveness. "
        "Revenue models must also consider the premium pricing achievable for green-certified and "
        "sustainably-produced nanomaterials in markets where environmental credentials command value. "
        "Government incentives including tax credits for green manufacturing, carbon pricing mechanisms, "
        "and preferential procurement policies for sustainable products can further improve the economic "
        "case for biological production routes by internalizing environmental costs currently borne by "
        "society rather than manufacturers."
    )
    
    # TABLE 2
    add_formatted_paragraph(doc,
        'Table 2: Economic and Sustainability Comparison of Nanomaterial Synthesis Methods',
        bold=True, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['Parameter', 'Microbial Biosynthesis', 'Chemical Reduction', 'Physical Methods']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "2E7D32")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table2_data = [
        ['Capital Investment', 'Moderate ($0.5-2M)', 'High ($2-5M)', 'Very High ($5-20M)'],
        ['Operating Costs ($/kg)', '$50-200', '$100-500', '$500-2000'],
        ['Energy Consumption', 'Low (ambient conditions)', 'Moderate (heating required)', 'High (extreme conditions)'],
        ['Hazardous Reagents', 'None/Minimal', 'Multiple toxic chemicals', 'Inert gases, high vacuum'],
        ['Waste Generation', 'Biodegradable, minimal', 'Toxic chemical waste', 'Minimal but energy-intensive'],
        ['Carbon Footprint (kg CO₂/kg)', '5-15', '20-80', '50-200'],
        ['Scalability Timeline', '2-4 years to industrial', '1-2 years (established)', 'Limited scalability'],
    ]
    
    for i, row_data in enumerate(table2_data):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "E8F5E9")
    
    add_formatted_paragraph(doc, '', space_after=6)
    
    add_body_text(doc,
        "Life-cycle assessment (LCA) provides a comprehensive methodology for evaluating the environmental "
        "impacts associated with bio-nanomaterial production across all stages from raw material extraction "
        "through end-of-life disposal [23]. Comparative LCA studies have demonstrated that microbial "
        "synthesis routes typically exhibit lower carbon footprints, reduced energy intensity, decreased "
        "hazardous waste generation, and improved resource efficiency compared to chemical and physical "
        "alternatives [24]. Resource efficiency can be further enhanced through integration with waste "
        "valorization strategies, utilizing agricultural residues, industrial effluents, or food processing "
        "wastes as both growth substrates and reducing agents for nanoparticle synthesis."
    )
    
    # =========================================================================
    # SECTION II
    # =========================================================================
    section_heading = doc.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(24)
    section_heading.paragraph_format.space_after = Pt(12)
    section_run = section_heading.add_run(
        'Section II: Regulatory, Safety, and Translational Considerations'
    )
    section_run.bold = True
    section_run.font.size = Pt(13)
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- 12.4 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.4 Regulatory Frameworks for Microbial-Derived Nanomaterials')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "The regulatory landscape governing microbial-derived nanomaterials is complex, fragmented, "
        "and continues to evolve as understanding of nanomaterial properties and biological interactions "
        "advances [25]. Different regulatory jurisdictions employ varying approaches to nanomaterial "
        "classification, with implications for required testing, documentation, and approval pathways. "
        "In the United States, the Environmental Protection Agency (EPA) regulates nanomaterials under "
        "the Toxic Substances Control Act (TSCA), while the Food and Drug Administration (FDA) oversees "
        "biomedical applications under existing pharmaceutical, medical device, or cosmetic regulatory "
        "frameworks [26]. The European Union applies the Registration, Evaluation, Authorization, and "
        "Restriction of Chemicals (REACH) regulation with specific provisions for nanomaterials, "
        "requiring detailed characterization and safety data [27]."
    )
    
    add_body_text(doc,
        "The regulatory challenge is compounded by the dual biological-nanomaterial nature of these "
        "products, which may simultaneously trigger requirements from both biotechnology and "
        "nanotechnology regulatory frameworks. In many jurisdictions, the use of living microorganisms "
        "in manufacturing processes invokes biosafety regulations governing genetically modified "
        "organisms, contained use requirements, and environmental release protocols, even when the "
        "final nanomaterial product contains no viable organisms. This regulatory complexity creates "
        "uncertainty for manufacturers and may require engagement with multiple regulatory agencies "
        "for a single product, significantly increasing time and cost for market authorization. "
        "Pre-submission meetings with regulatory authorities and early engagement strategies can help "
        "clarify requirements and avoid costly development pathway errors."
    )
    
    add_body_text(doc,
        "Product classification determines the applicable regulatory pathway and associated requirements "
        "for microbial-derived nanomaterials. Materials intended for drug delivery applications must "
        "navigate pharmaceutical regulatory frameworks requiring extensive preclinical and clinical "
        "testing [28]. Those designed for environmental remediation may fall under environmental "
        "protection regulations with different data requirements. The biological origin of these "
        "materials introduces additional considerations regarding microbial safety, sterility assurance, "
        "and potential immunogenicity of residual biological components. Documentation requirements "
        "encompass detailed manufacturing records, quality control specifications, stability data, and "
        "comprehensive physicochemical characterization using standardized methods."
    )
    
    add_body_text(doc,
        "International regulatory harmonization remains challenging due to differing national approaches "
        "to nanomaterial definition, risk assessment methodologies, and acceptable safety margins [29]. "
        "Organizations including the International Organization for Standardization (ISO), the Organisation "
        "for Economic Co-operation and Development (OECD), and the International Council for Harmonisation "
        "(ICH) are developing standardized testing protocols, characterization methods, and risk assessment "
        "frameworks to facilitate international regulatory convergence. The establishment of mutual "
        "recognition agreements and common technical standards would significantly reduce regulatory "
        "burden and accelerate market access for bio-nanomaterial products across multiple jurisdictions."
    )
    
    # --- 12.5 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.5 Toxicological, Biosafety, and Environmental Risk Assessment')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Comprehensive toxicological assessment of microbial-derived nanomaterials encompasses evaluation "
        "of cytotoxicity, genotoxicity, immunotoxicity, and chronic biological effects through both "
        "in vitro and in vivo testing paradigms [30]. Cytotoxicity studies employing standardized assays "
        "including MTT, WST-1, LDH release, and live/dead staining across relevant cell lines provide "
        "initial safety screening. Genotoxicity assessment through Ames tests, comet assays, micronucleus "
        "tests, and chromosomal aberration analysis evaluates potential mutagenic effects [31]. "
        "Immunotoxicity evaluation addresses concerns regarding nanoparticle-induced inflammatory "
        "responses, complement activation, and potential autoimmune effects."
    )
    
    add_body_text(doc,
        "The unique physicochemical properties of nanoparticles, including their high surface "
        "area-to-volume ratio, quantum effects, and ability to cross biological barriers, necessitate "
        "toxicological evaluation strategies that go beyond conventional chemical toxicity testing. "
        "Nanoparticle-specific considerations include the potential for lysosomal disruption, "
        "mitochondrial dysfunction, oxidative stress generation, and interference with cellular "
        "signaling pathways. Furthermore, standard toxicological assays may produce artifactual "
        "results due to nanoparticle interference with assay reagents, optical measurements, or "
        "enzymatic reactions, requiring careful validation of assay suitability for each specific "
        "nanomaterial system. Dose metrics for nanoparticle toxicology remain debated, with mass "
        "concentration, particle number, surface area, and particle volume each proposed as the most "
        "relevant dose descriptor depending on the mechanism of toxicity."
    )
    
    add_body_text(doc,
        "The biological origin of microbial-derived nanomaterials introduces unique biosafety "
        "considerations including potential microbial contamination, endotoxin presence, residual "
        "nucleic acids, and allergenic proteins [32]. Rigorous purification protocols and quality "
        "control testing must ensure removal of microbial components below acceptable safety thresholds. "
        "Nanoparticle stability under physiological conditions, potential for protein corona formation, "
        "and long-term degradation behavior influence biocompatibility and must be characterized under "
        "relevant exposure scenarios. The biogenic capping layers present on microbially-synthesized "
        "nanoparticles may actually enhance biocompatibility compared to chemically-synthesized "
        "counterparts, though this requires systematic validation [33]."
    )
    
    add_body_text(doc,
        "Environmental fate assessment evaluates the behavior of released bio-nanomaterials in aquatic, "
        "terrestrial, and atmospheric compartments [34]. Parameters including dissolution rate, aggregation "
        "behavior, sedimentation, and transformation under environmental conditions determine exposure "
        "concentrations for ecological receptors. Bioaccumulation potential through food web transfer and "
        "biomagnification in higher trophic levels requires investigation through standardized ecotoxicity "
        "testing with representative organisms. Risk mitigation strategies encompass exposure minimization "
        "through containment, safe-by-design approaches incorporating biodegradability, and end-of-life "
        "management protocols ensuring responsible disposal or recycling of nanomaterial-containing products."
    )
    
    add_body_text(doc,
        "Environmental transformation processes significantly alter the properties, behavior, and "
        "potential impacts of released bio-nanomaterials over time. Processes including oxidation, "
        "sulfidation, dissolution, and interaction with natural organic matter can fundamentally change "
        "nanoparticle surface chemistry, aggregation state, and bioavailability. For example, silver "
        "nanoparticles released into aquatic environments undergo sulfidation to form silver sulfide, "
        "dramatically reducing their dissolution rate and antimicrobial activity. Understanding these "
        "transformation pathways is essential for accurate environmental risk assessment, as the "
        "pristine manufactured form may differ substantially from the environmentally-relevant species "
        "to which organisms are actually exposed. Long-term mesocosm and field studies provide "
        "critical data on transformation kinetics and ultimate environmental fate under realistic "
        "conditions that cannot be replicated in short-term laboratory studies."
    )
    
    add_body_text(doc,
        "As depicted in Figure 2, the comprehensive risk assessment framework for microbial-derived "
        "bio-nanomaterials integrates hazard identification, exposure assessment, dose-response "
        "characterization, and risk management into a cohesive decision-making process."
    )
    
    # Insert Figure 2
    add_formatted_paragraph(doc, '', space_after=6)
    fig2_path = '/projects/sandbox/AMMAN/chapter12_figures/Figure_2_Risk_Assessment_Framework.png'
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig2_caption = doc.add_paragraph()
    fig2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig2_caption.paragraph_format.space_after = Pt(12)
    cap_run = fig2_caption.add_run(
        'Figure 2: Integrated risk assessment and regulatory decision framework for microbial-derived '
        'bio-nanomaterials, illustrating the pathway from hazard identification through regulatory '
        'decision-making to application-specific approval routes.'
    )
    cap_run.bold = True
    cap_run.font.size = Pt(9)
    
    # --- 12.6 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.6 Clinical and Environmental Translation')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Preclinical validation of microbial-derived nanomaterials for biomedical applications follows "
        "established pharmaceutical development paradigms adapted for nanomaterial-specific considerations [35]. "
        "This encompasses comprehensive physicochemical characterization, in vitro efficacy and safety "
        "assessment, pharmacokinetic and biodistribution studies in animal models, and demonstration of "
        "therapeutic efficacy in disease-relevant models. The transition from preclinical to clinical "
        "studies requires Investigational New Drug (IND) applications supported by manufacturing process "
        "validation, stability data, and toxicology packages meeting regulatory standards."
    )
    
    add_body_text(doc,
        "The preclinical development pathway for nanomaterial-based therapeutics typically requires "
        "demonstration of dose-response relationships, identification of maximum tolerated doses, "
        "characterization of absorption, distribution, metabolism, and excretion (ADME) properties, "
        "and evaluation of potential off-target effects in relevant organ systems. For nanomaterials "
        "intended for systemic administration, particular attention must be directed toward "
        "hemocompatibility, complement activation, and potential accumulation in reticuloendothelial "
        "system organs including liver and spleen. The development of appropriate in vitro-in vivo "
        "correlations (IVIVC) for nanomaterial pharmacokinetics remains an active area of research "
        "that would significantly accelerate preclinical development timelines."
    )
    
    add_body_text(doc,
        "Field validation of microbial-derived nanomaterials for environmental applications requires "
        "demonstration of performance under realistic environmental conditions including variable pH, "
        "temperature, ionic strength, natural organic matter, and competing contaminants [36]. Pilot-scale "
        "field trials evaluate long-term stability, sustained activity, potential ecological impacts, and "
        "practical implementation requirements. The development of application-specific formulations, "
        "delivery systems, and deployment strategies ensures effective performance while minimizing "
        "unintended environmental consequences."
    )
    
    add_body_text(doc,
        "Ethical considerations surrounding the development and deployment of microbial-derived "
        "nanomaterials encompass responsible innovation principles, stakeholder engagement, transparency, "
        "and equitable benefit sharing [37]. Public acceptance requires effective communication of both "
        "benefits and risks, addressing concerns regarding novel biological technologies and nanomaterial "
        "safety. Traceability systems enabling tracking of nanomaterials throughout their lifecycle "
        "support responsible stewardship and facilitate recall procedures if safety concerns emerge post-market."
    )
    
    add_body_text(doc,
        "The social license to operate for bio-nanomaterial technologies depends critically on "
        "transparent engagement with diverse stakeholders including patient advocacy groups, "
        "environmental organizations, local communities near production facilities, and the general "
        "public. Lessons from previous public controversies surrounding genetically modified organisms "
        "and nanotechnology highlight the importance of early, inclusive, and honest dialogue about "
        "both potential benefits and legitimate uncertainties. Citizen science initiatives, public "
        "deliberation forums, and accessible educational resources can build informed public understanding "
        "while ensuring that societal values and concerns are appropriately reflected in research "
        "priorities and governance frameworks."
    )
    
    # =========================================================================
    # SECTION III
    # =========================================================================
    section_heading = doc.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(24)
    section_heading.paragraph_format.space_after = Pt(12)
    section_run = section_heading.add_run(
        'Section III: Artificial Intelligence Integration and Market Potential'
    )
    section_run.bold = True
    section_run.font.size = Pt(13)
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- 12.7 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.7 AI-Driven Design and Optimization of Bio-Nanomaterials')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Artificial intelligence and machine learning have emerged as transformative tools for accelerating "
        "the design, optimization, and discovery of microbial-derived bio-nanomaterials [38]. Traditional "
        "experimental approaches to nanomaterial development rely on iterative trial-and-error optimization "
        "across multidimensional parameter spaces, requiring extensive time and resource investment. Machine "
        "learning algorithms, including random forests, support vector machines, neural networks, and "
        "gradient boosting methods, can learn complex relationships between synthesis parameters and "
        "nanoparticle properties from historical experimental data, enabling predictive design of materials "
        "with targeted characteristics [39]."
    )
    
    add_body_text(doc,
        "The application of AI to bio-nanomaterial development addresses a fundamental limitation "
        "of traditional approaches: the curse of dimensionality inherent in multi-parameter optimization "
        "problems. A typical microbial nanomaterial synthesis involves optimization of ten or more "
        "independent variables, each potentially interacting non-linearly with others. Full factorial "
        "experimental exploration of such parameter spaces would require thousands to millions of "
        "experiments, far exceeding practical resource constraints. Machine learning approaches overcome "
        "this limitation by learning from relatively sparse experimental data and generalizing to "
        "unexplored regions of parameter space, enabling efficient identification of optimal conditions "
        "with orders of magnitude fewer experiments than exhaustive search strategies."
    )
    
    add_body_text(doc,
        "The prediction of nanoparticle size, morphology, composition, and functionality from microbial "
        "synthesis conditions represents a primary application of machine learning in bio-nanomaterial "
        "research. Training datasets encompassing microbial species, growth conditions, metal precursor "
        "type and concentration, pH, temperature, reaction time, and resulting nanoparticle properties "
        "enable the development of quantitative structure-property relationship (QSPR) models [40]. "
        "These models can subsequently guide experimental design by identifying optimal synthesis "
        "conditions for achieving desired material specifications without exhaustive experimental screening."
    )
    
    add_body_text(doc,
        "AI-assisted optimization of microbial synthesis parameters employs advanced optimization "
        "algorithms including Bayesian optimization, genetic algorithms, and reinforcement learning "
        "to efficiently navigate complex parameter landscapes [41]. Bayesian optimization is particularly "
        "well-suited for bio-nanomaterial development as it balances exploration of unknown parameter "
        "regions with exploitation of promising conditions, achieving optimal solutions with minimal "
        "experimental iterations. The integration of active learning frameworks enables AI systems to "
        "autonomously design and request informative experiments, creating closed-loop optimization "
        "cycles that rapidly converge on optimal synthesis conditions."
    )
    
    add_body_text(doc,
        "Digital twin technology creates virtual representations of microbial biosynthesis processes, "
        "incorporating mechanistic models, empirical correlations, and real-time sensor data to simulate "
        "process behavior under varying conditions [42]. These digital replicas enable scenario analysis, "
        "process optimization, troubleshooting, and operator training without disrupting actual production. "
        "Predictive modeling combines thermodynamic, kinetic, and transport models with data-driven "
        "approaches to forecast nanoparticle nucleation, growth, and aggregation dynamics. Automated "
        "high-throughput platforms integrating robotic liquid handling, parallelized microbioreactors, "
        "and automated characterization enable rapid experimental exploration guided by AI-directed "
        "experimental design."
    )
    
    add_body_text(doc,
        "Figure 3 illustrates the integrated AI-driven workflow for bio-nanomaterial design, encompassing "
        "data collection, model training, optimization, and experimental validation in a closed-loop "
        "discovery cycle."
    )
    
    # Insert Figure 3
    add_formatted_paragraph(doc, '', space_after=6)
    fig3_path = '/projects/sandbox/AMMAN/chapter12_figures/Figure_3_AI_Workflow.png'
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig3_caption = doc.add_paragraph()
    fig3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig3_caption.paragraph_format.space_after = Pt(12)
    cap_run = fig3_caption.add_run(
        'Figure 3: AI-driven closed-loop workflow for design and optimization of microbial-derived '
        'bio-nanomaterials, showing the iterative cycle of data collection, model training, predictive '
        'design, automated synthesis, characterization, and Bayesian optimization.'
    )
    cap_run.bold = True
    cap_run.font.size = Pt(9)
    
    # --- 12.8 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.8 AI-Enabled Characterization, Quality Prediction, and Process Control')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Computer vision and deep learning techniques have revolutionized nanoparticle characterization "
        "by enabling automated analysis of microscopy images [43]. Convolutional neural networks (CNNs) "
        "trained on transmission electron microscopy (TEM) and scanning electron microscopy (SEM) images "
        "can automatically identify, segment, and measure individual nanoparticles, providing rapid and "
        "objective size distribution, morphology classification, and aggregation state assessment. These "
        "approaches dramatically reduce analysis time while eliminating operator bias inherent in manual "
        "image analysis [44]."
    )
    
    add_body_text(doc,
        "Beyond basic size and shape analysis, advanced deep learning architectures enable extraction "
        "of subtle morphological features, lattice fringe analysis from high-resolution TEM images, "
        "and automated identification of crystal defects, twin boundaries, and surface reconstruction "
        "patterns. Transfer learning approaches allow models pre-trained on large microscopy image "
        "datasets to be fine-tuned for specific nanoparticle systems with relatively small amounts "
        "of labeled training data, making these techniques accessible even to laboratories with "
        "limited computational resources and training datasets. The integration of computer vision "
        "with automated microscopy platforms enables high-throughput characterization of thousands "
        "of nanoparticles per sample, providing statistically robust descriptions of size distributions "
        "and morphological heterogeneity."
    )
    
    add_body_text(doc,
        "AI-based prediction of nanomaterial stability, toxicity, antimicrobial activity, and "
        "environmental performance leverages quantitative structure-activity relationship (QSAR) models "
        "trained on curated datasets [45]. Predictive toxicology models correlate physicochemical "
        "descriptors including size, surface charge, composition, and surface area with biological "
        "endpoints, enabling in silico safety screening prior to expensive experimental testing. "
        "Similarly, antimicrobial activity predictions based on nanoparticle characteristics can guide "
        "rational design of optimized antimicrobial formulations for specific pathogen targets."
    )
    
    add_body_text(doc,
        "The development of reliable predictive models requires high-quality, standardized datasets "
        "that currently remain limited in the bio-nanomaterial field. Data heterogeneity arising from "
        "different measurement protocols, characterization techniques, and reporting standards across "
        "research groups creates challenges for model training and validation. Initiatives to establish "
        "community-wide data standards, shared repositories, and minimum reporting requirements would "
        "significantly enhance the quality and utility of predictive models. Furthermore, the "
        "incorporation of uncertainty quantification into predictions ensures that model limitations "
        "are transparently communicated, enabling appropriate use of computational predictions in "
        "decision-making processes for material selection and process design."
    )
    
    add_body_text(doc,
        "Intelligent monitoring and process control systems employ sensor fusion, time-series analysis, "
        "and anomaly detection algorithms to maintain optimal production conditions and identify "
        "deviations requiring corrective action [46]. Real-time spectroscopic measurements including "
        "UV-visible absorption, dynamic light scattering, and fluorescence spectroscopy provide "
        "continuous information on nanoparticle formation kinetics. Machine learning models trained on "
        "historical process data and quality outcomes enable predictive quality control, forecasting "
        "final product characteristics from early-stage process measurements and enabling proactive "
        "process adjustments."
    )
    
    add_body_text(doc,
        "The implementation of Industry 4.0 concepts in bio-nanomaterial manufacturing creates "
        "smart production environments where interconnected sensors, actuators, and control systems "
        "operate as cyber-physical systems capable of autonomous decision-making. Edge computing "
        "deployed at the bioreactor level enables real-time processing of high-frequency sensor data "
        "without latency associated with cloud-based systems, supporting millisecond-level control "
        "responses to detected process disturbances. Digital records maintained through blockchain "
        "or distributed ledger technologies provide immutable batch records supporting regulatory "
        "compliance, traceability, and supply chain transparency. The integration of augmented "
        "reality interfaces enables operators to visualize real-time process data, AI-generated "
        "recommendations, and predictive maintenance alerts in intuitive formats that support "
        "rapid and informed decision-making during manufacturing operations."
    )
    
    # TABLE 3
    add_formatted_paragraph(doc,
        'Table 3: AI/ML Applications in Bio-Nanomaterial Research and Development',
        bold=True, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    
    table3 = doc.add_table(rows=8, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Application Domain', 'ML/AI Algorithm', 'Input Features', 'Performance/Outcome']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "4A148C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table3_data = [
        ['Size Prediction', 'Random Forest, ANN', 'pH, temp, concentration, time', 'R² = 0.91-0.96'],
        ['Morphology Classification', 'CNN, ResNet', 'TEM/SEM images', 'Accuracy >95%'],
        ['Toxicity Prediction', 'SVM, Gradient Boosting', 'Size, charge, composition', 'AUC = 0.88-0.94'],
        ['Synthesis Optimization', 'Bayesian Optimization', 'Process parameters (6-12)', '40-60% fewer experiments'],
        ['Antimicrobial Activity', 'Deep Neural Networks', 'NP properties, pathogen type', 'R² = 0.87-0.93'],
        ['Stability Prediction', 'Ensemble Methods', 'Zeta potential, pH, ionic strength', 'MAE < 5%'],
        ['Process Monitoring', 'LSTM, Autoencoders', 'Time-series sensor data', '99.2% anomaly detection'],
    ]
    
    for i, row_data in enumerate(table3_data):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EDE7F6")
    
    add_formatted_paragraph(doc, '', space_after=12)
    
    # --- 12.9 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.9 Commercialization and Market Opportunities')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "The global nanomaterials market continues to experience robust growth, with projections "
        "indicating expansion from approximately USD 11.3 billion in 2023 to over USD 36.4 billion "
        "by 2030, driven by increasing demand across healthcare, electronics, energy, and environmental "
        "sectors [47]. Bio-nanomaterials represent a rapidly growing segment within this market, "
        "benefiting from growing consumer preference for sustainable products, tightening environmental "
        "regulations, and demonstrated performance advantages in specific applications. The healthcare "
        "nanomedicine market alone is projected to exceed USD 350 billion by 2028, presenting substantial "
        "opportunities for microbial-derived nanomaterials in drug delivery, diagnostics, tissue "
        "engineering, and antimicrobial applications [48]."
    )
    
    add_body_text(doc,
        "Market segmentation analysis reveals distinct value propositions for microbial bio-nanomaterials "
        "across different application sectors. In the pharmaceutical industry, the premium pricing "
        "achievable for biocompatible, well-characterized nanomaterials with demonstrated safety "
        "profiles can offset higher production costs associated with biological manufacturing. In "
        "environmental applications, the lower unit costs enabled by ambient-condition biosynthesis "
        "and waste substrate utilization provide competitive advantages over chemically-synthesized "
        "alternatives for large-volume, lower-value applications such as wastewater treatment and "
        "soil remediation. The agricultural sector represents an emerging market where biosynthesized "
        "nanomaterials for crop protection, nutrient delivery, and soil health improvement align "
        "with growing demand for sustainable agricultural inputs and reduced chemical pesticide use."
    )
    
    add_body_text(doc,
        "Emerging biomedical applications of microbial-derived nanomaterials span diverse therapeutic "
        "and diagnostic domains. Silver nanoparticles biosynthesized by bacteria and fungi demonstrate "
        "potent antimicrobial activity against drug-resistant pathogens, addressing critical unmet medical "
        "needs in wound healing, medical device coatings, and infection control [49]. Gold nanoparticles "
        "produced by microorganisms exhibit excellent biocompatibility for photothermal therapy, drug "
        "delivery vehicles, and diagnostic biosensors. Magnetic iron oxide nanoparticles synthesized by "
        "magnetotactic bacteria find applications in magnetic resonance imaging contrast enhancement, "
        "targeted drug delivery, and hyperthermia cancer treatment."
    )
    
    add_body_text(doc,
        "Environmental applications represent equally significant commercial opportunities for "
        "microbial-derived nanomaterials [50]. Wastewater treatment applications utilize biosynthesized "
        "nanoparticles for heavy metal removal, organic pollutant degradation, and pathogen inactivation. "
        "Biosensing platforms incorporating biogenic nanomaterials enable rapid, sensitive, and "
        "cost-effective detection of environmental contaminants. Sustainable remediation technologies "
        "employing bio-nanomaterials for soil and groundwater cleanup offer advantages of enhanced "
        "biodegradability and reduced secondary pollution compared to conventional nanomaterials."
    )
    
    add_body_text(doc,
        "Market drivers for bio-nanomaterial commercialization include increasing regulatory pressure "
        "toward green chemistry, corporate sustainability commitments, consumer demand for eco-friendly "
        "products, and demonstrated cost advantages in specific applications [51]. However, "
        "commercialization barriers encompass limited production capacity, regulatory uncertainty, lack "
        "of standardized specifications, and insufficient long-term safety data. Intellectual property "
        "strategies must protect novel microbial strains, synthesis processes, and application-specific "
        "formulations while navigating complex patent landscapes in both biotechnology and nanotechnology domains."
    )
    
    add_body_text(doc,
        "Technology transfer mechanisms connecting academic research with industrial implementation "
        "require strengthened pathways including industry-sponsored research programs, technology "
        "licensing agreements, spin-off company creation, and collaborative development partnerships. "
        "Venture capital and government funding agencies increasingly recognize the commercial potential "
        "of bio-nanomaterial technologies, though investment levels remain below those needed to support "
        "the expensive pilot-scale and clinical validation studies required for market entry. Public "
        "funding instruments specifically targeting the scale-up valley of death between laboratory "
        "demonstration and commercial production would address this critical funding gap. Strategic "
        "partnerships between established nanomaterial companies and biotechnology firms offer "
        "complementary capabilities for accelerating commercialization, combining downstream market "
        "access and application expertise with biological production know-how."
    )
    
    add_body_text(doc,
        "As represented in Figure 1, the scale-up pathway directly influences commercial viability, "
        "while Figure 2 demonstrates how regulatory compliance enables market access through systematic "
        "risk management approaches."
    )
    
    # =========================================================================
    # SECTION IV
    # =========================================================================
    section_heading = doc.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(24)
    section_heading.paragraph_format.space_after = Pt(12)
    section_run = section_heading.add_run(
        'Section IV: Future Perspectives and Emerging Research Directions'
    )
    section_run.bold = True
    section_run.font.size = Pt(13)
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- 12.10 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.10 Next-Generation Microbial Bio-Nanomaterials')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Genetic engineering and synthetic biology approaches are creating next-generation microbial "
        "platforms specifically optimized for nanomaterial production [52]. Metabolic engineering strategies "
        "enhance relevant biosynthetic pathways, overexpress metal-reducing enzymes, and introduce novel "
        "functionalities not present in wild-type organisms. CRISPR-Cas9 genome editing enables precise "
        "modification of microbial genomes to improve nanoparticle yield, control morphology, and "
        "incorporate surface-displayed functional peptides [53]. Synthetic biology toolkits allow "
        "construction of engineered microbial consortia where different species perform complementary "
        "functions in nanoparticle synthesis, capping, and functionalization."
    )
    
    add_body_text(doc,
        "The rational design of microbial nanomaterial factories through synthetic biology represents "
        "a fundamental advance beyond traditional strain selection and optimization approaches. By "
        "understanding the genetic and biochemical basis of nanoparticle biosynthesis at molecular "
        "resolution, researchers can now engineer strains with precisely defined capabilities including "
        "controlled expression of specific reductase enzymes, programmed secretion of designer capping "
        "agents, and inducible production systems that decouple growth from nanomaterial synthesis phases. "
        "Modular genetic circuits enable construction of biosensor-coupled feedback systems that "
        "autonomously adjust synthesis conditions in response to real-time monitoring of nanoparticle "
        "formation, creating self-regulating production systems with enhanced consistency and reduced "
        "need for external process control interventions."
    )
    
    add_body_text(doc,
        "Smart, stimuli-responsive bio-nanomaterials represent an emerging class of advanced functional "
        "materials exhibiting triggered responses to environmental stimuli including pH, temperature, "
        "light, magnetic fields, and specific molecular signals [54]. Microbial synthesis can produce "
        "nanoparticles with inherent stimuli-responsiveness through biological capping molecules that "
        "undergo conformational changes under specific conditions. Multifunctional hybrid bio-nanomaterials "
        "combining therapeutic, diagnostic, and targeting capabilities within single nanostructures enable "
        "theranostic applications in personalized medicine. The integration of multiple metallic components "
        "within single nanoparticles through sequential or simultaneous microbial reduction creates "
        "bimetallic and multi-metallic structures with synergistic properties."
    )
    
    add_body_text(doc,
        "The integration of green synthesis principles with circular bioeconomy concepts creates "
        "sustainable manufacturing frameworks that maximize resource utilization and minimize environmental "
        "impact [55]. Waste-to-nanomaterial strategies convert agricultural residues, food processing "
        "wastes, industrial effluents, and electronic waste into value-added nanomaterials through "
        "microbial transformation. These approaches simultaneously address waste management challenges "
        "while producing high-value products, exemplifying circular economy principles. The coupling of "
        "microbial nanomaterial synthesis with other bioprocesses including biofuel production, "
        "bioremediation, and enzyme manufacture creates integrated biorefinery concepts with improved "
        "overall process economics."
    )
    
    # --- 12.11 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.11 Convergence of Biotechnology, Nanotechnology, and Artificial Intelligence')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "The convergence of artificial intelligence, microbiology, and nanotechnology is creating "
        "unprecedented opportunities for autonomous material discovery and accelerated innovation "
        "cycles [56]. Self-driving laboratories integrating AI-directed experimental design, robotic "
        "execution, automated characterization, and machine learning analysis can autonomously explore "
        "vast compositional and process spaces without human intervention. These platforms dramatically "
        "accelerate the pace of bio-nanomaterial discovery while generating comprehensive, standardized "
        "datasets that enhance predictive model accuracy [57]."
    )
    
    add_body_text(doc,
        "The concept of autonomous experimentation represents a paradigm shift from traditional "
        "researcher-directed investigation to AI-directed exploration of materials space. In this "
        "new paradigm, human researchers define high-level objectives and constraints while AI "
        "systems autonomously formulate hypotheses, design experiments, execute synthesis and "
        "characterization procedures through robotic platforms, interpret results, and iteratively "
        "refine their understanding through each experimental cycle. Early demonstrations of "
        "self-driving laboratory concepts in materials science have achieved discovery rates "
        "ten to one hundred times faster than traditional approaches, suggesting transformative "
        "potential when applied to the complex parameter spaces of microbial nanomaterial synthesis."
    )
    
    add_body_text(doc,
        "Multi-omics integration combining genomics, transcriptomics, proteomics, and metabolomics data "
        "with nanoparticle characterization provides holistic understanding of microbial nanomaterial "
        "biosynthesis mechanisms [58]. Systems biology approaches reveal the complex regulatory networks "
        "governing metal reduction, nanoparticle nucleation, and biological capping processes. "
        "Computational modeling at multiple scales, from molecular dynamics simulations of "
        "metal-biomolecule interactions through agent-based models of microbial populations to "
        "process-scale computational fluid dynamics, enables hierarchical understanding and optimization "
        "of biosynthesis processes."
    )
    
    add_body_text(doc,
        "The integration of multi-omics data with machine learning creates powerful predictive "
        "frameworks that connect genotype to nanomaterial phenotype through complex metabolic and "
        "regulatory intermediaries. Genome-scale metabolic models can predict which genetic "
        "modifications will enhance flux through metal-reducing pathways, while transcriptomic "
        "and proteomic data validate these predictions under actual production conditions. "
        "Metabolomic profiling identifies the specific biological reducing agents, capping molecules, "
        "and stabilizing compounds produced during nanoparticle biosynthesis, providing targets for "
        "metabolic engineering to enhance production of desired functionalities. This systems-level "
        "understanding represents a fundamental advance toward rational, predictive design of "
        "microbial nanomaterial production systems rather than empirical optimization approaches "
        "that have dominated the field to date."
    )
    
    add_body_text(doc,
        "Personalized biomedical nanomaterials represent a frontier application where AI-guided design "
        "creates patient-specific therapeutic formulations optimized for individual disease characteristics, "
        "genetic backgrounds, and treatment responses [59]. Application-specific environmental technologies "
        "leveraging AI-optimized bio-nanomaterials address site-specific contamination profiles, "
        "environmental conditions, and remediation objectives. The combination of real-time environmental "
        "sensing with adaptive nanomaterial deployment creates intelligent remediation systems capable "
        "of autonomous response to changing contamination conditions."
    )
    
    add_body_text(doc,
        "The realization of personalized nanomedicine requires integration of patient-specific data "
        "including genomic profiles, disease biomarkers, pharmacogenomic information, and treatment "
        "history with AI models capable of designing nanomaterial formulations optimized for individual "
        "therapeutic needs. This vision extends beyond simple dose adjustment to encompass customized "
        "nanoparticle size, surface functionalization, drug loading, and release kinetics tailored to "
        "each patient's unique biological context. While technically ambitious, rapid advances in "
        "both AI capabilities and automated nanomaterial synthesis platforms are bringing this vision "
        "progressively closer to practical realization within the coming decade."
    )
    
    # TABLE 4
    add_formatted_paragraph(doc,
        'Table 4: Future Research Directions and Technology Roadmap for Microbial Bio-Nanomaterials',
        bold=True, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    
    table4 = doc.add_table(rows=9, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['Research Direction', 'Enabling Technology', 'Timeline', 'Expected Impact']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "BF360C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table4_data = [
        ['CRISPR-engineered strains', 'Synthetic biology, genome editing', '1-3 years', 'High: 10x yield improvement'],
        ['Self-driving NP laboratories', 'AI, robotics, automation', '2-5 years', 'Very High: 100x discovery speed'],
        ['Digital twin bioprocesses', 'IoT, cloud computing, ML', '2-4 years', 'High: Real-time optimization'],
        ['Personalized nanomedicines', 'AI design, patient genomics', '5-10 years', 'Transformative: Precision medicine'],
        ['Smart remediation systems', 'Sensors, AI, responsive NPs', '3-7 years', 'High: Autonomous cleanup'],
        ['Circular bioeconomy integration', 'Waste valorization, LCA', '1-5 years', 'Medium-High: Zero-waste production'],
        ['Multi-omics guided synthesis', 'Systems biology, bioinformatics', '2-5 years', 'High: Mechanistic understanding'],
        ['Autonomous material discovery', 'AGI, quantum computing', '10-15 years', 'Transformative: New material classes'],
    ]
    
    for i, row_data in enumerate(table4_data):
        for j, cell_text in enumerate(row_data):
            cell = table4.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "FFEBEE")
    
    add_formatted_paragraph(doc, '', space_after=12)
    
    # --- 12.12 ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run('12.12 Future Challenges, Opportunities, and Research Roadmap')
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    
    add_body_text(doc,
        "Despite remarkable progress in microbial bio-nanomaterial research, several fundamental "
        "challenges must be addressed to realize the full potential of this field [60]. Scalability "
        "remains the foremost technical challenge, requiring development of production systems capable "
        "of manufacturing ton-scale quantities while maintaining the precise size, morphology, and "
        "functional characteristics achieved at laboratory scale. Reproducibility challenges arising "
        "from inherent biological variability necessitate robust process control strategies, standardized "
        "protocols, and comprehensive understanding of the biological mechanisms underlying nanoparticle "
        "formation [61]."
    )
    
    add_body_text(doc,
        "Addressing scalability requires parallel advancement across multiple fronts: development of "
        "high-productivity microbial strains through metabolic engineering, design of purpose-built "
        "bioreactor systems optimized for nanomaterial production, implementation of continuous "
        "manufacturing approaches that eliminate batch-to-batch variability, and establishment of "
        "robust supply chains for metal precursors and biological inputs at industrial scale. "
        "Furthermore, downstream processing must be scaled proportionally, with particular attention "
        "to maintaining nanoparticle quality during large-volume concentration, purification, and "
        "formulation operations. The development of platform technologies applicable across multiple "
        "nanoparticle types would significantly reduce the engineering effort required for each new "
        "product, accelerating the overall pace of commercialization across the field."
    )
    
    add_body_text(doc,
        "Regulatory uncertainty continues to impede commercialization efforts, as regulatory frameworks "
        "struggle to keep pace with rapid technological advancement in both nanotechnology and "
        "biotechnology [62]. The development of internationally harmonized regulatory standards "
        "specifically addressing bio-nanomaterials would reduce market entry barriers and provide clearer "
        "development pathways for industry. Toxicity concerns, while partially mitigated by the biological "
        "origin of these materials, require long-term safety studies addressing chronic exposure scenarios, "
        "environmental accumulation potential, and intergenerational effects that cannot be assessed "
        "through short-term testing alone."
    )
    
    add_body_text(doc,
        "The development of standardized protocols, comprehensive databases, performance benchmarks, "
        "and international guidelines represents essential infrastructure for advancing the field [63]. "
        "Open-access databases cataloging microbial synthesis conditions, resulting nanomaterial properties, "
        "and application performance would accelerate research progress and enable meta-analyses "
        "identifying universal principles. Standardized characterization protocols ensure comparability "
        "across studies, while performance benchmarks enable objective assessment of progress toward "
        "application-relevant targets. International collaboration through dedicated research networks, "
        "shared facilities, and coordinated funding programs would maximize research efficiency and impact."
    )
    
    add_body_text(doc,
        "The role of international standards organizations in establishing consensus measurement "
        "methods, reference materials, and terminology cannot be overstated. Currently, the absence "
        "of standardized methods for characterizing bio-nanomaterials makes comparison across studies "
        "extremely difficult and impedes regulatory assessment. Priority areas for standardization "
        "include methods for determining biogenic capping agent composition and quantity, protocols "
        "for assessing long-term stability under defined storage conditions, standardized antimicrobial "
        "susceptibility testing methods adapted for nanoparticle formulations, and harmonized "
        "environmental toxicity testing protocols that account for nanoparticle-specific behavior "
        "in test systems. The development of certified reference nanomaterials produced by biological "
        "synthesis would provide essential quality benchmarks for the entire research community."
    )
    
    add_body_text(doc,
        "Long-term prospects for sustainable industrial production of microbial-derived bio-nanomaterials "
        "are highly favorable, driven by converging trends in sustainability requirements, regulatory "
        "evolution, technological maturation, and market demand [64]. The integration of advanced "
        "biotechnology tools including synthetic biology and metabolic engineering with artificial "
        "intelligence and automation is expected to overcome current productivity and reproducibility "
        "limitations within the next decade. As illustrated in Figure 3, the AI-driven optimization "
        "workflow will be central to achieving these advances, while the risk assessment framework "
        "shown in Figure 4 ensures that accelerated development maintains appropriate safety standards."
    )
    
    add_body_text(doc,
        "The establishment of dedicated research centers and public-private partnerships focusing "
        "specifically on bio-nanomaterial scale-up and commercialization would provide critical "
        "infrastructure and expertise currently lacking in the field. Investment in shared pilot-scale "
        "manufacturing facilities would lower barriers to entry for academic researchers and small "
        "companies, enabling translation of promising laboratory discoveries without requiring each "
        "organization to independently develop expensive production capabilities. Training programs "
        "developing the specialized interdisciplinary workforce combining microbiology, nanotechnology, "
        "engineering, and data science skills would address current human capital constraints limiting "
        "the pace of progress in this rapidly evolving field."
    )
    
    # Insert Figure 4
    add_formatted_paragraph(doc, '', space_after=6)
    fig4_path = '/projects/sandbox/AMMAN/chapter12_figures/Figure_4_Research_Roadmap.png'
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig4_caption = doc.add_paragraph()
    fig4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig4_caption.paragraph_format.space_after = Pt(12)
    cap_run = fig4_caption.add_run(
        'Figure 4: Comprehensive research roadmap and technology convergence framework for '
        'next-generation microbial bio-nanomaterials, showing near-term to visionary timelines '
        'across biotechnology, AI, manufacturing, and application domains.'
    )
    cap_run.bold = True
    cap_run.font.size = Pt(9)
    
    add_body_text(doc,
        "The next generation of biomedical applications will increasingly leverage personalized, "
        "AI-designed nanomaterials optimized for individual patient needs, while environmental "
        "applications will employ intelligent, adaptive nanomaterial systems capable of autonomous "
        "response to contamination events [65]. The convergence of microbial nanotechnology with "
        "digital technologies, advanced manufacturing, and circular economy principles positions this "
        "field at the forefront of sustainable innovation with transformative potential across multiple "
        "sectors of the global economy."
    )
    
    add_body_text(doc,
        "In conclusion, the field of microbial-derived bio-nanomaterials stands at a critical "
        "inflection point where decades of fundamental research are converging with enabling "
        "technologies in artificial intelligence, synthetic biology, advanced manufacturing, and "
        "digital infrastructure to enable practical industrial implementation. The challenges of "
        "scalability, reproducibility, regulatory compliance, and market acceptance are substantial "
        "but increasingly tractable through the integrated approaches described in this chapter. "
        "Success will require sustained investment in both fundamental understanding and translational "
        "infrastructure, close collaboration between academic researchers, industry partners, and "
        "regulatory authorities, and commitment to responsible innovation principles that ensure "
        "the benefits of this technology are realized safely and equitably. The potential rewards—"
        "sustainable, cost-effective nanomaterial production supporting applications from personalized "
        "medicine to environmental restoration—justify the continued pursuit of these ambitious goals "
        "and position microbial bio-nanotechnology as a cornerstone of the emerging sustainable "
        "bioeconomy."
    )
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_page_break()
    
    ref_heading = doc.add_paragraph()
    ref_heading.paragraph_format.space_before = Pt(18)
    ref_heading.paragraph_format.space_after = Pt(12)
    ref_run = ref_heading.add_run('References')
    ref_run.bold = True
    ref_run.font.size = Pt(14)
    ref_run.font.color.rgb = RGBColor(0, 51, 102)
    
    references = [
        "[1] Iravani, S. (2014). Bacteria in nanoparticle synthesis: current status and future prospects. International Scholarly Research Notices, 2014, 359316.",
        "[2] Narayanan, K.B., & Sakthivel, N. (2010). Biological synthesis of metal nanoparticles by microbes. Advances in Colloid and Interface Science, 156(1-2), 1-13.",
        "[3] Li, X., Xu, H., Chen, Z.S., & Chen, G. (2011). Biosynthesis of nanoparticles by microorganisms and their applications. Journal of Nanomaterials, 2011, 270974.",
        "[4] Hulkoti, N.I., & Taranath, T.C. (2014). Biosynthesis of nanoparticles using microbes—a review. Colloids and Surfaces B: Biointerfaces, 121, 474-483.",
        "[5] Khandel, P., & Shahi, S.K. (2018). Mycogenic nanoparticles and their bio-prospective applications: current status and future challenges. Journal of Nanostructure in Chemistry, 8, 369-391.",
        "[6] Ovais, M., Khalil, A.T., Ayaz, M., Ahmad, I., Nethi, S.K., & Mukherjee, S. (2018). Biosynthesis of metal nanoparticles via microbial enzymes: a mechanistic approach. International Journal of Molecular Sciences, 19(12), 4100.",
        "[7] Mohd Yusof, H., Mohamad, R., Zaidan, U.H., & Abdul Rahman, N. (2019). Microbial synthesis of zinc oxide nanoparticles and their potential application as an antimicrobial agent. Journal of Animal Science and Biotechnology, 10, 57.",
        "[8] Singh, P., Kim, Y.J., Zhang, D., & Yang, D.C. (2016). Biological synthesis of nanoparticles from plants and microorganisms. Trends in Biotechnology, 34(7), 588-599.",
        "[9] Gahlawat, G., & Choudhury, A.R. (2019). A review on the biosynthesis of metal and metal salt nanoparticles by microbes. RSC Advances, 9(23), 12944-12967.",
        "[10] Fang, X., Wang, Y., Wang, Z., Jiang, Z., & Dong, M. (2019). Microorganism assisted synthesized nanoparticles for catalytic applications. Energies, 12(1), 190.",
        "[11] Srivastava, S.K., & Constanti, M. (2012). Room temperature biogenic synthesis of multiple nanoparticles by Pseudomonas aeruginosa SM1. Journal of Nanoparticle Research, 14, 831.",
        "[12] Deplanche, K., Caldelari, I., Mikheenko, I.P., Sargent, F., & Macaskie, L.E. (2010). Involvement of hydrogenases in the formation of highly catalytic Pd(0) nanoparticles. Microbiology, 156(9), 2630-2640.",
        "[13] Sriramulu, M., & Sumathi, S. (2018). Biosynthesis of palladium nanoparticles using Saccharomyces cerevisiae extract. Advances in Natural Sciences: Nanoscience and Nanotechnology, 9(2), 025018.",
        "[14] Gudikandula, K., & Charya Maringanti, S. (2016). Synthesis of silver nanoparticles by chemical and biological methods and their antimicrobial properties. Journal of Experimental Nanoscience, 11(9), 714-721.",
        "[15] Saxena, J., Sharma, P.K., Sharma, M.M., & Singh, A. (2016). Process optimization for green synthesis of silver nanoparticles by Sclerotinia sclerotiorum. SpringerPlus, 5, 861.",
        "[16] Moghaddam, A.B., Namvar, F., Moniri, M., et al. (2015). Nanoparticles biosynthesized by fungi and yeast: a review. Molecules, 20(9), 16540-16565.",
        "[17] Siddiqi, K.S., & Husen, A. (2016). Fabrication of metal nanoparticles from fungi and metal salts: scope and application. Nanoscale Research Letters, 11, 98.",
        "[18] Guilger-Casagrande, M., & Lima, R.D. (2019). Synthesis of silver nanoparticles mediated by fungi: a review. Frontiers in Bioengineering and Biotechnology, 7, 287.",
        "[19] Rathore, A.S., & Winkle, H. (2009). Quality by design for biopharmaceuticals. Nature Biotechnology, 27(1), 26-34.",
        "[20] Yu, L.X., Amidon, G., Khan, M.A., et al. (2014). Understanding pharmaceutical quality by design. AAPS Journal, 16(4), 771-783.",
        "[21] Kuppusamy, P., Yusoff, M.M., Maniam, G.P., & Govindan, N. (2016). Biosynthesis of metallic nanoparticles using plant derivatives. Saudi Pharmaceutical Journal, 24(4), 473-484.",
        "[22] Rajan, A., Vilas, V., & Philip, D. (2015). Studies on catalytic, antioxidant, antibacterial and anticancer activities of biogenic gold nanoparticles. Journal of Molecular Liquids, 212, 331-339.",
        "[23] Gavankar, S., Suh, S., & Keller, A.F. (2012). Life cycle assessment at nanoscale: review and recommendations. International Journal of Life Cycle Assessment, 17, 295-303.",
        "[24] Pourzahedi, L., & Eckelman, M.J. (2015). Comparative life cycle assessment of silver nanoparticle synthesis routes. Environmental Science: Nano, 2(4), 361-369.",
        "[25] Bleeker, E.A., de Jong, W.H., Geertsma, R.E., et al. (2013). Considerations on the EU definition of a nanomaterial. Regulatory Toxicology and Pharmacology, 65(1), 119-125.",
        "[26] FDA (2022). Nanotechnology guidance documents. U.S. Food and Drug Administration Center for Drug Evaluation and Research.",
        "[27] European Commission (2018). Commission Regulation (EU) 2018/1881 amending REACH Regulation as regards nanomaterials. Official Journal of the European Union, L308, 1-20.",
        "[28] Ventola, C.L. (2017). Progress in nanomedicine: approved and investigational nanodrugs. Pharmacy and Therapeutics, 42(12), 742-755.",
        "[29] Rasmussen, K., Rauscher, H., Mech, A., et al. (2018). Physico-chemical properties of manufactured nanomaterials. JRC Science for Policy Report.",
        "[30] Schins, R.P., & Knaapen, A.M. (2007). Genotoxicity of poorly soluble particles. Inhalation Toxicology, 19(sup1), 189-198.",
        "[31] Singh, N., Manshian, B., Jenkins, G.J., et al. (2009). NanoGenotoxicology: the DNA damaging potential of engineered nanomaterials. Biomaterials, 30(23-24), 3891-3914.",
        "[32] Dobrovolskaia, M.A., & McNeil, S.E. (2007). Immunological properties of engineered nanomaterials. Nature Nanotechnology, 2(8), 469-478.",
        "[33] Shankar, P.D., Shobana, S., Karuppusamy, I., et al. (2016). A review on the biosynthesis of metallic nanoparticles using bio-components of microalgae. Enzyme and Microbial Technology, 95, 28-44.",
        "[34] Lowry, G.V., Gregory, K.B., Apte, S.C., & Lead, J.R. (2012). Transformations of nanomaterials in the environment. Environmental Science & Technology, 46(13), 6893-6899.",
        "[35] Anselmo, A.C., & Mitragotri, S. (2019). Nanoparticles in the clinic: an update. Bioengineering & Translational Medicine, 4(3), e10143.",
        "[36] Gehrke, I., Geiser, A., & Somborn-Schulz, A. (2015). Innovations in nanotechnology for water treatment. Nanotechnology, Science and Applications, 8, 1-17.",
        "[37] Owen, R., Macnaghten, P., & Stilgoe, J. (2012). Responsible research and innovation. Science and Public Policy, 39(6), 751-760.",
        "[38] Barnard, A.S., & Motevalli, B. (2019). Using machine learning to predict the properties of nanoparticles. Nanoscale, 11(48), 23165-23172.",
        "[39] Tao, H., Wu, T., Aldeghi, M., et al. (2021). Nanoparticle synthesis assisted by machine learning. Nature Reviews Materials, 6(8), 701-716.",
        "[40] Yan, X., Sedykh, A., Wang, W., Yan, B., & Zhu, H. (2020). Construction of a web-based nanomaterial database. Nature Communications, 11, 2519.",
        "[41] Häse, F., Roch, L.M., & Aspuru-Guzik, A. (2019). Next-generation experimentation with self-driving laboratories. Trends in Chemistry, 1(3), 282-291.",
        "[42] Udugama, I.A., Gargalo, C.L., Yamashita, Y., et al. (2021). Digital twin in biomanufacturing: challenges and opportunities. Systems Microbiology and Biomanufacturing, 1, 257-274.",
        "[43] Modarres, M.H., et al. (2017). Neural network for nanoscience scanning electron microscope image recognition. Scientific Reports, 7, 13282.",
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_after = Pt(3)
        ref_para.paragraph_format.left_indent = Cm(1.0)
        ref_para.paragraph_format.first_line_indent = Cm(-1.0)
        ref_run = ref_para.add_run(ref)
        ref_run.font.size = Pt(9.5)
    
    # Continue references
    references2 = [
        "[44] Xu, M., Papageorgiou, D.P., Abidi, S.Z., et al. (2017). A deep convolutional neural network for classification of red blood cells. PLoS Computational Biology, 13(10), e1005746.",
        "[45] Puzyn, T., Rasulev, B., Gajewicz, A., et al. (2011). Using nano-QSAR to predict the cytotoxicity of metal oxide nanoparticles. Nature Nanotechnology, 6(3), 175-178.",
        "[46] Mowbray, M., Savage, T., Wu, C., et al. (2022). Machine learning for biochemical engineering: a review. Biochemical Engineering Journal, 172, 108054.",
        "[47] Grand View Research (2023). Nanomaterials Market Size, Share & Trends Analysis Report. Grand View Research Industry Report.",
        "[48] Allied Market Research (2023). Nanomedicine Market by Application, Modality, and Indication. Allied Market Research Report.",
        "[49] Durán, N., Durán, M., de Jesus, M.B., et al. (2016). Silver nanoparticles: a new view on mechanistic aspects on antimicrobial activity. Nanomedicine, 12(3), 789-799.",
        "[50] Cecchin, I., Reddy, K.R., Thomé, A., et al. (2017). Nanobioremediation: integration of nanoparticles and bioremediation. International Biodeterioration & Biodegradation, 119, 419-428.",
        "[51] Inshakova, E., & Inshakov, O. (2017). World market for nanomaterials: structure and trends. MATEC Web of Conferences, 129, 02013.",
        "[52] Chen, Y., Banerjee, D., Mukhopadhyay, A., & Petzold, C.J. (2020). Systems and synthetic biology tools for advanced bioproduction hosts. Current Opinion in Biotechnology, 64, 101-109.",
        "[53] Luo, C.H., Shanks, J.V., & Bravo, L. (2016). Genetically encoded biosensors for monitoring enzyme function. Analytical Chemistry, 88(17), 8381-8389.",
        "[54] Wei, M., Gao, Y., Li, X., & Serpe, M.J. (2017). Stimuli-responsive polymers and their applications. Polymer Chemistry, 8(1), 127-143.",
        "[55] Heimann, K. (2016). Novel approaches to microalgal and cyanobacterial cultivation for bioenergy. Current Opinion in Biotechnology, 38, 183-189.",
        "[56] Stach, E., DeCost, B., Kusne, A.G., et al. (2021). Autonomous experimentation systems for materials development. Matter, 4(9), 2702-2726.",
        "[57] Coley, C.W., Eyke, N.S., & Jensen, K.F. (2020). Autonomous discovery in the chemical sciences part II. Angewandte Chemie International Edition, 59(52), 23414-23436.",
        "[58] Kim, G.B., Kim, W.J., Kim, H.U., & Lee, S.Y. (2020). Machine learning applications in systems metabolic engineering. Current Opinion in Biotechnology, 64, 1-9.",
        "[59] Mitchell, M.J., Billingsley, M.M., Haley, R.M., et al. (2021). Engineering precision nanoparticles for drug delivery. Nature Reviews Drug Discovery, 20(2), 101-124.",
        "[60] Khan, I., Saeed, K., & Khan, I. (2019). Nanoparticles: properties, applications and toxicities. Arabian Journal of Chemistry, 12(7), 908-931.",
        "[61] Jeevanandam, J., Barhoum, A., Chan, Y.S., et al. (2018). Review on nanoparticles and nanostructured materials. Beilstein Journal of Nanotechnology, 9, 1050-1074.",
        "[62] Falkner, R., & Jaspers, N. (2012). Regulating nanotechnologies: risk, uncertainty and the global governance gap. Global Environmental Politics, 12(1), 30-55.",
        "[63] ISO (2020). ISO/TS 80004 Nanotechnologies—Vocabulary. International Organization for Standardization Technical Specification.",
        "[64] Saleh, T.A. (2020). Nanomaterials: classification, properties, and environmental toxicities. Environmental Technology & Innovation, 20, 101067.",
        "[65] Bayda, S., Adeel, M., Tuccinardi, T., et al. (2020). The history of nanoscience and nanotechnology. Molecules, 25(1), 112.",
    ]
    
    for ref in references2:
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_after = Pt(3)
        ref_para.paragraph_format.left_indent = Cm(1.0)
        ref_para.paragraph_format.first_line_indent = Cm(-1.0)
        ref_run = ref_para.add_run(ref)
        ref_run.font.size = Pt(9.5)
    
    # Save document
    output_path = '/projects/sandbox/AMMAN/Chapter_12_Industrial_Translation_AI_Integration.docx'
    doc.save(output_path)
    print(f"Document saved successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    create_document()
