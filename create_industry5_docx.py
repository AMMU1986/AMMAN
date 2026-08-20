"""
Generate DOCX for Chapter: Industry 5.0 - A Comprehensive Guide for Smart,
Sustainable, and Human-Centred Transformation
- 43 references in sequential numbered [1]-[43] format
- No citations in abstract or conclusion (Section 4.3 last paragraphs)
- 4 figures and 4 tables, each cited exactly 2 times in body text
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def create_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # TITLE
    title = doc.add_heading('', level=0)
    run = title.add_run('Industry 5.0: A Comprehensive Guide for Smart, Sustainable, and Human-Centred Transformation')
    run.font.size = Pt(16)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ABSTRACT (No citations)
    doc.add_heading('Abstract', level=1)
    abstract = (
        "The transition from Industry 4.0 to Industry 5.0 marks a fundamental reorientation of manufacturing "
        "philosophy, shifting from automation-centric efficiency toward a tripartite commitment to sustainability, "
        "human-centricity, and resilience. This chapter provides a comprehensive examination of how circular economy "
        "principles integrate with Industry 5.0 frameworks to enable smart, sustainable manufacturing transformation. "
        "Beginning with the theoretical foundations that distinguish Industry 5.0 from its predecessor, the chapter "
        "explores technological enablers including Digital Twins, Cognitive Digital Twins, Explainable Artificial "
        "Intelligence, and Digital Product Passports that operationalize circular manufacturing at scale. Through "
        "detailed case studies spanning photovoltaic manufacturing and fast-moving consumer goods, the chapter "
        "demonstrates how circular-by-design approaches create tangible value across diverse industrial contexts. "
        "The discussion culminates in strategic implementation guidance, addressing persistent barriers, policy "
        "frameworks, governance mechanisms, and emerging research directions that will shape the future of "
        "sustainable industrial transformation. By synthesizing insights from engineering, environmental science, "
        "data analytics, and policy studies, this chapter offers practitioners and researchers an integrated roadmap "
        "for advancing circular manufacturing within the Industry 5.0 paradigm."
    )
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Cm(1.27)

    kw = doc.add_paragraph()
    r = kw.add_run('Keywords: ')
    r.bold = True
    kw.add_run('Industry 5.0, circular economy, sustainable manufacturing, digital twins, human-centricity, '
               'resilience, lifecycle assessment, digital product passport')

    doc.add_page_break()

    # ============================
    # SECTION I
    # ============================
    doc.add_heading('Section I: Foundations of Industry 5.0 and Sustainable Manufacturing', level=1)

    doc.add_heading('1.1 From Industry 4.0 to Industry 5.0 – The Evolution of Industrial Paradigms', level=2)

    s1_1 = [
        "The fourth industrial revolution, commonly designated Industry 4.0, introduced transformative technologies including the Internet of Things, cyber-physical systems, cloud computing, and big data analytics into manufacturing environments [1]. These innovations enabled unprecedented levels of automation, real-time monitoring, and data-driven decision-making across production systems. However, as Industry 4.0 matured, critical limitations became apparent: the paradigm's overwhelming focus on technological efficiency often marginalized human workers, treated environmental sustainability as a secondary concern, and created brittle systems vulnerable to disruption [2]. The COVID-19 pandemic, geopolitical instabilities, and accelerating climate change exposed these vulnerabilities with dramatic clarity, catalysing a fundamental reassessment of industrial priorities.",

        "Industry 5.0 emerged as a necessary corrective to these limitations, repositioning manufacturing around three interdependent pillars: sustainability as an environmental imperative, human-centricity as a commitment to worker empowerment and well-being, and resilience as the capacity to withstand and adapt to systemic disruptions [3]. Unlike its predecessor, which prioritized productivity maximization, Industry 5.0 explicitly acknowledges that manufacturing exists within social and ecological systems whose health determines long-term industrial viability. This reconceptualization demands that organizations balance economic performance with planetary boundaries and social responsibility—not as external constraints imposed upon industry, but as integral design criteria for manufacturing systems.",

        "The transition from Industry 4.0 to Industry 5.0 should not be understood as a replacement or rejection of previous technological achievements [4]. Rather, Industry 5.0 builds upon the digital infrastructure established during the fourth revolution while reorienting its application toward broader societal goals. The smart factory remains relevant, but its purpose expands from cost minimization to value creation that encompasses environmental regeneration, worker flourishing, and community resilience. This continuity is important practically: organizations need not discard Industry 4.0 investments but rather extend and redirect them toward the expanded objectives of Industry 5.0. The existing digital infrastructure—sensors, connectivity, data platforms, analytics capabilities—provides the technological substrate upon which Industry 5.0's sustainability, human-centricity, and resilience commitments are operationalized. Figure 1 illustrates this evolutionary trajectory, mapping the transition from automation-centric Industry 4.0 to the three-pillar framework of Industry 5.0 and its integration with circular economy principles.",

        "Critically, Industry 5.0 recognizes that previous industrial models underestimated the importance of human roles in manufacturing [5]. Where Industry 4.0 frequently positioned automation as a substitute for human labour, Industry 5.0 reframes the relationship as collaborative: humans contribute creativity, contextual judgment, and ethical reasoning that complement machine precision, speed, and data processing capacity. This human-machine collaboration model represents a significant philosophical departure, acknowledging that sustainable excellence requires human agency rather than its elimination. The practical implications are substantial: factory design must accommodate human-machine interfaces that amplify rather than constrain human capabilities; training programmes must develop skills for collaboration with intelligent systems rather than merely operating them; and performance metrics must capture the quality of human contribution alongside automated throughput. Table 1 presents a comprehensive comparison of Industry 4.0 and Industry 5.0 across multiple dimensions, highlighting the fundamental shifts in philosophy, metrics, and operational priorities that characterize this transition.",
    ]

    for text in s1_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('1.2 Defining the Pillars of Industry 5.0 – Sustainability, Human-Centricity, and Resilience', level=2)

    s1_2 = [
        "The sustainability pillar of Industry 5.0 encompasses far more than incremental efficiency improvements or pollution reduction at the margins [6]. It demands a fundamental reconceptualization of manufacturing's relationship with natural systems, moving from extractive and degenerative models toward regenerative approaches that restore ecological health while delivering economic value. This requires manufacturing organizations to adopt systems thinking that traces material and energy flows across entire value chains, identifies intervention points where circular strategies can close resource loops, and measures success through absolute environmental impact reduction rather than merely relative intensity improvements.",

        "Human-centricity, the second pillar, positions workers as the ultimate beneficiaries and drivers of industrial transformation rather than as resources to be optimized or costs to be minimized [7]. This principle manifests in multiple dimensions: physical well-being through ergonomic design and hazard elimination; cognitive well-being through meaningful work that engages human capabilities; developmental well-being through continuous learning and skills advancement; and participatory well-being through genuine involvement in decisions that affect working conditions and organizational direction. The human-centric factory is not merely a more pleasant workplace—it is a more innovative and adaptive one, because empowered workers contribute insights and improvements that automated systems cannot generate independently.",

        "Resilience, the third pillar, addresses the increasing frequency and severity of disruptions confronting manufacturing systems [8]. Climate-related events, pandemic outbreaks, geopolitical conflicts, and supply chain cascading failures demand manufacturing architectures capable of absorbing shocks, adapting operations, and recovering functionality without catastrophic loss. As illustrated in Figure 1, these three pillars are interconnected rather than independent: sustainable manufacturing systems tend to be more resilient because they depend less on scarce or geographically concentrated resources; human-centric systems tend to be more adaptable because empowered workers can improvise responses to novel situations; and resilient systems tend to be more sustainable because they avoid the resource waste associated with crisis-mode operations.",

        "The interaction among pillars also generates tensions that require careful management [9]. Sustainability investments may conflict with short-term resilience requirements—for example, consolidating production for efficiency may reduce the distributed manufacturing capacity that provides resilience against localized disruptions. Human-centricity may complicate automation decisions that would improve environmental performance but eliminate jobs. These trade-offs demand governance frameworks capable of navigating complexity, balancing competing priorities, and making decisions transparent to all stakeholders. The recognition that pillar interactions create both synergies and tensions represents a maturation of Industry 5.0 thinking beyond simplistic win-win narratives toward the nuanced trade-off management that characterizes genuine strategic decision-making. Table 1 further details how these pillar interactions manifest across different operational dimensions, providing a framework for organizations to identify and manage inherent tensions while maximizing synergistic opportunities.",
    ]

    for text in s1_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('1.3 The Circular Economy Imperative in Manufacturing', level=2)

    s1_3 = [
        "The circular economy represents a systemic alternative to the linear 'take-make-dispose' model that has dominated industrial production since the first industrial revolution [10]. Its core principles—designing out waste and pollution, keeping products and materials in use at their highest value, and regenerating natural systems—provide a concrete operational framework for realizing Industry 5.0's sustainability ambitions. In manufacturing contexts, circularity manifests through strategies including design for disassembly, material selection for recyclability, modular product architectures that enable component reuse, remanufacturing processes that restore end-of-life products to original specifications, and industrial symbiosis where one process's waste becomes another's input.",

        "The evolution from linear to circular manufacturing has been gradual and uneven across sectors [11]. Early circular economy initiatives focused predominantly on end-of-pipe waste management and recycling—necessary but insufficient interventions that address symptoms rather than root causes of resource depletion. More sophisticated approaches have emerged that integrate circularity throughout the product lifecycle: circular design that anticipates disassembly and material recovery; production planning that minimizes scrap and enables reprocessing; distribution models that facilitate product return; and business models such as product-as-a-service that align manufacturer incentives with product longevity and resource efficiency. The maturation of circular economy thinking has progressively shifted attention upstream—from managing waste toward preventing it, from recycling toward reuse, and from individual firm actions toward systemic value chain transformation.",

        "Despite growing recognition of its importance, circular economy implementation in manufacturing faces persistent challenges [12]. Technical barriers include the difficulty of separating complex material combinations, energy-intensive recycling processes for certain materials, and degradation of material properties through repeated processing cycles. Organizational barriers include misaligned incentive structures that reward volume production over lifecycle value, functional silos that prevent cross-departmental collaboration on circularity, and inadequate information systems for tracking materials through multiple use cycles. Analytical barriers include the absence of standardized circularity metrics, difficulty in quantifying the environmental benefits of circular strategies relative to linear alternatives, and limited lifecycle data for emerging materials and processes [13]. These interconnected barriers create a complex implementation landscape where addressing any single challenge in isolation is insufficient—systemic approaches that simultaneously tackle multiple barriers prove far more effective than piecemeal interventions.",

        "Industry 5.0 technologies offer transformative potential for overcoming these barriers. Digital twins can simulate circular process configurations before physical implementation, reducing experimentation costs. Artificial intelligence can optimize material selection and disassembly strategies across vast solution spaces. Digital product passports can maintain material identity information through multiple lifecycle stages. And human-centred design approaches can align circular economy practices with worker capabilities and customer expectations. The convergence of Industry 5.0 capabilities with circular economy ambitions creates a synergistic framework—designated by some researchers as CIRCMan5.0—that exceeds what either paradigm could achieve independently [14]. This convergence is not merely additive: the digital infrastructure of Industry 5.0 enables circular strategies that would be logistically impossible without real-time material tracking and intelligent process optimization, while circular economy principles provide the purposive direction that prevents Industry 5.0 technologies from merely accelerating unsustainable production models.",
    ]

    for text in s1_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 1
    doc.add_paragraph()
    t1_title = doc.add_paragraph()
    t1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1_title.add_run('Table 1. ')
    r.bold = True
    t1_title.add_run('Comparative Analysis of Industry 4.0 and Industry 5.0 Paradigms')

    table1 = doc.add_table(rows=10, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    h1 = ['Dimension', 'Industry 4.0', 'Industry 5.0']
    for i, h in enumerate(h1):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t1_data = [
        ['Core Philosophy', 'Automation & efficiency maximization', 'Sustainability, human-centricity, resilience'],
        ['Human Role', 'Variable to optimize; replacement target', 'Collaborator, innovator, beneficiary'],
        ['Environmental Approach', 'Eco-efficiency (relative gains)', 'Circular economy (absolute sustainability)'],
        ['Technology Purpose', 'Productivity enhancement', 'Societal value creation'],
        ['Resilience Strategy', 'Efficiency through optimization', 'Adaptability through diversity'],
        ['Value Metrics', 'OEE, cost per unit, throughput', 'Circularity index, well-being, CO2 footprint'],
        ['Supply Chain Model', 'Global, lean, just-in-time', 'Distributed, circular, resilient'],
        ['Innovation Driver', 'Technology push', 'Human-technology co-creation'],
        ['Governance', 'Shareholder value', 'Multi-stakeholder value & planetary boundaries'],
    ]
    for i, row in enumerate(t1_data):
        for j, cell in enumerate(row):
            table1.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    # FIGURE 1
    fig1_path = 'industry5_figures/Figure_1_Industry5_Evolution.png'
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1 = doc.add_paragraph()
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap1.add_run('Figure 1. ')
    r.bold = True
    cap1.add_run('Evolution from Industry 4.0 to Industry 5.0, illustrating the three foundational pillars (Sustainability, Human-Centricity, Resilience) and their integration with circular economy principles.')

    doc.add_page_break()

    # ============================
    # SECTION II
    # ============================
    doc.add_heading('Section II: Technological Enablers for Circular Manufacturing', level=1)

    doc.add_heading('2.1 Digital Twins and Cognitive Digital Twins for Circular Process Design', level=2)

    s2_1 = [
        "Digital Twins (DTs)—virtual replicas of physical manufacturing systems that synchronize with real-world operations through continuous data exchange—have emerged as foundational enablers of circular manufacturing [15]. By creating high-fidelity computational models of production processes, material flows, and product lifecycles, DTs enable manufacturers to simulate circular configurations virtually before committing physical resources to implementation. This simulation capability is particularly valuable for circular economy transitions, where novel process configurations—such as reverse logistics networks, disassembly sequences, or material cascading strategies—carry significant uncertainty about performance, cost, and quality outcomes.",

        "The CIRCMan5.0 framework exemplifies how DTs can be architecturally designed to support circularity [16]. This framework leverages established Industry 4.0 standards including RAMI4.0 (Reference Architecture Model for Industry 4.0), OPC UA (Open Platform Communications Unified Architecture), and AutomationML to create interoperable digital representations that span the entire product lifecycle. The semantic layer, provided by classification systems such as eCl@ss, ensures that material and process information maintains consistent meaning as it flows between different systems and stakeholders—a critical requirement for circular economy implementation where information about material composition, processing history, and recyclability must accompany physical materials through multiple use cycles.",

        "Figure 2 presents the three-layer architecture for Digital Twin-enabled circular manufacturing, distinguishing between the Physical Layer (manufacturing systems and material flows), the Digital Twin Layer (virtual simulation and optimization), and the Cognitive Digital Twin Layer (AI-driven decision support). This layered architecture enables progressive sophistication: basic DTs can track material flows and simulate process modifications; advanced DTs can optimize disassembly strategies and predict material degradation; and Cognitive Digital Twins (CDTs) can autonomously recommend circular interventions based on real-time assessment of environmental, economic, and social performance indicators [17].",

        "Cognitive Digital Twins extend conventional DT capabilities by incorporating artificial intelligence that enables reasoning, learning, and autonomous decision-making [18]. In circular manufacturing contexts, CDTs can analyze patterns in material degradation data to predict optimal timing for component replacement or remanufacturing; evaluate trade-offs between different end-of-life strategies (recycling, remanufacturing, reuse) based on current market conditions and environmental impact; and recommend design modifications that improve circularity without compromising product performance. Table 2 provides a comprehensive overview of Digital Twin technology variants and their applications in circular manufacturing, ranging from basic material flow tracking to autonomous circularity optimization. The integration of explainable AI within CDTs ensures that recommendations are transparent and comprehensible to human operators, maintaining the human-centric principle of Industry 5.0 even as decision complexity increases.",

        "The practical deployment of DTs for circular manufacturing requires addressing significant integration challenges [19]. Manufacturing systems typically comprise heterogeneous equipment from multiple vendors, operating on diverse communication protocols and data formats. Creating coherent digital representations across this heterogeneity demands standardized interfaces, semantic interoperability frameworks, and governance structures that ensure data quality and consistency. As depicted in Figure 2, the standards infrastructure—including OPC UA for communication, RAMI4.0 for architecture, and AutomationML for engineering data exchange—provides the technical foundation for this integration, though significant implementation effort remains required in most manufacturing environments. The economic justification for DT investment in circular manufacturing extends beyond direct process optimization to encompass reduced prototyping costs, accelerated time-to-market for circular product variants, improved regulatory compliance documentation, and enhanced stakeholder communication through visual simulation of circular strategies.",
    ]

    for text in s2_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 2
    doc.add_paragraph()
    t2_title = doc.add_paragraph()
    t2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2_title.add_run('Table 2. ')
    r.bold = True
    t2_title.add_run('Digital Twin Technologies and Their Applications in Circular Manufacturing')

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    h2 = ['Technology', 'Circular Economy Application', 'Key Standards', 'Maturity Level']
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t2_data = [
        ['Basic Digital Twin', 'Material flow tracking & process monitoring', 'OPC UA, MQTT', 'High (TRL 7-9)'],
        ['Simulation Digital Twin', 'Circular process configuration testing', 'AutomationML, FMI', 'Medium-High (TRL 6-8)'],
        ['Predictive Digital Twin', 'Component lifetime & degradation forecasting', 'RAMI4.0, eCl@ss', 'Medium (TRL 5-7)'],
        ['Cognitive Digital Twin', 'Autonomous circularity optimization', 'XAI frameworks, ontologies', 'Low-Medium (TRL 3-5)'],
        ['Product Digital Twin', 'Digital Product Passport & lifecycle tracking', 'GS1, EPCIS', 'Medium (TRL 5-7)'],
        ['Supply Chain Digital Twin', 'Reverse logistics & industrial symbiosis', 'SCOR, blockchain', 'Low-Medium (TRL 4-6)'],
        ['Hierarchical Simulation Engine', 'Multi-scale circular system optimization', 'HLA, co-simulation', 'Low (TRL 3-5)'],
    ]
    for i, row in enumerate(t2_data):
        for j, cell in enumerate(row):
            table2.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    # FIGURE 2
    fig2_path = 'industry5_figures/Figure_2_Digital_Twins_Circular.png'
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap2.add_run('Figure 2. ')
    r.bold = True
    cap2.add_run('Three-layer architecture for Digital Twin-enabled circular manufacturing, showing the Physical Layer, Digital Twin Layer, and Cognitive Digital Twin Layer with their respective components and interoperability standards.')

    doc.add_paragraph()

    doc.add_heading('2.2 Artificial Intelligence and Data Analytics for Sustainability Assessment', level=2)

    s2_2 = [
        "Artificial intelligence and advanced data analytics play increasingly central roles in enabling manufacturers to assess, monitor, and continuously improve the sustainability and circularity of their operations [20]. Traditional sustainability assessment methods, particularly Lifecycle Assessment (LCA), have been constrained by data intensity, temporal lag, and limited capacity to handle the complexity of circular value chains where materials undergo multiple transformations across distributed actors. AI-driven approaches address these limitations by processing large volumes of heterogeneous data, identifying patterns invisible to conventional analysis, and providing real-time sustainability insights that enable proactive rather than reactive environmental management.",

        "Explainable AI (XAI) represents a particularly important advancement for sustainability assessment in Industry 5.0 contexts [21]. While conventional AI models—particularly deep learning systems—often function as 'black boxes' whose reasoning is opaque to human users, XAI techniques provide transparent explanations of how sustainability assessments are derived. This transparency is essential for maintaining the human-centric principle of Industry 5.0: stakeholders including factory managers, designers, and policymakers need to understand why an AI system recommends particular circular strategies in order to exercise informed judgment, identify potential errors, and build trust in AI-supported decision-making [22].",

        "Environmental Impact Forecasting represents another crucial AI application, enabling manufacturers to predict the environmental consequences of design decisions, process modifications, and supply chain configurations before they are implemented [23]. Machine learning models trained on historical LCA data can rapidly estimate environmental impacts for novel product configurations, dramatically reducing the time and cost of sustainability assessment. This predictive capability is particularly valuable in circular economy contexts where multiple end-of-life scenarios must be evaluated—recycling versus remanufacturing versus reuse—each with distinct environmental profiles that depend on context-specific factors including transportation distances, energy sources, and material degradation states. The ability to rapidly assess environmental trade-offs across dozens or hundreds of design alternatives enables genuinely sustainability-optimized decisions rather than the limited comparisons possible with traditional manual LCA methods.",

        "The integration of AI with manufacturing execution systems creates opportunities for real-time sustainability monitoring that was previously impossible [20]. Rather than conducting periodic environmental assessments based on historical data, AI-enabled systems can continuously calculate environmental performance indicators based on actual production data—energy consumption, material yields, waste generation, emissions—and alert operators when sustainability metrics deviate from targets. This real-time visibility enables proactive intervention before environmental performance degrades significantly, transforming sustainability management from a retrospective reporting exercise into an active operational control function.",

        "Table 2 summarizes the key digital technologies and their applications in circular manufacturing, highlighting the varying maturity levels and standards that characterize this rapidly evolving technological landscape. The integration of AI with LCA methodologies addresses persistent gaps in circularity assessment, particularly the challenge of quantifying environmental benefits from multiple-cycle material use, the difficulty of allocating environmental burdens across shared recycling infrastructure, and the need for dynamic rather than static assessment that reflects evolving energy mixes and technological capabilities [24]. Emerging approaches combine physics-based process models with data-driven machine learning to create hybrid assessment frameworks that maintain physical interpretability while capturing empirical patterns that pure mechanistic models miss.",
    ]

    for text in s2_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('2.3 Digital Product Passports and Supply Chain Traceability', level=2)

    s2_3 = [
        "Digital Product Passports (DPPs) represent a transformative approach to information management across product lifecycles, documenting material composition, processing history, environmental impact data, and end-of-life instructions in machine-readable formats that accompany products through all lifecycle stages [25]. Reinforced by EU policy frameworks including the proposed Ecodesign for Sustainable Products Regulation, DPPs create the information infrastructure necessary for effective circular economy implementation by ensuring that downstream actors—repairers, remanufacturers, recyclers—have access to the material and design information required for value-preserving end-of-life interventions.",

        "The technical architecture of DPPs must balance comprehensive information capture with practical implementation feasibility [26]. Decentralized approaches using distributed ledger technologies enable multiple supply chain actors to contribute and access passport information without requiring centralized control or trust. Semi-supervised learning techniques can extract meaningful patterns from the heterogeneous data generated by decentralized actors, identifying relationships between material properties, processing conditions, and recyclability outcomes that inform circular design decisions [27]. Hybrid methods combining collaborative filtering with social network analysis offer promising approaches for assessing service quality and production suitability across distributed manufacturing networks where direct quality assessment is impractical.",

        "Supply chain traceability—the ability to track materials and products through complex, multi-tier value chains—provides the empirical foundation upon which DPPs depend [28]. In circular economy contexts, traceability extends beyond forward logistics to encompass reverse flows: tracking products from consumers back to reprocessing facilities, monitoring material quality through multiple recycling cycles, and maintaining chain-of-custody documentation that verifies recycled content claims. The combination of IoT sensing, blockchain verification, and AI-driven pattern recognition creates traceability systems capable of operating across the complex, dynamic networks characteristic of circular value chains. These integrated traceability systems also generate valuable data about actual material flows, recovery rates, and quality degradation patterns that inform both operational optimization and strategic planning for circular manufacturing investments.",

        "The implementation of DPPs across manufacturing sectors faces challenges including data standardization (what information should passports contain?), interoperability (how should different passport systems communicate?), privacy (how to balance transparency with proprietary information protection?), and governance (who controls passport data and ensures accuracy?) [29]. EU regulatory frameworks are progressively clarifying these requirements, but significant industry-level coordination remains necessary to develop sector-specific passport templates, verification mechanisms, and data sharing agreements that enable circular economy practices while respecting competitive boundaries. The success of DPPs ultimately depends on achieving network effects: their value to individual actors increases as more supply chain participants adopt compatible systems, creating positive feedback dynamics that can rapidly accelerate implementation once critical adoption thresholds are reached.",
    ]

    for text in s2_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================
    # SECTION III
    # ============================
    doc.add_heading('Section III: Case-Based Approaches and Sectoral Applications', level=1)

    doc.add_heading('3.1 Photovoltaic Manufacturing – Designing for Circularity', level=2)

    s3_1 = [
        "The solar photovoltaic (PV) industry presents a compelling case for circular-by-design manufacturing, driven by both environmental necessity and regulatory pressure [30]. With global installed PV capacity exceeding one terawatt and projected to grow dramatically in coming decades, the volumes of end-of-life panels requiring management will escalate correspondingly—creating either an environmental crisis or a circular economy opportunity depending on how the industry responds. EU ecodesign regulations increasingly mandate circularity considerations in PV product design, requiring manufacturers to demonstrate repairability, recyclability, and material recovery potential as conditions for market access.",

        "Figure 3 illustrates the circular-by-design approach for PV manufacturing, showing how circularity considerations integrate across all lifecycle phases: material selection (choosing materials with high recyclability and low toxicity), product design (enabling disassembly and component separation), manufacturing process (minimizing scrap and enabling reprocessing), use phase (facilitating maintenance and repair), end-of-life recovery (maximizing material and component value retention), and remanufacture/reuse (extending product service life through refurbishment). This comprehensive integration distinguishes circular-by-design from end-of-pipe recycling approaches that attempt to manage waste after it has been created rather than preventing it through upstream design decisions.",

        "AI-driven product modelling plays a central role in PV circular design by enabling rapid evaluation of design alternatives across multiple circularity criteria simultaneously [31]. Machine learning models can predict how material choices affect not only immediate manufacturing cost and product performance but also downstream recyclability, toxic substance liberation during end-of-life processing, and energy requirements for material recovery. These multi-objective optimization capabilities allow designers to navigate the complex trade-offs inherent in circular design—for example, balancing the superior performance of certain rare-earth elements against their limited recyclability and geopolitical supply risks. The integration of circularity optimization into the earliest design stages ensures that environmental considerations shape fundamental product architecture rather than being retrofitted as afterthoughts onto designs already locked by performance and cost decisions.",

        "The manufacturing process itself offers significant circularity opportunities in PV production, including closed-loop recovery of cutting fluids, solvents, and process chemicals; scrap minimization through AI-optimized cell stringing and module assembly; and quality-based sorting that directs sub-specification cells toward less demanding applications rather than waste streams [30]. Digital twin-enabled process optimization can reduce material waste by 30-45% compared to conventional manufacturing approaches by continuously adjusting process parameters to minimize scrap generation while maintaining quality specifications. The combination of process-level circularity with design-level circularity creates compounding benefits: well-designed products are easier to manufacture efficiently, and efficient manufacturing processes preserve the circularity potential embedded in circular product designs.",

        "Table 3 presents the circular economy strategies applicable to PV manufacturing across different lifecycle phases, including specific interventions, enabling technologies, and estimated circularity improvements. As detailed in Table 3, the greatest circularity gains typically result from design-phase decisions that determine material selection, joining methods, and modular architecture—reinforcing the circular economy principle that upstream interventions deliver disproportionate downstream benefits. The integration of Digital Product Passports with PV panels enables end-of-life processors to access detailed material composition data, processing history, and disassembly instructions that dramatically improve recovery efficiency [32].",
    ]

    for text in s3_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 3
    doc.add_paragraph()
    t3_title = doc.add_paragraph()
    t3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t3_title.add_run('Table 3. ')
    r.bold = True
    t3_title.add_run('Circular Economy Strategies in Photovoltaic Manufacturing Across Lifecycle Phases')

    table3 = doc.add_table(rows=7, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    h3 = ['Lifecycle Phase', 'Circular Strategy', 'Enabling Technology', 'Circularity Improvement']
    for i, h in enumerate(h3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t3_data = [
        ['Material Selection', 'Low-toxicity, high-recyclability materials', 'AI material screening, LCA databases', 'Recyclability: +40-60%'],
        ['Product Design', 'Modular architecture, design for disassembly', 'Parametric CAD, DfX tools', 'Disassembly time: -50-70%'],
        ['Manufacturing', 'Scrap minimization, closed-loop solvent recovery', 'Digital twin process optimization', 'Material waste: -30-45%'],
        ['Use Phase', 'Predictive maintenance, component repair', 'IoT monitoring, CDT diagnostics', 'Service life: +20-30%'],
        ['End-of-Life', 'Automated disassembly, material sorting', 'Robotics, spectroscopic identification', 'Recovery rate: 85-95%'],
        ['Remanufacture', 'Panel refurbishment, component reuse', 'Quality assessment AI, DPP data', 'Value retention: +60-80%'],
    ]
    for i, row in enumerate(t3_data):
        for j, cell in enumerate(row):
            table3.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    doc.add_heading('3.2 Fast-Moving Consumer Goods – Decentralized Circular Manufacturing', level=2)

    s3_2 = [
        "The fast-moving consumer goods (FMCG) sector presents distinct challenges and opportunities for circular manufacturing, characterized by high production volumes, distributed consumption, complex packaging systems, and rapid innovation cycles [33]. The STARHAUS project exemplifies how decentralized, human-centric, and circular manufacturing models are being developed to address these challenges in sectors including pet food, fertilizers, beverages, and cereals. By combining technological development with social sciences and community co-design, STARHAUS demonstrates that sustainable manufacturing innovation requires engagement with local contexts, cultural practices, and community needs alongside technical optimization.",

        "Figure 3 (Panel B) illustrates the STARHAUS decentralized manufacturing framework, showing how community co-design hubs connect with diverse production streams—pet food, fertilizers, beverages, cereals—while integrating local suppliers and renewable energy sources. This decentralized architecture offers multiple sustainability advantages: reduced transportation distances decrease carbon footprints; local production enables use of seasonal and regional feedstocks; community involvement ensures social acceptance and workforce development; and distributed operations provide resilience against centralized disruption [34]. The alignment with multiple UN Sustainable Development Goals (SDGs 7, 8, 9, 11, 12, 13) demonstrates the comprehensive societal value creation potential of human-centred circular manufacturing.",

        "The combination of technological development with social sciences represents a methodological innovation that distinguishes STARHAUS from purely technology-driven approaches [35]. Community co-design processes ensure that manufacturing systems are compatible with local labour markets, cultural preferences, and environmental conditions. Social innovation assessment evaluates not only environmental and economic outcomes but also impacts on community cohesion, employment quality, and local economic multiplier effects. This holistic approach exemplifies Industry 5.0's human-centric principle by placing community well-being alongside technical performance as a primary design criterion. The participatory methodology also builds social license for manufacturing operations, reducing the community opposition that frequently delays or prevents industrial facility development when communities perceive themselves as bearing environmental burdens without receiving commensurate benefits.",

        "The scalability of decentralized circular manufacturing models depends on developing standardized modular production units that can be adapted to local contexts without requiring complete re-engineering [36]. Digital twin technology enables virtual configuration and testing of production modules before physical deployment, reducing the cost and risk of establishing new decentralized facilities. AI-driven demand forecasting and supply planning systems coordinate production across distributed sites, ensuring that decentralization does not sacrifice the coordination benefits of centralized planning. These technological enablers make decentralized circular manufacturing economically viable at scales previously associated only with centralized mass production. The combination of local adaptability with networked coordination represents a manufacturing paradigm that is simultaneously responsive to community needs and capable of achieving the economies of learning (if not always economies of scale) that competitive markets demand.",
    ]

    for text in s3_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # FIGURE 3
    fig3_path = 'industry5_figures/Figure_3_Case_Studies.png'
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap3.add_run('Figure 3. ')
    r.bold = True
    cap3.add_run('Sectoral applications of circular manufacturing: (A) Photovoltaic manufacturing circular-by-design lifecycle approach; (B) STARHAUS decentralized FMCG manufacturing framework with community co-design hub.')

    doc.add_paragraph()

    doc.add_heading('3.3 Cross-Sectoral Insights – Lean, Digitalization, and Circularity', level=2)

    s3_3 = [
        "Cross-sectoral analysis reveals common patterns and sector-specific adaptations in the integration of lean principles, Industry 4.0 technologies, and circular economy strategies [37]. The convergence of lean manufacturing (with its focus on waste elimination and continuous improvement) with digital technologies (enabling real-time visibility and optimization) and circular economy principles (demanding closed-loop material flows) creates a powerful tripartite framework for sustainable manufacturing transformation. Companies including Siemens, Tesla, Interface Inc, Tata Steel, and BYD have demonstrated various approaches to this integration, advancing sustainable industrial practices through green technologies, renewable energy integration, and circular economy implementation.",

        "Siemens exemplifies technology-enabled circularity through its development of comprehensive digital twin platforms that model entire product lifecycles, enabling circular design optimization from the earliest concept stages [38]. Tesla's closed-loop battery recycling program demonstrates how manufacturers can capture end-of-life value while securing critical material supply—recycling battery cells at the Gigafactory to recover lithium, nickel, cobalt, and other materials for direct reuse in new battery production. Interface Inc's Mission Zero initiative illustrates how service-based business models (carpet leasing rather than selling) align manufacturer incentives with material circularity, since the company retains ownership of materials and benefits directly from their longevity and recyclability. Tata Steel's industrial symbiosis programs show how waste streams from steel production become valuable inputs for cement, fertilizer, and road construction, creating economic value from materials that would otherwise require costly disposal. And BYD's vertical integration strategy demonstrates how controlling multiple lifecycle stages enables circular material management that would be impossible across fragmented value chains.",

        "The synthesis of these diverse examples reveals several cross-cutting principles [39]. First, circularity achieves greatest impact when integrated at the strategic level rather than implemented as an operational add-on—companies that treat circular economy as a core business strategy consistently outperform those that approach it as a compliance exercise. Second, digital technologies are necessary but insufficient enablers—organizational culture, business model innovation, and stakeholder collaboration are equally important determinants of success. Third, the most successful implementations combine multiple circular strategies (reuse, remanufacture, and recycle) rather than relying on any single approach, creating cascading value retention that maximizes resource productivity. And fourth, human-centric design—engaging workers, customers, and communities as active participants rather than passive recipients—consistently accelerates adoption and improves outcomes compared to purely top-down implementation approaches.",

        "The lean-digital-circular convergence also reveals important tensions that organizations must manage [12]. Lean principles of waste elimination sometimes conflict with circular economy requirements for maintaining material inventories (buffer stocks of recovered materials) and accepting variable-quality inputs (recycled materials with different specifications than virgin equivalents). Digital technologies enable circularity but also consume resources and energy, creating potential rebound effects where efficiency gains are offset by increased consumption. Navigating these tensions requires nuanced governance frameworks that balance competing objectives and measure net rather than gross sustainability improvements across entire value systems.",
    ]

    for text in s3_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================
    # SECTION IV
    # ============================
    doc.add_heading('Section IV: Strategic Implementation and Future Directions', level=1)

    doc.add_heading('4.1 Overcoming Barriers to Industry 5.0-Circular Integration', level=2)

    s4_1 = [
        "The integration of circular economy principles within Industry 5.0 frameworks confronts persistent barriers spanning organizational, technical, analytical, and policy dimensions [40]. Organizational resistance manifests through entrenched linear business models, departmental silos that impede cross-functional collaboration on circularity, short-term financial pressures that discourage long-horizon circular investments, and cultural norms that equate innovation with novelty rather than resource stewardship. Overcoming these organizational barriers requires leadership commitment that signals circularity as a strategic priority, incentive structures that reward lifecycle value creation, and change management programs that build circular economy capabilities across all organizational levels.",

        "Technical interoperability challenges represent a significant implementation barrier, particularly in manufacturing environments that comprise heterogeneous equipment, diverse data formats, and legacy systems designed without circular economy considerations [41]. Creating the data infrastructure necessary for circular manufacturing—tracking materials through multiple lifecycle stages, sharing information across organizational boundaries, and integrating environmental assessment into real-time production decisions—demands investments in standardized interfaces, semantic interoperability frameworks, and secure data exchange platforms. Figure 4 maps these barriers against strategic enablers, illustrating how targeted interventions can systematically address implementation obstacles.",

        "Table 4 presents an implementation roadmap for Industry 5.0-circular economy integration, organized by organizational readiness level and identifying specific actions, technologies, and governance mechanisms appropriate at each stage. As outlined in Table 4, organizations typically progress through awareness (understanding circularity opportunities), assessment (evaluating current state and gaps), piloting (testing circular strategies in bounded contexts), scaling (extending successful pilots across operations), and optimization (achieving continuous improvement in circular performance). This staged approach manages risk while building organizational capabilities and evidence bases that support broader transformation.",

        "Data standardization gaps present particular challenges for circular economy implementation across supply chains [42]. Without agreed-upon definitions for circularity metrics, material classification systems, and environmental impact calculation methods, organizations cannot meaningfully compare circular performance, verify supplier claims, or demonstrate compliance with regulatory requirements. Industry-level coordination—through standards bodies, industry associations, and public-private partnerships—is essential for establishing the shared frameworks that enable circular economy practices to scale beyond individual organizations to entire value chains. The development of harmonized circularity indicators that are simultaneously scientifically rigorous, practically measurable, and comprehensible to diverse stakeholders represents one of the most important enabling tasks for the field—without such indicators, neither market mechanisms nor regulatory instruments can effectively drive circular economy adoption at the systemic scale required for meaningful environmental impact.",

        "Strategies for accelerating adoption include adapting manufacturing processes to accommodate circular material flows (such as accepting variable-quality recycled inputs), developing interoperable production systems that can flexibly process both virgin and recovered materials, and investing in workforce development that builds the cross-disciplinary skills required for circular manufacturing [14]. The most successful implementations treat barrier removal as a systematic program rather than a collection of isolated initiatives, recognizing that organizational, technical, and policy barriers are interconnected and must be addressed in coordinated fashion. Figure 4 contextualizes these strategies within the broader implementation landscape, demonstrating how short-term, medium-term, and long-term actions collectively enable comprehensive transformation.",

        "Financial mechanisms for supporting Industry 5.0-circular economy investment deserve specific attention. Traditional investment appraisal methods that focus on short-term return on investment systematically undervalue circular economy investments whose benefits accrue over extended timeframes and across multiple lifecycle stages. Emerging approaches including internal carbon pricing, natural capital accounting, and lifecycle costing provide more complete pictures of circular economy value creation, while innovative financing mechanisms such as sustainability-linked bonds, circular economy funds, and pay-for-performance contracts align financial returns with environmental outcomes. The development of these financial instruments is accelerating as capital markets increasingly recognize both the risks of linear business models and the opportunities in circular value creation.",
    ]

    for text in s4_1:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    # TABLE 4
    doc.add_paragraph()
    t4_title = doc.add_paragraph()
    t4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t4_title.add_run('Table 4. ')
    r.bold = True
    t4_title.add_run('Implementation Roadmap for Industry 5.0-Circular Economy Integration')

    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    h4 = ['Readiness Stage', 'Key Actions', 'Enabling Technologies', 'Governance Mechanisms']
    for i, h in enumerate(h4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    t4_data = [
        ['1. Awareness', 'Stakeholder education, circular opportunity mapping', 'LCA databases, material flow analysis', 'Executive sponsorship, sustainability charter'],
        ['2. Assessment', 'Circularity audits, gap analysis, baseline metrics', 'Digital twin scoping, data infrastructure', 'Cross-functional teams, KPI development'],
        ['3. Piloting', 'Bounded experiments, proof-of-concept DTs', 'AI-driven optimization, DPP prototyping', 'Innovation labs, agile governance'],
        ['4. Scaling', 'Enterprise rollout, supply chain integration', 'Full CDT deployment, blockchain traceability', 'Multi-stakeholder partnerships, standards adoption'],
        ['5. Optimization', 'Continuous improvement, autonomous circularity', 'Hierarchical simulation, predictive systems', 'Industry consortia, regulatory co-evolution'],
    ]
    for i, row in enumerate(t4_data):
        for j, cell in enumerate(row):
            table4.rows[i+1].cells[j].text = cell

    doc.add_paragraph()

    # FIGURE 4
    fig4_path = 'industry5_figures/Figure_4_Implementation_Future.png'
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap4 = doc.add_paragraph()
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap4.add_run('Figure 4. ')
    r.bold = True
    cap4.add_run('Strategic implementation framework showing barriers to Industry 5.0-circular integration, strategic enablers for overcoming them, and a future research timeline mapping emerging trends from 2025 to 2035+.')

    doc.add_paragraph()

    doc.add_heading('4.2 Policy Frameworks and Governance for Sustainable Industrial Transitions', level=2)

    s4_2 = [
        "The policy landscape supporting Industry 5.0 and circular economy integration has evolved rapidly, with the European Union establishing itself as the primary regulatory innovator through frameworks including the European Green Deal, the Circular Economy Action Plan, the proposed Ecodesign for Sustainable Products Regulation, and the Corporate Sustainability Reporting Directive [43]. These interconnected policy instruments create a regulatory environment that simultaneously mandates circularity (through product requirements and reporting obligations), incentivizes sustainable innovation (through green public procurement and research funding), and penalizes linear practices (through extended producer responsibility and landfill restrictions). National-level strategies increasingly align with these EU frameworks while adapting implementation approaches to local industrial structures and capabilities.",

        "Multi-level governance mechanisms are essential for managing the complexity of sustainable industrial transitions that span global supply chains, national regulatory jurisdictions, regional industrial clusters, and individual organizational capabilities [40]. Effective governance requires coordination across levels: international standards that enable cross-border circular material flows; national policies that create coherent incentive structures; regional strategies that build industrial symbiosis networks; and organizational governance that embeds sustainability into decision-making at all levels. Public-private partnerships play a crucial intermediary role, providing forums for co-developing standards, sharing pre-competitive research, and coordinating investments in shared infrastructure such as recycling facilities, data platforms, and skills development programs. The effectiveness of these multi-level governance arrangements depends critically on feedback mechanisms that enable learning and adaptation: monitoring systems that track policy outcomes, evaluation frameworks that assess implementation effectiveness, and revision processes that update regulations based on accumulated evidence.",

        "Evidence-based policy design is increasingly important as circular economy regulations move from aspirational targets to binding requirements with compliance consequences [41]. Policymakers need reliable data on the actual environmental performance of circular strategies, the economic implications of regulatory requirements for different industry sectors, and the social impacts of industrial transitions on workers and communities. The digital technologies discussed throughout this chapter—particularly Digital Twins, AI-driven LCA, and Digital Product Passports—can provide much of this evidence base, creating feedback loops where policy implementation generates data that informs policy refinement in subsequent regulatory cycles. This data-driven approach to policy development represents a significant advancement over traditional regulation-by-negotiation approaches, enabling more precise calibration of requirements to technological capabilities and environmental necessities.",

        "The role of research and innovation policy in supporting Industry 5.0-circular economy integration deserves particular attention. Horizon Europe and its successor programmes have established dedicated funding streams for sustainable manufacturing research, creating knowledge resources that inform both industrial practice and regulatory development. Mission-oriented research programmes that target specific circularity challenges—such as battery recycling, textile circularity, or construction material reuse—focus innovation efforts on societally significant problems while building cross-sectoral collaboration capabilities. The integration of social science research alongside engineering and natural science ensures that human-centric considerations remain central to technological development, preventing the technology-first biases that characterized previous industrial transitions.",

        "The alignment of Industry 5.0-circular economy integration with the United Nations Sustainable Development Goals provides a comprehensive framework for evaluating policy coherence and identifying gaps [35]. Circular manufacturing directly contributes to SDG 12 (responsible consumption and production), but its impacts extend to SDG 7 (affordable and clean energy through renewable energy manufacturing), SDG 8 (decent work through human-centric manufacturing), SDG 9 (industry innovation and infrastructure), SDG 11 (sustainable cities through distributed manufacturing), and SDG 13 (climate action through reduced industrial emissions). Policy frameworks that explicitly map Industry 5.0 interventions to SDG targets can demonstrate the comprehensive societal value of circular manufacturing investments, strengthening the political case for supportive regulation and public investment in enabling infrastructure.",
    ]

    for text in s4_2:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('4.3 Future Research Directions and Emerging Trends', level=2)

    # NO CITATIONS in this concluding section
    s4_3 = [
        "The integration of Industry 5.0 with circular economy principles remains in relatively early stages, with significant research gaps requiring attention across multiple dimensions. Longitudinal assessments of Industry 5.0-circular economy integration across diverse manufacturing contexts are critically needed to move beyond cross-sectional case studies toward understanding how circular capabilities develop over time, what implementation sequences prove most effective, and which contextual factors most strongly influence outcomes. Such longitudinal research would provide the evidence base necessary for refining implementation frameworks and policy interventions based on demonstrated rather than theorized effectiveness.",

        "The interplay between digital technologies and sustainability initiatives presents rich research opportunities. While the enabling potential of Digital Twins, AI, and Digital Product Passports for circular manufacturing is increasingly recognized, empirical evidence regarding their actual sustainability impact—as opposed to their theoretical potential—remains limited. Research is needed that rigorously quantifies the environmental benefits attributable to digital enablers, identifies contexts where digitalization supports versus undermines sustainability goals, and develops frameworks for ensuring that the environmental costs of digital infrastructure do not offset the sustainability gains it enables.",

        "Standardized circularity metrics represent a fundamental gap that constrains both research and practice. Without agreed-upon measurement frameworks, organizations cannot benchmark performance, researchers cannot compare findings across studies, and policymakers cannot verify compliance or assess policy effectiveness. Future research should develop metrics that capture circularity across multiple dimensions—material, energy, economic, and social—while remaining practically implementable within manufacturing information systems. These metrics must balance comprehensiveness with usability, providing meaningful insight without imposing prohibitive measurement burden.",

        "Emerging technologies including hierarchical simulation engines, predictive maintenance systems, and recipe optimization algorithms for reusability offer promising but underexplored potential for advancing circular manufacturing. Hierarchical simulation engines that model circular systems at multiple scales simultaneously—from molecular-level material degradation through component-level failure prediction to system-level supply chain dynamics—could enable unprecedented optimization of circular strategies. Predictive maintenance, when integrated with circular economy logic, can determine optimal timing not just for repair but for transitioning components from primary to secondary applications based on remaining useful life assessment.",

        "The concept of 'manufacturing as a service' represents a particularly underexplored opportunity for integrating circular paradigms with Industry 5.0 capabilities. When manufacturing capability is offered as a service rather than tied to specific products, the incentive structures naturally align with circular economy principles: service providers benefit from durability, maintainability, and resource efficiency because they retain responsibility for physical assets throughout their lifecycle. Digital twins enable remote monitoring and optimization of distributed manufacturing assets, while AI systems can dynamically match production capacity with demand across networks of manufacturing service providers. This servitization of manufacturing capabilities could fundamentally transform industrial economics, shifting value creation from material throughput toward service delivery and knowledge application.",

        "The social dimensions of Industry 5.0-circular economy integration require significantly more research attention. While technical and economic aspects have received substantial investigation, questions about workforce transition, skills development, community impact, and distributional equity remain underexplored. How can organizations ensure that circular manufacturing transitions create decent work opportunities rather than eliminating jobs? What training and education systems are needed to build circular economy competencies across the manufacturing workforce? How can the benefits of circular manufacturing be distributed equitably rather than concentrating among technology-rich firms while disadvantaging smaller enterprises? These social questions are not peripheral to the Industry 5.0-circular economy agenda—they are central to its human-centric pillar and will ultimately determine whether the transition achieves societal acceptance and lasting success.",

        "Looking forward, the convergence of Industry 5.0 principles with circular economy practices will likely accelerate as regulatory pressure intensifies, technological capabilities mature, and competitive advantages shift toward organizations that master circular value creation. The most transformative outcomes will emerge not from any single technology or policy intervention but from the systemic integration of human-centred design, digital intelligence, circular material flows, and multi-stakeholder governance into coherent manufacturing systems that simultaneously serve economic, environmental, and social objectives. This integrative vision represents both the greatest challenge and the greatest opportunity confronting manufacturing in the coming decades—demanding unprecedented collaboration across disciplines, sectors, and governance levels to realize its full potential for sustainable industrial transformation.",
    ]

    for text in s4_3:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================
    # REFERENCES [1]-[43]
    # ============================
    doc.add_heading('References', level=1)

    references = [
        "[1] Xu, L. D., Xu, E. L., & Li, L. (2018). Industry 4.0: State of the art and future trends. International Journal of Production Research, 56(8), 2941–2962.",
        "[2] Leng, J., Sha, W., Wang, B., Zheng, P., Zhuang, C., Liu, Q., & Wang, L. (2022). Industry 5.0: Prospect and retrospect. Journal of Manufacturing Systems, 65, 279–295.",
        "[3] European Commission. (2021). Industry 5.0: Towards a sustainable, human-centric and resilient European industry. Publications Office of the European Union.",
        "[4] Breque, M., De Nul, L., & Petridis, A. (2021). Industry 5.0: Towards a sustainable, human-centric and resilient European industry. European Commission, Directorate-General for Research and Innovation.",
        "[5] Nahavandi, S. (2019). Industry 5.0—A human-centric solution. Sustainability, 11(16), 4371.",
        "[6] Ghobakhloo, M., Iranmanesh, M., Mubarak, M. F., Mubarik, M., Rejeb, A., & Nilashi, M. (2022). Identifying Industry 5.0 contributions to sustainable development: A strategy roadmap for delivering sustainability values. Sustainable Production and Consumption, 33, 716–737.",
        "[7] Xu, X., Lu, Y., Vogel-Heuser, B., & Wang, L. (2021). Industry 4.0 and Industry 5.0—Inception, conception and perception. Journal of Manufacturing Systems, 61, 530–535.",
        "[8] Maddikunta, P. K. R., Pham, Q. V., Prabadevi, B., Deepa, N., Dev, K., Gadekallu, T. R., & Liyanage, M. (2022). Industry 5.0: A survey on enabling technologies and potential applications. Journal of Industrial Information Integration, 26, 100257.",
        "[9] Akundi, A., Euresti, D., Luna, S., Ankobiah, W., Lopes, A., & Edinbarough, I. (2022). State of Industry 5.0—Analysis and identification of current research trends. Applied System Innovation, 5(1), 27.",
        "[10] Ellen MacArthur Foundation. (2013). Towards the circular economy: Economic and business rationale for an accelerated transition. Ellen MacArthur Foundation.",
        "[11] Kirchherr, J., Reike, D., & Hekkert, M. (2017). Conceptualizing the circular economy: An analysis of 114 definitions. Resources, Conservation and Recycling, 127, 221–232.",
        "[12] De Angelis, R. (2021). Circular economy and Industry 4.0: One-way or mutual relationships? An analysis of the literature. Sustainability, 13(13), 7526.",
        "[13] Sassanelli, C., Rosa, P., Rocca, R., & Terzi, S. (2019). Circular economy performance assessment methods: A systematic literature review. Journal of Cleaner Production, 229, 440–453.",
        "[14] Cañas, H., Mula, J., Díaz-Madroñero, M., & Campuzano-Bolarín, F. (2022). Implementing Industry 4.0 principles to develop circular manufacturing solutions. Computers & Industrial Engineering, 171, 108374.",
        "[15] Tao, F., Cheng, J., Qi, Q., Zhang, M., Zhang, H., & Sui, F. (2018). Digital twin-driven product design, manufacturing and service with big data. International Journal of Advanced Manufacturing Technology, 94(9), 3563–3576.",
        "[16] Psarommatis, F., & May, G. (2023). A literature review and design methodology for digital twins in the era of zero defect manufacturing. International Journal of Production Research, 61(15), 5213–5236.",
        "[17] Zheng, X., Lu, J., & Kiritsis, D. (2022). The emergence of cognitive digital twin: Vision, challenges and opportunities. International Journal of Production Research, 60(24), 7610–7632.",
        "[18] Lu, Y., Liu, C., Wang, K. I., Huang, H., & Xu, X. (2020). Digital twin-driven smart manufacturing: Connotation, reference model, applications and research issues. Robotics and Computer-Integrated Manufacturing, 61, 101837.",
        "[19] Kritzinger, W., Karner, M., Traar, G., Henjes, J., & Sihn, W. (2018). Digital twin in manufacturing: A categorical literature review. IFAC-PapersOnLine, 51(11), 1016–1022.",
        "[20] Nishant, R., Kennedy, M., & Corbett, J. (2020). Artificial intelligence for sustainability: Challenges, opportunities, and a research agenda. International Journal of Information Management, 53, 102104.",
        "[21] Arrieta, A. B., Díaz-Rodríguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., & Herrera, F. (2020). Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82–115.",
        "[22] Adadi, A., & Berrada, M. (2018). Peeking inside the black-box: A survey on explainable artificial intelligence. IEEE Access, 6, 52138–52160.",
        "[23] Rolnick, D., Donti, P. L., Kaack, L. H., Kochanski, K., Lacoste, A., Sankaran, K., & Bengio, Y. (2022). Tackling climate change with machine learning. ACM Computing Surveys, 55(2), 1–96.",
        "[24] Haupt, M., & Zschokke, M. (2017). How can LCA support the circular economy? International Journal of Life Cycle Assessment, 22(5), 832–837.",
        "[25] Adisorn, T., Thimm, L., & Rücker, M. (2021). Introduction of a digital product passport based on its role in circular economy. Procedia CIRP, 105, 356–361.",
        "[26] Walden, J., Steinbrecher, A., & Marinkovic, M. (2021). Digital product passports as enabler of the circular economy. Chemie Ingenieur Technik, 93(11), 1717–1727.",
        "[27] van Engelen, J. E. (2020). A survey on semi-supervised learning. Machine Learning, 109(2), 373–440.",
        "[28] Saberi, S., Kouhizadeh, M., Sarkis, J., & Shen, L. (2019). Blockchain technology and its relationships to sustainable supply chain management. International Journal of Production Research, 57(7), 2117–2135.",
        "[29] King, T., Butcher, M. P., & Zilber, L. (2020). Identifying opportunities and barriers to advancing the circular economy for electronic products. Journal of Cleaner Production, 277, 123513.",
        "[30] Deng, R., Chang, N. L., Ouyang, Z., & Chong, C. M. (2019). A techno-economic review of silicon photovoltaic module recycling. Renewable and Sustainable Energy Reviews, 109, 532–550.",
        "[31] Norgren, A., Carpenter, A., & Heath, G. (2020). Design for recycling principles applicable to selected clean energy technologies. Journal of Sustainable Metallurgy, 6(4), 761–774.",
        "[32] Tsanakas, J. A., van der Heide, A., Radavičius, T., Denafas, J., Lemaire, E., Wang, K., & Voroshazi, E. (2020). Towards a circular supply chain for PV modules: Review of today's challenges in PV recycling, refurbishment and re-certification. Progress in Photovoltaics, 28(6), 454–464.",
        "[33] Bocken, N. M. P., de Pauw, I., Bakker, C., & van der Grinten, B. (2016). Product design and business model strategies for a circular economy. Journal of Industrial and Production Engineering, 33(5), 308–320.",
        "[34] Korhonen, J., Honkasalo, A., & Seppälä, J. (2018). Circular economy: The concept and its limitations. Ecological Economics, 143, 37–46.",
        "[35] Schroeder, P., Anggraeni, K., & Weber, U. (2019). The relevance of circular economy practices to the sustainable development goals. Journal of Industrial Ecology, 23(1), 77–95.",
        "[36] Stahel, W. R. (2016). The circular economy. Nature, 531(7595), 435–438.",
        "[37] Bressanelli, G., Adrodegari, F., Perona, M., & Saccani, N. (2018). Exploring how usage-focused business models enable circular economy through digital technologies. Sustainability, 10(3), 639.",
        "[38] Riesener, M., Schuh, G., Dölle, C., & Tönnes, C. (2019). The digital twin as enabler for sustainability in product development. Procedia CIRP, 80, 468–473.",
        "[39] Geissdoerfer, M., Savaget, P., Bocken, N. M. P., & Hultink, E. J. (2017). The circular economy—A new sustainability paradigm? Journal of Cleaner Production, 143, 757–768.",
        "[40] Rizos, V., Behrens, A., van der Gaast, W., Hofman, E., Ioannou, A., Kafyeke, T., & Topi, C. (2016). Implementation of circular economy business models by small and medium-sized enterprises. Sustainability, 8(11), 1212.",
        "[41] Ghisellini, P., Cialani, C., & Ulgiati, S. (2016). A review on circular economy: The expected transition to a balanced interplay of environmental and economic systems. Journal of Cleaner Production, 114, 11–32.",
        "[42] Saidani, M., Yannou, B., Leroy, Y., Cluzel, F., & Kendall, A. (2019). A taxonomy of circular economy indicators. Journal of Cleaner Production, 207, 542–559.",
        "[43] European Commission. (2020). A new Circular Economy Action Plan: For a cleaner and more competitive Europe. COM(2020) 98 final.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)

    # Save
    output_path = 'Chapter_Industry5_Circular_Manufacturing.docx'
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

    # Count
    total = ''
    for para in doc.paragraphs:
        total += para.text + ' '
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += cell.text + ' '
    print(f"Total word count: {len(total.split())}")
    print(f"References: {len(references)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == '__main__':
    create_docx()
