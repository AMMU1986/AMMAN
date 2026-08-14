"""
Generate Word Document for Chapter: Reframing Design Thinking for Business Strategy
Includes all text content, 4 tables, 4 figures, and 43 references.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for key, val in kwargs[edge].items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=None, space_after=None, space_before=None, alignment=None):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
    if alignment:
        para.alignment = alignment
    return para

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ==================== TITLE PAGE ====================
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Reframing Design Thinking for Business Strategy:\nA Comparative Analysis')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Chapter: Design Thinking and Other Approaches')
    run.italic = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add page break
    doc.add_page_break()
    
    # ==================== ABSTRACT ====================
    abstract_heading = doc.add_heading('Abstract', level=1)
    
    abstract_text = (
        "The contemporary business environment is characterised by unprecedented complexity, rapid technological change, "
        "and evolving stakeholder expectations. Traditional strategic frameworks, while providing essential analytical rigour, "
        "often fall short in addressing the multifaceted challenges that modern organisations face. This chapter presents a "
        "comprehensive comparative analysis of four foundational approaches to strategic decision-making: Analytical "
        "Decision-Making, Systems Thinking, Creative Thinking, and Design Thinking. By examining the theoretical "
        "underpinnings, methodological tools, and strategic applications of each approach, this chapter argues that Design "
        "Thinking serves as a powerful integrative framework capable of synthesising the strengths of analytical rigour, "
        "systemic awareness, and creative ideation into a cohesive, human-centred strategy. Drawing on real-world "
        "applications from product development, digital transformation, healthcare, entrepreneurship, and manufacturing, "
        "the analysis demonstrates how organisations that adopt an integrated framework achieve superior innovation outcomes, "
        "enhanced customer loyalty, and sustainable competitive advantage. The chapter further examines key challenges in "
        "implementation, including cultural resistance, scalability concerns, and process integration barriers, before "
        "concluding with future directions that emphasise the role of artificial intelligence, digital tools, and hybrid "
        "strategic frameworks in advancing integrated innovation methodologies."
    )
    
    para = doc.add_paragraph(abstract_text)
    para.paragraph_format.first_line_indent = Cm(0)
    for run in para.runs:
        run.font.size = Pt(11)
    
    # Keywords
    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    run = kw_para.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(11)
    run = kw_para.add_run('Design Thinking, Business Strategy, Systems Thinking, Analytical Decision-Making, '
                          'Creative Thinking, Innovation, Human-Centred Design, Strategic Integration')
    run.font.size = Pt(11)
    run.italic = True
    
    doc.add_page_break()
    
    # ==================== SECTION 1 ====================
    doc.add_heading('1. The Landscape of Strategic Decision-Making', level=1)
    
    # 1.1
    doc.add_heading('1.1 The Evolution of Business Strategy', level=2)
    
    doc.add_paragraph(
        'The discipline of business strategy has undergone a profound transformation over the past five decades. '
        'The era of classical strategic planning, characterised by top-down, hierarchical models emphasising '
        'long-term forecasting and resource allocation, dominated corporate thinking from the 1960s through the '
        '1980s [1]. Pioneering frameworks such as Porter\'s Five Forces and the BCG Growth-Share Matrix provided '
        'managers with structured tools for competitive analysis and portfolio management [2]. These models assumed '
        'a relatively stable and predictable business environment where past performance could reliably inform '
        'future direction.'
    )
    
    doc.add_paragraph(
        'However, the advent of globalisation, digital disruption, and the knowledge economy has fundamentally '
        'altered the strategic landscape. The concept of VUCA (Volatility, Uncertainty, Complexity, and Ambiguity), '
        'originally developed by the U.S. Army War College, has become a defining descriptor of the modern business '
        'environment [3]. In this context, organisations can no longer rely solely on historical data extrapolation '
        'or linear planning processes. The emergence of agile methodologies, lean startup principles, and '
        'innovation-driven models reflects a paradigm shift toward more adaptive, iterative, and human-centric '
        'approaches to strategy formulation [4].'
    )
    
    doc.add_paragraph(
        'The implications of this shift extend beyond methodology to encompass fundamental questions about the '
        'nature of strategic advantage in the twenty-first century. Where once competitive advantage could be '
        'sustained through superior resource positioning or cost leadership, contemporary markets increasingly '
        'reward organisations that demonstrate superior capability in understanding and responding to human needs. '
        'The proliferation of digital platforms, the democratisation of production technologies, and the increasing '
        'transparency of markets have collectively eroded traditional barriers to entry and shortened the lifespan '
        'of product-based advantages. In this environment, the capacity for continuous innovation — rooted in deep '
        'understanding of human experience — emerges as perhaps the most durable source of competitive '
        'differentiation available to modern organisations.'
    )
    
    doc.add_paragraph(
        'This evolution is not merely a theoretical exercise; it has tangible implications for organisational '
        'performance. Research indicates that companies embracing human-centred innovation practices outperform '
        'their peers by significant margins in revenue growth, shareholder returns, and market capitalisation [5]. '
        'The Design Management Institute\'s Design Value Index, for instance, demonstrated that design-led companies '
        'outperformed the S&P 500 by 219% over a ten-year period [6]. Such findings underscore the strategic '
        'imperative of integrating diverse thinking approaches into the fabric of organisational decision-making.'
    )
    
    doc.add_paragraph(
        'The shift from purely analytical frameworks to more holistic approaches also reflects a deeper '
        'epistemological change in how organisations conceptualise problems. Rather than viewing challenges as '
        'well-defined puzzles amenable to optimisation, progressive organisations increasingly recognise that their '
        'most pressing strategic issues are "wicked problems" — complex, interconnected, and resistant to '
        'straightforward solutions [7]. This recognition necessitates a multi-faceted strategic toolkit that '
        'combines the precision of analysis with the breadth of systemic understanding, the spark of creativity, '
        'and the empathy of human-centred design.'
    )
    
    doc.add_paragraph(
        'Moreover, the acceleration of technological change has compressed strategic planning horizons dramatically. '
        'Where five-year strategic plans were once standard practice, many organisations now operate on quarterly or '
        'even monthly planning cycles, recognising that the assumptions underlying longer-term plans become obsolete '
        'before they can be fully implemented. This temporal compression demands strategic approaches that enable '
        'rapid learning, flexible adaptation, and continuous recalibration — qualities that are inherent in the '
        'iterative, prototype-driven approaches discussed in this chapter. The organisations that successfully '
        'navigate this compressed timeline are those that have developed the capability to sense, interpret, and '
        'respond to environmental changes with speed and precision, drawing simultaneously on analytical insight, '
        'systemic understanding, creative imagination, and empathic connection with stakeholders.'
    )
    
    # 1.2
    doc.add_heading('1.2 Defining the Core Approaches', level=2)
    
    doc.add_paragraph(
        'This chapter examines four distinct yet complementary approaches to strategic decision-making, each '
        'offering unique perspectives and methodological tools for navigating business complexity.'
    )
    
    doc.add_paragraph(
        'Analytical Decision-Making represents the most established tradition in strategic management. Rooted in '
        'the principles of rationality, evidence-based reasoning, and quantitative optimisation, this approach '
        'seeks to reduce uncertainty through systematic data collection, statistical analysis, and logical '
        'inference [8]. It provides the essential foundation of rigour and objectivity upon which sound strategic '
        'decisions are built.'
    )
    
    doc.add_paragraph(
        'Systems Thinking offers a holistic lens for understanding organisations as complex, dynamic entities '
        'embedded within broader environmental contexts. Drawing on cybernetics, general systems theory, and '
        'complexity science, this approach emphasises interconnectedness, feedback loops, emergence, and non-linear '
        'causality [9]. It enables strategists to see beyond immediate cause-and-effect relationships and '
        'understand the deeper structural forces shaping organisational behaviour.'
    )
    
    doc.add_paragraph(
        'The systems perspective is particularly valuable in an era of increasing organisational complexity, where '
        'decisions in one domain inevitably produce ripple effects across others. Supply chain disruptions, for '
        'instance, cannot be understood in isolation from geopolitical dynamics, environmental constraints, and '
        'technological dependencies. Similarly, human resource strategies interact with innovation capabilities, '
        'customer relationships, and financial performance in ways that linear analytical models often fail to '
        'capture. Systems thinking provides the conceptual vocabulary and analytical tools necessary to navigate '
        'this web of interdependencies.'
    )
    
    doc.add_paragraph(
        'Creative Thinking encompasses the cognitive processes and structured techniques that enable the generation '
        'of novel ideas, solutions, and perspectives. Grounded in research on divergent thinking, associative '
        'cognition, and cognitive flexibility, creative thinking provides the generative engine for innovation [10]. '
        'It enables organisations to break free from established mental models and explore previously unconsidered '
        'possibilities.'
    )
    
    doc.add_paragraph(
        'Design Thinking (DT) represents a human-centred, iterative, and empathetic problem-solving framework '
        'that has gained substantial traction in business contexts since its popularisation by IDEO and the '
        'Stanford d.school [11]. As illustrated in Figure 1, DT integrates elements of empathy, experimentation, '
        'and iteration into a cohesive methodology that bridges the gap between understanding user needs and '
        'delivering viable solutions.'
    )
    
    # Insert Figure 1
    doc.add_paragraph()
    if os.path.exists('dt_chapter_figures/Figure_1_Integrated_Framework.png'):
        doc.add_picture('dt_chapter_figures/Figure_1_Integrated_Framework.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig1_caption = doc.add_paragraph()
    fig1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig1_caption.add_run('Figure 1: ')
    run.bold = True
    run.font.size = Pt(10)
    run = fig1_caption.add_run('The Integrated Strategic Framework — Positioning Design Thinking at the Centre '
                               'of Analytical, Systems, and Creative Approaches')
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()
    
    # 1.3
    doc.add_heading('1.3 The Need for Integration', level=2)
    
    doc.add_paragraph(
        'While each of these approaches offers significant value individually, their true strategic power emerges '
        'through thoughtful integration. Analytical decision-making, when used in isolation, can lead to "analysis '
        'paralysis" and an overemphasis on quantifiable metrics at the expense of human insight and creative '
        'possibility [12]. Systems thinking, while providing valuable holistic understanding, may struggle to '
        'translate systemic insights into actionable solutions at the operational level [13]. Creative thinking, '
        'unbounded by constraints, can generate ideas that are brilliant in concept but impractical in '
        'implementation [14].'
    )
    
    doc.add_paragraph(
        'The limitations of siloed approaches are well documented in the strategic management literature. '
        'Mintzberg\'s critique of "the design school" of strategy formation highlighted how overly rational '
        'planning processes could disconnect strategy from the emergent, learning-based nature of organisational '
        'adaptation [15]. Similarly, Christensen\'s work on disruptive innovation demonstrated how established '
        'firms\' reliance on analytical models optimised for current customers could blind them to transformative '
        'market shifts [16].'
    )
    
    doc.add_paragraph(
        'This chapter argues that Design Thinking, with its inherent emphasis on integration, iteration, and '
        'human-centredness, provides a natural framework for synthesising these complementary approaches. As '
        'demonstrated in Table 1, each approach contributes distinct strengths that, when combined through a '
        'DT-led integration, create a more robust and comprehensive strategic capability. The following sections '
        'will elaborate on this thesis through detailed examination of each approach and their synergistic potential.'
    )
    
    # ==================== TABLE 1 ====================
    doc.add_paragraph()
    table1_title = doc.add_paragraph()
    table1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table1_title.add_run('Table 1: ')
    run.bold = True
    run.font.size = Pt(10)
    run = table1_title.add_run('Comparative Overview of Four Strategic Approaches')
    run.font.size = Pt(10)
    run.italic = True
    
    table1 = doc.add_table(rows=5, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Dimension', 'Analytical\nDecision-Making', 'Systems\nThinking', 'Creative\nThinking', 'Design\nThinking']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1A237E')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    table1_data = [
        ['Core Focus', 'Data, logic,\noptimisation', 'Interconnectedness,\nfeedback loops', 'Ideation, novelty,\ndivergent thinking', 'Human needs,\nempathy, iteration'],
        ['Primary Strength', 'Risk reduction,\nevidence-based', 'Holistic understanding,\nunintended consequences', 'Novel solutions,\nbreakthrough ideas', 'User-centred solutions,\nrapid validation'],
        ['Key Limitation', 'Analysis paralysis,\nmisses qualitative', 'Difficulty translating\nto action', 'May lack feasibility\nand grounding', 'Scalability,\ncultural resistance'],
        ['Strategic Role', 'Validation and\nmeasurement', 'Contextual awareness\nand sustainability', 'Innovation engine\nand disruption', 'Integration and\nhuman-centred strategy'],
    ]
    
    for row_idx, row_data in enumerate(table1_data, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table1.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if col_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
        if row_idx % 2 == 0:
            for col_idx in range(5):
                set_cell_shading(table1.rows[row_idx].cells[col_idx], 'E8EAF6')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ==================== SECTION 2 ====================
    doc.add_heading('2. A Deep Dive into Complementary Approaches', level=1)
    
    # 2.1
    doc.add_heading('2.1 Analytical Decision-Making Models', level=2)
    doc.add_heading('2.1.1 Theory and Principles', level=3)
    
    doc.add_paragraph(
        'Analytical decision-making has its intellectual roots in the Enlightenment tradition of rational inquiry, '
        'subsequently formalised through the development of operations research during World War II and the '
        'emergence of management science in the post-war era [17]. The foundational assumption is that optimal '
        'decisions can be derived through systematic application of logic, mathematics, and empirical evidence '
        'to well-defined problem spaces.'
    )
    
    doc.add_paragraph(
        'The theoretical framework rests on several key principles. First, the principle of bounded rationality, '
        'introduced by Herbert Simon, acknowledges that while decision-makers aspire to rational optimisation, '
        'cognitive limitations and information constraints mean that "satisficing" — finding a satisfactory rather '
        'than optimal solution — is often the practical reality [18]. Second, evidence-based management extends '
        'the principles of evidence-based medicine to organisational contexts, arguing that managerial decisions '
        'should be grounded in the best available empirical evidence rather than tradition, intuition, or '
        'anecdote [19].'
    )
    
    doc.add_paragraph(
        'The analytical tradition also encompasses Bayesian reasoning, which provides a formal framework for '
        'updating beliefs and predictions in light of new evidence [20]. This probabilistic approach to '
        'decision-making is particularly relevant in strategic contexts characterised by uncertainty, where '
        'managers must continuously revise their assumptions based on emerging market signals and competitive '
        'developments.'
    )
    
    doc.add_paragraph(
        'A further important principle is the distinction between programmed and non-programmed decisions. '
        'Programmed decisions are routine, repetitive, and amenable to standardised analytical procedures, while '
        'non-programmed decisions are novel, unstructured, and require judgment that extends beyond purely '
        'quantitative analysis. Strategic decisions, by their nature, tend toward the non-programmed end of this '
        'spectrum, suggesting that while analytical tools provide essential input, they cannot fully determine '
        'strategic choices without being complemented by other modes of thinking. This recognition of the inherent '
        'limitations of purely rational models has led to the development of behavioural decision theory, which '
        'integrates insights from psychology regarding cognitive biases, heuristics, and the affective dimensions '
        'of judgment into our understanding of how strategic decisions are actually made in practice.'
    )
    
    doc.add_heading('2.1.2 Key Methodologies', level=3)
    
    doc.add_paragraph(
        'The analytical toolkit available to strategic decision-makers is extensive and continually evolving. '
        'Traditional methods include cost-benefit analysis, which provides systematic comparison of the expected '
        'costs and benefits of alternative courses of action; SWOT analysis, which maps internal strengths and '
        'weaknesses against external opportunities and threats; and financial modelling, which projects future '
        'cash flows, returns, and risk profiles under various scenarios [21].'
    )
    
    doc.add_paragraph(
        'More contemporary analytical methods include predictive analytics, which leverages machine learning '
        'algorithms and large datasets to forecast future outcomes; decision tree analysis, which maps complex '
        'decision paths and their probabilistic outcomes; and Monte Carlo simulation, which models the impact of '
        'uncertainty by running thousands of randomised scenarios [22]. The emergence of big data and artificial '
        'intelligence has dramatically expanded the analytical capabilities available to organisations, enabling '
        'real-time analysis of customer behaviour, market trends, and operational performance at unprecedented '
        'scale and granularity [23].'
    )
    
    doc.add_paragraph(
        'As shown in Table 2, these methodologies can be categorised by their primary function, complexity level, '
        'and typical application context within strategic decision-making processes.'
    )
    
    # ==================== TABLE 2 ====================
    doc.add_paragraph()
    table2_title = doc.add_paragraph()
    table2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table2_title.add_run('Table 2: ')
    run.bold = True
    run.font.size = Pt(10)
    run = table2_title.add_run('Analytical Decision-Making Methodologies — Classification and Applications')
    run.font.size = Pt(10)
    run.italic = True
    
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers2 = ['Methodology', 'Primary Function', 'Complexity', 'Application Context']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2E7D32')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    table2_data = [
        ['Cost-Benefit Analysis', 'Evaluation', 'Low-Medium', 'Project selection, investment decisions'],
        ['SWOT Analysis', 'Assessment', 'Low', 'Strategic positioning, competitive analysis'],
        ['Predictive Analytics', 'Forecasting', 'High', 'Market trends, customer behaviour'],
        ['Decision Tree Analysis', 'Decision mapping', 'Medium', 'Complex decisions with multiple paths'],
        ['Monte Carlo Simulation', 'Risk modelling', 'High', 'Financial planning, scenario analysis'],
        ['Balanced Scorecard', 'Performance tracking', 'Medium', 'Strategy execution, KPI monitoring'],
        ['Regression Analysis', 'Relationship modelling', 'Medium-High', 'Demand forecasting, causal analysis'],
    ]
    
    for row_idx, row_data in enumerate(table2_data, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        if row_idx % 2 == 0:
            for col_idx in range(4):
                set_cell_shading(table2.rows[row_idx].cells[col_idx], 'E8F5E9')
    
    doc.add_paragraph()
    
    doc.add_heading('2.1.3 Role in Strategy', level=3)
    
    doc.add_paragraph(
        'Analytical decision-making plays an indispensable role in strategic management by providing the empirical '
        'foundation for informed action. It reduces decision risk by quantifying uncertainty, identifying patterns '
        'in historical data, and projecting likely outcomes of alternative strategies [24]. In capital-intensive '
        'industries such as aerospace, energy, and pharmaceuticals, where the cost of strategic error can be '
        'measured in billions, analytical rigour is not merely desirable but essential for organisational survival.'
    )
    
    doc.add_paragraph(
        'Furthermore, analytical approaches provide the metrics and measurement frameworks necessary for strategic '
        'accountability. Balanced scorecards, key performance indicators, and return-on-investment calculations '
        'enable organisations to track strategic progress, identify deviations from planned trajectories, and '
        'implement corrective actions [25]. Without this analytical infrastructure, strategic management becomes '
        'an exercise in wishful thinking rather than disciplined execution.'
    )
    
    doc.add_paragraph(
        'However, the limitations of purely analytical approaches become apparent when organisations face genuinely '
        'novel challenges — situations where historical data provides limited guidance, where the relevant variables '
        'are difficult to quantify, or where the problem itself is poorly defined. It is precisely in these contexts '
        'that the complementary perspectives of systems thinking, creative thinking, and design thinking become essential.'
    )
    
    doc.add_paragraph(
        'The growing recognition of these limitations has fuelled interest in what scholars term "decision quality" '
        'frameworks that acknowledge the multi-dimensional nature of good strategic decisions. A high-quality '
        'strategic decision requires not only analytical soundness but also creative framing of alternatives, '
        'systemic awareness of consequences, and alignment with the values and needs of key stakeholders. This '
        'multi-dimensional view of decision quality provides a natural bridge to the integrative framework '
        'advocated in this chapter, suggesting that the most effective strategic decisions emerge from processes '
        'that systematically engage multiple cognitive modes rather than privileging any single approach.'
    )
    
    # 2.2
    doc.add_heading('2.2 Systems Thinking', level=2)
    doc.add_heading('2.2.1 Theory and Principles', level=3)
    
    doc.add_paragraph(
        'Systems thinking emerged from the interdisciplinary ferment of the mid-twentieth century, drawing on '
        'cybernetics, general systems theory, and the study of complex adaptive systems [26]. At its core, systems '
        'thinking is founded on the recognition that organisations are not collections of independent parts but '
        'rather integrated wholes whose behaviour emerges from the interactions and relationships among their '
        'constituent elements [27].'
    )
    
    doc.add_paragraph(
        'The key principles of systems thinking include interconnectedness — the recognition that all elements '
        'within a system are linked, often in non-obvious ways; feedback loops — both reinforcing (amplifying) and '
        'balancing (stabilising) loops that drive system behaviour; emergence — the appearance of properties and '
        'behaviours at the system level that cannot be predicted from the characteristics of individual components; '
        'and non-linearity — the recognition that small changes can sometimes produce disproportionately large '
        'effects, while large interventions may have surprisingly limited impact [28].'
    )
    
    doc.add_paragraph(
        'Peter Senge\'s influential work on the "learning organisation" brought systems thinking into mainstream '
        'management discourse, arguing that the inability to see systemic patterns is one of the primary barriers '
        'to organisational learning and adaptation [9]. Senge identified several "systems archetypes" — recurring '
        'patterns of behaviour such as "limits to growth," "shifting the burden," and "tragedy of the commons" — '
        'that help managers recognise familiar systemic dynamics in novel contexts.'
    )
    
    doc.add_heading('2.2.2 Key Methodologies', level=3)
    
    doc.add_paragraph(
        'The methodological toolkit of systems thinking includes several distinctive approaches. Causal loop '
        'diagrams provide visual representations of the feedback structures underlying system behaviour, enabling '
        'stakeholders to map the complex web of cause-and-effect relationships within and around an '
        'organisation [29]. Stock-and-flow models add quantitative rigour by modelling the accumulation and '
        'depletion of resources over time and the rates of change that govern system dynamics.'
    )
    
    doc.add_paragraph(
        'System dynamics modelling, pioneered by Jay Forrester at MIT, enables simulation of complex systems over '
        'extended time horizons, revealing counter-intuitive behaviours and unintended consequences that would be '
        'invisible to linear analytical approaches [30]. More recent developments include agent-based modelling, '
        'which simulates the behaviour of complex systems by modelling the decisions and interactions of individual '
        'agents, and network analysis, which maps and analyses the structure of relationships within and between '
        'organisations.'
    )
    
    doc.add_paragraph(
        'The concept of "leverage points" — places within a complex system where a small intervention can produce '
        'significant and lasting change — is particularly valuable for strategic decision-making [31]. Donella '
        'Meadows identified a hierarchy of leverage points ranging from relatively superficial interventions (such '
        'as adjusting numerical parameters) to profound systemic changes (such as shifting the goals or paradigms '
        'that govern system behaviour).'
    )
    
    doc.add_heading('2.2.3 Role in Strategy', level=3)
    
    doc.add_paragraph(
        'Systems thinking contributes to strategy by enabling leaders to understand the broader context within '
        'which their organisations operate and to anticipate the ripple effects of strategic decisions across '
        'multiple dimensions [32]. It guards against the common strategic pitfall of "fixing" one problem only '
        'to create or exacerbate others — a phenomenon known as "policy resistance" in systems dynamics literature.'
    )
    
    doc.add_paragraph(
        'In practice, systems thinking helps organisations identify root causes rather than symptoms, design '
        'strategies that account for delayed effects and unintended consequences, and recognise when short-term '
        'gains come at the expense of long-term systemic health [33]. For example, a systems perspective might '
        'reveal that aggressive cost-cutting, while improving short-term profitability, undermines employee morale, '
        'reduces innovation capacity, and ultimately erodes competitive position — a reinforcing loop that can lead '
        'to organisational decline.'
    )
    
    doc.add_paragraph(
        'The strategic value of systems thinking is particularly evident in contexts characterised by high complexity '
        'and stakeholder interdependence, such as sustainability initiatives, digital ecosystem strategies, and '
        'organisational transformation programmes [34]. In these contexts, the ability to see the "whole board" '
        'rather than isolated pieces provides a crucial strategic advantage, as illustrated in Figure 2.'
    )
    
    # Insert Figure 2
    doc.add_paragraph()
    if os.path.exists('dt_chapter_figures/Figure_2_Systems_Thinking.png'):
        doc.add_picture('dt_chapter_figures/Figure_2_Systems_Thinking.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig2_caption = doc.add_paragraph()
    fig2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig2_caption.add_run('Figure 2: ')
    run.bold = True
    run.font.size = Pt(10)
    run = fig2_caption.add_run('Systems Thinking in Strategic Context — Feedback Loops and Leverage Points '
                               'in Organisational Decision-Making')
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()
    
    # 2.3
    doc.add_heading('2.3 Creative Thinking', level=2)
    doc.add_heading('2.3.1 Theory and Principles', level=3)
    
    doc.add_paragraph(
        'Creative thinking as a field of study has its origins in Guilford\'s seminal 1950 address to the American '
        'Psychological Association, which called for systematic research into the nature of creative cognitive '
        'processes [35]. Since then, a rich body of research has illuminated the cognitive, social, and environmental '
        'factors that enable or inhibit creative ideation in organisational contexts.'
    )
    
    doc.add_paragraph(
        'The fundamental distinction between divergent thinking (the generation of multiple novel ideas) and '
        'convergent thinking (the evaluation and selection of the most promising ideas) provides the conceptual '
        'backbone of creative thinking methodology [10]. Effective creative processes require both modes — the '
        'uninhibited generation of possibilities followed by critical evaluation against relevant criteria — but '
        'these modes require different cognitive stances and environmental conditions.'
    )
    
    doc.add_paragraph(
        'Additional theoretical concepts include lateral thinking, introduced by Edward de Bono, which involves '
        'approaching problems from unconventional angles rather than following established logical sequences [36]; '
        'associative thinking, which generates novel ideas through unexpected combinations of previously unrelated '
        'concepts; and cognitive flexibility, the ability to shift between different mental frameworks and '
        'perspectives. Research on creative cognition has also highlighted the importance of incubation — the '
        'unconscious processing that occurs when attention is diverted from a problem — and the role of constraints '
        'in stimulating creative solutions [37].'
    )
    
    doc.add_heading('2.3.2 Key Methodologies', level=3)
    
    doc.add_paragraph(
        'The methodological landscape of creative thinking encompasses a diverse array of structured techniques '
        'designed to overcome cognitive fixedness and stimulate novel ideation. Classic brainstorming, as originally '
        'formulated by Alex Osborn, remains widely used despite ongoing debate about its effectiveness relative to '
        'individual ideation [38]. Contemporary variations include brainwriting, reverse brainstorming, and '
        'electronic brainstorming, each addressing specific limitations of the traditional approach.'
    )
    
    doc.add_paragraph(
        'Mind mapping provides a visual technique for exploring the associative connections between ideas, enabling '
        'non-linear exploration of problem spaces [39]. The SCAMPER technique (Substitute, Combine, Adapt, Modify, '
        'Put to other uses, Eliminate, Reverse) offers a structured framework for systematically transforming '
        'existing concepts into novel variations. Other notable methods include analogical reasoning, which draws '
        'insights from distant domains; morphological analysis, which systematically explores combinations of '
        'problem dimensions; and provocation techniques, which use deliberately unrealistic statements to trigger '
        'new thinking directions.'
    )
    
    doc.add_paragraph(
        'More recently, computational creativity tools powered by artificial intelligence have expanded the creative '
        'toolkit, enabling the generation of novel designs, texts, and solutions through generative adversarial '
        'networks, evolutionary algorithms, and large language models [40]. These tools represent a new frontier in '
        'creative thinking methodology, augmenting human creative capacity rather than replacing it.'
    )
    
    doc.add_heading('2.3.3 Role in Strategy', level=3)
    
    doc.add_paragraph(
        'Creative thinking serves as the generative engine of strategic innovation, enabling organisations to '
        'envision possibilities beyond the constraints of current assumptions and established practices [41]. In '
        'strategic contexts, creative thinking contributes at multiple levels: at the product level, it drives the '
        'development of novel offerings; at the business model level, it enables the reconceptualisation of value '
        'creation and capture mechanisms; and at the industry level, it can catalyse the emergence of entirely new '
        'market categories.'
    )
    
    doc.add_paragraph(
        'The strategic importance of creative thinking is amplified in hyper-competitive environments where '
        'sustainable advantage increasingly derives from the ability to innovate continuously rather than to '
        'exploit existing positions [42]. Research on ambidextrous organisations — those capable of simultaneously '
        'exploiting current capabilities while exploring new possibilities — highlights the critical role of '
        'creative thinking in maintaining the exploratory capacity essential for long-term organisational vitality.'
    )
    
    doc.add_paragraph(
        'The relationship between creative thinking and organisational culture deserves particular attention in '
        'strategic contexts. Organisations that systematically cultivate creative capabilities — through diverse '
        'hiring practices, cross-functional collaboration, dedicated innovation time, and tolerance of productive '
        'failure — develop what researchers term "creative capital" that compounds over time. This accumulated '
        'creative capacity becomes a strategic asset that enables rapid response to market disruptions and the '
        'identification of opportunities invisible to less creatively capable competitors. However, building and '
        'maintaining this creative capital requires sustained investment and leadership commitment, as creative '
        'capabilities are easily eroded by short-term efficiency pressures or risk-averse management practices.'
    )
    
    doc.add_paragraph(
        'However, creativity in isolation does not constitute strategy. Unbounded ideation without the discipline '
        'of analytical validation, systemic awareness, or user-centred grounding can produce solutions that are '
        'novel but impractical, exciting but unsustainable, or innovative but irrelevant to actual human needs. '
        'It is this recognition that motivates the integrative framework presented in the following section.'
    )
    
    doc.add_page_break()
    
    # ==================== SECTION 3 ====================
    doc.add_heading('3. The Unifying Power of Design Thinking', level=1)
    
    # 3.1
    doc.add_heading('3.1 The Human-Centred DT Process', level=2)
    
    doc.add_paragraph(
        'Design Thinking, as a formalised methodology for innovation and problem-solving, has its roots in the '
        'design profession\'s long-standing practices of user research, iterative prototyping, and solution '
        'refinement [11]. However, its application has expanded far beyond traditional design contexts to encompass '
        'business strategy, organisational development, public policy, healthcare, and education [43]. The '
        'methodology\'s appeal lies in its ability to make the tacit practices of expert designers accessible and '
        'applicable to non-designers facing complex challenges.'
    )
    
    doc.add_paragraph(
        'The five-stage DT model — Empathize, Define, Ideate, Prototype, and Test — provides a structured yet '
        'flexible framework for navigating from problem understanding to solution implementation [6]. Each stage '
        'serves a distinct purpose while maintaining iterative connections to the others:'
    )
    
    doc.add_paragraph(
        'Empathize involves deep immersion in the experiences, perspectives, and needs of the people for whom '
        'solutions are being designed. This stage employs ethnographic methods, contextual inquiry, stakeholder '
        'interviews, and observational research to develop a rich, nuanced understanding of human needs that goes '
        'beyond surface-level preferences to uncover latent needs and unarticulated desires.'
    )
    
    doc.add_paragraph(
        'Define synthesises the insights gathered during empathy research into a clear, actionable problem '
        'statement — often expressed as a "How Might We" question that frames the challenge in a way that invites '
        'creative solutions while maintaining focus on genuine human needs.'
    )
    
    doc.add_paragraph(
        'Ideate is the generative phase where diverse teams produce a wide range of potential solutions without '
        'premature judgment. This stage deliberately leverages creative thinking techniques to maximise the breadth '
        'and novelty of ideas generated.'
    )
    
    doc.add_paragraph(
        'Prototype translates selected ideas into tangible, low-fidelity representations that can be experienced '
        'and evaluated. The emphasis is on rapid, inexpensive creation of "thinking tools" that externalise ideas '
        'and enable learning through making.'
    )
    
    doc.add_paragraph(
        'Test involves presenting prototypes to users and stakeholders, gathering feedback, and iterating based on '
        'what is learned. This stage closes the loop between solution development and user needs, ensuring that '
        'final solutions genuinely address the problems identified during the Empathize phase.'
    )
    
    doc.add_paragraph(
        'The iterative nature of this process is crucial — DT practitioners expect to cycle through these stages '
        'multiple times, with each iteration refining understanding and improving solutions. This tolerance for '
        'ambiguity and willingness to "fail forward" through rapid experimentation represents a fundamental '
        'departure from traditional linear planning approaches [4].'
    )
    
    doc.add_paragraph(
        'Importantly, the DT process is not merely a sequence of steps but an embodiment of a particular '
        'philosophical orientation toward problem-solving. It privileges action over deliberation, experimentation '
        'over prediction, collaboration over hierarchy, and learning over certainty. These values are not merely '
        'methodological preferences but reflect a fundamental reconceptualisation of how innovation happens in '
        'complex, uncertain environments. The DT practitioner accepts that perfect information is unavailable, '
        'that user needs are often tacit and evolving, and that the best path to a successful solution is not '
        'linear reasoning but iterative exploration. This orientation makes DT particularly well-suited to the '
        'VUCA environments that characterise contemporary business, where the ability to learn and adapt rapidly '
        'is often more valuable than the ability to plan comprehensively.'
    )
    
    # 3.2
    doc.add_heading('3.2 DT as a Synergistic Integrator', level=2)
    
    doc.add_paragraph(
        'The central thesis of this chapter is that Design Thinking functions not merely as an alternative to '
        'analytical, systemic, and creative approaches but as a meta-framework capable of integrating their '
        'respective strengths into a cohesive strategic methodology. This integrative capacity arises from DT\'s '
        'fundamental characteristics: its iterative structure naturally accommodates different thinking modes at '
        'different stages; its human-centred focus provides a unifying criterion for evaluating diverse inputs; '
        'and its emphasis on prototyping and testing creates a pragmatic bridge between abstract thinking and '
        'concrete action.'
    )
    
    doc.add_heading('3.2.1 DT + Analytical Decision-Making: Informed Empathy', level=3)
    
    doc.add_paragraph(
        'The integration of analytical approaches with Design Thinking creates what might be termed "informed '
        'empathy" — a combination of quantitative rigour and qualitative depth that produces richer, more reliable '
        'insights than either approach alone [5]. Data analytics can identify patterns in user behaviour at scale, '
        'revealing needs and opportunities that qualitative research alone might miss. Simultaneously, empathic '
        'research provides the contextual understanding necessary to interpret data meaningfully and avoid the '
        'trap of mistaking correlation for causation.'
    )
    
    doc.add_paragraph(
        'In the Prototype and Test stages, analytical methods provide essential validation mechanisms. A/B testing, '
        'statistical significance analysis, and quantitative user metrics enable rigorous evaluation of solution '
        'effectiveness [22]. The analytical approach transforms the inherently subjective process of design '
        'evaluation into a more systematic and reliable assessment, while DT\'s emphasis on human outcomes ensures '
        'that the metrics chosen reflect genuine value creation rather than arbitrary quantification.'
    )
    
    doc.add_paragraph(
        'This synthesis is particularly powerful in contexts such as digital product development, where large-scale '
        'behavioural data can be combined with deep user research to create solutions that are both empirically '
        'validated and emotionally resonant. Companies like Netflix, Spotify, and Airbnb exemplify this integration, '
        'using sophisticated data analytics to inform and validate design decisions while maintaining a relentless '
        'focus on user experience [16].'
    )
    
    doc.add_heading('3.2.2 DT + Systems Thinking: Sustainable Solutions', level=3)
    
    doc.add_paragraph(
        'The integration of systems thinking with Design Thinking addresses one of the most common criticisms of '
        'DT — its potential to generate solutions that work at the individual user level but fail to account for '
        'broader organisational, environmental, or social contexts [13]. By incorporating systemic awareness into '
        'the DT process, practitioners can ensure that human-centred solutions are also systemically viable and '
        'sustainable.'
    )
    
    doc.add_paragraph(
        'This integration is particularly evident in the Define stage, where systems thinking helps frame problems '
        'within their broader context. Rather than defining a challenge narrowly in terms of individual user needs, '
        'a systems-informed DT approach considers how the problem is embedded within larger organisational processes, '
        'stakeholder networks, and environmental dynamics [27]. This broader framing leads to solutions that address '
        'root causes rather than symptoms and that anticipate potential unintended consequences.'
    )
    
    doc.add_paragraph(
        'Figure 3 illustrates how the integration of systems thinking and Design Thinking creates a multi-level '
        'framework for strategic innovation, connecting individual user needs with organisational capabilities and '
        'broader ecosystem dynamics.'
    )
    
    # Insert Figure 3
    doc.add_paragraph()
    if os.path.exists('dt_chapter_figures/Figure_3_MultiLevel_Integration.png'):
        doc.add_picture('dt_chapter_figures/Figure_3_MultiLevel_Integration.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig3_caption = doc.add_paragraph()
    fig3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig3_caption.add_run('Figure 3: ')
    run.bold = True
    run.font.size = Pt(10)
    run = fig3_caption.add_run('The Multi-Level Integration Framework — Connecting User Needs, Organisational '
                               'Systems, and Ecosystem Dynamics through Design Thinking')
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()
    
    doc.add_paragraph(
        'In the Prototype and Test stages, systems thinking provides tools for evaluating whether proposed solutions '
        'will function effectively within the larger system. Causal loop diagrams can be used to model how a new '
        'product, service, or process will interact with existing organisational elements, identify potential '
        'resistance or unintended consequences, and design implementation strategies that account for systemic '
        'dynamics [29].'
    )
    
    doc.add_heading('3.2.3 DT + Creative Thinking: Grounded Innovation', level=3)
    
    doc.add_paragraph(
        'The relationship between Design Thinking and creative thinking is perhaps the most intuitive of the three '
        'integrations, as the Ideate stage of DT explicitly draws on creative thinking methodologies [38]. However, '
        'the integration is more nuanced than simple incorporation. DT provides creative thinking with something it '
        'often lacks when used in isolation: a clear set of human-centred constraints that channel creative energy '
        'toward genuinely valuable solutions rather than merely novel ones.'
    )
    
    doc.add_paragraph(
        'The DT framework grounds creativity in two critical ways. First, the insights generated during the '
        'Empathize and Define stages provide a rich understanding of user needs that serves as both inspiration '
        'and constraint for creative ideation [35]. Second, the Prototype and Test stages provide rapid feedback '
        'mechanisms that help creative teams distinguish between ideas that are merely novel and those that '
        'represent genuine innovation — defined as novelty combined with value.'
    )
    
    doc.add_paragraph(
        'This grounding effect addresses the common organisational frustration with creative processes that generate '
        'exciting ideas but fail to deliver practical results. By embedding creative thinking within the larger DT '
        'framework, organisations can harness the generative power of divergent thinking while maintaining the '
        'discipline necessary to translate creative possibilities into implemented innovations [41].'
    )
    
    doc.add_paragraph(
        'Table 3 presents a detailed mapping of how Design Thinking integrates with each complementary approach '
        'across the five stages of the DT process.'
    )
    
    # ==================== TABLE 3 ====================
    doc.add_paragraph()
    table3_title = doc.add_paragraph()
    table3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table3_title.add_run('Table 3: ')
    run.bold = True
    run.font.size = Pt(10)
    run = table3_title.add_run('Integration Matrix — How Design Thinking Synthesises Analytical, Systems, and '
                               'Creative Approaches Across the Five DT Stages')
    run.font.size = Pt(10)
    run.italic = True
    
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers3 = ['DT Stage', 'Analytical\nContribution', 'Systems\nContribution', 'Creative\nContribution']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1565C0')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    table3_data = [
        ['Empathize', 'Quantitative user data,\nbehavioural analytics', 'Stakeholder mapping,\ncontext analysis', 'Open observation,\nassumption challenging'],
        ['Define', 'Data-driven problem\nscoping, metrics', 'Root cause analysis,\nsystemic framing', 'Reframing problems,\nnew perspectives'],
        ['Ideate', 'Evidence-based\nconstraints', 'Ecosystem-aware\nsolution space', 'Divergent thinking,\nnovel combinations'],
        ['Prototype', 'Testable hypotheses,\nmeasurement criteria', 'System interaction\nmodelling', 'Rapid visualisation,\nexpressive forms'],
        ['Test', 'Statistical validation,\nA/B testing', 'Impact assessment,\nfeedback analysis', 'Iterative refinement,\nalternative exploration'],
    ]
    
    for row_idx, row_data in enumerate(table3_data, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table3.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if col_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
        if row_idx % 2 == 0:
            for col_idx in range(4):
                set_cell_shading(table3.rows[row_idx].cells[col_idx], 'E3F2FD')
    
    doc.add_paragraph()
    
    # 3.3
    doc.add_heading('3.3 The Strategic Outcomes of DT', level=2)
    
    doc.add_paragraph(
        'The strategic outcomes of implementing an integrated DT approach extend across multiple dimensions of '
        'organisational performance. Research consistently demonstrates positive associations between DT adoption '
        'and innovation outcomes, including increased patent activity, accelerated time-to-market, higher new '
        'product success rates, and greater customer satisfaction [5, 6].'
    )
    
    doc.add_paragraph(
        'Enhanced Customer Loyalty and Experience: By grounding strategic decisions in deep empathic understanding '
        'of customer needs, organisations create offerings that resonate at both functional and emotional levels. '
        'This emotional connection translates into stronger brand loyalty, higher customer lifetime value, and more '
        'positive word-of-mouth [43]. The iterative testing inherent in DT ensures that solutions are continuously '
        'refined based on actual user feedback rather than internal assumptions.'
    )
    
    doc.add_paragraph(
        'The depth of customer understanding enabled by DT-integrated approaches extends beyond traditional market '
        'research to encompass the full spectrum of human experience — emotional, social, cultural, and contextual '
        'dimensions that quantitative surveys alone cannot capture. This rich understanding enables organisations '
        'to create not merely functional solutions but meaningful experiences that build deep emotional bonds with '
        'customers. In an era where products are increasingly commoditised, this experiential differentiation '
        'represents one of the most powerful and sustainable sources of competitive advantage available.'
    )
    
    doc.add_paragraph(
        'Increased Organisational Agility: The iterative, experimental nature of DT cultivates organisational '
        'capabilities — such as comfort with ambiguity, rapid prototyping, and learning from failure — that are '
        'essential for agility in volatile environments [4]. Organisations that internalise DT principles develop '
        'a "muscle" for rapid adaptation that serves them well beyond individual innovation projects.'
    )
    
    doc.add_paragraph(
        'Reduced Risk Through Rapid Prototyping: Traditional approaches to strategy often involve large upfront '
        'investments in untested assumptions. DT\'s emphasis on low-fidelity prototyping and early user testing '
        'enables organisations to validate strategic hypotheses at minimal cost, reducing the risk of expensive '
        'failures and enabling more efficient resource allocation [11].'
    )
    
    doc.add_paragraph(
        'Sustainable Competitive Advantage Through Innovation: Perhaps most significantly, the integrated DT '
        'approach creates a renewable source of competitive advantage. While individual products or services can '
        'be imitated, the organisational capability for continuous, human-centred innovation is far more difficult '
        'for competitors to replicate [42]. This capability advantage is particularly valuable in knowledge-intensive '
        'industries where the pace of change renders product-based advantages increasingly transient.'
    )
    
    doc.add_page_break()
    
    # ==================== SECTION 4 ====================
    doc.add_heading('4. Practice, Challenges, and Future Opportunities', level=1)
    
    # 4.1
    doc.add_heading('4.1 Applications and Industry Examples', level=2)
    
    doc.add_paragraph(
        'The practical application of the integrated DT framework spans diverse industries and organisational '
        'contexts, demonstrating its versatility and adaptability as a strategic methodology.'
    )
    
    doc.add_paragraph(
        'Product Development: In the consumer technology sector, companies such as Apple, Samsung, and Dyson have '
        'demonstrated the power of combining deep user empathy with analytical market research, systemic '
        'understanding of technology ecosystems, and creative design excellence [6]. Apple\'s product development '
        'process exemplifies the integration thesis, combining ethnographic user research, rigorous engineering '
        'analysis, awareness of ecosystem dynamics, and creative industrial design into a cohesive innovation '
        'methodology that has produced some of the most successful products in commercial history.'
    )
    
    doc.add_paragraph(
        'Digital Transformation: Organisations undertaking digital transformation initiatives have increasingly '
        'adopted DT as a guiding methodology for navigating the complex interplay of technology, process, and '
        'human factors [23]. IBM\'s Enterprise Design Thinking framework, for example, explicitly integrates user '
        'research, systemic process analysis, and creative ideation to guide the design of digital products and '
        'services at enterprise scale. The framework\'s emphasis on "Hills" (desired user outcomes), "Sponsor Users" '
        '(real people whose needs guide design), and "Playbacks" (iterative feedback sessions) demonstrates the '
        'practical application of DT principles in complex organisational contexts.'
    )
    
    doc.add_paragraph(
        'Healthcare: The healthcare sector presents particularly compelling examples of integrated DT application, '
        'given its combination of complex systemic challenges, critical human needs, and strict regulatory '
        'requirements [7]. The Mayo Clinic\'s Center for Innovation has pioneered the application of DT to '
        'healthcare delivery, using patient journey mapping, rapid prototyping of care processes, and systemic '
        'analysis of clinical workflows to develop innovations that improve both patient experience and clinical '
        'outcomes.'
    )
    
    doc.add_paragraph(
        'Entrepreneurship: The lean startup methodology, which has become the dominant framework for technology '
        'entrepreneurship, shares significant conceptual overlap with DT [4]. Both emphasise iterative development, '
        'early user feedback, and willingness to pivot based on learning. Successful startups such as Airbnb, which '
        'famously used DT techniques including immersive user research and rapid prototyping to transform its '
        'business model, demonstrate the strategic power of human-centred approaches in entrepreneurial contexts.'
    )
    
    doc.add_paragraph(
        'The entrepreneurial application of integrated DT is particularly instructive because it demonstrates the '
        'framework\'s value under conditions of extreme uncertainty and resource constraint. Entrepreneurs cannot '
        'afford the luxury of extensive analytical studies or comprehensive systemic modelling before acting; they '
        'must learn by doing, iterating rapidly based on market feedback. DT provides a structured yet lightweight '
        'framework for this rapid learning, enabling entrepreneurs to validate assumptions about customer needs '
        'through low-cost experiments rather than expensive product launches. The integration of analytical thinking '
        '(in the form of metrics-driven validation), systems thinking (in the form of business model canvas '
        'analysis), and creative thinking (in the form of rapid ideation and pivoting) within the DT framework '
        'creates a comprehensive yet agile approach to entrepreneurial strategy that has proven remarkably effective '
        'in practice.'
    )
    
    doc.add_paragraph(
        'Manufacturing: In manufacturing contexts, the integration of DT with systems thinking and analytical '
        'methods has driven advances in areas such as lean production, smart factory design, and sustainable '
        'manufacturing [34]. Toyota\'s production system, often cited as a pinnacle of manufacturing excellence, '
        'demonstrates many of the integrative principles discussed in this chapter — combining rigorous analytical '
        'methods (statistical process control), systemic thinking (value stream mapping), creative problem-solving '
        '(kaizen), and human-centred design (respect for people) into a cohesive operational philosophy.'
    )
    
    doc.add_paragraph(
        'Figure 4 provides a comparative visualisation of DT application outcomes across these five industry '
        'sectors, highlighting the common patterns of improvement in innovation metrics, customer satisfaction, '
        'and operational efficiency.'
    )
    
    # Insert Figure 4
    doc.add_paragraph()
    if os.path.exists('dt_chapter_figures/Figure_4_Comparative_Outcomes.png'):
        doc.add_picture('dt_chapter_figures/Figure_4_Comparative_Outcomes.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig4_caption = doc.add_paragraph()
    fig4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig4_caption.add_run('Figure 4: ')
    run.bold = True
    run.font.size = Pt(10)
    run = fig4_caption.add_run('Comparative Outcomes of Integrated Design Thinking Application Across Five '
                               'Industry Sectors')
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()
    
    # 4.2
    doc.add_heading('4.2 Key Challenges and Barriers', level=2)
    
    doc.add_paragraph(
        'Despite the compelling case for integrated DT approaches, significant challenges remain in their effective '
        'implementation within established organisations.'
    )
    
    doc.add_heading('4.2.1 Cultural and Leadership Challenges', level=3)
    
    doc.add_paragraph(
        'The adoption of DT requires fundamental shifts in organisational culture that often encounter significant '
        'resistance. Traditional corporate cultures that prize certainty, efficiency, and predictability may find '
        'the ambiguity, iteration, and experimentation inherent in DT uncomfortable or even threatening [15]. '
        'Leaders accustomed to "command and control" management styles must develop new capabilities in facilitation, '
        'empowerment, and comfort with uncertainty.'
    )
    
    doc.add_paragraph(
        'Psychological safety — the shared belief that team members will not be punished for taking interpersonal '
        'risks — has been identified as a critical enabler of DT implementation [12]. Without psychological safety, '
        'the vulnerability required for genuine empathy, the risk-taking necessary for creative ideation, and the '
        'honesty essential for constructive testing feedback are all inhibited. Creating and maintaining psychological '
        'safety requires sustained leadership commitment and often represents a multi-year cultural transformation '
        'rather than a quick programmatic fix.'
    )
    
    doc.add_paragraph(
        'The "innovation theatre" phenomenon, where organisations adopt the superficial trappings of DT (sticky '
        'notes, workshops, design sprints) without the deeper cultural and structural changes necessary for genuine '
        'impact, represents a significant barrier to realising DT\'s strategic potential [14]. Effective DT '
        'implementation requires not just methodological training but fundamental shifts in how organisations '
        'define success, allocate resources, tolerate failure, and reward learning.'
    )
    
    doc.add_heading('4.2.2 Process and Practice Challenges', level=3)
    
    doc.add_paragraph(
        'At the operational level, several challenges complicate the integration of DT with existing organisational '
        'processes. The tension between the open-ended, exploratory nature of DT and the structured, milestone-driven '
        'requirements of traditional project management creates friction in many organisational contexts [25]. '
        'Finding the appropriate balance between creative freedom and operational discipline remains an ongoing '
        'challenge for DT practitioners.'
    )
    
    doc.add_paragraph(
        'Scalability represents another significant concern. While DT has demonstrated effectiveness at the project '
        'and team levels, scaling its principles across large, complex organisations presents substantial '
        'challenges [43]. Questions of governance, resource allocation, cross-functional coordination, and quality '
        'assurance become increasingly complex as DT initiatives expand beyond pilot projects to enterprise-wide '
        'transformation programmes.'
    )
    
    doc.add_paragraph(
        'The integration challenge itself — combining analytical, systemic, creative, and design thinking in '
        'practice — requires individuals and teams with broad cognitive repertoires and the ability to shift '
        'fluidly between different thinking modes [32]. Such "T-shaped" professionals — combining deep expertise '
        'in one domain with broad capability across multiple domains — remain relatively rare, creating talent '
        'constraints for organisations seeking to implement fully integrated approaches.'
    )
    
    doc.add_paragraph(
        'Table 4 summarises the key challenges and barriers to integrated DT implementation, categorised by domain, '
        'along with recommended mitigation strategies drawn from successful implementation cases.'
    )
    
    # ==================== TABLE 4 ====================
    doc.add_paragraph()
    table4_title = doc.add_paragraph()
    table4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table4_title.add_run('Table 4: ')
    run.bold = True
    run.font.size = Pt(10)
    run = table4_title.add_run('Key Challenges in Integrated Design Thinking Implementation and Mitigation Strategies')
    run.font.size = Pt(10)
    run.italic = True
    
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers4 = ['Challenge Domain', 'Specific Barrier', 'Impact Level', 'Mitigation Strategy']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'E65100')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    table4_data = [
        ['Cultural', 'Resistance to ambiguity\nand experimentation', 'High', 'Leadership modelling, safe-to-fail\nexperiments, incremental adoption'],
        ['Leadership', 'Command-and-control\nmindset', 'High', 'Executive coaching, DT immersion\nprogrammes, visible sponsorship'],
        ['Process', 'Tension with existing\nproject management', 'Medium', 'Hybrid governance frameworks,\nadaptive milestone structures'],
        ['Scalability', 'Difficulty expanding\nbeyond pilot projects', 'Medium-High', 'Centres of excellence, DT\nchampion networks, toolkits'],
        ['Talent', 'Shortage of T-shaped\nprofessionals', 'Medium', 'Cross-functional rotations,\ninterdisciplinary education'],
        ['Measurement', 'Difficulty quantifying\nDT impact', 'Medium', 'Leading indicators, innovation\naccounting, portfolio metrics'],
    ]
    
    for row_idx, row_data in enumerate(table4_data, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table4.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if col_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
        if row_idx % 2 == 0:
            for col_idx in range(4):
                set_cell_shading(table4.rows[row_idx].cells[col_idx], 'FFF3E0')
    
    doc.add_paragraph()
    
    # 4.3
    doc.add_heading('4.3 Future Directions and Hybrid Frameworks', level=2)
    doc.add_heading('4.3.1 The Role of Technology', level=3)
    
    doc.add_paragraph(
        'The accelerating development of digital technologies, particularly artificial intelligence, is creating new '
        'possibilities for strengthening and extending integrated innovation frameworks [40]. AI-powered analytics '
        'can process vast quantities of user data to identify needs and patterns at scales impossible for human '
        'researchers alone, augmenting the Empathize stage of DT with quantitative breadth while human researchers '
        'provide qualitative depth.'
    )
    
    doc.add_paragraph(
        'Generative AI tools are beginning to augment the Ideate stage by producing diverse concept variations '
        'based on specified constraints, enabling design teams to explore solution spaces more rapidly and '
        'comprehensively than manual ideation alone allows [22]. Virtual and augmented reality technologies are '
        'transforming the Prototype and Test stages by enabling the creation of immersive, interactive prototypes '
        'that can be experienced and evaluated by users before physical production.'
    )
    
    doc.add_paragraph(
        'Digital twin technology, which creates virtual replicas of physical systems, represents a particularly '
        'promising integration point between systems thinking and DT [30]. Digital twins enable the simulation of '
        'how proposed design solutions will interact with complex organisational or physical systems, providing '
        'systemic validation before implementation and enabling rapid iteration on system-level design decisions.'
    )
    
    doc.add_paragraph(
        'Furthermore, the emergence of collaborative platforms that integrate user research tools, data analytics, '
        'systems modelling, creative ideation support, and prototyping capabilities into unified digital environments '
        'is reducing the practical barriers to integrated approaches [23]. These platforms enable cross-functional '
        'teams to move fluidly between analytical, systemic, creative, and design thinking modes within a shared '
        'digital workspace, supporting the kind of integrated practice that this chapter advocates.'
    )
    
    doc.add_paragraph(
        'The convergence of these technologies with integrated innovation methodologies represents more than '
        'incremental improvement — it has the potential to fundamentally transform how organisations approach '
        'strategic innovation. Real-time data streams can feed directly into empathy research, enabling continuous '
        'monitoring of user needs and behaviours. Machine learning algorithms can identify patterns across vast '
        'solution spaces, augmenting human creative capacity. Network simulation tools can model the systemic '
        'implications of proposed interventions before they are implemented. And collaborative digital environments '
        'can enable geographically distributed teams to engage in rich, multi-modal innovation processes that were '
        'previously possible only through co-located interaction. Together, these technological developments promise '
        'to make integrated innovation more accessible, more rigorous, and more effective than ever before.'
    )
    
    doc.add_heading('4.3.2 Research Agenda', level=3)
    
    doc.add_paragraph(
        'The academic and practitioner communities face several important research challenges in advancing the '
        'theory and practice of integrated strategic innovation. First, there is a need for more rigorous empirical '
        'research on the effectiveness of integrated approaches relative to single-methodology implementations [19]. '
        'While case studies and design-led performance indices provide suggestive evidence, controlled studies that '
        'isolate the contribution of integration are rare and methodologically challenging.'
    )
    
    doc.add_paragraph(
        'Second, the development of formal frameworks for "strategic integration intelligence" — the organisational '
        'capability to determine which combination of approaches is most appropriate for a given strategic '
        'challenge — represents a significant research opportunity [8]. Such frameworks would help organisations '
        'move beyond generic prescriptions to context-sensitive application of integrated methods.'
    )
    
    doc.add_paragraph(
        'Third, research on the cognitive and team-level processes that enable effective integration is needed to '
        'inform education and professional development programmes [37]. Understanding how individuals develop the '
        'cognitive flexibility to move between analytical, systemic, creative, and empathic thinking modes — and '
        'how teams coordinate these different modes effectively — would provide crucial foundations for scaling '
        'integrated approaches.'
    )
    
    doc.add_paragraph(
        'Fourth, the intersection of AI and integrated innovation methodologies presents a rich research '
        'frontier [40]. Questions include how AI tools can best augment human capabilities at each stage of the '
        'integrated process, what risks are associated with AI-mediated design decisions, and how the balance '
        'between human judgment and algorithmic recommendation should be calibrated in different contexts.'
    )
    
    doc.add_paragraph(
        'Finally, as the reference to Figure 1 earlier indicated, the theoretical model positioning DT as a '
        'central integrating framework requires ongoing refinement and validation across diverse organisational '
        'and cultural contexts. The framework presented in this chapter draws primarily on Western organisational '
        'contexts, and further research is needed to understand how cultural factors influence the effectiveness '
        'and appropriate adaptation of integrated approaches in different global settings [3].'
    )
    
    doc.add_page_break()
    
    # ==================== SECTION 5: CONCLUSION ====================
    doc.add_heading('5. Conclusion', level=1)
    
    doc.add_paragraph(
        'This chapter has presented a comprehensive comparative analysis of four foundational approaches to '
        'strategic decision-making — Analytical Decision-Making, Systems Thinking, Creative Thinking, and Design '
        'Thinking — and argued for their thoughtful integration under the unifying umbrella of Design Thinking. '
        'The analysis has demonstrated that while each approach offers distinct and valuable contributions to '
        'strategic capability, their true power emerges through synergistic combination.'
    )
    
    doc.add_paragraph(
        'Design Thinking\'s unique characteristics — its human-centred orientation, iterative structure, emphasis '
        'on tangible prototyping, and inherent tolerance for ambiguity — position it as an effective meta-framework '
        'for integration. When combined with the empirical rigour of analytical approaches, the holistic perspective '
        'of systems thinking, and the generative power of creative thinking, DT enables organisations to develop '
        'strategies that are simultaneously desirable (meeting genuine human needs), viable (sustainable within '
        'systemic constraints), and feasible (technically and operationally achievable).'
    )
    
    doc.add_paragraph(
        'The tripartite criterion of desirability, viability, and feasibility — long recognised in design discourse '
        '— provides a particularly useful lens for understanding the integrative value of DT. Desirability is '
        'ensured through the empathic, human-centred orientation of DT itself. Viability is supported through the '
        'integration of systems thinking, which ensures that solutions are sustainable within broader organisational '
        'and environmental contexts. Feasibility is validated through the integration of analytical approaches, '
        'which provide rigorous assessment of technical and operational constraints. And the creative thinking '
        'dimension ensures that the solution space explored is sufficiently broad and novel to yield genuinely '
        'innovative outcomes rather than incremental improvements to existing offerings. This comprehensive coverage '
        'of the innovation value space represents one of the most compelling arguments for the integrated framework '
        'advocated in this chapter.'
    )
    
    doc.add_paragraph(
        'The practical applications reviewed across product development, digital transformation, healthcare, '
        'entrepreneurship, and manufacturing demonstrate both the versatility and the effectiveness of integrated '
        'approaches. However, significant challenges remain, particularly in the areas of cultural transformation, '
        'scalability, and talent development. As referenced in Figure 4, the outcomes across industries consistently '
        'show positive returns when integration is effectively achieved.'
    )
    
    doc.add_paragraph(
        'Looking forward, the convergence of advanced digital technologies, artificial intelligence, and evolving '
        'organisational models creates unprecedented opportunities for strengthening and extending integrated '
        'innovation frameworks. The organisations that will thrive in an increasingly complex and uncertain future '
        'will be those that develop robust capabilities across all four approaches and, crucially, the integration '
        'intelligence to combine them effectively in response to specific strategic challenges.'
    )
    
    doc.add_paragraph(
        'The vision this chapter advances is ultimately one of strategic pluralism — not the dominance of any '
        'single approach, but the mature, context-sensitive integration of multiple perspectives in service of '
        'creating value for people, organisations, and society. Design Thinking, with its deeply humanistic '
        'orientation and inherent integrative capacity, provides both the philosophical foundation and the '
        'practical methodology for realising this vision.'
    )
    
    doc.add_paragraph(
        'As organisations continue to grapple with increasing complexity, accelerating change, and rising '
        'stakeholder expectations, the ability to draw upon multiple strategic thinking approaches — and to '
        'integrate them effectively — will become an increasingly critical organisational capability. The framework '
        'presented in this chapter provides a conceptual and practical foundation for developing this capability, '
        'offering both the theoretical justification for integration and the methodological guidance for its '
        'implementation. The challenge for practitioners and scholars alike is to continue refining, testing, and '
        'extending this framework in response to the ever-evolving demands of the business environment, ensuring '
        'that our strategic thinking keeps pace with the complexity of the challenges we seek to address.'
    )
    
    doc.add_page_break()
    
    # ==================== REFERENCES ====================
    doc.add_heading('References', level=1)
    
    references = [
        '[1] Ansoff, H.I. (1965). Corporate Strategy: An Analytical Approach to Business Policy for Growth and Expansion. McGraw-Hill, New York.',
        '[2] Porter, M.E. (1985). Competitive Advantage: Creating and Sustaining Superior Performance. Free Press, New York.',
        '[3] Bennett, N. and Lemoine, G.J. (2014). What a difference a word makes: Understanding threats to performance in a VUCA world. Business Horizons, 57(3), pp.311-317.',
        '[4] Ries, E. (2011). The Lean Startup: How Today\'s Entrepreneurs Use Continuous Innovation to Create Radically Successful Businesses. Crown Business, New York.',
        '[5] Liedtka, J. (2018). Why Design Thinking Works. Harvard Business Review, 96(5), pp.72-79.',
        '[6] Brown, T. (2009). Change by Design: How Design Thinking Transforms Organizations and Inspires Innovation. HarperBusiness, New York.',
        '[7] Buchanan, R. (1992). Wicked Problems in Design Thinking. Design Issues, 8(2), pp.5-21.',
        '[8] Davenport, T.H. (2009). Make Better Decisions. Harvard Business Review, 87(11), pp.117-123.',
        '[9] Senge, P.M. (1990). The Fifth Discipline: The Art and Practice of the Learning Organization. Doubleday, New York.',
        '[10] Guilford, J.P. (1967). The Nature of Human Intelligence. McGraw-Hill, New York.',
        '[11] Kelley, T. and Kelley, D. (2013). Creative Confidence: Unleashing the Creative Potential Within Us All. Crown Business, New York.',
        '[12] Edmondson, A.C. (2019). The Fearless Organization: Creating Psychological Safety in the Workplace. Wiley, New York.',
        '[13] Meadows, D.H. (2008). Thinking in Systems: A Primer. Chelsea Green Publishing, Vermont.',
        '[14] Verganti, R. (2009). Design Driven Innovation: Changing the Rules of Competition by Radically Innovating What Things Mean. Harvard Business Press, Boston.',
        '[15] Mintzberg, H. (1994). The Rise and Fall of Strategic Planning. Free Press, New York.',
        '[16] Christensen, C.M. (1997). The Innovator\'s Dilemma: When New Technologies Cause Great Firms to Fail. Harvard Business School Press, Boston.',
        '[17] Simon, H.A. (1960). The New Science of Management Decision. Harper & Row, New York.',
        '[18] Simon, H.A. (1997). Administrative Behavior: A Study of Decision-Making Processes in Administrative Organizations. 4th ed. Free Press, New York.',
        '[19] Pfeffer, J. and Sutton, R.I. (2006). Hard Facts, Dangerous Half-Truths, and Total Nonsense: Profiting from Evidence-Based Management. Harvard Business School Press, Boston.',
        '[20] Kahneman, D. (2011). Thinking, Fast and Slow. Farrar, Straus and Giroux, New York.',
        '[21] Grant, R.M. (2016). Contemporary Strategy Analysis. 9th ed. Wiley, Chichester.',
        '[22] Provost, F. and Fawcett, T. (2013). Data Science for Business. O\'Reilly Media, Sebastopol.',
        '[23] McAfee, A. and Brynjolfsson, E. (2017). Machine, Platform, Crowd: Harnessing Our Digital Future. W.W. Norton, New York.',
        '[24] Bazerman, M.H. and Moore, D.A. (2013). Judgment in Managerial Decision Making. 8th ed. Wiley, New York.',
        '[25] Kaplan, R.S. and Norton, D.P. (1996). The Balanced Scorecard: Translating Strategy into Action. Harvard Business School Press, Boston.',
        '[26] von Bertalanffy, L. (1968). General System Theory: Foundations, Development, Applications. George Braziller, New York.',
        '[27] Ackoff, R.L. (1981). Creating the Corporate Future: Plan or Be Planned For. Wiley, New York.',
        '[28] Sterman, J.D. (2000). Business Dynamics: Systems Thinking and Modeling for a Complex World. McGraw-Hill, Boston.',
        '[29] Kim, D.H. (1999). Introduction to Systems Thinking. Pegasus Communications, Waltham.',
        '[30] Forrester, J.W. (1961). Industrial Dynamics. MIT Press, Cambridge.',
        '[31] Meadows, D.H. (1999). Leverage Points: Places to Intervene in a System. The Sustainability Institute, Hartland.',
        '[32] Jackson, M.C. (2003). Systems Thinking: Creative Holism for Managers. Wiley, Chichester.',
        '[33] Gharajedaghi, J. (2011). Systems Thinking: Managing Chaos and Complexity. 3rd ed. Morgan Kaufmann, Burlington.',
        '[34] Geissdoerfer, M., Savaget, P. and Evans, S. (2017). The Cambridge Business Model Innovation Process. Procedia Manufacturing, 8, pp.262-269.',
        '[35] Guilford, J.P. (1950). Creativity. American Psychologist, 5(9), pp.444-454.',
        '[36] De Bono, E. (1970). Lateral Thinking: Creativity Step by Step. Harper & Row, New York.',
        '[37] Amabile, T.M. (1996). Creativity in Context: Update to the Social Psychology of Creativity. Westview Press, Boulder.',
        '[38] Osborn, A.F. (1963). Applied Imagination: Principles and Procedures of Creative Problem-Solving. 3rd ed. Scribner\'s, New York.',
        '[39] Buzan, T. and Buzan, B. (1993). The Mind Map Book. BBC Books, London.',
        '[40] Colton, S. and Wiggins, G.A. (2012). Computational Creativity: The Final Frontier? Proceedings of the 20th European Conference on Artificial Intelligence, pp.21-26.',
        '[41] Robinson, K. (2011). Out of Our Minds: Learning to be Creative. Capstone, Chichester.',
        '[42] Teece, D.J. (2010). Business Models, Business Strategy and Innovation. Long Range Planning, 43(2-3), pp.172-194.',
        '[43] Kolko, J. (2015). Design Thinking Comes of Age. Harvard Business Review, 93(9), pp.66-71.',
    ]
    
    for ref in references:
        para = doc.add_paragraph(ref)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.first_line_indent = Cm(-1.27)
        para.paragraph_format.left_indent = Cm(1.27)
        for run in para.runs:
            run.font.size = Pt(10)
    
    # Save document
    output_path = 'Chapter_Design_Thinking_Comparative_Analysis.docx'
    doc.save(output_path)
    print(f"\nWord document saved successfully: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
    
    # Word count estimate
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count: {word_count}")

if __name__ == '__main__':
    create_document()
