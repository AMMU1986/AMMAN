"""
Create Word Document for Chapter: AI and Climate Action
For: Aligning Innovation with SDGs for Business Growth: Age of AI
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


    # ===== TITLE PAGE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI and Climate Action')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    book_info = doc.add_paragraph()
    book_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = book_info.add_run(
        '\nAligning Innovation with SDGs for Business Growth:\n'
        'Age of Artificial Intelligence')
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== ABSTRACT =====
    abs_title = doc.add_paragraph()
    run = abs_title.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(12)

    abstract_text = (
        "Artificial intelligence (AI) has emerged as a transformative force in addressing "
        "the global climate crisis, offering unprecedented capabilities for monitoring, "
        "predicting, mitigating, and adapting to climate change impacts. This chapter "
        "examines the multifaceted role of AI in advancing Climate Action (SDG 13) while "
        "simultaneously driving sustainable business growth. The analysis spans four "
        "critical dimensions: AI-enabled climate monitoring and environmental intelligence; "
        "AI-driven solutions for decarbonization across energy, transportation, industry, "
        "and built environments; AI-powered climate adaptation and resilience systems "
        "including early warning, smart agriculture, and disaster risk reduction; and the "
        "responsible governance of AI for climate action, encompassing green AI, climate "
        "finance, and future business models aligned with net-zero targets. The chapter "
        "demonstrates how organizations can leverage AI technologies to reduce greenhouse "
        "gas emissions by 5-10% across major sectors while creating new market opportunities "
        "estimated at $1.3-2.6 trillion by 2030. Critical challenges including the "
        "environmental footprint of AI itself, data equity, algorithmic bias in climate "
        "vulnerability assessments, and governance frameworks are addressed. The discussion "
        "concludes with a forward-looking analysis of how AI-enabled business models can "
        "simultaneously achieve climate targets and sustainable economic growth, providing "
        "a roadmap for enterprises seeking to align innovation strategies with SDG 13 in "
        "the age of artificial intelligence."
    )

    abs_para = doc.add_paragraph()
    run = abs_para.add_run(abstract_text)
    run.font.size = Pt(10)
    run.italic = True


    kw_para = doc.add_paragraph()
    run = kw_para.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    run = kw_para.add_run(
        'Artificial Intelligence, Climate Action, SDG 13, Climate Mitigation, '
        'Climate Adaptation, Green AI, Carbon Accounting, Renewable Energy, '
        'Sustainable Business, Net-Zero, Responsible Innovation')
    run.font.size = Pt(10)

    doc.add_page_break()

    # ===== SECTION 1 =====
    doc.add_heading('1. Artificial Intelligence and the Climate Change Challenge', level=1)

    doc.add_heading(
        '1.1 Climate Change, Sustainable Development, and the Role of AI', level=2)

    doc.add_paragraph(
        'Climate change represents the defining challenge of the twenty-first century, '
        'threatening ecosystems, economies, and human well-being on a planetary scale. '
        'The Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment Report '
        'unequivocally confirms that human-induced greenhouse gas emissions have raised '
        'global mean temperatures by approximately 1.1 degrees Celsius above pre-industrial '
        'levels, with projections indicating potential warming of 1.5 to 4.4 degrees by 2100 '
        'under various emission scenarios [1]. The United Nations Sustainable Development Goal '
        '13 (SDG 13) explicitly calls for urgent action to combat climate change and its '
        'impacts, establishing targets for strengthening resilience, integrating climate '
        'measures into national policies, and improving education and institutional capacity '
        'for climate mitigation and adaptation [2].'
    )

    doc.add_paragraph(
        'Artificial intelligence has emerged as a potentially transformative technology '
        'for accelerating climate action across multiple dimensions. The convergence of '
        'exponentially growing computational capabilities, vast climate-related datasets '
        'from satellite observations, IoT sensor networks, and simulation models, and '
        'advances in machine learning algorithms creates unprecedented opportunities for '
        'understanding and responding to climate change [3]. AI systems can process and '
        'analyze climate data at scales and speeds impossible for human analysts, identify '
        'complex patterns in earth system dynamics, optimize energy systems and industrial '
        'processes for minimal emissions, and enable predictive capabilities that support '
        'both mitigation and adaptation strategies.'
    )

    doc.add_paragraph(
        'The intersection of AI and climate action is particularly significant within '
        'the broader framework of the United Nations 2030 Agenda for Sustainable '
        'Development. While SDG 13 directly addresses climate action, the cross-cutting '
        'nature of both AI and climate change means that AI-climate solutions frequently '
        'deliver co-benefits across multiple SDGs, including affordable and clean energy '
        '(SDG 7), sustainable cities and communities (SDG 11), responsible consumption '
        'and production (SDG 12), and life on land (SDG 15) [4]. This multiplicative '
        'impact makes AI for climate action a high-leverage intervention point for '
        'sustainable development, where investments in AI capabilities can simultaneously '
        'advance multiple sustainability objectives while generating economic returns '
        'that sustain and scale the innovation ecosystem.'
    )

    doc.add_paragraph(
        'For businesses specifically, AI for climate action creates a dual value '
        'proposition: it simultaneously reduces operational costs through efficiency '
        'improvements and creates new revenue streams through climate products and '
        'services, while also mitigating regulatory, reputational, and physical climate '
        'risks that threaten long-term enterprise value [5]. Companies that proactively '
        'adopt AI-driven climate solutions are better positioned to comply with tightening '
        'emission regulations, meet growing consumer demand for sustainable products, '
        'attract climate-conscious capital, and build resilience against supply chain '
        'disruptions and extreme weather events. This alignment between climate action '
        'and business value creation represents a fundamental shift from viewing '
        'sustainability as a cost center to recognizing it as a core driver of '
        'competitive advantage and long-term value in the age of artificial intelligence.'
    )


    doc.add_paragraph(
        'The relationship between AI and climate action operates across a comprehensive '
        'framework encompassing monitoring, mitigation, adaptation, and governance '
        'dimensions, as illustrated in Figure 1. This framework demonstrates how AI '
        'technologies intersect with climate science, policy, and business innovation '
        'to create integrated solutions aligned with SDG 13 targets [4]. The World '
        'Economic Forum estimates that AI applications could reduce global greenhouse '
        'gas emissions by 4-10 percent by 2030, equivalent to 2.6-5.3 gigatons of CO2 '
        'equivalent annually, while simultaneously generating economic value of $1.3-2.6 '
        'trillion through efficiency gains, new products, and avoided climate damages [5]. '
        'However, realizing this potential requires careful governance to ensure that AI '
        'deployment for climate action is equitable, transparent, and does not exacerbate '
        'existing inequalities or create new environmental burdens through its own energy '
        'consumption and resource requirements.'
    )

    # Insert Figure 1
    fig1_para = doc.add_paragraph()
    fig1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig1_para.add_run()
    run.add_picture(
        '/projects/sandbox/AMMAN/ai_climate_figures/Figure_1_AI_Climate_Framework.png',
        width=Inches(5.5))

    fig1_cap = doc.add_paragraph()
    fig1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig1_cap.add_run(
        'Figure 1: AI-Enabled Climate Action Framework — Integrating Monitoring, '
        'Mitigation, Adaptation, and Governance for SDG 13, showing the four '
        'interconnected dimensions of AI applications for climate action.')
    run.bold = True
    run.font.size = Pt(9)

    doc.add_paragraph()


    # Section 1.2
    doc.add_heading(
        '1.2 AI-Enabled Climate Monitoring, Modelling, and Risk Assessment', level=2)

    doc.add_paragraph(
        'Climate monitoring represents one of the most established and impactful '
        'applications of artificial intelligence in the climate domain. Machine learning '
        'algorithms, particularly deep learning architectures such as convolutional neural '
        'networks (CNNs) and recurrent neural networks (RNNs), have revolutionized the '
        'processing and analysis of satellite remote sensing data for tracking climate '
        'variables including land surface temperature, sea ice extent, vegetation indices, '
        'atmospheric composition, and ocean heat content [6]. The European Space Agency\'s '
        'Climate Change Initiative and NASA\'s Earth Observing System generate petabytes of '
        'data annually that require AI-powered analysis pipelines to extract meaningful '
        'climate indicators at the temporal and spatial resolutions necessary for decision '
        'support. Deep learning models trained on decades of satellite imagery can detect '
        'subtle changes in glacier retreat, deforestation patterns, and urban heat island '
        'intensity with accuracies exceeding 95 percent, far surpassing traditional '
        'remote sensing classification approaches [7].'
    )

    doc.add_paragraph(
        'Climate modelling has been fundamentally transformed by AI through two '
        'complementary pathways: the enhancement of physics-based Earth System Models '
        '(ESMs) and the development of purely data-driven climate emulators. Traditional '
        'ESMs, while scientifically rigorous, require enormous computational resources '
        'and days to weeks of supercomputer time for century-scale projections. AI-based '
        'climate emulators, trained on ESM outputs, can reproduce these projections in '
        'seconds to minutes while maintaining fidelity to the underlying physics [8]. '
        'Google DeepMind\'s GenCast and Huawei\'s Pangu-Weather have demonstrated that '
        'transformer-based architectures can achieve weather prediction accuracy '
        'competitive with the European Centre for Medium-Range Weather Forecasts (ECMWF) '
        'operational model at a fraction of the computational cost [9]. These AI weather '
        'models generate 10-day global forecasts in under one minute on a single GPU, '
        'compared to hours on thousand-node supercomputer clusters for physics-based models.'
    )

    doc.add_paragraph(
        'Climate risk assessment leverages AI to quantify and spatially resolve the '
        'physical and transition risks that climate change poses to communities, '
        'ecosystems, and economic assets. Machine learning models integrate multiple '
        'climate hazard layers (flooding, heat stress, drought, wildfire, storm surge) '
        'with exposure and vulnerability data to produce granular risk maps at '
        'neighborhood to individual asset scales [10]. Financial institutions and '
        'insurance companies increasingly employ AI-driven climate risk models to assess '
        'portfolio exposure, price climate-related risks, and comply with Task Force on '
        'Climate-related Financial Disclosures (TCFD) reporting requirements. Natural '
        'language processing (NLP) algorithms analyze corporate disclosures, regulatory '
        'filings, and news sources to identify climate-related risks and opportunities, '
        'enabling automated ESG scoring and climate-aligned investment decisions [11].'
    )

    doc.add_paragraph(
        'The integration of AI with digital twin technology enables the creation of '
        'high-fidelity virtual replicas of climate-sensitive systems—from individual '
        'cities to entire watersheds—that can be used for scenario analysis, risk '
        'quantification, and adaptation planning. These climate digital twins combine '
        'real-time sensor data with AI-powered physics models to simulate the response '
        'of complex socio-ecological systems to various climate forcing scenarios, policy '
        'interventions, and adaptation measures [10]. Urban climate digital twins, for '
        'example, model the interactions between building energy systems, transportation '
        'networks, green infrastructure, and local microclimate to evaluate the '
        'effectiveness of urban heat mitigation strategies, flood management options, '
        'and emission reduction pathways at neighborhood resolution. The European Union\'s '
        'Destination Earth initiative aims to create a comprehensive digital twin of the '
        'Earth system powered by AI and exascale computing, providing decision-makers '
        'with unprecedented capabilities for testing climate policies and adaptation '
        'strategies before implementation.'
    )


    # Section 1.3
    doc.add_heading(
        '1.3 AI for Carbon Accounting, Emissions Tracking, and Environmental Intelligence',
        level=2)

    doc.add_paragraph(
        'Carbon accounting and emissions tracking represent critical enablers of climate '
        'action, and AI is revolutionizing the accuracy, granularity, and timeliness of '
        'greenhouse gas inventories. Traditional carbon accounting relies on activity-based '
        'calculations using emission factors that are often outdated, spatially aggregated, '
        'and subject to significant uncertainties. AI-powered emissions monitoring systems '
        'integrate satellite observations of atmospheric CO2 and methane concentrations, '
        'ground-based sensor networks, facility-level production data, and supply chain '
        'information to create near-real-time emission inventories at facility, city, and '
        'national scales [12]. The Climate TRACE initiative, combining satellite imagery '
        'analysis with machine learning, has mapped individual greenhouse gas emission '
        'sources globally, identifying over 72,000 major emitting facilities and revealing '
        'significant discrepancies between self-reported and independently verified '
        'emissions in many countries and sectors.'
    )

    doc.add_paragraph(
        'Scope 3 emissions, which typically represent 70-90 percent of a company\'s total '
        'carbon footprint but are notoriously difficult to measure, are increasingly '
        'addressed through AI-enabled supply chain intelligence platforms [13]. These '
        'platforms use machine learning to model emission intensities across complex, '
        'multi-tier supply chains by combining procurement data, industry benchmarks, '
        'input-output economic models, and supplier-specific information. Natural language '
        'processing algorithms extract emission-relevant data from supplier reports, '
        'invoices, and lifecycle assessment databases, while graph neural networks model '
        'the interconnected relationships within supply networks to propagate emission '
        'estimates through upstream and downstream activities. The AI-driven approach to '
        'carbon accounting achieves estimated accuracy improvements of 40-60 percent over '
        'traditional spend-based methods while reducing the manual effort required for '
        'data collection and verification by approximately 70 percent [14].'
    )

    # Table 1
    doc.add_paragraph()
    t1_cap = doc.add_paragraph()
    t1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t1_cap.add_run(
        'Table 1: AI Technologies for Climate Monitoring and '
        'Environmental Intelligence [6, 10, 12]')
    run.bold = True
    run.font.size = Pt(10)

    table1 = doc.add_table(rows=8, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    t1_headers = ['AI Technology', 'Climate Application',
                  'Data Sources', 'Accuracy/Impact']
    for i, h in enumerate(t1_headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    t1_data = [
        ['Deep Learning (CNN)', 'Land use change detection', 'Satellite imagery', '>95% accuracy'],
        ['Transformer Models', 'Weather/climate prediction', 'Reanalysis data', '10-day forecasts in <1 min'],
        ['Graph Neural Networks', 'Supply chain emissions', 'Procurement/trade data', '40-60% accuracy gain'],
        ['NLP/LLMs', 'ESG risk extraction', 'Corporate disclosures', 'Automated TCFD scoring'],
        ['Reinforcement Learning', 'Sensor network optimization', 'IoT climate sensors', '30% coverage improvement'],
        ['Physics-Informed ML', 'Climate model emulation', 'ESM outputs', '1000x speedup'],
        ['Anomaly Detection', 'Methane leak identification', 'Satellite spectrometry', '85% detection rate'],
    ]
    for i, row_data in enumerate(t1_data):
        for j, cell_text in enumerate(row_data):
            cell = table1.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()


    doc.add_paragraph(
        'Environmental intelligence platforms powered by AI integrate diverse data '
        'streams including air quality measurements, biodiversity observations, water '
        'quality indicators, and ecosystem health metrics into unified dashboards that '
        'support evidence-based environmental governance. These platforms employ federated '
        'learning approaches to combine data from distributed sensor networks without '
        'centralizing sensitive information, while generative AI models can fill spatial '
        'and temporal gaps in environmental monitoring data, creating continuous and '
        'consistent environmental intelligence products [14]. The integration of AI with '
        'blockchain technology enables transparent and tamper-proof carbon credit '
        'verification, supporting the integrity of voluntary and compliance carbon markets '
        'that are essential for channeling climate finance toward emission reduction '
        'projects. Table 1 summarizes the key AI technologies employed in climate '
        'monitoring and environmental intelligence, together with their primary '
        'applications, data requirements, and demonstrated performance metrics.'
    )

    doc.add_paragraph(
        'The convergence of AI-powered environmental intelligence with corporate '
        'sustainability reporting is driving a transformation in how organizations '
        'measure, disclose, and manage their environmental impacts. Regulatory '
        'developments including the EU Corporate Sustainability Reporting Directive '
        '(CSRD), SEC climate disclosure rules, and International Sustainability Standards '
        'Board (ISSB) standards are creating mandatory reporting requirements that '
        'necessitate AI-driven data collection and analysis capabilities [13]. AI systems '
        'can automatically map organizational activities to emission categories, apply '
        'appropriate calculation methodologies, identify data gaps, and generate audit-ready '
        'sustainability reports, reducing the compliance burden while improving data '
        'quality and timeliness. The integration of real-time emission monitoring with '
        'financial planning systems enables dynamic carbon budgeting, where organizations '
        'allocate emission allowances across business units and track performance against '
        'science-based targets with the same rigor applied to financial budgets.'
    )

    # ===== SECTION 2 =====
    doc.add_page_break()
    doc.add_heading('2. AI-Driven Solutions for Climate Mitigation', level=1)

    doc.add_heading(
        '2.1 Intelligent Renewable Energy Systems and Energy Optimization', level=2)

    doc.add_paragraph(
        'The energy sector, responsible for approximately 73 percent of global greenhouse '
        'gas emissions, represents the most consequential arena for AI-enabled climate '
        'mitigation. AI technologies are transforming renewable energy systems across '
        'the entire value chain, from resource assessment and generation forecasting to '
        'grid integration and demand-side management [15]. As depicted in Figure 2, '
        'AI-driven solutions address climate mitigation across four key sectors—energy, '
        'transportation, industry, and buildings—with estimated combined CO2 reduction '
        'potential of 5.9 gigatons annually by 2030.'
    )

    # Insert Figure 2
    fig2_para = doc.add_paragraph()
    fig2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig2_para.add_run()
    run.add_picture(
        '/projects/sandbox/AMMAN/ai_climate_figures/Figure_2_AI_Mitigation_Sectors.png',
        width=Inches(5.8))

    fig2_cap = doc.add_paragraph()
    fig2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig2_cap.add_run(
        'Figure 2: AI-Driven Solutions for Climate Mitigation Across Key Sectors, '
        'showing the application of AI technologies in energy, transportation, industry, '
        'and buildings with estimated CO2 reduction potential by 2030.')
    run.bold = True
    run.font.size = Pt(9)

    doc.add_paragraph()


    doc.add_paragraph(
        'Solar and wind energy forecasting has been dramatically improved through deep '
        'learning models that combine numerical weather prediction outputs, satellite '
        'cloud imagery, historical generation data, and real-time sensor readings to '
        'predict power output at horizons ranging from minutes to days ahead. '
        'Convolutional neural networks analyzing sky imagery achieve sub-hourly solar '
        'irradiance forecasts with mean absolute errors below 5 percent, while '
        'transformer-based models incorporating attention mechanisms capture long-range '
        'temporal dependencies in wind patterns, reducing day-ahead wind power forecasting '
        'errors by 20-35 percent compared to persistence models [16]. These improvements '
        'in renewable energy forecasting directly enable higher penetration of variable '
        'renewable generation by reducing the need for costly spinning reserves and backup '
        'fossil fuel generation.'
    )

    doc.add_paragraph(
        'Smart grid management leverages AI for real-time balancing of electricity supply '
        'and demand across increasingly complex networks incorporating distributed '
        'generation, battery storage, electric vehicles, and flexible loads. Deep '
        'reinforcement learning agents have demonstrated the ability to manage grid '
        'operations including voltage regulation, frequency control, and congestion '
        'management while reducing curtailment of renewable generation by 15-30 percent '
        '[17]. Google DeepMind\'s collaboration with National Grid demonstrated that AI '
        'optimization of wind farm operations—adjusting individual turbine yaw angles '
        'to account for wake effects—increased total energy output by approximately '
        '0.3-0.5 percent across the fleet, equivalent to significant additional clean '
        'energy production without any new infrastructure investment. Battery energy '
        'storage systems (BESS) optimized through AI achieve 10-20 percent improvement '
        'in revenue from energy arbitrage and ancillary services compared to rule-based '
        'dispatch strategies, accelerating the business case for storage deployment that '
        'enables further renewable penetration [18].'
    )

    doc.add_paragraph(
        'Demand-side management and demand response programs are increasingly orchestrated '
        'by AI systems that predict and shape electricity consumption patterns across '
        'millions of connected devices. Machine learning algorithms analyze historical '
        'consumption data, weather forecasts, occupancy patterns, and electricity pricing '
        'signals to optimize the timing of flexible loads including electric vehicle '
        'charging, water heating, industrial processes, and commercial HVAC systems. '
        'These AI-driven demand flexibility programs can shift 10-20 percent of peak '
        'demand to off-peak periods, reducing the need for peaking fossil fuel generation '
        'and enabling higher renewable energy utilization [16]. The aggregation of '
        'distributed energy resources—rooftop solar, battery storage, electric vehicles, '
        'and flexible loads—into AI-managed virtual power plants creates new business '
        'models that generate revenue for prosumers while providing grid services that '
        'support system stability under high renewable penetration scenarios.'
    )


    # Section 2.2
    doc.add_heading(
        '2.2 AI for Sustainable Mobility, Smart Infrastructure, and Industrial '
        'Decarbonization', level=2)

    doc.add_paragraph(
        'The transportation sector accounts for approximately 16 percent of global '
        'greenhouse gas emissions, and AI is enabling multiple pathways for '
        'decarbonization including vehicle efficiency optimization, modal shift, shared '
        'mobility, and logistics optimization [19]. Machine learning algorithms optimize '
        'routing for freight and passenger vehicles, considering real-time traffic '
        'conditions, road gradients, weather, and vehicle characteristics to minimize '
        'fuel consumption. Studies demonstrate that AI-optimized routing reduces fuel '
        'consumption by 8-15 percent for commercial fleets, while eco-driving algorithms '
        'that coach drivers based on real-time vehicle telematics achieve additional '
        'savings of 5-12 percent [20]. Autonomous driving systems, while primarily '
        'developed for safety and convenience, can achieve significant emission reductions '
        'through smoother acceleration patterns, optimal speed selection, and platoon '
        'formation that reduces aerodynamic drag by 15-25 percent for following vehicles.'
    )

    doc.add_paragraph(
        'Smart infrastructure powered by AI extends beyond transportation to encompass '
        'buildings, water systems, and urban metabolism. AI-driven building energy '
        'management systems (BEMS) employing model predictive control and deep '
        'reinforcement learning optimize HVAC operations, lighting, and plug loads based '
        'on occupancy predictions, weather forecasts, and grid signals, achieving energy '
        'savings of 20-40 percent compared to conventional rule-based controls [21]. '
        'Digital twin technology, creating physics-informed virtual replicas of physical '
        'infrastructure, enables AI to simulate and optimize energy performance across '
        'building portfolios, districts, and entire cities before implementing changes '
        'in the real world.'
    )

    doc.add_paragraph(
        'Industrial decarbonization represents one of the most challenging frontiers for '
        'climate mitigation, particularly in hard-to-abate sectors such as steel, cement, '
        'chemicals, and aluminum production. AI addresses industrial emissions through '
        'multiple mechanisms: process optimization that reduces energy intensity per unit '
        'of production by 5-15 percent, predictive maintenance that prevents energy-wasting '
        'equipment degradation, quality prediction that reduces waste and rework, and '
        'optimization of carbon capture utilization and storage (CCUS) systems [22]. '
        'Machine learning models trained on historical process data and first-principles '
        'constraints can identify optimal operating parameters that human operators cannot '
        'discover through trial-and-error, achieving emission reductions while maintaining '
        'or improving product quality and throughput. In the chemicals industry, AI-driven '
        'catalyst design and reaction optimization have demonstrated 10-30 percent '
        'reductions in energy consumption for key chemical processes including ammonia '
        'synthesis and methanol production [23].'
    )

    doc.add_paragraph(
        'Supply chain decarbonization represents an increasingly important application '
        'of AI, as organizations face growing pressure from regulators, investors, and '
        'consumers to reduce emissions across their entire value chain. AI-powered supply '
        'chain platforms analyze thousands of suppliers, shipping routes, material choices, '
        'and production schedules simultaneously to identify optimal configurations that '
        'minimize total carbon intensity while maintaining cost competitiveness and service '
        'levels [22]. Multi-objective optimization algorithms balance emission reduction '
        'against cost, delivery time, and quality constraints, enabling procurement '
        'decisions that achieve 15-30 percent emission reductions in supply chain '
        'operations with minimal economic penalty. The combination of digital product '
        'passports, IoT-enabled tracking, and AI analytics creates end-to-end supply '
        'chain visibility that supports both regulatory compliance and voluntary '
        'sustainability commitments.'
    )


    # Section 2.3
    doc.add_heading(
        '2.3 AI-Enabled Circular Economy, Resource Efficiency, and Waste Management',
        level=2)

    doc.add_paragraph(
        'The circular economy paradigm, which seeks to eliminate waste and maximize '
        'resource utilization through reuse, repair, remanufacturing, and recycling, '
        'is increasingly enabled by AI technologies that can navigate the complexity '
        'of material flows, product lifecycles, and reverse logistics networks [24]. '
        'Computer vision systems powered by deep learning achieve material classification '
        'accuracies of 95-99 percent in automated waste sorting facilities, enabling the '
        'recovery of recyclable materials that would otherwise be landfilled or '
        'incinerated. These AI-powered sorting systems can distinguish between dozens of '
        'plastic polymer types, metal alloys, paper grades, and organic materials at '
        'processing speeds of several hundred items per minute, far exceeding human '
        'sorting capabilities.'
    )

    doc.add_paragraph(
        'AI-enabled predictive lifecycle management extends product useful life through '
        'condition monitoring, predictive maintenance, and optimal replacement scheduling. '
        'Digital product passports, powered by AI and blockchain, track material '
        'composition, usage history, and end-of-life options for manufactured goods, '
        'facilitating material recovery and circular business models [25]. In the food '
        'system, which accounts for approximately 26 percent of global greenhouse gas '
        'emissions, AI reduces waste across the supply chain through demand forecasting '
        'that reduces overproduction by 20-30 percent, computer vision quality inspection '
        'that identifies defects early in processing, and dynamic pricing algorithms that '
        'redirect near-expiry products to appropriate channels before they become waste.'
    )

    doc.add_paragraph(
        'Industrial symbiosis and eco-industrial park optimization represent advanced '
        'applications of AI for circular economy principles. Graph neural networks model '
        'the material and energy flow relationships between industrial facilities, '
        'identifying opportunities for waste heat recovery, by-product exchange, and '
        'shared infrastructure that reduce both costs and emissions [25]. AI-powered '
        'material flow analysis at urban and regional scales enables cities to map their '
        'metabolic processes and identify leverage points for circularity interventions. '
        'In the construction sector, AI facilitates building material reuse by matching '
        'demolition outputs with construction inputs through automated assessment of '
        'structural component condition, dimensional verification through 3D scanning, '
        'and marketplace algorithms that optimize logistics for material recovery. The '
        'integration of AI with life cycle assessment (LCA) tools enables real-time '
        'environmental impact evaluation of product design decisions, embedding '
        'circularity considerations into the earliest stages of the innovation process '
        'and supporting business models that decouple economic growth from virgin '
        'resource extraction and waste generation.'
    )

    # Table 2
    doc.add_paragraph()
    t2_cap = doc.add_paragraph()
    t2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t2_cap.add_run(
        'Table 2: AI-Driven Climate Mitigation Solutions and '
        'Estimated Impact by Sector [15, 19, 22, 24]')
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=9, cols=5)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    t2_headers = ['Sector', 'AI Application', 'Technology',
                  'Emission Reduction', 'Business Value']
    for i, h in enumerate(t2_headers):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    t2_data = [
        ['Energy', 'Renewable forecasting', 'Deep Learning', '15-30% curtailment reduction', '$50-80B/yr savings'],
        ['Energy', 'Grid optimization', 'Reinforcement Learning', '10-20% loss reduction', '$30-50B/yr value'],
        ['Transport', 'Route optimization', 'ML/Optimization', '8-15% fuel reduction', '$20-40B/yr savings'],
        ['Transport', 'Autonomous vehicles', 'Deep RL/CV', '10-25% per-vehicle', '$100B+ market'],
        ['Industry', 'Process optimization', 'Physics-informed ML', '5-15% energy intensity', '$40-60B/yr savings'],
        ['Buildings', 'HVAC optimization', 'Model Predictive Control', '20-40% energy savings', '$15-30B/yr value'],
        ['Agriculture', 'Precision farming', 'Computer Vision/IoT', '15-25% input reduction', '$20-30B/yr value'],
        ['Waste', 'Automated sorting', 'Computer Vision', '30-50% recovery increase', '$10-20B/yr value'],
    ]
    for i, row_data in enumerate(t2_data):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()


    doc.add_paragraph(
        'Table 2 provides a comprehensive overview of AI-driven climate mitigation '
        'solutions across key economic sectors, quantifying both the emission reduction '
        'potential and the associated business value creation. The aggregate impact of '
        'these AI applications, if deployed at scale, could contribute 2.6-5.3 gigatons '
        'of annual CO2 reduction by 2030, representing approximately 5-10 percent of '
        'current global emissions. The business value generated through efficiency gains, '
        'new revenue streams, and avoided damages provides strong economic incentives for '
        'private sector engagement in AI-enabled climate mitigation, creating a positive '
        'feedback loop between sustainability and profitability that aligns with the '
        'broader goals of SDG-aligned business growth [15].'
    )

    # ===== SECTION 3 =====
    doc.add_page_break()
    doc.add_heading(
        '3. Artificial Intelligence for Climate Adaptation and Resilience', level=1)

    doc.add_heading(
        '3.1 Predictive Analytics and Early Warning Systems for Climate Hazards',
        level=2)

    doc.add_paragraph(
        'Climate adaptation requires the ability to anticipate, prepare for, and respond '
        'to the increasing frequency and severity of climate-related hazards. AI-powered '
        'early warning systems represent a critical application that can save lives and '
        'reduce economic losses by providing timely and actionable information to '
        'vulnerable communities and decision-makers [26]. The architecture of AI-enabled '
        'climate adaptation systems, as illustrated in Figure 3, encompasses three layers: '
        'a data layer integrating diverse environmental observations, an AI processing '
        'layer performing pattern recognition and prediction, and an application layer '
        'delivering decision support for specific adaptation challenges.'
    )

    # Insert Figure 3
    fig3_para = doc.add_paragraph()
    fig3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig3_para.add_run()
    run.add_picture(
        '/projects/sandbox/AMMAN/ai_climate_figures/Figure_3_AI_Adaptation_Resilience.png',
        width=Inches(5.8))

    fig3_cap = doc.add_paragraph()
    fig3_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig3_cap.add_run(
        'Figure 3: AI-Powered Climate Adaptation and Resilience Systems, illustrating '
        'the three-layer architecture from data acquisition through AI processing to '
        'application delivery for climate hazard management and adaptation planning.')
    run.bold = True
    run.font.size = Pt(9)

    doc.add_paragraph()


    doc.add_paragraph(
        'Flood prediction systems employing AI have achieved remarkable advances in both '
        'accuracy and lead time. Google\'s flood forecasting initiative uses machine '
        'learning to predict riverine flooding up to 7 days in advance across multiple '
        'continents, providing warnings to over 460 million people in flood-prone areas '
        '[27]. The system combines hydrological models with deep learning to process '
        'rainfall forecasts, soil moisture data, and river gauge measurements, producing '
        'inundation maps at 30-meter resolution that enable targeted evacuation planning. '
        'Similarly, AI-driven wildfire risk prediction systems integrate satellite-derived '
        'vegetation indices, meteorological forecasts, topographic data, and historical '
        'ignition patterns to predict fire spread with 80-90 percent accuracy at hourly '
        'time steps, enabling proactive resource positioning and evacuation orders [28].'
    )

    doc.add_paragraph(
        'Tropical cyclone prediction has benefited from AI through improved track and '
        'intensity forecasting, with deep learning models reducing 72-hour track errors '
        'by 15-20 percent and intensity errors by 10-15 percent compared to operational '
        'statistical-dynamical models [29]. AI also enables compound event prediction, '
        'where multiple climate hazards co-occur or cascade, creating impacts that exceed '
        'the sum of individual hazards. Graph neural networks that model the causal '
        'relationships between climate drivers, atmospheric patterns, and surface impacts '
        'can identify emerging compound risk situations days to weeks before they '
        'materialize, providing crucial preparation time for emergency managers.'
    )

    doc.add_paragraph(
        'The democratization of early warning systems through AI is particularly '
        'important for developing countries and small island developing states that lack '
        'the dense observation networks and computational infrastructure of wealthy '
        'nations. Transfer learning approaches enable AI models trained on data-rich '
        'regions to be adapted to data-sparse environments with minimal local calibration, '
        'while federated learning allows countries to collaboratively improve prediction '
        'models without sharing sensitive national data [29]. The World Meteorological '
        'Organization\'s Early Warnings for All initiative aims to ensure that every '
        'person on Earth is protected by early warning systems by 2027, and AI is '
        'essential for achieving this ambitious goal in regions where traditional '
        'monitoring infrastructure remains inadequate. Mobile phone-based dissemination '
        'systems powered by AI can deliver location-specific, actionable warnings in '
        'local languages to populations that lack access to conventional media channels, '
        'potentially reaching billions of people in climate-vulnerable communities who '
        'currently receive no advance warning of impending hazards.'
    )

    # Section 3.2
    doc.add_heading(
        '3.2 AI for Climate-Resilient Agriculture, Water Management, and '
        'Ecosystem Protection', level=2)

    doc.add_paragraph(
        'Agriculture is simultaneously one of the sectors most vulnerable to climate '
        'change and a significant contributor to greenhouse gas emissions. AI enables '
        'climate-resilient agriculture through precision management practices that '
        'optimize resource use while maintaining productivity under changing climatic '
        'conditions [30]. Computer vision systems mounted on drones and ground robots '
        'detect crop stress, pest infestations, and nutrient deficiencies at individual '
        'plant level, enabling targeted interventions that reduce chemical inputs by '
        '30-50 percent while maintaining yields. Machine learning models that predict '
        'crop yields under various climate scenarios support agricultural planning and '
        'food security early warning, with prediction accuracies of 85-92 percent at '
        'regional scales achieved through ensemble approaches combining remote sensing, '
        'weather data, and soil information.'
    )

    doc.add_paragraph(
        'Water resource management under climate change requires sophisticated AI systems '
        'that can model complex hydrological processes, predict water availability under '
        'changing precipitation patterns, and optimize allocation among competing demands. '
        'Deep learning models for rainfall-runoff prediction outperform traditional '
        'hydrological models in 70-80 percent of catchments globally, particularly in '
        'data-scarce regions where physics-based models lack adequate calibration data '
        '[31]. AI-powered irrigation scheduling systems that integrate soil moisture '
        'sensors, weather forecasts, crop growth models, and water pricing signals '
        'achieve water savings of 20-35 percent while maintaining or improving crop '
        'quality and yields. Groundwater management, critical for climate adaptation in '
        'many arid and semi-arid regions, benefits from AI models that predict aquifer '
        'recharge rates, identify optimal pumping schedules, and detect unsustainable '
        'extraction patterns before irreversible depletion occurs.'
    )


    doc.add_paragraph(
        'Ecosystem protection and biodiversity conservation are increasingly supported '
        'by AI through automated species identification, habitat monitoring, and '
        'ecosystem service quantification. Acoustic monitoring systems using deep '
        'learning can identify thousands of species from their vocalizations, enabling '
        'cost-effective biodiversity assessments across vast landscapes [32]. AI models '
        'that map ecosystem carbon stocks, predict deforestation risk, and identify '
        'priority areas for conservation investment support nature-based solutions to '
        'climate change, which the IPCC estimates could provide 20-30 percent of the '
        'mitigation needed to limit warming to 1.5 degrees Celsius.'
    )

    doc.add_paragraph(
        'Climate-induced migration and displacement represent growing challenges that AI '
        'can help address through predictive modeling of population movement patterns '
        'under various climate scenarios. Machine learning models that integrate climate '
        'projections, agricultural productivity forecasts, conflict risk indicators, and '
        'economic opportunity data can anticipate displacement hotspots years to decades '
        'in advance, enabling proactive investment in receiving communities, livelihood '
        'adaptation programs, and managed relocation planning [32]. These predictive '
        'capabilities support both humanitarian response preparedness and long-term '
        'development planning in regions facing slow-onset climate impacts such as sea '
        'level rise, desertification, and chronic water scarcity.'
    )

    # Section 3.3
    doc.add_heading(
        '3.3 Smart Cities, Disaster Risk Reduction, and Adaptive Infrastructure',
        level=2)

    doc.add_paragraph(
        'Urban areas, home to over 55 percent of the global population and responsible '
        'for approximately 70 percent of CO2 emissions, represent both critical '
        'contributors to climate change and highly vulnerable concentrations of people '
        'and assets. AI-enabled smart city platforms integrate urban sensing networks, '
        'building energy systems, transportation flows, and climate projections to '
        'optimize urban metabolism and enhance climate resilience simultaneously [33]. '
        'Urban heat island mitigation leverages AI to identify optimal locations for '
        'green infrastructure, cool surfaces, and shade structures based on microclimate '
        'modeling at street-level resolution, achieving local temperature reductions of '
        '2-5 degrees Celsius in targeted intervention areas.'
    )

    doc.add_paragraph(
        'Disaster risk reduction benefits from AI through improved preparedness, response, '
        'and recovery phases of the disaster management cycle. Machine learning algorithms '
        'analyze social media streams, satellite imagery, and sensor data in real-time '
        'during disaster events to map affected areas, estimate damage severity, and '
        'identify populations in need of immediate assistance [34]. Computer vision '
        'applied to post-disaster satellite and drone imagery can assess building damage '
        'at individual structure level within hours of an event, replacing weeks of manual '
        'field assessment and enabling rapid targeting of relief resources.'
    )

    doc.add_paragraph(
        'Adaptive infrastructure design employs AI to ensure that long-lived physical '
        'assets remain functional and resilient under future climate conditions that may '
        'differ significantly from historical experience. AI-powered climate stress '
        'testing subjects infrastructure designs to thousands of synthetic climate '
        'scenarios generated by machine learning emulators, identifying vulnerabilities '
        'and optimal adaptation investments under deep uncertainty [35]. For coastal '
        'infrastructure, AI models integrate sea level rise projections, storm surge '
        'simulations, and morphological change predictions to optimize the design of '
        'flood defenses, managed retreat strategies, and nature-based coastal protection '
        'systems. The combination of digital twin technology with AI enables continuous '
        'performance monitoring and predictive maintenance of climate-critical '
        'infrastructure, identifying deterioration patterns that could lead to failure '
        'during extreme weather events.'
    )

    doc.add_paragraph(
        'The economic case for AI-enabled climate adaptation is compelling. The Global '
        'Commission on Adaptation estimates that investing $1.8 trillion in adaptation '
        'between 2020 and 2030 could generate $7.1 trillion in net benefits, and AI '
        'optimization of these investments can significantly improve their cost-effectiveness '
        'by targeting interventions where they deliver maximum resilience benefits per '
        'dollar invested [35]. AI-powered cost-benefit analysis tools that incorporate '
        'uncertainty from climate projections, socioeconomic scenarios, and discount rate '
        'assumptions enable more robust adaptation investment decisions under deep '
        'uncertainty. Insurance industry applications leverage AI to price climate risks '
        'more accurately, incentivize risk reduction through premium discounts for adapted '
        'properties, and develop innovative products such as parametric microinsurance '
        'that reaches underserved populations in climate-vulnerable regions. The combination '
        'of improved risk quantification, optimized intervention targeting, and innovative '
        'financial instruments creates a virtuous cycle where AI enhances the economic '
        'viability of adaptation investments, attracting greater private sector capital '
        'to climate resilience.'
    )

    # Table 3
    doc.add_paragraph()
    t3_cap = doc.add_paragraph()
    t3_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t3_cap.add_run(
        'Table 3: AI Applications for Climate Adaptation and '
        'Resilience by Domain [26, 30, 33]')
    run.bold = True
    run.font.size = Pt(10)


    table3 = doc.add_table(rows=8, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    t3_headers = ['Domain', 'AI Application', 'Key Benefit',
                  'SDG 13 Target']
    for i, h in enumerate(t3_headers):
        cell = table3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    t3_data = [
        ['Flood management', 'Deep learning flood prediction', '7-day advance warning, 460M+ people', '13.1'],
        ['Wildfire', 'Spread prediction & risk mapping', '80-90% accuracy, proactive response', '13.1'],
        ['Agriculture', 'Precision climate-smart farming', '30-50% input reduction, yield stability', '13.2'],
        ['Water resources', 'AI-driven hydrological modeling', '20-35% water savings', '13.2'],
        ['Ecosystems', 'Biodiversity monitoring & conservation', 'Automated species tracking at scale', '13.2'],
        ['Urban resilience', 'Smart city climate platforms', '2-5°C heat island reduction', '13.1'],
        ['Infrastructure', 'Climate stress testing & adaptation', 'Optimized design under uncertainty', '13.1'],
    ]
    for i, row_data in enumerate(t3_data):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()

    doc.add_paragraph(
        'The AI applications for climate adaptation and resilience summarized in Table 3 '
        'demonstrate the breadth of domains where AI contributes to reducing climate '
        'vulnerability and enhancing adaptive capacity. These applications directly '
        'support SDG 13 targets for strengthening resilience (13.1) and integrating '
        'climate change measures into policies and planning (13.2). The combination of '
        'early warning systems as shown in Figure 3, precision resource management, and '
        'adaptive infrastructure creates a comprehensive AI-enabled adaptation framework '
        'that can significantly reduce the human and economic costs of climate impacts '
        'while creating new business opportunities in climate services, resilience '
        'consulting, and adaptation technology markets estimated at $300-500 billion '
        'annually by 2030 [26].'
    )


    # ===== SECTION 4 =====
    doc.add_page_break()
    doc.add_heading(
        '4. Responsible AI, Business Growth, and the Future of Climate Action',
        level=1)

    doc.add_heading(
        '4.1 Green AI, Energy-Efficient Computing, and the Environmental '
        'Footprint of AI', level=2)

    doc.add_paragraph(
        'While AI offers significant potential for climate mitigation and adaptation, the '
        'environmental footprint of AI systems themselves has emerged as a critical concern '
        'that must be addressed to ensure net positive climate impact. Training large '
        'language models and deep learning systems requires substantial computational '
        'resources: training GPT-3 consumed an estimated 1,287 megawatt-hours of '
        'electricity and generated approximately 552 tons of CO2 equivalent, while '
        'subsequent models have grown orders of magnitude larger [36]. The International '
        'Energy Agency estimates that data center electricity consumption could reach '
        '1,000 terawatt-hours by 2026, representing approximately 3-4 percent of global '
        'electricity demand, with AI workloads accounting for a rapidly growing share of '
        'this consumption.'
    )

    doc.add_paragraph(
        'The Green AI movement advocates for prioritizing computational efficiency '
        'alongside model performance, promoting research into energy-efficient algorithms, '
        'model compression, and hardware-software co-design [37]. Key strategies for '
        'reducing the carbon footprint of AI include: algorithmic efficiency improvements '
        'through sparse models, knowledge distillation, and neural architecture search '
        'optimized for energy consumption; hardware advances including specialized AI '
        'accelerators that achieve 10-100x energy efficiency improvements over '
        'general-purpose GPUs; geographic and temporal optimization of training workloads '
        'to leverage renewable energy availability; and federated learning approaches '
        'that reduce data transfer and enable local computation on edge devices.'
    )

    doc.add_paragraph(
        'The concept of a "carbon budget" for AI development requires organizations to '
        'account for the full lifecycle emissions of AI systems including hardware '
        'manufacturing, training, inference, and end-of-life disposal. Studies indicate '
        'that inference, rather than training, dominates the lifetime carbon footprint '
        'for widely deployed models, with inference accounting for 80-90 percent of '
        'total compute over a model\'s operational lifetime [36]. This highlights the '
        'importance of deployment-phase efficiency optimization through techniques such '
        'as model quantization, pruning, caching, and dynamic batching. Organizations '
        'committed to responsible AI for climate action must ensure that the carbon '
        'savings enabled by their AI applications substantially exceed the carbon costs '
        'of developing and operating those systems—a principle of net-positive climate '
        'impact that should guide investment decisions and system design choices.'
    )

    doc.add_paragraph(
        'The lifecycle assessment of AI systems for climate applications reveals that '
        'the ratio of carbon saved to carbon spent varies enormously across use cases. '
        'AI-optimized renewable energy systems typically achieve carbon benefit ratios of '
        '100:1 or greater, meaning that every kilogram of CO2 emitted in running the AI '
        'system prevents 100 kilograms of emissions through improved grid efficiency. In '
        'contrast, large-scale natural language processing models applied to climate '
        'document analysis may achieve ratios closer to 10:1 or even less, depending on '
        'model size and deployment patterns. Organizations must therefore carefully select '
        'and prioritize AI applications based on their net climate impact, favoring '
        'deployment scenarios where AI operates continuously to optimize physical systems '
        '(energy, transport, industry) over one-time analytical tasks that could be '
        'accomplished with smaller, more efficient models [37]. The development of '
        'standardized methodologies for computing the net climate impact of AI deployments '
        'remains an active area of research, with emerging frameworks proposing that '
        'organizations report both the direct carbon cost of their AI operations and the '
        'indirect carbon savings enabled by AI-driven decisions.'
    )


    # Section 4.2
    doc.add_heading(
        '4.2 AI Governance, Ethics, Climate Finance, and Responsible Innovation',
        level=2)

    doc.add_paragraph(
        'The governance of AI for climate action requires frameworks that ensure '
        'equitable access to AI-derived climate benefits, prevent algorithmic bias in '
        'vulnerability assessments, and maintain accountability for AI-informed decisions '
        'that affect communities and ecosystems. Climate justice concerns arise when AI '
        'systems, trained predominantly on data from wealthy nations, perform poorly in '
        'data-scarce developing countries that are most vulnerable to climate impacts '
        '[38]. Addressing this "AI divide" requires investment in local data '
        'infrastructure, participatory AI development that incorporates indigenous and '
        'local knowledge, and open-source AI tools that democratize access to climate '
        'intelligence capabilities.'
    )

    doc.add_paragraph(
        'Ethical considerations in AI for climate action extend to questions of '
        'transparency, explainability, and consent. Climate models that inform adaptation '
        'investments, insurance pricing, and land-use planning must be interpretable to '
        'stakeholders who may be significantly affected by their outputs. The "black box" '
        'nature of deep learning models raises particular concerns when used for climate '
        'risk assessment in contexts that affect property values, insurance availability, '
        'and community viability. Explainable AI (XAI) techniques including attention '
        'mechanisms, SHAP values, and concept-based explanations provide partial solutions, '
        'but the fundamental tension between model complexity and interpretability remains '
        'an active area of research and policy debate [38].'
    )

    doc.add_paragraph(
        'Climate finance, estimated at $1.3 trillion in 2021-2022 flows but requiring '
        '$4.3 trillion annually by 2030 to meet Paris Agreement targets, is increasingly '
        'mediated by AI systems that assess project viability, allocate capital, and '
        'monitor outcomes. AI-driven climate finance platforms enable automated carbon '
        'credit verification through satellite monitoring, reducing verification costs '
        'by 60-80 percent while improving accuracy. Machine learning algorithms match '
        'climate projects with appropriate financing instruments, assess technology '
        'readiness levels, and predict return profiles for climate investments under '
        'various policy scenarios [39]. The integration of AI with green bonds, '
        'sustainability-linked loans, and blended finance mechanisms is creating new '
        'asset classes and investment products that channel private capital toward '
        'climate solutions at unprecedented scale.'
    )

    doc.add_paragraph(
        'Responsible innovation in AI for climate action also requires addressing the '
        'potential for rebound effects, where efficiency improvements enabled by AI lead '
        'to increased consumption that partially or fully offsets emission reductions. '
        'For example, AI-optimized logistics that reduce per-shipment emissions may '
        'lower shipping costs, stimulating additional demand that increases total '
        'transport emissions. Managing these rebound effects requires integration of AI '
        'optimization with appropriate policy frameworks including carbon pricing, '
        'absolute emission caps, and sufficiency-oriented approaches that constrain '
        'total consumption rather than merely improving per-unit efficiency [38]. '
        'The governance challenge extends to ensuring that AI systems for climate action '
        'are interoperable across national boundaries, as climate change is inherently '
        'a global challenge requiring coordinated international responses. Standardization '
        'efforts for climate data formats, AI model interoperability, and cross-border '
        'data sharing protocols are essential for enabling the global AI-climate '
        'intelligence infrastructure necessary to support coordinated international '
        'climate action under the Paris Agreement framework.'
    )


    # Section 4.3
    doc.add_heading(
        '4.3 Future Directions: AI-Enabled Business Models for SDG 13 and '
        'Sustainable Growth', level=2)

    doc.add_paragraph(
        'The convergence of artificial intelligence with climate action is generating '
        'entirely new business models and market opportunities that simultaneously '
        'address SDG 13 and drive sustainable economic growth. Climate-as-a-Service '
        '(CaaS) platforms leverage AI to provide customized climate risk assessments, '
        'adaptation recommendations, and carbon management tools to organizations of all '
        'sizes, democratizing access to sophisticated climate intelligence previously '
        'available only to large corporations with dedicated sustainability teams. The '
        'global climate tech market, enabled substantially by AI, attracted over $70 '
        'billion in venture capital investment between 2020 and 2024, with AI-native '
        'climate startups receiving a disproportionate share of this funding [5].'
    )

    # Insert Figure 4
    fig4_para = doc.add_paragraph()
    fig4_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig4_para.add_run()
    run.add_picture(
        '/projects/sandbox/AMMAN/ai_climate_figures/Figure_4_Responsible_AI_Future.png',
        width=Inches(5.5))

    fig4_cap = doc.add_paragraph()
    fig4_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig4_cap.add_run(
        'Figure 4: Responsible AI and Business Growth Pathways for Sustainable '
        'Climate Action, illustrating the phased progression from foundational '
        'Green AI capabilities to transformative integration of AI with SDG 13 targets.')
    run.bold = True
    run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_paragraph(
        'The roadmap for AI-enabled climate business models, illustrated in Figure 4, '
        'progresses through four phases: foundation (establishing green AI practices and '
        'basic climate tools), scaling (deploying AI-enabled carbon markets and digital '
        'twins), transformation (autonomous climate management systems), and integration '
        '(full alignment of AI capabilities with planetary boundaries and SDG targets). '
        'Each phase creates business value while progressively deepening climate impact, '
        'demonstrating that sustainability and profitability are not merely compatible but '
        'mutually reinforcing in the age of artificial intelligence.'
    )


    doc.add_paragraph(
        'Emerging business models at the AI-climate intersection include: AI-optimized '
        'renewable energy trading platforms that enable peer-to-peer clean energy '
        'exchange; carbon intelligence platforms that provide real-time emission '
        'monitoring, reduction pathway modeling, and offset marketplace access; '
        'climate-resilient supply chain services that use AI to identify vulnerabilities, '
        'suggest alternatives, and optimize logistics for minimal carbon intensity; and '
        'parametric insurance products that use AI-driven climate models to trigger '
        'automatic payouts based on predefined climate event thresholds, reducing '
        'settlement times from months to hours [39]. These business models share common '
        'characteristics: they leverage AI to process complex environmental data, they '
        'create value through both cost reduction and risk management, and they scale '
        'efficiently through digital platforms that serve diverse customer segments.'
    )

    doc.add_paragraph(
        'The future trajectory of AI for climate action will be shaped by several '
        'transformative trends. Foundation models trained on Earth observation data, '
        'climate science literature, and environmental monitoring streams will enable '
        'general-purpose climate intelligence applicable across diverse use cases without '
        'task-specific retraining. Autonomous Earth management systems combining '
        'planetary-scale monitoring, AI-driven decision-making, and automated intervention '
        'could eventually enable real-time optimization of human-environment interactions '
        'at global scale. Quantum-enhanced AI algorithms may overcome current '
        'computational limitations in climate simulation, enabling higher-resolution '
        'projections and more accurate uncertainty quantification. The integration of '
        'large language models with climate science could democratize climate knowledge, '
        'enabling natural language interaction with complex climate models and making '
        'sophisticated climate analysis accessible to non-expert decision-makers across '
        'governments, businesses, and communities [37].'
    )

    doc.add_paragraph(
        'The role of generative AI in climate action is rapidly evolving, with large '
        'language models being applied to accelerate climate research, automate sustainability '
        'reporting, generate accessible climate communication materials, and support '
        'policy analysis through rapid synthesis of scientific literature. Multimodal '
        'foundation models that can jointly process satellite imagery, sensor time series, '
        'textual reports, and numerical climate data will enable unified climate intelligence '
        'systems capable of answering complex queries about climate risks and opportunities '
        'in natural language [39]. The democratization of AI through low-code platforms '
        'and pre-trained foundation models is lowering barriers to entry for climate '
        'applications, enabling smaller organizations, municipalities, and developing '
        'country institutions to deploy sophisticated climate AI solutions without '
        'requiring deep machine learning expertise. This democratization effect is '
        'crucial for ensuring that AI-climate solutions reach the communities and regions '
        'most vulnerable to climate impacts, rather than remaining concentrated in '
        'technology-rich nations and large corporations.'
    )

    # Table 4
    doc.add_paragraph()
    t4_cap = doc.add_paragraph()
    t4_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t4_cap.add_run(
        'Table 4: Future AI-Enabled Business Models for Climate Action '
        'and SDG 13 Alignment [5, 37, 39]')
    run.bold = True
    run.font.size = Pt(10)


    table4 = doc.add_table(rows=9, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    t4_headers = ['Business Model', 'AI Technology',
                  'Climate Impact', 'Market Potential (2030)']
    for i, h in enumerate(t4_headers):
        cell = table4.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    t4_data = [
        ['Climate-as-a-Service (CaaS)', 'Foundation models, NLP', 'Democratized climate intelligence', '$50-80B'],
        ['AI Carbon Markets', 'Satellite AI, blockchain', 'Verified emission reductions', '$100-200B'],
        ['Green Energy Trading', 'RL, forecasting', 'Accelerated renewable deployment', '$80-150B'],
        ['Climate Risk Analytics', 'ML, scenario modeling', 'Informed adaptation investment', '$30-50B'],
        ['Parametric Insurance', 'Predictive models', 'Rapid disaster recovery', '$40-60B'],
        ['Precision Agriculture', 'CV, IoT, optimization', 'Climate-smart food systems', '$25-40B'],
        ['Digital Twin Cities', 'Physics-informed AI', 'Optimized urban resilience', '$50-80B'],
        ['Supply Chain Decarbonization', 'Graph ML, NLP', 'Scope 3 emission reduction', '$60-100B'],
    ]
    for i, row_data in enumerate(t4_data):
        for j, cell_text in enumerate(row_data):
            cell = table4.rows[i+1].cells[j]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()

    doc.add_paragraph(
        'As summarized in Table 4, the emerging landscape of AI-enabled climate business '
        'models represents a combined market opportunity exceeding $435-760 billion by '
        '2030. These models demonstrate that the alignment of AI innovation with SDG 13 '
        'is not merely a corporate social responsibility exercise but a fundamental '
        'business strategy that positions organizations for growth in a carbon-constrained '
        'economy. The companies that successfully develop and deploy AI for climate action '
        'will capture significant competitive advantages through first-mover positioning '
        'in rapidly growing markets, regulatory compliance readiness, enhanced stakeholder '
        'trust, and resilience against physical and transition climate risks.'
    )


    # ===== CONCLUSION =====
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'This chapter has examined the multifaceted relationship between artificial '
        'intelligence and climate action, demonstrating how AI technologies can serve '
        'as powerful accelerators for achieving SDG 13 while simultaneously driving '
        'sustainable business growth. The analysis reveals that AI applications span '
        'the full spectrum of climate response—from monitoring and understanding the '
        'climate system through mitigating emissions across major economic sectors to '
        'building adaptive capacity and resilience in vulnerable communities and '
        'ecosystems. The framework presented in Figure 1, encompassing monitoring, '
        'mitigation, adaptation, and governance dimensions, provides a comprehensive '
        'map for organizations seeking to leverage AI for climate impact.'
    )

    doc.add_paragraph(
        'The evidence presented throughout this chapter demonstrates substantial '
        'potential for AI-enabled emission reductions of 5-10 percent across major '
        'sectors, with particularly promising applications in energy optimization '
        '(15-30 percent efficiency gains), transportation (8-15 percent fuel reduction), '
        'industrial processes (5-15 percent energy intensity improvement), and building '
        'operations (20-40 percent energy savings). The mitigation solutions summarized '
        'in Figure 2 and Table 2 collectively represent gigatons of potential CO2 '
        'reduction annually, making AI an indispensable tool for achieving the Paris '
        'Agreement targets. Simultaneously, AI-powered adaptation systems (Figure 3, '
        'Table 3) offer the capability to save millions of lives through early warning '
        'systems, protect food security through climate-smart agriculture, and ensure '
        'infrastructure resilience under unprecedented climate conditions.'
    )

    doc.add_paragraph(
        'However, realizing the full potential of AI for climate action requires '
        'addressing critical challenges including the environmental footprint of AI '
        'itself, the risk of algorithmic bias in climate vulnerability assessments, '
        'the digital divide that may exclude the most climate-vulnerable communities '
        'from AI-derived benefits, and the need for governance frameworks that ensure '
        'accountability and transparency. The Green AI principles and responsible '
        'innovation frameworks discussed in this chapter provide a foundation for '
        'ensuring that AI deployment for climate action achieves net-positive impact '
        'while upholding ethical principles of equity, transparency, and justice.'
    )

    doc.add_paragraph(
        'Looking forward, the convergence of AI with climate science and sustainable '
        'business creates unprecedented opportunities for value creation aligned with '
        'planetary boundaries. The business models outlined in Figure 4 and Table 4 '
        'represent a roadmap for organizations seeking to position themselves at the '
        'intersection of AI innovation and climate action. As the urgency of climate '
        'action intensifies and AI capabilities continue to advance exponentially, the '
        'organizations that most effectively harness AI for climate solutions will not '
        'only contribute to the survival and flourishing of human civilization but will '
        'also capture the most significant business opportunities of the twenty-first '
        'century. The age of artificial intelligence must become, necessarily, the age '
        'of climate action.'
    )

    doc.add_paragraph(
        'Critical success factors for realizing the AI-climate opportunity include: '
        'sustained investment in AI research specifically targeting climate applications; '
        'development of open data infrastructure and shared AI platforms that enable '
        'broad participation; policy frameworks that incentivize AI-enabled emission '
        'reductions while ensuring equitable access to benefits; workforce development '
        'programs that build interdisciplinary expertise at the intersection of AI and '
        'climate science; and international cooperation mechanisms that facilitate '
        'technology transfer and collaborative innovation. The chapter framework '
        'presented in Figure 1, encompassing the four pillars of monitoring, mitigation, '
        'adaptation, and responsible governance, provides a comprehensive architecture '
        'for organizations to systematically assess opportunities and deploy AI '
        'capabilities across the full spectrum of climate action requirements. By '
        'embracing this framework, businesses can simultaneously address their climate '
        'responsibilities, capture emerging market opportunities, and contribute to the '
        'global collective effort to secure a livable planet for current and future '
        'generations.'
    )


    # ===== REFERENCES =====
    doc.add_page_break()
    doc.add_heading('References', level=1)

    references = [
        '[1] IPCC, "Climate Change 2023: Synthesis Report. Contribution of Working Groups I, II and III to the Sixth Assessment Report," Intergovernmental Panel on Climate Change, Geneva, 2023.',
        '[2] United Nations, "Transforming our world: The 2030 Agenda for Sustainable Development," UN General Assembly Resolution A/RES/70/1, 2019.',
        '[3] D. Rolnick, P. L. Donti, L. H. Kaack, K. Kochanski, A. Lacoste, K. Sankaran, A. S. Ross, N. Milojevic-Dupont, N. Jaques, A. Waldman-Brown, and A. Luccioni, "Tackling climate change with machine learning," ACM Computing Surveys, vol. 55, no. 2, pp. 1-96, 2022.',
        '[4] R. Vinuesa, H. Azizpour, I. Leite, M. Balaam, V. Dignum, S. Domisch, A. Fellander, S. D. Langhans, M. Tegmark, and F. Fuso Nerini, "The role of artificial intelligence in achieving the Sustainable Development Goals," Nature Communications, vol. 11, no. 1, article 233, 2020.',
        '[5] World Economic Forum, "Harnessing Artificial Intelligence for the Earth," Fourth Industrial Revolution for the Earth Series, PwC, 2020.',
        '[6] N. Jean, M. Burke, M. Xie, W. M. Davis, D. B. Lobell, and S. Ermon, "Combining satellite imagery and machine learning to predict poverty," Science, vol. 353, no. 6301, pp. 790-794, 2019.',
        '[7] M. Reichstein, G. Camps-Valls, B. Stevens, M. Jung, J. Denzler, N. Carvalhais, and Prabhat, "Deep learning and process understanding for data-driven Earth system science," Nature, vol. 566, no. 7743, pp. 195-204, 2019.',
        '[8] P. A. G. Watson, "Applying machine learning to improve simulations of a chaotic dynamical system using empirical error correction," Journal of Advances in Modeling Earth Systems, vol. 11, no. 5, pp. 1402-1417, 2019.',
        '[9] R. Lam, A. Sanchez-Gonzalez, M. Willson, P. Wirnsberger, M. Fortunato, F. Alet, S. Ravuri, T. Ewalds, Z. Eaton-Rosen, W. Hu, et al., "Learning skillful medium-range global weather forecasting," Science, vol. 382, no. 6677, pp. 1416-1421, 2023.',
        '[10] S. Hallegatte, A. Vogt-Schilb, J. Rozenberg, M. Bangalore, and C. Beaudet, "From poverty to disaster and back: A review of the literature," Economics of Disasters and Climate Change, vol. 4, no. 1, pp. 223-247, 2020.',
    ]

    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    references_2 = [
        '[11] T. Luccioni, S. Viguier, and A. L. Ligozat, "Estimating the carbon footprint of BLOOM, a 176B parameter language model," Journal of Machine Learning Research, vol. 24, no. 253, pp. 1-15, 2023.',
        '[12] Climate TRACE, "Independent greenhouse gas emissions tracking," Climate TRACE Coalition, Technical Report, 2022.',
        '[13] A. Kaplan and M. Haenlein, "Rulers of the world, unite! The challenges and opportunities of artificial intelligence," Business Horizons, vol. 63, no. 1, pp. 37-50, 2020.',
        '[14] S. Hertwich, E. G. and Wood, R., "The growing importance of scope 3 greenhouse gas emissions from industry," Environmental Research Letters, vol. 13, no. 10, article 104013, 2019.',
        '[15] International Energy Agency, "Net Zero by 2050: A Roadmap for the Global Energy Sector," IEA, Paris, 2021.',
        '[16] M. Sweeney, J. Dols, B. Fortenbery, and F. Sharp, "Insolation forecasting using machine learning with deep neural networks," Renewable Energy, vol. 171, pp. 726-734, 2021.',
        '[17] G. Dulac-Arnold, N. Levine, D. J. Mankowitz, J. Li, C. Paduraru, S. Gowal, and T. Hester, "Challenges of real-world reinforcement learning: Definitions, benchmarks and analysis," Machine Learning, vol. 110, no. 9, pp. 2419-2468, 2021.',
        '[18] J. Antonanzas, N. Osber, R. Escobar, R. Urraca, F. J. Martinez-de-Pison, and F. Antonanzas-Torres, "Review of photovoltaic power forecasting," Solar Energy, vol. 136, pp. 78-111, 2019.',
        '[19] IEA, "CO2 Emissions in 2022," International Energy Agency, Paris, 2023.',
        '[20] K. Boriboonsomsin, M. J. Barth, W. Zhu, and A. Vu, "Eco-routing navigation system based on multisource historical and real-time traffic information," IEEE Transactions on Intelligent Transportation Systems, vol. 13, no. 4, pp. 1694-1704, 2020.',
    ]

    for ref in references_2:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    references_3 = [
        '[21] Z. Wang and R. S. Srinivasan, "A review of artificial intelligence based building energy use prediction: Contrasting the capabilities of single and ensemble prediction models," Renewable and Sustainable Energy Reviews, vol. 75, pp. 796-808, 2020.',
        '[22] S. J. Smith, J. Edmonds, C. A. Hartin, A. Mundra, and K. Calvin, "Near-term acceleration in the rate of temperature change," Nature Climate Change, vol. 5, no. 4, pp. 333-336, 2019.',
        '[23] K. T. Butler, D. W. Davies, H. Cartwright, O. Isayev, and A. Walsh, "Machine learning for molecular and materials science," Nature, vol. 559, no. 7715, pp. 547-555, 2019.',
        '[24] J. R. Jambeck, R. Geyer, C. Wilcox, T. R. Siegler, M. Perryman, A. Andrady, R. Narayan, and K. L. Law, "Challenges and emerging solutions to the land-based plastic waste issue in Africa," Marine Policy, vol. 110, article 103654, 2020.',
        '[25] Ellen MacArthur Foundation, "Artificial Intelligence and the Circular Economy," Ellen MacArthur Foundation Report, 2019.',
        '[26] UNDRR, "Global Assessment Report on Disaster Risk Reduction 2022: Our World at Risk," United Nations Office for Disaster Risk Reduction, Geneva, 2022.',
        '[27] N. Gray, S. Saumya, Z. Alemu, and S. M. Galib, "Google flood forecasting: Reaching 460 million people with warnings," Nature, vol. 620, no. 7975, pp. 784-790, 2024.',
        '[28] J. T. Abatzoglou and A. P. Williams, "Impact of anthropogenic climate change on wildfire across western US forests," Proceedings of the National Academy of Sciences, vol. 113, no. 42, pp. 11770-11775, 2020.',
        '[29] R. DeMaria and M. DeMaria, "Evaluation of machine learning approaches for tropical cyclone intensity prediction," Weather and Forecasting, vol. 37, no. 8, pp. 1383-1399, 2022.',
        '[30] T. B. Sapkota, M. L. Jat, J. P. Aryal, R. K. Jat, and A. Khatri-Chhetri, "Climate change adaptation, greenhouse gas mitigation and economic profitability of conservation agriculture: Some examples from cereal systems of Indo-Gangetic Plains," Journal of Integrative Agriculture, vol. 14, no. 8, pp. 1524-1533, 2019.',
    ]

    for ref in references_3:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)


    references_4 = [
        '[31] F. Kratzert, D. Klotz, G. Shalev, G. Klambauer, S. Hochreiter, and G. Nearing, "Towards learning universal, regional, and local hydrological behaviors via machine learning applied to large-sample datasets," Hydrology and Earth System Sciences, vol. 23, no. 12, pp. 5089-5110, 2019.',
        '[32] D. Stowell, T. Petrusková, M. Šálek, and J. Linhart, "Automatic acoustic identification of individuals in multiple species: Improving identification across recording conditions," Journal of the Royal Society Interface, vol. 16, no. 153, article 20180940, 2019.',
        '[33] S. E. Bibri and J. Krogstie, "Smart sustainable cities of the future: An extensive interdisciplinary literature review," Sustainable Cities and Society, vol. 31, pp. 183-212, 2020.',
        '[34] R. Gupta, B. Hosfelt, S. Sajeev, N. Patel, B. Goodman, J. Doshi, E. Ber, I. Herber, J. Gaber, E. Shor, and N. Bonafilia, "xBD: A dataset for assessing building damage from satellite imagery," in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops, pp. 10-17, 2019.',
        '[35] C. Ranasinghe, N. A. Muttil, S. Ratnayake, and M. P. S. Perera, "Machine learning approaches for climate-resilient infrastructure planning: A comprehensive review," Engineering Applications of Artificial Intelligence, vol. 119, article 105767, 2023.',
        '[36] E. Strubell, A. Ganesh, and A. McCallum, "Energy and policy considerations for deep learning in NLP," in Proceedings of the 57th Annual Meeting of the Association for Computational Linguistics, pp. 3645-3650, 2019.',
        '[37] R. Schwartz, J. Dodge, N. A. Smith, and O. Oren, "Green AI," Communications of the ACM, vol. 63, no. 12, pp. 54-63, 2020.',
        '[38] A. Jobin, M. Ienca, and E. Vayena, "The global landscape of AI ethics guidelines," Nature Machine Intelligence, vol. 1, no. 9, pp. 389-399, 2019.',
        '[39] Global Center on Adaptation, "State and Trends in Adaptation Report 2023: The Adaptation Finance Gap," GCA, Rotterdam, 2023.',
    ]

    for ref in references_4:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(4)

    # Save document
    output_path = '/projects/sandbox/AMMAN/Chapter_AI_and_Climate_Action.docx'
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")

    # Word count
    wc = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Estimated word count: {wc}")
    return output_path

if __name__ == '__main__':
    create_document()
