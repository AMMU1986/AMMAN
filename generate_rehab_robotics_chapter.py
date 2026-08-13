"""
Generate Complete Word Document: Rehabilitation Robotics Book Chapter
~8300 words, 47 references, 4 tables, 4 figures
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=11, space_after=6, alignment=None):
    """Add a formatted paragraph."""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return para

def add_section_heading(doc, text, level=1):
    """Add a section heading."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_body_text(doc, text):
    """Add body text paragraph."""
    para = doc.add_paragraph(text, style='Normal')
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    run = title.add_run('REHABILITATION ROBOTICS')
    run.bold = True
    run.font.size = Pt(22)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Foundations, Technologies, Clinical Applications, and Future Directions')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('A Comprehensive Book Chapter')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_page_break()
    
    # ============================================================
    # ABSTRACT
    # ============================================================
    add_section_heading(doc, 'Abstract', level=1)
    
    abstract_text = (
        "Rehabilitation robotics represents a rapidly evolving interdisciplinary field at the intersection of "
        "mechanical engineering, biomedical engineering, neuroscience, and clinical rehabilitation science. "
        "This chapter provides a comprehensive overview of the foundations, enabling technologies, clinical "
        "applications, and future directions of rehabilitation robotic systems. Beginning with the historical "
        "evolution from rudimentary mechanical orthoses to modern intelligent systems, the chapter explores "
        "the fundamental principles of human-robot interaction, biomechanics, and therapeutic requirements "
        "that underpin effective robotic rehabilitation. A detailed taxonomy of rehabilitation robots is "
        "presented, encompassing end-effector-based systems, exoskeletons, robotic prostheses, and socially "
        "assistive platforms designed for diverse patient populations including stroke survivors, spinal cord "
        "injury patients, and individuals with neurodegenerative disorders. The chapter examines enabling "
        "technologies including advanced sensing systems, adaptive control algorithms, artificial intelligence, "
        "virtual and augmented reality interfaces, brain-computer interfaces, and Internet of Medical Things "
        "connectivity. Current challenges related to cost, accessibility, clinical acceptance, and regulatory "
        "compliance are critically analyzed alongside emerging solutions. Finally, future directions including "
        "soft robotics, digital twins, collaborative robots, and personalized AI-driven rehabilitation are "
        "discussed, emphasizing the trajectory toward patient-centred, home-based, and autonomous "
        "rehabilitation ecosystems."
    )
    add_body_text(doc, abstract_text)
    
    keywords = doc.add_paragraph()
    run = keywords.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(11)
    run = keywords.add_run(
        'Rehabilitation robotics; Exoskeletons; Human-robot interaction; Artificial intelligence; '
        'Brain-computer interfaces; Telerehabilitation; Soft robotics; Clinical rehabilitation; '
        'Assistive technology; Digital twins'
    )
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION I: FOUNDATIONS AND EVOLUTION
    # ============================================================
    add_section_heading(doc, 'Section I: Foundations and Evolution of Rehabilitation Robotics', level=1)
    
    # Chapter 1
    add_section_heading(doc, '1. Introduction, History, and Evolution of Rehabilitation Robots', level=2)
    
    add_section_heading(doc, '1.1 Historical Development of Rehabilitation Robotics', level=3)
    
    text = (
        "The history of rehabilitation robotics traces its origins to the mid-twentieth century, when the "
        "convergence of advances in mechanical engineering, electronics, and clinical medicine created the "
        "foundations for automated therapeutic systems [1]. The earliest rehabilitation devices were passive "
        "mechanical orthoses designed to support limb positioning and prevent contractures in paralyzed "
        "patients. These rudimentary systems, while lacking active actuation, represented the conceptual "
        "predecessors of modern rehabilitation robots by establishing the principle that mechanical systems "
        "could augment human physical capabilities [2]."
    )
    add_body_text(doc, text)
    
    text = (
        "The philosophical underpinnings of rehabilitation robotics emerged from the recognition that "
        "human motor recovery following neurological injury requires sustained, intensive, and precisely "
        "controlled physical practice that exceeds the capacity of manual therapeutic delivery. Early "
        "investigators observed that the repetitive nature of rehabilitative exercise, combined with "
        "the growing evidence for activity-dependent neural plasticity, created an ideal application "
        "domain for robotic automation. The precision, repeatability, and tirelessness of robotic "
        "systems offered theoretical advantages over purely manual therapy in delivering the high-dose, "
        "consistent training necessary for optimal neural reorganization [1]."
    )
    add_body_text(doc, text)
    
    text = (
        "The 1960s witnessed the emergence of the first electrically powered orthotic devices and robotic "
        "manipulators adapted for individuals with severe physical disabilities. Pioneering work at the "
        "Rancho Los Amigos Hospital in California led to the development of powered arm orthoses that "
        "could assist patients with high-level spinal cord injuries in performing basic activities of daily "
        "living [3]. Simultaneously, researchers at Case Western Reserve University and the Massachusetts "
        "Institute of Technology began exploring the application of industrial robotic technology to "
        "assistive and therapeutic contexts [4]. These early efforts were characterized by bulky, "
        "stationary systems with limited degrees of freedom, but they established critical proof-of-concept "
        "demonstrations that mechanical assistance could meaningfully improve functional outcomes for "
        "individuals with severe motor disabilities."
    )
    add_body_text(doc, text)
    
    text = (
        "The 1970s and 1980s brought incremental but important advances, including the development of "
        "functional electrical stimulation systems that could be combined with mechanical orthoses, "
        "the introduction of microprocessor control for adaptive device behaviour, and growing clinical "
        "evidence supporting technology-assisted rehabilitation approaches. The Veterans Administration "
        "in the United States played a significant role in funding early research, motivated by the "
        "rehabilitation needs of injured service members. Academic-clinical partnerships established "
        "during this era created the interdisciplinary research infrastructure that would later enable "
        "the rapid development of sophisticated rehabilitation robotic systems [4]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '1.2 Major Milestones in Robotics, Biomechanics, and Assistive Technology', level=3)
    
    text = (
        "The development of the MIT-Manus system in the early 1990s marked a transformative milestone "
        "in rehabilitation robotics, establishing the paradigm of robot-assisted therapy for neurological "
        "rehabilitation [5]. This planar robot demonstrated that intensive, repetitive, and goal-directed "
        "robotic training could facilitate neuroplasticity and motor recovery in stroke survivors. The "
        "clinical success of MIT-Manus catalyzed widespread research interest and led to the development "
        "of numerous subsequent systems [6]. As illustrated in Figure 1, the evolution of rehabilitation "
        "robotics reflects a progression from simple mechanical assistance to sophisticated intelligent "
        "systems integrating artificial intelligence, advanced sensing, and cloud connectivity."
    )
    add_body_text(doc, text)
    
    text = (
        "The late 1990s and early 2000s saw the introduction of lower-limb rehabilitation systems, most "
        "notably the Lokomat (Hocoma AG), which combined a body-weight support system with motorized "
        "leg orthoses on a treadmill to provide intensive gait training [7]. This period also witnessed "
        "the emergence of end-effector-based gait training systems such as the Gait Trainer GT I, which "
        "guided foot movement patterns without constraining the entire limb kinematic chain [8]. The "
        "parallel development of robotic prosthetic limbs, exemplified by the DEKA Arm System and the "
        "Modular Prosthetic Limb program, demonstrated the feasibility of restoring complex motor "
        "functions through advanced mechatronic design [9]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '1.3 Evolution from Manually Assisted Therapy to Intelligent Robotic Rehabilitation', level=3)
    
    text = (
        "The transition from manually assisted therapy to intelligent robotic rehabilitation has been "
        "driven by several converging factors: the growing evidence for dose-dependent neuroplasticity, "
        "the clinical workforce limitations in providing high-intensity therapy, and rapid advances in "
        "computing, sensing, and actuation technologies [10]. Early rehabilitation robots functioned "
        "primarily as motorized guides that moved patients' limbs along predetermined trajectories. "
        "Modern systems, in contrast, employ sophisticated control strategies that adapt in real-time "
        "to patient performance, providing assistance only as needed through impedance-based and "
        "assist-as-needed control paradigms [11]. The historical evolution and key milestones "
        "of rehabilitation robotics are depicted in Figure 1, illustrating the progressive integration "
        "of intelligence and adaptivity into therapeutic systems."
    )
    add_body_text(doc, text)
    
    text = (
        "This evolutionary trajectory reflects a fundamental shift in therapeutic philosophy from "
        "passive mobilization to active engagement. Contemporary intelligent rehabilitation systems "
        "continuously assess patient intent, effort, and performance, dynamically modulating their "
        "behavior to challenge patients at their optimal training intensity. The concept of the "
        "challenge point framework, where learning is maximized when task difficulty matches the "
        "learner's current capability, has become a guiding principle for adaptive rehabilitation "
        "robot control design [11]. Furthermore, the integration of gamification elements, performance "
        "feedback, and motivational strategies has transformed rehabilitation robots from mere mechanical "
        "assistants into engaging therapeutic partners that address both physical and psychological "
        "dimensions of recovery."
    )
    add_body_text(doc, text)
    
    # Insert Figure 1
    doc.add_paragraph()
    if os.path.exists('/projects/sandbox/AMMAN/rehab_figures/Figure_1_Evolution_Timeline.png'):
        doc.add_picture('/projects/sandbox/AMMAN/rehab_figures/Figure_1_Evolution_Timeline.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption = doc.add_paragraph()
    fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_caption.add_run('Figure 1: Historical Evolution and Major Milestones in Rehabilitation Robotics')
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph()
    
    # Chapter 2
    add_section_heading(doc, '2. Fundamentals of Rehabilitation Robotics', level=2)
    
    add_section_heading(doc, '2.1 Biomechanics and Human-Robot Interaction', level=3)
    
    text = (
        "The design and operation of rehabilitation robots requires deep understanding of human "
        "biomechanics and the complex dynamics of human-robot physical interaction. Human joints exhibit "
        "nonlinear viscoelastic properties, variable stiffness characteristics, and complex multi-axial "
        "movement patterns that rehabilitation robots must accommodate [12]. The mechanical interface "
        "between robot and patient represents a critical design challenge, as improperly aligned or "
        "overly constrained connections can generate harmful interaction forces, cause discomfort, and "
        "reduce therapeutic efficacy [13]. The human musculoskeletal system possesses inherent redundancy, "
        "with multiple joint configurations capable of achieving a given hand or foot position, and "
        "rehabilitation robots must respect this kinematic freedom rather than imposing artificially "
        "constrained movement paths."
    )
    add_body_text(doc, text)
    
    text = (
        "Biomechanical modeling of the human body segments, including accurate representation of joint "
        "centres of rotation, segment inertial properties, and muscle force generation capabilities, "
        "is essential for proper rehabilitation robot design. Misalignment between robot joint axes "
        "and anatomical joint axes generates parasitic forces at the human-robot interface that increase "
        "with movement amplitude and can cause skin pressure injuries, joint stress, or patient "
        "discomfort [12]. Advanced exoskeleton designs address this challenge through self-aligning "
        "mechanisms, passive alignment degrees of freedom, and redundant kinematic structures that "
        "accommodate inter-individual anatomical variability."
    )
    add_body_text(doc, text)
    
    text = (
        "Effective human-robot interaction in rehabilitation contexts requires consideration of both "
        "physical and cognitive dimensions. Physical interaction encompasses force exchange, kinematic "
        "compatibility, and ergonomic design, while cognitive interaction involves the patient's "
        "understanding of and engagement with the robotic therapy [14]. Modern rehabilitation robots "
        "increasingly incorporate transparent control strategies that allow patients to feel in control "
        "of their movements, thereby promoting active participation and motor learning [15]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '2.2 Principles of Therapeutic and Assistive Robotics', level=3)
    
    text = (
        "Rehabilitation robots serve two distinct but complementary functions: therapeutic and assistive. "
        "Therapeutic robots are designed to facilitate motor recovery through intensive, repetitive, and "
        "goal-directed training, leveraging principles of neuroplasticity and motor learning [16]. These "
        "systems are typically used in clinical settings under professional supervision and aim to restore "
        "lost motor functions. Assistive robots, conversely, compensate for permanent functional deficits "
        "by providing ongoing support for activities of daily living, mobility, and social participation [17]. "
        "The distinction between therapeutic and assistive functions is not always clear-cut, as assistive "
        "use of robotic devices during daily activities may also contribute to motor learning and "
        "functional improvement through repeated practice in meaningful contexts."
    )
    add_body_text(doc, text)
    
    text = (
        "The therapeutic principles underlying rehabilitation robotics derive from neuroscience research "
        "demonstrating that the adult nervous system retains significant capacity for reorganization "
        "following injury. Activity-dependent plasticity, Hebbian learning rules, and use-dependent "
        "cortical reorganization provide the neurobiological basis for robot-assisted motor rehabilitation "
        "[16]. Effective therapeutic robots must therefore promote active patient engagement, provide "
        "appropriate sensory feedback to reinforce successful movement patterns, and deliver training "
        "at intensities sufficient to drive neural adaptation. The principles of motor learning theory, "
        "including practice variability, contextual interference, and knowledge of results, inform "
        "the design of robotic therapy protocols that optimize skill acquisition and retention."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '2.3 Classification and Working Principles of Rehabilitation Robots', level=3)
    
    text = (
        "Rehabilitation robots can be classified along multiple dimensions including anatomical target "
        "(upper limb, lower limb, full body), mechanical architecture (end-effector-based, exoskeleton-based), "
        "functional purpose (therapeutic, assistive, prosthetic), and level of intelligence (passive, "
        "active-adaptive, intelligent) [18]. Figure 2 presents a comprehensive taxonomy of rehabilitation "
        "robotic systems organized by these classification criteria. End-effector-based systems interact "
        "with the patient only at the distal point of contact, allowing natural joint coordination patterns "
        "while providing forces at the hand or foot. Exoskeleton-based systems, in contrast, align their "
        "joints and links with the human skeletal structure, enabling independent control of individual "
        "joints but requiring precise kinematic matching [19]."
    )
    add_body_text(doc, text)
    
    # Insert Figure 2
    doc.add_paragraph()
    if os.path.exists('/projects/sandbox/AMMAN/rehab_figures/Figure_2_Classification.png'):
        doc.add_picture('/projects/sandbox/AMMAN/rehab_figures/Figure_2_Classification.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption = doc.add_paragraph()
    fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_caption.add_run('Figure 2: Classification of Rehabilitation Robotic Systems')
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph()
    
    # Chapter 3
    add_section_heading(doc, '3. Clinical Rehabilitation and Therapeutic Requirements', level=2)
    
    add_section_heading(doc, '3.1 Physiotherapy and Occupational Therapy', level=3)
    
    text = (
        "Rehabilitation robotics must align with established principles of physiotherapy and occupational "
        "therapy to achieve clinically meaningful outcomes. Physiotherapy emphasizes restoration of "
        "movement, strength, and functional mobility through progressive exercise programs, manual "
        "techniques, and modality-based interventions [20]. Occupational therapy focuses on enabling "
        "participation in meaningful daily activities through task-specific training, environmental "
        "adaptation, and compensatory strategies. Rehabilitation robots must support both therapeutic "
        "philosophies by providing versatile platforms that can deliver movement-based training while "
        "also enabling practice of functional tasks relevant to patients' daily lives [21]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '3.2 Repetitive, Intensive, and Task-Oriented Rehabilitation', level=3)
    
    text = (
        "Compelling evidence from neuroscience demonstrates that motor recovery following neurological "
        "injury depends critically on the intensity, repetition, and task-specificity of rehabilitation "
        "training [22]. Studies have shown that stroke survivors require hundreds to thousands of "
        "movement repetitions per session to drive meaningful neuroplastic reorganization, far exceeding "
        "what can be practically delivered through manual therapy alone [23]. Rehabilitation robots "
        "address this challenge by enabling sustained, high-repetition training sessions without "
        "therapist fatigue, while maintaining consistent movement quality and providing objective "
        "performance measurement. Research indicates that conventional therapy sessions typically provide "
        "only 30 to 50 active movement repetitions, whereas robotic therapy can deliver several hundred "
        "repetitions within the same time frame, potentially accelerating the recovery process."
    )
    add_body_text(doc, text)
    
    text = (
        "Task-oriented training, where patients practice functional movements relevant to their daily "
        "activities, has demonstrated superior outcomes compared to abstract movement exercises. "
        "Rehabilitation robots facilitate task-oriented training by providing virtual or physical "
        "task environments that simulate reaching, grasping, lifting, walking, and other functional "
        "activities [22]. The ability to precisely control task difficulty, adjust environmental "
        "constraints, and provide augmented sensory feedback enables systematic progression through "
        "increasingly challenging functional tasks as patient capability improves. Furthermore, "
        "the objective measurement capabilities of robotic systems provide quantitative outcome data "
        "that support evidence-based clinical decision-making and enable precise monitoring of "
        "patient progress over time."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '3.3 Patient-Centred and Personalized Rehabilitation Approaches', level=3)
    
    text = (
        "Contemporary rehabilitation philosophy emphasizes patient-centred care, where therapeutic goals "
        "and interventions are individualized based on each patient's specific impairments, functional "
        "limitations, personal goals, and contextual factors [24]. Advanced rehabilitation robots support "
        "personalization through adaptive algorithms that automatically adjust difficulty, assistance "
        "levels, and exercise parameters based on real-time assessment of patient performance and "
        "progress. Machine learning techniques enable these systems to identify optimal training "
        "parameters for individual patients, potentially outperforming standardized protocols [25]."
    )
    add_body_text(doc, text)
    
    # TABLE 1
    doc.add_paragraph()
    table_title = doc.add_paragraph()
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table_title.add_run('Table 1: Comparison of Major Rehabilitation Robotic Systems')
    run.bold = True
    run.font.size = Pt(10)
    
    table1 = doc.add_table(rows=7, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['System', 'Type', 'Target Limb', 'Control Strategy', 'Clinical Application']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1a5276')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    data = [
        ['MIT-Manus/InMotion', 'End-effector', 'Upper limb', 'Impedance control', 'Stroke rehabilitation'],
        ['Lokomat (Hocoma)', 'Exoskeleton', 'Lower limb', 'Position/impedance', 'Gait training (SCI, stroke)'],
        ['ARMin III', 'Exoskeleton', 'Upper limb', 'Assist-as-needed', 'Stroke, TBI rehabilitation'],
        ['ALEX (Active Leg Exo)', 'Exoskeleton', 'Lower limb', 'Force-field control', 'Gait rehabilitation'],
        ['Armeo Spring/Power', 'End-effector/Exo', 'Upper limb', 'Gravity compensation', 'Neurorehabilitation'],
        ['ReWalk/Ekso', 'Exoskeleton', 'Lower limb', 'Triggered stepping', 'SCI mobility assistance'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table1.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'ebf5fb')
    
    doc.add_paragraph()
    
    text = (
        "Table 1 presents a comparison of major rehabilitation robotic systems currently in clinical use "
        "or advanced research stages, highlighting the diversity of mechanical architectures, control "
        "strategies, and clinical applications within the field."
    )
    add_body_text(doc, text)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION II: TYPES AND CLINICAL APPLICATIONS
    # ============================================================
    add_section_heading(doc, 'Section II: Types and Clinical Applications of Rehabilitation Robots', level=1)
    
    # Chapter 4
    add_section_heading(doc, '4. Robotic Systems for Physical Rehabilitation', level=2)
    
    add_section_heading(doc, '4.1 End-Effector-Based Rehabilitation Systems', level=3)
    
    text = (
        "End-effector-based rehabilitation robots interact with the patient through a single distal "
        "attachment point, typically at the hand or foot, without constraining the intermediate joints "
        "of the limb [26]. This architectural approach offers several advantages: simplified mechanical "
        "design, reduced setup time, no requirement for precise kinematic alignment with human joints, "
        "and the freedom for patients to use natural coordination patterns. The MIT-Manus, subsequently "
        "commercialized as the InMotion ARM robot, exemplifies this approach, providing planar reaching "
        "exercises through a two-degree-of-freedom haptic interface. Clinical trials demonstrated that "
        "chronic stroke survivors receiving robot-assisted therapy with MIT-Manus showed significant "
        "improvements in motor function compared to conventional therapy controls [27]. The system's "
        "ability to record detailed kinematic and kinetic data during each therapy session has also "
        "proven valuable for research purposes, generating large datasets that have informed our "
        "understanding of motor recovery patterns and optimal training parameters."
    )
    add_body_text(doc, text)
    
    text = (
        "More recent end-effector systems have expanded beyond planar movement to provide three-dimensional "
        "workspace training, including the InMotion WRIST module and the KINARM robotic platform. "
        "These systems can assess and train complex upper extremity movements including reaching in "
        "three-dimensional space, wrist pronation-supination, and bimanual coordination tasks [26]. "
        "The simplicity of end-effector designs also makes them particularly amenable to portable and "
        "home-based configurations, as the reduced mechanical complexity translates to lower cost, "
        "lighter weight, and simplified maintenance requirements compared to full exoskeleton systems."
    )
    add_body_text(doc, text)
    
    text = (
        "For lower-limb rehabilitation, end-effector-based systems such as the Gait Trainer GT I and "
        "the Haptic Walker guide the patient's feet along prescribed trajectories, simulating gait "
        "patterns including level walking, stair climbing, and terrain variations. These systems "
        "can accommodate patients with varying anthropometric dimensions without mechanical adjustment, "
        "making them particularly practical for clinical deployment [28]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '4.2 Exoskeletons and Wearable Rehabilitation Robots', level=3)
    
    text = (
        "Exoskeleton-based rehabilitation robots feature an anthropomorphic mechanical structure with "
        "joints aligned to the patient's anatomical joints, enabling independent control of individual "
        "joint movements [29]. Upper-limb exoskeletons such as the ARMin III and the HARMONY system "
        "provide six or more degrees of freedom, allowing training of complex three-dimensional movements "
        "including reaching, grasping, and manipulation tasks. Lower-limb exoskeletons such as the Lokomat, "
        "ReWalk, and Ekso GT provide powered hip and knee actuation for gait training and ambulation "
        "assistance. The classification of these systems within the broader taxonomy of rehabilitation "
        "robots is presented in Figure 2, demonstrating the relationship between mechanical architecture "
        "and functional application."
    )
    add_body_text(doc, text)
    
    text = (
        "Wearable exoskeletons represent a significant advancement, enabling rehabilitation training beyond "
        "the confines of clinical settings. Over-ground walking exoskeletons allow individuals with complete "
        "spinal cord injury to stand and walk independently, providing both therapeutic benefits through "
        "weight-bearing exercise and assistive function for community mobility [30]. Recent developments "
        "in lightweight materials, compact actuators, and efficient power systems have substantially "
        "improved the wearability and practical utility of these devices. Carbon fiber composite "
        "structures have reduced device mass while maintaining structural integrity, and advances in "
        "battery technology have extended operating duration to several hours of continuous use. "
        "The development of series elastic actuators and quasi-direct drive mechanisms has improved "
        "the backdrivability and force transparency of exoskeleton joints, enhancing comfort and "
        "enabling more natural human-robot interaction during daily activities."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '4.3 Upper- and Lower-Limb Rehabilitation Platforms', level=3)
    
    text = (
        "The development of comprehensive rehabilitation platforms that address both upper and lower "
        "extremity deficits reflects the holistic nature of neurological rehabilitation. Upper-limb "
        "platforms typically combine robotic guidance with virtual reality environments that present "
        "engaging and motivating task-oriented exercises [31]. The Armeo Spring system, for example, "
        "integrates a passive spring-based arm support with interactive computer games, enabling "
        "patients with moderate-to-severe upper extremity impairment to practice reaching and grasping "
        "movements in a motivating context. Lower-limb platforms increasingly incorporate body-weight "
        "support, automated gait pattern generation, and real-time biofeedback to optimize gait "
        "retraining outcomes [32]."
    )
    add_body_text(doc, text)
    
    # Chapter 5
    add_section_heading(doc, '5. Assistive and Specialized Robotic Technologies', level=2)
    
    add_section_heading(doc, '5.1 Robotic Prostheses', level=3)
    
    text = (
        "Robotic prostheses represent a distinct category of rehabilitation robotics focused on replacing "
        "lost limb function through advanced mechatronic devices. Modern myoelectric upper-limb prostheses "
        "utilize surface electromyography signals from residual muscles to control multi-articulated "
        "robotic hands with individual finger actuation [33]. Pattern recognition algorithms and "
        "machine learning classifiers enable intuitive control of multiple grip patterns and wrist "
        "movements, approaching the dexterity of the natural hand. Lower-limb powered prostheses "
        "incorporate microprocessor-controlled knee and ankle joints that adapt their behaviour to "
        "terrain, walking speed, and activity context, significantly improving functional mobility "
        "and reducing metabolic cost compared to passive devices [34]. Targeted muscle reinnervation "
        "surgery and osseointegrated implants have further expanded the control and mechanical "
        "interface options for prosthetic users, enabling more intuitive motor control and improved "
        "proprioceptive feedback through direct bone-anchored connections."
    )
    add_body_text(doc, text)
    
    text = (
        "The development of sensory feedback systems for robotic prostheses represents a critical "
        "research frontier, addressing the significant limitation of absent somatosensory information "
        "in conventional prosthetic devices. Tactile sensor arrays integrated into prosthetic fingertips "
        "can detect contact forces, object slip, and surface texture, with this sensory information "
        "communicated to the user through vibrotactile, electrotactile, or neural stimulation feedback "
        "channels [33]. Direct peripheral nerve interfaces and cortical implants offer the possibility "
        "of restoring naturalistic sensory perception, potentially enabling prosthetic users to feel "
        "objects as if through their own hand."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '5.2 Balance and Gait-Training Systems', level=3)
    
    text = (
        "Specialized robotic systems for balance and gait training address the complex sensorimotor "
        "requirements of upright stance and locomotion. These systems range from instrumented treadmills "
        "with perturbation capabilities for balance training to overground robotic gait trainers that "
        "provide pelvic support and leg guidance during walking [35]. The KineAssist and Balance Master "
        "systems exemplify this category, providing controlled postural challenges that systematically "
        "train balance reactions and improve postural stability in populations at risk of falls. "
        "Split-belt treadmills with independent speed control of each leg enable asymmetric gait training "
        "that can address step length asymmetries common in stroke survivors, while perturbation platforms "
        "deliver unexpected balance disturbances that train reactive postural responses essential for "
        "fall prevention in elderly populations."
    )
    add_body_text(doc, text)
    
    text = (
        "The integration of biofeedback mechanisms with balance and gait training systems enhances "
        "therapeutic outcomes by providing patients with real-time information about their movement "
        "patterns, weight distribution, and postural alignment. Visual, auditory, and haptic feedback "
        "channels convey performance information that helps patients develop internal movement "
        "representations and refine their motor strategies [35]. Advanced systems employ machine "
        "learning algorithms to identify optimal perturbation parameters and progression criteria "
        "for individual patients, personalizing the balance training protocol to maximize improvement "
        "while maintaining safety margins appropriate for each patient's current stability level."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '5.3 Socially Assistive Robots for Cognitive and Neurological Rehabilitation', level=3)
    
    text = (
        "Socially assistive robots (SARs) represent an emerging paradigm that leverages social interaction "
        "to motivate and guide patients through rehabilitation exercises without physical contact [36]. "
        "These robots use verbal communication, gestural cues, and facial expressions to provide "
        "encouragement, instruction, and feedback during therapy sessions. Research has demonstrated "
        "the effectiveness of SARs in maintaining patient engagement during repetitive exercises, "
        "reducing feelings of isolation in home-based rehabilitation, and providing cognitive stimulation "
        "for patients with dementia and other neurocognitive disorders [37]."
    )
    add_body_text(doc, text)
    
    # TABLE 2
    doc.add_paragraph()
    table_title = doc.add_paragraph()
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table_title.add_run('Table 2: Clinical Applications of Rehabilitation Robots by Patient Population')
    run.bold = True
    run.font.size = Pt(10)
    
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['Patient Population', 'Primary Impairments', 'Robotic Interventions', 'Evidence Level']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1a5276')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    data2 = [
        ['Stroke (acute/subacute)', 'Hemiparesis, spasticity', 'Upper/lower limb robots, exoskeletons', 'Level I (RCTs)'],
        ['Stroke (chronic)', 'Persistent motor deficit', 'Intensive robot-assisted training', 'Level I-II'],
        ['Spinal cord injury', 'Paraplegia/tetraplegia', 'Gait exoskeletons, arm supports', 'Level II'],
        ['Traumatic brain injury', 'Motor/cognitive deficits', 'Multi-modal robotic therapy', 'Level II-III'],
        ['Cerebral palsy', 'Spastic motor patterns', 'Adaptive paediatric robots', 'Level II-III'],
        ["Parkinson's disease", 'Gait freezing, bradykinesia', 'Rhythmic gait training robots', 'Level II'],
        ['Multiple sclerosis', 'Fatigue, weakness, ataxia', 'Assist-as-needed platforms', 'Level III'],
    ]
    
    for i, row_data in enumerate(data2):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'ebf5fb')
    
    doc.add_paragraph()
    
    # Chapter 6
    add_section_heading(doc, '6. Clinical Applications and Patient Populations', level=2)
    
    add_section_heading(doc, '6.1 Stroke and Spinal Cord Injury Rehabilitation', level=3)
    
    text = (
        "Stroke is the most extensively studied clinical application of rehabilitation robotics, "
        "with numerous randomized controlled trials demonstrating the efficacy of robot-assisted "
        "therapy for upper and lower extremity motor recovery [38]. Meta-analyses have shown that "
        "robot-assisted arm training improves motor function and muscle strength in stroke survivors, "
        "particularly when combined with conventional therapy and delivered at high intensity. For "
        "lower-limb rehabilitation, robotic gait training with body-weight support has shown superior "
        "outcomes compared to conventional overground training in non-ambulatory acute stroke patients, "
        "increasing the probability of independent walking [39]. Table 2 summarizes the clinical "
        "applications of rehabilitation robots across major patient populations, highlighting the "
        "diversity of conditions amenable to robotic intervention."
    )
    add_body_text(doc, text)
    
    text = (
        "The timing of robotic intervention relative to stroke onset significantly influences outcomes. "
        "Evidence suggests that robot-assisted therapy initiated in the acute and subacute phases "
        "(within the first three months post-stroke) may capitalize on enhanced neural plasticity "
        "during this critical recovery window. However, studies also demonstrate meaningful improvements "
        "in chronic stroke survivors (more than six months post-onset), indicating that the brain "
        "retains capacity for reorganization well beyond the traditional recovery plateau [38]. "
        "The optimal dosage, frequency, and duration of robot-assisted therapy remain active areas "
        "of research, with emerging evidence suggesting that higher-intensity protocols may yield "
        "superior outcomes, particularly for patients with moderate-to-severe impairments."
    )
    add_body_text(doc, text)
    
    text = (
        "Spinal cord injury rehabilitation represents another major application domain, where robotic "
        "exoskeletons serve both therapeutic and assistive functions. Locomotor training with robotic "
        "assistance provides intensive, repetitive stepping practice that may promote neural plasticity "
        "in incomplete injuries, while powered exoskeletons enable individuals with complete injuries "
        "to stand and walk independently, with demonstrated benefits for cardiovascular health, bone "
        "density, bowel function, and psychological well-being [40]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '6.2 Traumatic Brain Injury and Cerebral Palsy', level=3)
    
    text = (
        "Traumatic brain injury (TBI) rehabilitation presents unique challenges due to the heterogeneous "
        "nature of cognitive and motor impairments. Rehabilitation robots for TBI populations must "
        "accommodate cognitive limitations, behavioral issues, and variable attention spans while "
        "providing appropriate physical training [41]. Paediatric rehabilitation robotics for cerebral "
        "palsy requires specially designed systems that account for smaller body dimensions, growth "
        "patterns, developmental motor learning principles, and the need for playful, engaging "
        "therapeutic experiences. Systems such as the WREX (Wilmington Robotic Exoskeleton) and "
        "the Lokomat Paediatric have been specifically developed for this population."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '6.3 Neurodegenerative Disorders and Age-Related Mobility Loss', level=3)
    
    text = (
        "Parkinson's disease, multiple sclerosis, and age-related mobility decline represent growing "
        "application areas for rehabilitation robotics. For Parkinson's disease, robotic treadmill "
        "training with external rhythmic cueing has shown promise in addressing gait festination "
        "and freezing episodes [42]. Multiple sclerosis rehabilitation benefits from the ability of "
        "robotic systems to provide graded assistance that accommodates the fatigue sensitivity "
        "characteristic of this condition. The aging population worldwide is driving increasing "
        "interest in robotic systems that can maintain mobility, prevent falls, and support "
        "independent living in older adults through both therapeutic and assistive functions."
    )
    add_body_text(doc, text)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION III: ENABLING TECHNOLOGIES
    # ============================================================
    add_section_heading(doc, 'Section III: Enabling Technologies and Intelligent Rehabilitation', level=1)
    
    # Chapter 7
    add_section_heading(doc, '7. Sensing, Control, and Human-Robot Interaction', level=2)
    
    add_section_heading(doc, '7.1 Force and Motion Sensing', level=3)
    
    text = (
        "Accurate sensing of human-robot interaction forces and patient movement is fundamental to "
        "effective rehabilitation robotics. Multi-axis force/torque sensors embedded at the human-robot "
        "interface measure interaction forces, enabling force-controlled therapy modes and safety "
        "monitoring [43]. Inertial measurement units (IMUs), optical encoders, and motion capture systems "
        "provide real-time measurement of joint angles, velocities, and movement trajectories. "
        "Electromyography (EMG) sensors detect muscle activation patterns, enabling intent estimation "
        "and myoelectric control. The integration of these sensing modalities within the layered "
        "architecture of intelligent rehabilitation systems is illustrated in Figure 3, showing how "
        "sensor data flows through processing, control, and application layers."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '7.2 Force Feedback and Adaptive Control', level=3)
    
    text = (
        "The control architecture of rehabilitation robots has evolved from simple position control "
        "to sophisticated adaptive strategies that respond to patient behavior in real-time. "
        "Impedance control, first applied to rehabilitation by Hogan and colleagues, allows robots "
        "to behave as programmable mechanical environments with adjustable stiffness and damping "
        "properties [5]. Assist-as-needed control algorithms continuously monitor patient effort "
        "and reduce robotic assistance when the patient demonstrates capability, encouraging active "
        "participation and preventing learned non-use. Adaptive controllers using model reference "
        "adaptive control (MRAC) and iterative learning control (ILC) progressively optimize therapy "
        "parameters based on patient response patterns across multiple training sessions [44]."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '7.3 Multimodal Sensing and Patient-Performance Monitoring', level=3)
    
    text = (
        "Advanced rehabilitation systems integrate multiple sensing modalities to create comprehensive "
        "patient performance profiles. Multimodal fusion of kinematic, kinetic, physiological, and "
        "behavioral data enables nuanced assessment of patient status, effort level, and recovery "
        "trajectory. Wearable sensor networks incorporating accelerometers, gyroscopes, EMG electrodes, "
        "heart rate monitors, and galvanic skin response sensors provide continuous monitoring during "
        "and between therapy sessions, supporting longitudinal tracking of recovery progress [43]."
    )
    add_body_text(doc, text)
    
    # Insert Figure 3
    doc.add_paragraph()
    if os.path.exists('/projects/sandbox/AMMAN/rehab_figures/Figure_3_Enabling_Technologies.png'):
        doc.add_picture('/projects/sandbox/AMMAN/rehab_figures/Figure_3_Enabling_Technologies.png', width=Inches(5.8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption = doc.add_paragraph()
    fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_caption.add_run('Figure 3: Layered Architecture of Enabling Technologies for Intelligent Rehabilitation')
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph()
    
    # TABLE 3
    table_title = doc.add_paragraph()
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table_title.add_run('Table 3: Sensing Technologies in Rehabilitation Robotics')
    run.bold = True
    run.font.size = Pt(10)
    
    table3 = doc.add_table(rows=7, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Sensor Type', 'Measurement', 'Application', 'Key Advantage']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1a5276')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    data3 = [
        ['Force/Torque Sensor', 'Interaction forces', 'Force-controlled therapy', 'Direct force measurement'],
        ['IMU', 'Orientation, acceleration', 'Motion tracking', 'Wearable, wireless'],
        ['EMG', 'Muscle activation', 'Intent detection', 'Volitional control signal'],
        ['EEG', 'Brain electrical activity', 'BCI control', 'Non-invasive neural input'],
        ['Optical Encoder', 'Joint angles', 'Position control', 'High resolution'],
        ['Pressure Sensor', 'Contact pressure', 'Gait analysis, safety', 'Interface monitoring'],
    ]
    
    for i, row_data in enumerate(data3):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'ebf5fb')
    
    doc.add_paragraph()
    
    text = (
        "Table 3 summarizes the principal sensing technologies employed in rehabilitation robotics, "
        "their measurement capabilities, and their specific applications within the rehabilitation "
        "technology architecture shown in Figure 3."
    )
    add_body_text(doc, text)
    
    # Chapter 8
    add_section_heading(doc, '8. Artificial Intelligence and Digital Technologies', level=2)
    
    add_section_heading(doc, '8.1 Artificial Intelligence and Machine Learning', level=3)
    
    text = (
        "Artificial intelligence (AI) and machine learning (ML) are transforming rehabilitation robotics "
        "by enabling systems that learn from data, adapt to individual patients, and make intelligent "
        "decisions about therapy delivery [25]. Supervised learning algorithms classify patient movement "
        "patterns, predict recovery trajectories, and identify optimal therapy parameters from historical "
        "data. Reinforcement learning enables rehabilitation robots to autonomously discover effective "
        "therapy strategies through trial-and-error interaction with patients, optimizing long-term "
        "outcomes without explicit programming [45]. Deep learning architectures process high-dimensional "
        "sensor data for movement quality assessment, anomaly detection, and real-time adaptation "
        "of therapy difficulty."
    )
    add_body_text(doc, text)
    
    text = (
        "Transfer learning and domain adaptation techniques address the challenge of limited patient-specific "
        "data by leveraging knowledge from related tasks or populations to improve model performance for "
        "individual patients. Federated learning approaches enable collaborative model training across "
        "multiple clinical sites without sharing sensitive patient data, addressing privacy concerns "
        "while benefiting from larger and more diverse training datasets [45]. Natural language processing "
        "technologies are being integrated into socially assistive rehabilitation systems, enabling "
        "conversational interaction between patients and robotic therapists that can provide verbal "
        "instruction, encouragement, and feedback during therapy sessions. The combination of these "
        "AI technologies creates increasingly autonomous rehabilitation systems capable of conducting "
        "therapy sessions with minimal human supervision, while maintaining safety and clinical "
        "effectiveness."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '8.2 Computer Vision and Data Analytics', level=3)
    
    text = (
        "Computer vision technologies enable markerless motion analysis, facial expression recognition "
        "for pain and engagement assessment, and gesture-based interaction with rehabilitation systems. "
        "Deep learning-based pose estimation algorithms such as OpenPose and MediaPipe can track body "
        "joint positions from standard camera images, providing affordable and accessible movement "
        "analysis for home-based rehabilitation [46]. Data analytics platforms aggregate longitudinal "
        "patient data across institutions, enabling population-level analyses of rehabilitation outcomes, "
        "identification of prognostic factors, and development of evidence-based clinical decision "
        "support systems."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '8.3 Internet of Medical Things, Cloud Systems, and Intelligent Decision Support', level=3)
    
    text = (
        "The Internet of Medical Things (IoMT) creates connected ecosystems of rehabilitation devices, "
        "wearable sensors, and clinical information systems that enable seamless data exchange and "
        "coordinated care delivery. Cloud-based platforms provide centralized data storage, computational "
        "resources for AI algorithms, and remote access for clinicians and patients [47]. Intelligent "
        "clinical decision support systems synthesize multimodal patient data to recommend optimal "
        "therapy protocols, predict outcomes, identify patients at risk of complications, and support "
        "clinical decision-making through evidence-based recommendations."
    )
    add_body_text(doc, text)
    
    text = (
        "The architecture of IoMT-enabled rehabilitation systems typically comprises edge computing "
        "devices for real-time local processing, secure communication protocols for data transmission, "
        "cloud storage and computing infrastructure, and application layers for clinical interfaces "
        "and patient portals. Interoperability standards such as HL7 FHIR and IEEE 11073 enable "
        "integration of rehabilitation robot data with electronic health records, creating unified "
        "patient information systems that support coordinated multidisciplinary care [47]. Edge-cloud "
        "hybrid architectures balance the requirements for low-latency real-time control with the "
        "computational demands of AI-based analysis and personalization, ensuring responsive robot "
        "behavior while leveraging powerful cloud resources for complex decision-making tasks."
    )
    add_body_text(doc, text)
    
    # Chapter 9
    add_section_heading(doc, '9. Immersive and Neurotechnology-Based Rehabilitation', level=2)
    
    add_section_heading(doc, '9.1 Virtual Reality and Augmented Reality', level=3)
    
    text = (
        "Virtual reality (VR) and augmented reality (AR) technologies provide immersive, engaging "
        "environments for rehabilitation training that enhance motivation, provide meaningful task "
        "contexts, and enable systematic manipulation of training parameters [31]. VR-based rehabilitation "
        "combines robotic movement assistance with interactive virtual environments where patients "
        "practice reaching for objects, navigating virtual spaces, or playing therapeutic games. "
        "AR systems overlay digital information onto the real world, providing visual feedback about "
        "movement quality, target positions, and performance metrics during physical activities. "
        "The gamification of rehabilitation through VR and AR has demonstrated significant improvements "
        "in patient adherence and motivation, with studies reporting increased therapy duration and "
        "more positive attitudes toward rehabilitation in patients using immersive technologies "
        "compared to conventional approaches."
    )
    add_body_text(doc, text)
    
    text = (
        "Head-mounted VR displays create fully immersive virtual environments that can transport "
        "patients to engaging scenarios while performing therapeutic movements, such as navigating "
        "a virtual kitchen to practice reaching and grasping, or walking through a virtual park "
        "for gait rehabilitation. Haptic feedback integrated with VR provides tactile sensations "
        "that reinforce interaction with virtual objects, enhancing the sense of presence and "
        "ecological validity of the training environment [31]. Mixed reality systems that blend "
        "physical robotic devices with virtual enhancements represent a particularly promising "
        "approach, combining the physical support and measurement capabilities of robots with "
        "the motivational and contextual benefits of immersive virtual environments."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '9.2 Brain-Computer Interfaces', level=3)
    
    text = (
        "Brain-computer interfaces (BCIs) enable direct neural control of rehabilitation robots, "
        "creating a pathway for motor rehabilitation in patients with severe paralysis who cannot "
        "generate voluntary movements. EEG-based BCIs detect motor imagery or motor intention signals "
        "from the brain and translate them into commands for robotic devices, enabling patients to "
        "actively participate in therapy through mental effort alone [36]. This neural engagement "
        "during robot-assisted movement is hypothesized to strengthen residual neural pathways and "
        "promote cortical reorganization. Recent advances in implantable BCI technology have "
        "demonstrated remarkable control fidelity, enabling complex multi-degree-of-freedom robotic "
        "manipulation through intracortical recordings."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '9.3 Telerehabilitation and Remote Robotic Therapy', level=3)
    
    text = (
        "Telerehabilitation leverages telecommunications technology to deliver rehabilitation services "
        "remotely, addressing barriers of geographic distance, transportation limitations, and clinician "
        "availability. Robotic telerehabilitation systems enable patients to receive robot-assisted "
        "therapy at home while clinicians remotely monitor performance, adjust therapy parameters, "
        "and provide guidance through videoconferencing [47]. The COVID-19 pandemic accelerated "
        "adoption of telerehabilitation approaches, demonstrating both feasibility and clinical "
        "effectiveness when in-person services were disrupted. Advanced haptic telerehabilitation "
        "systems transmit force feedback between remote locations, enabling therapists to physically "
        "guide patient movements from a distance."
    )
    add_body_text(doc, text)
    
    text = (
        "The technical architecture of robotic telerehabilitation systems must address significant "
        "challenges including network latency, bandwidth limitations, data security, and the need "
        "for robust local safety systems that function independently of network connectivity. "
        "Time-delay compensation algorithms and predictive control strategies mitigate the effects "
        "of communication latency on haptic interaction quality, while local safety controllers "
        "ensure that the robotic system remains safe even during network interruptions. Encrypted "
        "data transmission, secure authentication protocols, and compliance with healthcare data "
        "regulations such as HIPAA and GDPR are essential for protecting sensitive patient "
        "information in networked rehabilitation systems. The integration of artificial intelligence "
        "with telerehabilitation enables semi-autonomous therapy delivery where the robot conducts "
        "routine exercises independently, with clinician oversight focused on clinical decision-making, "
        "progress evaluation, and therapeutic goal adjustment during periodic virtual consultations."
    )
    add_body_text(doc, text)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION IV: CHALLENGES AND FUTURE
    # ============================================================
    add_section_heading(doc, 'Section IV: Challenges, Emerging Trends, and Future Directions', level=1)
    
    # Chapter 10
    add_section_heading(doc, '10. Current Challenges and Barriers', level=2)
    
    add_section_heading(doc, '10.1 Cost, Accessibility, and Ergonomic Limitations', level=3)
    
    text = (
        "Despite demonstrated clinical benefits, rehabilitation robots face significant barriers to "
        "widespread adoption. The high cost of rehabilitation robotic systems, typically ranging from "
        "$50,000 to over $500,000 for clinical-grade devices, restricts access to well-funded "
        "rehabilitation centres and limits availability for patients in resource-constrained settings "
        "[10]. Ergonomic limitations including device weight, donning/doffing time, limited "
        "adjustability for different body sizes, and restricted workspace continue to affect patient "
        "comfort and therapy efficiency. The complexity of setup and operation often requires "
        "specialized technical staff, adding operational costs and limiting deployment in smaller "
        "clinical facilities."
    )
    add_body_text(doc, text)
    
    text = (
        "The economic sustainability of rehabilitation robotics programmes depends on demonstrating "
        "cost-effectiveness relative to conventional therapy approaches. While the per-session cost "
        "of robot-assisted therapy may exceed that of conventional physiotherapy due to equipment "
        "amortization, the ability to provide higher-intensity training, treat multiple patients "
        "with reduced therapist time, and potentially accelerate recovery timelines may offer "
        "net economic benefits from a health system perspective [10]. However, robust health-economic "
        "analyses comparing the lifetime costs and outcomes of robotic versus conventional rehabilitation "
        "remain limited, hindering evidence-based reimbursement decisions and institutional "
        "investment planning."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '10.2 Clinical Acceptance, Interoperability, and Regulatory Requirements', level=3)
    
    text = (
        "Clinical acceptance of rehabilitation robots varies considerably across institutions and "
        "geographic regions, influenced by factors including clinician training, evidence awareness, "
        "reimbursement policies, and cultural attitudes toward technology-assisted care. "
        "Interoperability challenges arise from proprietary data formats, incompatible communication "
        "protocols, and lack of standardized outcome measures across different robotic systems [18]. "
        "Regulatory requirements for medical robotic devices impose substantial development costs "
        "and time, with certification processes varying significantly between jurisdictions. The "
        "classification of rehabilitation robots as medical devices subjects them to rigorous safety "
        "standards including IEC 60601 for medical electrical equipment and ISO 13482 for personal "
        "care robots."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '10.3 Safety, Reliability, and Clinician Training', level=3)
    
    text = (
        "Safety remains paramount in rehabilitation robotics, as these systems physically interact "
        "with vulnerable patient populations. Hardware safety mechanisms including force limiters, "
        "emergency stops, and mechanical compliance must be complemented by software-based safety "
        "monitoring that detects abnormal forces, velocities, or patient physiological responses "
        "[29]. Reliability requirements demand extensive testing and validation, as system failures "
        "during therapy could result in patient injury or loss of clinical confidence. Comprehensive "
        "clinician training programs are essential to ensure safe and effective operation, yet "
        "training infrastructure and standardized curricula for rehabilitation robotics remain "
        "underdeveloped in many regions."
    )
    add_body_text(doc, text)
    
    # TABLE 4
    doc.add_paragraph()
    table_title = doc.add_paragraph()
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table_title.add_run('Table 4: Current Challenges and Proposed Solutions in Rehabilitation Robotics')
    run.bold = True
    run.font.size = Pt(10)
    
    table4 = doc.add_table(rows=7, cols=3)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['Challenge Category', 'Specific Barriers', 'Proposed Solutions']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1a5276')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    data4 = [
        ['Economic', 'High device cost ($50K-$500K+)', 'Modular designs, shared-use models, soft robotics'],
        ['Accessibility', 'Limited to specialized centres', 'Home-based systems, telerehabilitation'],
        ['Clinical', 'Low adoption, training gaps', 'Standardized curricula, evidence dissemination'],
        ['Technical', 'Interoperability, data silos', 'Open standards, cloud platforms, APIs'],
        ['Regulatory', 'Complex certification paths', 'Harmonized international standards'],
        ['Safety', 'Risk of interaction injuries', 'Compliant actuators, AI-based monitoring'],
    ]
    
    for i, row_data in enumerate(data4):
        for j, cell_text in enumerate(row_data):
            cell = table4.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'ebf5fb')
    
    doc.add_paragraph()
    
    text = (
        "Table 4 presents a structured analysis of current challenges facing the rehabilitation "
        "robotics field alongside proposed solutions derived from ongoing research and development "
        "efforts, as discussed in the context of Table 1 comparative systems analysis."
    )
    add_body_text(doc, text)
    
    # Chapter 11
    add_section_heading(doc, '11. Emerging Technologies and Future Research', level=2)
    
    add_section_heading(doc, '11.1 Soft Robotics and Lightweight Rehabilitation Systems', level=3)
    
    text = (
        "Soft robotics represents a paradigm shift in rehabilitation robot design, replacing rigid "
        "metallic structures with compliant, lightweight materials that inherently conform to the human "
        "body and provide safe physical interaction [30]. Soft pneumatic actuators, shape memory alloys, "
        "and dielectric elastomer actuators enable the creation of wearable rehabilitation devices that "
        "are lightweight, comfortable, and inherently safe. Soft robotic gloves for hand rehabilitation "
        "demonstrate this approach, providing finger flexion and extension assistance through pneumatic "
        "bladders or cable-driven mechanisms integrated into a textile glove. These systems represent "
        "the emerging future directions illustrated in Figure 4, where bio-inspired and compliant "
        "systems converge with intelligent control."
    )
    add_body_text(doc, text)
    
    text = (
        "The inherent compliance of soft robotic systems provides passive safety through mechanical "
        "properties that limit interaction forces without requiring active force monitoring or "
        "electronic safety systems. This property makes soft rehabilitation robots particularly "
        "suitable for home-based applications where professional supervision may be intermittent "
        "or absent [30]. Furthermore, the textile-like properties of many soft robotic designs "
        "enhance social acceptability, as these devices can be worn under clothing without visible "
        "mechanical apparatus, reducing the stigma sometimes associated with assistive technology. "
        "However, soft robotic systems face challenges including limited force output, imprecise "
        "position control, and reduced durability compared to rigid systems, requiring continued "
        "research in materials science and control engineering to realize their full potential."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '11.2 Digital Twins and Cloud Robotics', level=3)
    
    text = (
        "Digital twin technology creates virtual replicas of physical rehabilitation systems and "
        "patient biomechanical models, enabling simulation-based therapy planning, predictive "
        "maintenance, and personalized treatment optimization [45]. Cloud robotics architectures "
        "distribute computational intelligence across local robotic devices and remote cloud servers, "
        "enabling resource-constrained rehabilitation devices to access powerful AI capabilities "
        "through network connectivity. This distributed architecture supports fleet learning, where "
        "experience from multiple rehabilitation robots across institutions is aggregated to "
        "continuously improve therapy algorithms."
    )
    add_body_text(doc, text)
    
    text = (
        "Patient-specific digital twins incorporate musculoskeletal models calibrated with individual "
        "biomechanical parameters, enabling simulation of treatment outcomes before physical implementation. "
        "Clinicians can test different therapy protocols, predict patient responses, and optimize "
        "intervention strategies in the virtual domain before applying them to the actual patient, "
        "reducing trial-and-error in clinical practice and potentially improving outcomes [45]. "
        "The integration of digital twin technology with real-time sensor data creates continuously "
        "updated models that track patient progress, predict future recovery trajectories, and "
        "alert clinicians to deviations from expected patterns that may warrant therapeutic adjustment. "
        "Cloud-based platforms hosting these digital twins enable collaborative clinical decision-making, "
        "where multidisciplinary teams can access and contribute to the patient's rehabilitation plan "
        "regardless of geographic location."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '11.3 Collaborative Robots and Personalized Adaptive Rehabilitation', level=3)
    
    text = (
        "Collaborative robots (cobots) designed for safe human interaction without protective barriers "
        "are increasingly adapted for rehabilitation applications. These systems leverage advanced "
        "force sensing, compliant actuation, and sophisticated collision detection to ensure safety "
        "during close physical interaction with patients [44]. Personalized adaptive rehabilitation "
        "represents the convergence of AI-driven assessment, individualized goal setting, and "
        "real-time therapy adaptation, creating systems that continuously optimize the rehabilitation "
        "experience for each unique patient. The integration of these emerging technologies within "
        "the broader landscape of future rehabilitation robotics is depicted in Figure 4."
    )
    add_body_text(doc, text)
    
    text = (
        "The cobot approach offers particular advantages for rehabilitation settings where robots "
        "must share workspace with patients, therapists, and other clinical staff. Unlike industrial "
        "robots that operate in segregated environments, rehabilitation cobots must anticipate and "
        "respond to unpredictable human movements, detecting incipient collisions and responding "
        "with appropriate force limiting before harm occurs [44]. Machine learning algorithms trained "
        "on human motion data enable predictive safety systems that anticipate patient movements "
        "and pre-emptively adjust robot behavior, providing a proactive rather than merely reactive "
        "safety paradigm. The combination of collaborative robotic hardware with AI-driven "
        "personalization creates systems that can serve as autonomous therapy assistants, conducting "
        "prescribed exercise sessions while continuously adapting to patient responses and maintaining "
        "safety margins appropriate for each individual's physical capabilities and limitations."
    )
    add_body_text(doc, text)
    
    # Insert Figure 4
    doc.add_paragraph()
    if os.path.exists('/projects/sandbox/AMMAN/rehab_figures/Figure_4_Future_Directions.png'):
        doc.add_picture('/projects/sandbox/AMMAN/rehab_figures/Figure_4_Future_Directions.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption = doc.add_paragraph()
    fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_caption.add_run('Figure 4: Emerging Trends and Future Directions in Rehabilitation Robotics')
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph()
    
    # Chapter 12
    add_section_heading(doc, '12. Towards Patient-Centred Intelligent Rehabilitation', level=2)
    
    add_section_heading(doc, '12.1 Home-Based and Autonomous Rehabilitation', level=3)
    
    text = (
        "The future of rehabilitation robotics is increasingly oriented toward home-based systems that "
        "enable patients to receive intensive therapy in their own environments without requiring "
        "frequent visits to clinical facilities. Home-based rehabilitation robots must balance "
        "therapeutic capability with practical requirements including compact size, simple setup, "
        "autonomous operation, and robust safety systems that function without professional supervision "
        "[47]. Cloud-connected home rehabilitation systems can transmit therapy data to clinicians "
        "for remote monitoring, receive updated therapy protocols, and alert healthcare providers "
        "to concerning changes in patient performance or safety events."
    )
    add_body_text(doc, text)
    
    text = (
        "The design philosophy for home-based rehabilitation robots differs fundamentally from "
        "clinical systems, prioritizing user-friendliness, aesthetic integration with home environments, "
        "and autonomous operation over maximum therapeutic versatility. These systems must accommodate "
        "patients with cognitive impairments who may struggle with complex setup procedures, and must "
        "function safely in uncontrolled home environments that lack the protective infrastructure "
        "of clinical settings [47]. Voice-activated control, automatic calibration routines, and "
        "simplified user interfaces based on touchscreen tablets or gesture recognition reduce "
        "the technical burden on patients and caregivers. Remote firmware updates, predictive "
        "maintenance algorithms, and automated performance verification ensure that home-based "
        "systems maintain optimal function over extended deployment periods without requiring "
        "technical service visits."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '12.2 Data-Driven Personalization and Continuous Monitoring', level=3)
    
    text = (
        "Data-driven personalization leverages machine learning algorithms to analyze individual patient "
        "data and optimize therapy parameters for maximum recovery. Continuous monitoring through "
        "wearable sensors and connected rehabilitation devices provides unprecedented longitudinal "
        "data about patient function, activity levels, and recovery trajectories [46]. This data "
        "enables predictive models that forecast recovery outcomes, identify patients who may benefit "
        "from therapy modifications, and support shared decision-making between patients and clinicians. "
        "The integration of electronic health records, robotic therapy data, and wearable sensor "
        "information creates comprehensive patient digital profiles that inform personalized "
        "rehabilitation planning."
    )
    add_body_text(doc, text)
    
    text = (
        "Advanced analytics applied to continuous monitoring data can detect subtle changes in motor "
        "function that may not be apparent during periodic clinical assessments, enabling earlier "
        "intervention when recovery plateaus or complications arise. Pattern recognition algorithms "
        "identify temporal trends, diurnal variations, and environmental factors that influence patient "
        "performance, providing insights that inform both clinical decisions and self-management "
        "strategies [46]. The concept of ecological momentary assessment, where patient function is "
        "measured in real-world contexts rather than controlled clinical environments, provides "
        "more ecologically valid outcome data that better reflects patients' actual functional "
        "capabilities and challenges in their daily lives."
    )
    add_body_text(doc, text)
    
    add_section_heading(doc, '12.3 Interdisciplinary Collaboration and the Future of Rehabilitation Robotics', level=3)
    
    text = (
        "The continued advancement of rehabilitation robotics requires sustained interdisciplinary "
        "collaboration among engineers, clinicians, neuroscientists, computer scientists, and patients "
        "themselves. User-centred design methodologies that actively involve patients and clinicians "
        "in the development process are essential to creating systems that address real clinical needs "
        "and integrate effectively into existing care workflows [21]. The convergence of advances in "
        "materials science, artificial intelligence, neurotechnology, and telecommunications is creating "
        "unprecedented opportunities for rehabilitation robotics to deliver personalized, accessible, "
        "and effective rehabilitation at scale. As the global population ages and the burden of "
        "neurological and musculoskeletal disability increases, rehabilitation robotics will play "
        "an increasingly central role in healthcare systems worldwide, transforming how motor "
        "recovery is achieved and functional independence is maintained across the lifespan."
    )
    add_body_text(doc, text)
    
    text = (
        "The vision for the next decade encompasses fully autonomous home rehabilitation ecosystems, "
        "where intelligent robotic systems seamlessly integrate with smart home infrastructure, "
        "wearable health monitoring, and telemedicine platforms to provide continuous, personalized "
        "rehabilitation support. These systems will leverage digital twin technology for predictive "
        "therapy planning, federated learning for privacy-preserving improvement from collective "
        "patient data, and advanced soft robotic designs that are comfortable enough for all-day "
        "wear and powerful enough for effective therapy delivery."
    )
    add_body_text(doc, text)
    
    text = (
        "The democratization of rehabilitation robotics through reduced costs, simplified designs, "
        "and telehealth integration has the potential to dramatically improve access to intensive "
        "rehabilitation services, particularly for underserved populations in rural areas and "
        "developing countries. Open-source hardware and software platforms for rehabilitation "
        "robotics are emerging, enabling local adaptation and cost-effective manufacturing in "
        "resource-limited settings. The combination of low-cost sensor technologies, affordable "
        "computing platforms, and cloud-based AI services creates opportunities for developing "
        "effective rehabilitation robots at price points accessible to individual consumers "
        "and small community clinics, not merely large academic medical centres."
    )
    add_body_text(doc, text)
    
    text = (
        "In conclusion, rehabilitation robotics stands at a transformative juncture where the "
        "convergence of mechanical engineering, artificial intelligence, neuroscience, materials "
        "science, and telecommunications is enabling unprecedented capabilities for motor "
        "rehabilitation and functional assistance. The field has progressed from simple motorized "
        "guides to sophisticated intelligent systems capable of personalized, adaptive, and "
        "autonomous therapeutic intervention. While significant challenges remain in cost, "
        "accessibility, and clinical integration, the trajectory of technological advancement "
        "and growing clinical evidence base suggest that rehabilitation robots will become "
        "increasingly central to healthcare delivery for individuals with motor disabilities. "
        "Continued interdisciplinary collaboration, patient-centred design, and rigorous clinical "
        "validation will be essential to realizing the full potential of these transformative "
        "technologies in improving quality of life for millions of individuals worldwide affected "
        "by neurological and musculoskeletal conditions."
    )
    add_body_text(doc, text)
    
    doc.add_page_break()
    
    # ============================================================
    # REFERENCES
    # ============================================================
    add_section_heading(doc, 'References', level=1)
    
    references = [
        "[1] Krebs, H.I., Hogan, N., Aisen, M.L., and Volpe, B.T. (1998). Robot-aided neurorehabilitation. IEEE Transactions on Rehabilitation Engineering, 6(1), 75-87.",
        "[2] Colombo, G., Joerg, M., Schreier, R., and Dietz, V. (2000). Treadmill training of paraplegic patients using a robotic orthosis. Journal of Rehabilitation Research and Development, 37(6), 693-700.",
        "[3] Lum, P.S., Burgar, C.G., Shor, P.C., Majmundar, M., and Van der Loos, M. (2002). Robot-assisted movement training compared with conventional therapy techniques for the rehabilitation of upper-limb motor function after stroke. Archives of Physical Medicine and Rehabilitation, 83(7), 952-959.",
        "[4] Prange, G.B., Jannink, M.J., Groothuis-Oudshoorn, C.G., Hermens, H.J., and IJzerman, M.J. (2006). Systematic review of the effect of robot-aided therapy on recovery of the hemiparetic arm after stroke. Journal of Rehabilitation Research and Development, 43(2), 171-184.",
        "[5] Hogan, N., Krebs, H.I., Charnnarong, J., Srikrishna, P., and Sharon, A. (1992). MIT-MANUS: A workstation for manual therapy and training. Proceedings of the IEEE International Workshop on Robot and Human Communication, 161-165.",
        "[6] Volpe, B.T., Krebs, H.I., Hogan, N., Edelstein, L., Diels, C., and Aisen, M. (2000). A novel approach to stroke rehabilitation: Robot-aided sensorimotor stimulation. Neurology, 54(10), 1938-1944.",
        "[7] Colombo, G., Wirz, M., and Dietz, V. (2001). Driven gait orthosis for improvement of locomotor training in paraplegic patients. Spinal Cord, 39(5), 252-255.",
        "[8] Hesse, S., and Uhlenbrock, D. (2000). A mechanized gait trainer for restoration of gait. Journal of Rehabilitation Research and Development, 37(6), 701-708.",
        "[9] Resnik, L., Klinger, S.L., and Etter, K. (2014). The DEKA Arm: Its features, functionality, and evolution during the Veterans Affairs Study to optimize the DEKA Arm. Prosthetics and Orthotics International, 38(6), 492-504.",
        "[10] Maciejasz, P., Eschweiler, J., Gerlach-Hahn, K., Jansen-Troy, A., and Leonhardt, S. (2014). A survey on robotic devices for upper limb rehabilitation. Journal of NeuroEngineering and Rehabilitation, 11(1), 3.",
        "[11] Basteris, A., Nijenhuis, S.M., Stienen, A.H., Buurke, J.H., Prange, G.B., and Amirabdollahian, F. (2014). Training modalities in robot-mediated upper limb rehabilitation in stroke: A framework for classification based on a systematic review. Journal of NeuroEngineering and Rehabilitation, 11(1), 111.",
        "[12] Flash, T., and Hogan, N. (1985). The coordination of arm movements: An experimentally confirmed mathematical model. Journal of Neuroscience, 5(7), 1688-1703.",
        "[13] Jarrasse, N., Tagliabue, M., Robertson, J.V., Maiza, A., Crocher, V., Roby-Brami, A., and Morel, G. (2010). A methodology to quantify alterations in human upper limb movement during co-manipulation with an exoskeleton. IEEE Transactions on Neural Systems and Rehabilitation Engineering, 18(4), 389-397.",
        "[14] Krakauer, J.W. (2006). Motor learning: Its relevance to stroke recovery and neurorehabilitation. Current Opinion in Neurology, 19(1), 84-90.",
        "[15] Marchal-Crespo, L., and Reinkensmeyer, D.J. (2009). Review of control strategies for robotic movement training after neurologic injury. Journal of NeuroEngineering and Rehabilitation, 6(1), 20.",
        "[16] Langhorne, P., Bernhardt, J., and Kwakkel, G. (2011). Stroke rehabilitation. The Lancet, 377(9778), 1693-1702.",
        "[17] Turchetti, G., Vitiello, N., Triber, L., Scattolini, M., Lenzi, T., and Carrozza, M.C. (2014). Technology and innovative strategies for rehabilitation. In Lupu, R.G. et al. (Eds.), Advanced Technologies for Enhanced Quality of Life, IEEE, 68-72.",
        "[18] Siciliano, B., and Khatib, O. (Eds.) (2016). Springer Handbook of Robotics (2nd ed.). Springer International Publishing.",
        "[19] Nef, T., Guidali, M., and Riener, R. (2009). ARMin III—arm therapy exoskeleton with an ergonomic shoulder actuation. Applied Bionics and Biomechanics, 6(2), 127-142.",
        "[20] Veerbeek, J.M., van Wegen, E., van Peppen, R., van der Wees, P.J., Hendriks, E., Rietberg, M., and Kwakkel, G. (2014). What is the evidence for physical therapy poststroke? A systematic review and meta-analysis. PLoS ONE, 9(2), e87987.",
        "[21] Stephenson, A., and Stephens, J. (2018). An exploration of physiotherapists' experiences of robotic therapy in upper limb rehabilitation within a stroke rehabilitation centre. Disability and Rehabilitation: Assistive Technology, 13(3), 245-252.",
        "[22] Kwakkel, G., van Peppen, R., Wagenaar, R.C., Wood Dauphinee, S., Richards, C., Ashburn, A., and Langhorne, P. (2004). Effects of augmented exercise therapy time after stroke: A meta-analysis. Stroke, 35(11), 2529-2539.",
        "[23] Lang, C.E., Macdonald, J.R., Reisman, D.S., Boyd, L., Jacobson Kimberley, T., Schindler-Ivens, S.M., and Bhatt, E. (2009). Observation of amounts of movement practice provided during stroke rehabilitation. Archives of Physical Medicine and Rehabilitation, 90(10), 1692-1698.",
        "[24] Mehrholz, J., Pohl, M., Platz, T., Kugler, J., and Elsner, B. (2018). Electromechanical and robot-assisted arm training for improving activities of daily living, arm function, and arm muscle strength after stroke. Cochrane Database of Systematic Reviews, (9), CD006876.",
        "[25] Qian, Z., and Bi, Z. (2015). Recent development of rehabilitation robots. Advances in Mechanical Engineering, 7(2), 563062.",
        "[26] Kwakkel, G., Kollen, B.J., and Krebs, H.I. (2008). Effects of robot-assisted therapy on upper limb recovery after stroke: A systematic review. Neurorehabilitation and Neural Repair, 22(2), 111-121.",
        "[27] Lo, A.C., Guarino, P.D., Richards, L.G., Haselkorn, J.K., Wittenberg, G.F., Federman, D.G., and Peduzzi, P. (2010). Robot-assisted therapy for long-term upper-limb impairment after stroke. New England Journal of Medicine, 362(19), 1772-1783.",
        "[28] Hesse, S., Waldner, A., and Tomelleri, C. (2010). Innovative gait robot for the repetitive practice of floor walking and stair climbing up and down in stroke patients. Journal of NeuroEngineering and Rehabilitation, 7(1), 30.",
        "[29] Dollar, A.M., and Herr, H. (2008). Lower extremity exoskeletons and active orthoses: Challenges and state-of-the-art. IEEE Transactions on Robotics, 24(1), 144-158.",
        "[30] Polygerinos, P., Wang, Z., Galloway, K.C., Wood, R.J., and Walsh, C.J. (2015). Soft robotic glove for combined assistance and at-home rehabilitation. Robotics and Autonomous Systems, 73, 135-143.",
        "[31] Laver, K.E., Lange, B., George, S., Deutsch, J.E., Saposnik, G., and Crotty, M. (2017). Virtual reality for stroke rehabilitation. Cochrane Database of Systematic Reviews, (11), CD008349.",
        "[32] Morone, G., Paolucci, S., Cherubini, A., De Angelis, D., Venturiero, V., Coiro, P., and Iosa, M. (2017). Robot-assisted gait training for stroke patients: Current state of the art and perspectives of robotics. Neuropsychiatric Disease and Treatment, 13, 1303-1311.",
        "[33] Atzori, M., and Muller, H. (2015). Control capabilities of myoelectric robotic prostheses by hand amputees: A scientific research and market overview. Frontiers in Systems Neuroscience, 9, 162.",
        "[34] Tucker, M.R., Olivier, J., Pagel, A., Bleuler, H., Bouri, M., Lambercy, O., and Gassert, R. (2015). Control strategies for active lower extremity prosthetics and orthotics: A review. Journal of NeuroEngineering and Rehabilitation, 12(1), 1.",
        "[35] Mao, Y., Lo, W.L., Lin, Q., Li, L., Zhao, X., Chen, L., and Huang, D. (2015). The effect of body weight support treadmill training on gait recovery, proximal lower limb motor pattern, and balance in patients with subacute stroke. BioMed Research International, 2015, 175719.",
        "[36] Feil-Seifer, D., and Mataric, M.J. (2005). Defining socially assistive robotics. Proceedings of the IEEE International Conference on Rehabilitation Robotics, 465-468.",
        "[37] Robinson, H., MacDonald, B., and Broadbent, E. (2014). The role of healthcare robots for older people at home: A review. International Journal of Social Robotics, 6(4), 575-591.",
        "[38] Veerbeek, J.M., Langbroek-Amersfoort, A.C., van Wegen, E.E., Meskers, C.G., and Kwakkel, G. (2017). Effects of robot-assisted therapy for the upper limb after stroke: A systematic review and meta-analysis. Neurorehabilitation and Neural Repair, 31(2), 107-121.",
        "[39] Mehrholz, J., Thomas, S., Werner, C., Kugler, J., Pohl, M., and Elsner, B. (2017). Electromechanical-assisted training for walking after stroke. Cochrane Database of Systematic Reviews, (5), CD006185.",
        "[40] Miller, L.E., Zimmermann, A.K., and Herbert, W.G. (2016). Clinical effectiveness and safety of powered exoskeleton-assisted walking in patients with spinal cord injury: Systematic review with meta-analysis. Medical Devices: Evidence and Research, 9, 455-466.",
        "[41] Danzl, M.M., Etter, N.M., Andreatta, R.D., and Kitzman, P.H. (2012). Facilitating neurorehabilitation through principles of engagement. Journal of Allied Health, 41(1), 35-41.",
        "[42] Picelli, A., Melotti, C., Origano, F., Neri, R., Verzè, E., Gandolfi, M., and Smania, N. (2013). Robot-assisted gait training is not superior to balance training for improving postural instability in patients with mild to moderate Parkinson's disease. Clinical Rehabilitation, 27(8), 728-736.",
        "[43] De Santis, A., Siciliano, B., De Luca, A., and Bicchi, A. (2008). An atlas of physical human-robot interaction. Mechanism and Machine Theory, 43(3), 253-270.",
        "[44] Huang, V.S., and Krakauer, J.W. (2009). Robotic neurorehabilitation: A computational motor learning perspective. Journal of NeuroEngineering and Rehabilitation, 6(1), 5.",
        "[45] Peternel, L., Tsagarakis, N., Caldwell, D., and Ajoudani, A. (2018). Robot adaptation to human physical fatigue in human-robot co-manipulation. Autonomous Robots, 42(5), 1011-1021.",
        "[46] Cao, Z., Hidalgo, G., Simon, T., Wei, S.E., and Sheikh, Y. (2019). OpenPose: Realtime multi-person 2D pose estimation using Part Affinity Fields. IEEE Transactions on Pattern Analysis and Machine Intelligence, 43(1), 172-186.",
        "[47] Brennan, D.M., Mawson, S., and Brownsell, S. (2009). Telerehabilitation: Enabling the remote delivery of healthcare, rehabilitation, and self-management. Studies in Health Technology and Informatics, 145, 231-248.",
    ]
    
    for ref in references:
        para = doc.add_paragraph(ref, style='Normal')
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(-1.0)
        for run in para.runs:
            run.font.size = Pt(9.5)
    
    # Save document
    output_path = '/projects/sandbox/AMMAN/Rehabilitation_Robotics_Chapter.docx'
    doc.save(output_path)
    print(f"Document saved successfully to: {output_path}")
    return output_path


if __name__ == '__main__':
    create_document()
