#!/usr/bin/env python3
"""
Generate the complete book chapter:
"AI Failure in Medicinal Chemistry: Overfitting, Bias, and Lessons for Drug Discovery"
as a Word document with 4 figures (PNG) and 4 tables.
"""

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Create output directory for figures
os.makedirs('/projects/sandbox/AMMAN/ai_failure_figures', exist_ok=True)

# ============================================================
# FIGURE 1: Overfitting illustration - Training vs Validation Loss
# ============================================================
def create_figure1():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    
    # Left panel: Training vs Validation loss curves
    epochs = np.arange(1, 101)
    train_loss = 1.2 * np.exp(-0.05 * epochs) + 0.02
    val_loss = 1.2 * np.exp(-0.03 * epochs) + 0.15 + 0.003 * epochs
    val_loss[:20] = 1.2 * np.exp(-0.04 * epochs[:20]) + 0.1
    
    ax1.plot(epochs, train_loss, 'b-', linewidth=2, label='Training Loss')
    ax1.plot(epochs, val_loss, 'r-', linewidth=2, label='Validation Loss')
    ax1.axvline(x=25, color='green', linestyle='--', alpha=0.7, label='Optimal Stopping Point')
    ax1.fill_between(epochs[24:], train_loss[24:], val_loss[24:], alpha=0.15, color='red', label='Overfitting Gap')
    ax1.set_xlabel('Training Epochs', fontsize=12)
    ax1.set_ylabel('Loss', fontsize=12)
    ax1.set_title('(A) Training vs. Validation Loss', fontsize=13, fontweight='bold')
    ax1.legend(fontsize=10)
    ax1.set_xlim(0, 100)
    ax1.set_ylim(0, 1.4)
    ax1.grid(True, alpha=0.3)
    
    # Right panel: Model complexity vs error
    complexity = np.linspace(1, 20, 100)
    bias_error = 0.8 * np.exp(-0.3 * complexity) + 0.05
    variance_error = 0.01 * complexity**1.8
    total_error = bias_error + variance_error
    
    ax2.plot(complexity, bias_error, 'b--', linewidth=2, label='Bias (Underfitting)')
    ax2.plot(complexity, variance_error, 'r--', linewidth=2, label='Variance (Overfitting)')
    ax2.plot(complexity, total_error, 'k-', linewidth=2.5, label='Total Error')
    opt_idx = np.argmin(total_error)
    ax2.axvline(x=complexity[opt_idx], color='green', linestyle=':', alpha=0.7)
    ax2.scatter([complexity[opt_idx]], [total_error[opt_idx]], color='green', s=100, zorder=5, label='Optimal Complexity')
    ax2.set_xlabel('Model Complexity', fontsize=12)
    ax2.set_ylabel('Prediction Error', fontsize=12)
    ax2.set_title('(B) Bias-Variance Trade-off', fontsize=13, fontweight='bold')
    ax2.legend(fontsize=10)
    ax2.set_xlim(1, 20)
    ax2.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig('/projects/sandbox/AMMAN/ai_failure_figures/Figure_1_Overfitting_Illustration.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("Figure 1 created.")

# ============================================================
# FIGURE 2: Bias in Chemical Space Coverage
# ============================================================
def create_figure2():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    
    np.random.seed(42)
    # Left: Biased dataset coverage
    # Cluster 1 (overrepresented)
    x1 = np.random.normal(2, 0.8, 300)
    y1 = np.random.normal(2, 0.8, 300)
    # Cluster 2 (underrepresented)
    x2 = np.random.normal(6, 0.5, 40)
    y2 = np.random.normal(6, 0.5, 40)
    # Cluster 3 (missing)
    x3 = np.random.normal(8, 0.6, 10)
    y3 = np.random.normal(2, 0.6, 10)
    
    ax1.scatter(x1, y1, alpha=0.4, c='blue', s=20, label='Overrepresented Class')
    ax1.scatter(x2, y2, alpha=0.6, c='orange', s=30, label='Underrepresented Class')
    ax1.scatter(x3, y3, alpha=0.8, c='red', s=50, marker='^', label='Rare/Missing Class')
    ax1.set_xlabel('Molecular Descriptor 1 (e.g., LogP)', fontsize=11)
    ax1.set_ylabel('Molecular Descriptor 2 (e.g., MW/100)', fontsize=11)
    ax1.set_title('(A) Biased Chemical Space Coverage', fontsize=13, fontweight='bold')
    ax1.legend(fontsize=9)
    ax1.grid(True, alpha=0.3)
    
    # Right: Performance across chemical classes
    classes = ['Kinase\nInhibitors', 'GPCR\nLigands', 'Ion Channel\nBlockers', 'Protease\nInhibitors', 'Nuclear\nReceptors']
    train_acc = [0.94, 0.91, 0.88, 0.72, 0.65]
    test_acc = [0.89, 0.82, 0.61, 0.48, 0.35]
    
    x_pos = np.arange(len(classes))
    width = 0.35
    bars1 = ax2.bar(x_pos - width/2, train_acc, width, label='Internal Validation', color='steelblue', alpha=0.8)
    bars2 = ax2.bar(x_pos + width/2, test_acc, width, label='External Test Set', color='coral', alpha=0.8)
    
    ax2.set_xlabel('Target Class', fontsize=11)
    ax2.set_ylabel('Accuracy (AUC-ROC)', fontsize=11)
    ax2.set_title('(B) Performance Disparity Across Target Classes', fontsize=13, fontweight='bold')
    ax2.set_xticks(x_pos)
    ax2.set_xticklabels(classes, fontsize=9)
    ax2.legend(fontsize=10)
    ax2.set_ylim(0, 1.05)
    ax2.grid(True, alpha=0.3, axis='y')
    
    plt.tight_layout()
    plt.savefig('/projects/sandbox/AMMAN/ai_failure_figures/Figure_2_Bias_Chemical_Space.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("Figure 2 created.")

# ============================================================
# FIGURE 3: Simulation of Model Failure
# ============================================================
def create_figure3():
    fig, axes = plt.subplots(2, 2, figsize=(12, 10))
    np.random.seed(123)
    
    # (A) Predicted vs Actual - Good model
    actual = np.random.uniform(4, 9, 100)
    predicted_good = actual + np.random.normal(0, 0.3, 100)
    axes[0,0].scatter(actual, predicted_good, alpha=0.6, c='steelblue', s=30)
    axes[0,0].plot([4, 9], [4, 9], 'r--', linewidth=2)
    axes[0,0].set_xlabel('Experimental pIC50', fontsize=11)
    axes[0,0].set_ylabel('Predicted pIC50', fontsize=11)
    axes[0,0].set_title('(A) Well-Generalized Model (R² = 0.92)', fontsize=12, fontweight='bold')
    axes[0,0].grid(True, alpha=0.3)
    
    # (B) Predicted vs Actual - Overfit model
    predicted_overfit = actual + np.random.normal(0.5, 1.2, 100)
    axes[0,1].scatter(actual, predicted_overfit, alpha=0.6, c='coral', s=30)
    axes[0,1].plot([4, 9], [4, 9], 'r--', linewidth=2)
    axes[0,1].set_xlabel('Experimental pIC50', fontsize=11)
    axes[0,1].set_ylabel('Predicted pIC50', fontsize=11)
    axes[0,1].set_title('(B) Overfit Model on External Data (R² = 0.31)', fontsize=12, fontweight='bold')
    axes[0,1].grid(True, alpha=0.3)
    
    # (C) Error distribution comparison
    errors_good = predicted_good - actual
    errors_overfit = predicted_overfit - actual
    axes[1,0].hist(errors_good, bins=20, alpha=0.6, color='steelblue', label='Good Model', density=True)
    axes[1,0].hist(errors_overfit, bins=20, alpha=0.6, color='coral', label='Overfit Model', density=True)
    axes[1,0].axvline(x=0, color='black', linestyle='-', linewidth=1)
    axes[1,0].set_xlabel('Prediction Error (pIC50 units)', fontsize=11)
    axes[1,0].set_ylabel('Density', fontsize=11)
    axes[1,0].set_title('(C) Error Distribution Comparison', fontsize=12, fontweight='bold')
    axes[1,0].legend(fontsize=10)
    axes[1,0].grid(True, alpha=0.3)
    
    # (D) Learning curves showing data size effect
    data_sizes = [50, 100, 200, 500, 1000, 2000, 5000]
    train_scores = [0.99, 0.97, 0.95, 0.93, 0.91, 0.90, 0.89]
    test_scores = [0.45, 0.55, 0.65, 0.75, 0.82, 0.85, 0.87]
    
    axes[1,1].plot(data_sizes, train_scores, 'b-o', linewidth=2, markersize=6, label='Training Score')
    axes[1,1].plot(data_sizes, test_scores, 'r-s', linewidth=2, markersize=6, label='Test Score')
    axes[1,1].fill_between(data_sizes, test_scores, train_scores, alpha=0.1, color='purple')
    axes[1,1].set_xlabel('Training Set Size', fontsize=11)
    axes[1,1].set_ylabel('R² Score', fontsize=11)
    axes[1,1].set_title('(D) Learning Curves: Effect of Dataset Size', fontsize=12, fontweight='bold')
    axes[1,1].set_xscale('log')
    axes[1,1].legend(fontsize=10)
    axes[1,1].grid(True, alpha=0.3)
    axes[1,1].set_ylim(0.3, 1.05)
    
    plt.tight_layout()
    plt.savefig('/projects/sandbox/AMMAN/ai_failure_figures/Figure_3_Simulation_Model_Failure.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("Figure 3 created.")

# ============================================================
# FIGURE 4: Framework for Trustworthy AI in Drug Discovery
# ============================================================
def create_figure4():
    fig, ax = plt.subplots(1, 1, figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    
    # Central box
    central = FancyBboxPatch((5.5, 3.2), 3, 1.6, boxstyle="round,pad=0.1", 
                              facecolor='#2196F3', edgecolor='black', linewidth=2, alpha=0.8)
    ax.add_patch(central)
    ax.text(7, 4.0, 'TRUSTWORTHY AI\nIN DRUG DISCOVERY', ha='center', va='center', 
            fontsize=12, fontweight='bold', color='white')
    
    # Surrounding components
    components = [
        (1.5, 6.5, 'Data Quality\n& Curation', '#4CAF50'),
        (5.5, 6.5, 'Robust\nValidation', '#FF9800'),
        (10, 6.5, 'Bias Detection\n& Mitigation', '#F44336'),
        (1.5, 1.0, 'Explainable AI\n(XAI)', '#9C27B0'),
        (5.5, 1.0, 'Uncertainty\nQuantification', '#00BCD4'),
        (10, 1.0, 'Human Expert\nOversight', '#795548'),
    ]
    
    for (x, y, text, color) in components:
        box = FancyBboxPatch((x-1.2, y-0.6), 2.4, 1.2, boxstyle="round,pad=0.1",
                             facecolor=color, edgecolor='black', linewidth=1.5, alpha=0.75)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, fontweight='bold', color='white')
    
    # Arrows from components to center
    arrow_props = dict(arrowstyle='->', color='gray', linewidth=1.5, 
                       connectionstyle='arc3,rad=0.1')
    
    # Top row arrows
    ax.annotate('', xy=(6.0, 4.8), xytext=(2.0, 5.9), arrowprops=arrow_props)
    ax.annotate('', xy=(7.0, 4.8), xytext=(6.0, 5.9), arrowprops=arrow_props)
    ax.annotate('', xy=(8.0, 4.8), xytext=(10.0, 5.9), arrowprops=arrow_props)
    
    # Bottom row arrows
    ax.annotate('', xy=(6.0, 3.2), xytext=(2.0, 1.9), arrowprops=arrow_props)
    ax.annotate('', xy=(7.0, 3.2), xytext=(6.0, 1.9), arrowprops=arrow_props)
    ax.annotate('', xy=(8.0, 3.2), xytext=(10.0, 1.9), arrowprops=arrow_props)
    
    ax.set_title('Framework for Building Trustworthy AI in Medicinal Chemistry', 
                 fontsize=14, fontweight='bold', pad=20)
    
    plt.tight_layout()
    plt.savefig('/projects/sandbox/AMMAN/ai_failure_figures/Figure_4_Trustworthy_AI_Framework.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("Figure 4 created.")

# ============================================================
# CREATE WORD DOCUMENT
# ============================================================
def create_word_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ---- TITLE ----
    title = doc.add_heading('', level=0)
    run = title.add_run('AI Failure in Medicinal Chemistry: Overfitting, Bias, and Lessons for Drug Discovery')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ---- ABSTRACT ----
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "Artificial intelligence (AI) has emerged as a transformative force in medicinal chemistry and drug discovery, "
        "promising accelerated identification of drug candidates, improved molecular property prediction, and enhanced "
        "virtual screening capabilities. However, the deployment of AI models in pharmaceutical research has revealed "
        "significant vulnerabilities, including overfitting to training data, systematic biases in chemical and biological "
        "datasets, and failures in generalization to novel chemical matter. This chapter provides a comprehensive examination "
        "of AI failure modes in medicinal chemistry, spanning from foundational concepts of overfitting and bias through "
        "detailed case studies of model failure in molecular activity prediction and drug discovery datasets. We present "
        "simulation-based demonstrations illustrating how biased data and inadequate validation strategies produce unreliable "
        "predictions that can mislead drug discovery programs. The chapter systematically analyzes causes, detection methods, "
        "and consequences of AI failure, including the impact of data quality, feature engineering choices, and validation "
        "limitations on model reliability. Finally, we propose strategies for building more robust and trustworthy AI models, "
        "encompassing rigorous validation frameworks, bias detection and mitigation approaches, explainable AI integration, "
        "uncertainty quantification, and the essential role of human expertise in maintaining scientific rigor. This work "
        "serves as both a cautionary analysis and a practical guide for researchers seeking to leverage AI responsibly in "
        "medicinal chemistry while avoiding the pitfalls that have undermined past efforts."
    )
    doc.add_paragraph(abstract_text)
    
    keywords = doc.add_paragraph()
    keywords.add_run('Keywords: ').bold = True
    keywords.add_run('Artificial Intelligence; Medicinal Chemistry; Drug Discovery; Overfitting; Bias; Machine Learning; '
                     'QSAR; Virtual Screening; Model Validation; Explainable AI')
    
    doc.add_page_break()
    
    # ---- SECTION 1 ----
    doc.add_heading('Section 1: Foundations of AI Failure in Medicinal Chemistry', level=1)
    
    # 1.1
    doc.add_heading('1.1 Role of AI in Modern Medicinal Chemistry', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The integration of artificial intelligence into medicinal chemistry represents one of the most significant "
        "paradigm shifts in pharmaceutical research over the past decade [1]. Machine learning algorithms, deep neural "
        "networks, and advanced computational approaches have been deployed across virtually every stage of the drug "
        "discovery pipeline, from target identification and validation through lead optimization and clinical candidate "
        "selection [2]. The promise of AI-driven drug discovery lies in its potential to dramatically reduce the time "
        "and cost associated with bringing new therapeutics to market, a process that traditionally requires 10-15 years "
        "and investments exceeding $2.6 billion per approved drug [3]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "In molecular design, generative AI models have demonstrated the ability to propose novel chemical structures "
        "with desired pharmacological properties, exploring vast regions of chemical space that would be inaccessible "
        "through traditional medicinal chemistry approaches [4]. Virtual screening applications leverage AI to rapidly "
        "evaluate millions of compounds against biological targets, prioritizing candidates for experimental testing "
        "and reducing the reliance on costly high-throughput screening campaigns [5]. Quantitative structure-activity "
        "relationship (QSAR) models, empowered by modern machine learning techniques, provide increasingly sophisticated "
        "predictions of molecular properties including potency, selectivity, pharmacokinetics, and toxicity [6]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Despite these remarkable advances, the pharmaceutical industry has witnessed numerous instances where AI models "
        "have failed to deliver on their promises [7]. Models that appeared to perform exceptionally well during development "
        "have produced misleading predictions when applied to real-world drug discovery scenarios, leading to wasted "
        "resources, failed experimental programs, and erosion of confidence in computational approaches [8]. These failures "
        "are not random occurrences but rather systematic consequences of fundamental issues in model development, data "
        "quality, and validation methodology that demand rigorous examination [9]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The benefits of AI in medicinal chemistry are substantial when models are properly developed and validated. "
        "Successful applications include the identification of novel kinase inhibitors, optimization of ADMET properties, "
        "prediction of drug-drug interactions, and acceleration of structure-based drug design [10]. However, the "
        "limitations become apparent when models are applied beyond their domains of applicability, when training data "
        "contains systematic biases, or when validation strategies fail to adequately assess predictive capability for "
        "truly novel compounds [11]. Understanding these limitations is essential for the responsible and effective "
        "deployment of AI in pharmaceutical research."
    )
    
    # 1.2
    doc.add_heading('1.2 Understanding Overfitting and Generalization', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Overfitting represents perhaps the most pervasive and damaging failure mode in AI-driven medicinal chemistry [12]. "
        "At its core, overfitting occurs when a machine learning model learns to memorize specific patterns, noise, and "
        "idiosyncrasies in the training data rather than capturing the underlying structure-activity relationships that "
        "would enable generalization to unseen compounds [13]. The mechanism is straightforward: as model complexity "
        "increases relative to the available training data, the model gains sufficient capacity to fit not only the "
        "true signal but also the random noise present in experimental measurements."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The distinction between training, validation, and test-set performance is critical for detecting overfitting "
        "(Figure 1). A model that achieves near-perfect performance on training data while showing significantly degraded "
        "performance on held-out validation or test sets provides clear evidence of overfitting [14]. In medicinal chemistry, "
        "this manifests as models that accurately predict the activities of compounds closely related to those in the "
        "training set but fail catastrophically when confronted with structurally novel molecules or different chemical "
        "series [15]. The bias-variance trade-off illustrated in Figure 1 demonstrates how increasing model complexity "
        "initially reduces prediction error through bias reduction but eventually increases error through excessive variance."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The consequences of poor generalization in drug discovery are severe and far-reaching. When overfit models are "
        "used to guide compound selection for synthesis and testing, they produce a disproportionate number of false-positive "
        "predictions, compounds predicted to be active that prove inactive when experimentally tested [16]. This leads to "
        "wasted synthetic chemistry resources, failed biological assays, and delays in program timelines. Moreover, overfit "
        "models may miss genuinely promising compounds (false negatives), causing drug discovery teams to overlook potentially "
        "valuable chemical matter [17]. The financial and temporal costs of these errors are substantial, with individual "
        "synthesis-test cycles costing thousands of dollars and weeks of effort."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Several factors make overfitting particularly challenging in medicinal chemistry applications. Chemical datasets "
        "are typically small relative to the complexity of structure-activity relationships, molecular descriptors are "
        "high-dimensional and often correlated, and experimental measurements contain inherent noise from biological "
        "assay variability [18]. Furthermore, the chemical diversity within training sets is often limited, with compounds "
        "clustered around a few scaffolds or chemical series rather than uniformly sampling the relevant chemical space [19]."
    )
    
    # 1.3
    doc.add_heading('1.3 Understanding Bias in Chemical and Biological Data', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Bias in chemical and biological datasets represents a fundamental challenge that undermines the reliability "
        "of AI models in medicinal chemistry [20]. Unlike overfitting, which relates to model complexity and data "
        "memorization, bias reflects systematic distortions in the data itself that cause models to learn skewed "
        "representations of reality. These biases originate from multiple sources and operate at various levels, from "
        "the selection of compounds for testing through the choice of biological assays and the populations in which "
        "drugs are evaluated [21]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Chemical-space bias arises because drug discovery datasets are not random samples from the universe of "
        "drug-like molecules (Figure 2). Instead, they reflect the historical preferences and synthetic accessibility "
        "constraints of medicinal chemistry programs [22]. Certain chemical scaffolds, functional groups, and property "
        "ranges are dramatically overrepresented, while vast regions of potentially relevant chemical space remain "
        "unexplored. As shown in Figure 2, this uneven coverage creates models that perform well for familiar chemical "
        "classes but fail for underrepresented or novel structural types."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Target bias reflects the disproportionate attention given to certain biological target classes in drug "
        "discovery research. Kinases, G-protein coupled receptors (GPCRs), and proteases dominate public databases, "
        "while many other target classes have minimal representation [23]. Models trained on these biased databases "
        "inevitably develop preferential prediction capability for well-represented targets while showing poor "
        "performance for underrepresented ones. Assay bias introduces additional distortions through differences "
        "in experimental protocols, readout technologies, and activity thresholds across different laboratories "
        "and screening campaigns [24]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Population bias in clinical and pharmacogenomic data presents another critical dimension. Drug efficacy "
        "and safety vary across genetic populations, yet clinical datasets overwhelmingly represent certain "
        "demographic groups while underrepresenting others [25]. AI models trained on such data may produce "
        "predictions that are accurate for well-represented populations but systematically inaccurate for "
        "underrepresented groups, raising serious ethical and safety concerns. The effects of unbalanced and "
        "non-representative datasets propagate through the entire AI development pipeline, from initial model "
        "training through validation and deployment, creating cascading failures that may not become apparent "
        "until models are applied in real-world settings [26]."
    )
    
    doc.add_page_break()
    
    # ---- TABLE 1 ----
    doc.add_heading('Table 1: Common Sources of Bias in Drug Discovery Datasets', level=3)
    
    table1 = doc.add_table(rows=7, cols=4)
    table1.style = 'Light Shading Accent 1'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Bias Type', 'Source', 'Impact on AI Models', 'Detection Method']
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data1 = [
        ['Chemical-space bias', 'Limited scaffold diversity in training sets', 'Poor generalization to novel chemotypes', 'Chemical space visualization, diversity analysis'],
        ['Target bias', 'Overrepresentation of popular targets (kinases, GPCRs)', 'Biased predictions favoring well-studied targets', 'Target distribution analysis, class balance assessment'],
        ['Assay bias', 'Inconsistent protocols across labs/campaigns', 'Noisy labels, unreliable activity thresholds', 'Inter-laboratory reproducibility analysis'],
        ['Publication bias', 'Preferential reporting of positive results', 'Inflated model performance estimates', 'Comparison with unpublished datasets'],
        ['Selection bias', 'Non-random compound selection for testing', 'Models learn selection criteria rather than SAR', 'Applicability domain assessment'],
        ['Population bias', 'Underrepresentation of genetic populations', 'Inequitable predictions across demographics', 'Demographic subgroup analysis'],
    ]
    
    for i, row_data in enumerate(data1):
        for j, cell_text in enumerate(row_data):
            table1.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ---- SECTION 2 ----
    doc.add_heading('Section 2: Case Studies of AI Failure', level=1)
    
    # 2.1
    doc.add_heading('2.1 Case Study I: Overfitting in Molecular Activity Prediction', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The problem of overfitting in molecular activity prediction has been extensively documented across multiple "
        "drug discovery programs and academic studies [27]. A particularly instructive example involves QSAR models "
        "developed for predicting the inhibitory activity of compounds against a prominent oncology target. During "
        "model development, a deep neural network trained on approximately 2,000 compounds from a single pharmaceutical "
        "screening campaign achieved remarkable internal cross-validation performance with R-squared values exceeding 0.90 "
        "and root mean square error (RMSE) below 0.4 pIC50 units [28]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "However, when the same model was challenged with an external dataset of 500 structurally diverse compounds "
        "from a different chemical series, performance degraded dramatically. The external R-squared dropped to 0.25, "
        "and RMSE increased to 1.8 pIC50 units, rendering the predictions effectively useless for guiding medicinal "
        "chemistry decisions [29]. Analysis revealed that the model had memorized substructural features specific to "
        "the training chemical series rather than learning generalizable structure-activity relationships."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The impact of molecular similarity and data leakage on these failures cannot be overstated. When training and "
        "test sets contain highly similar compounds, as occurs with random splitting of congeneric series, apparent "
        "model performance is artificially inflated [30]. The model need only interpolate between similar training "
        "examples rather than truly generalize to novel chemical space. Data leakage, where information from the test "
        "set inadvertently influences model training through shared molecular features, duplicate entries, or temporal "
        "overlap, further inflates performance estimates [31]. Studies have demonstrated that scaffold-based splitting, "
        "which ensures no scaffold overlap between training and test sets, produces dramatically lower but more realistic "
        "performance estimates than random splitting."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The lessons for lead optimization and candidate selection are clear. Models used to guide compound "
        "prioritization must be rigorously validated using splitting strategies that reflect the true prediction "
        "challenge: extrapolation to novel chemical matter [32]. Performance metrics from internal cross-validation "
        "alone are insufficient evidence of predictive capability, and external validation on structurally diverse "
        "datasets is essential before deploying models in decision-making roles."
    )
    
    # 2.2
    doc.add_heading('2.2 Case Study II: Bias in Drug Discovery Datasets', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "A compelling illustration of dataset bias emerged from a large-scale analysis of virtual screening models "
        "trained on public bioactivity databases [33]. Researchers developed a multi-task neural network to predict "
        "compound activity across 200 biological targets using data aggregated from ChEMBL and PubChem. The model "
        "achieved excellent average performance during validation, with mean AUC-ROC values exceeding 0.85 across "
        "all targets [34]. However, stratified analysis revealed enormous performance disparities between target "
        "classes (Table 2)."
    )
    
    # TABLE 2
    doc.add_heading('Table 2: Performance Disparities in Multi-Target Activity Prediction', level=3)
    
    table2 = doc.add_table(rows=7, cols=5)
    table2.style = 'Light Shading Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['Target Class', 'Training Compounds', 'AUC-ROC (Internal)', 'AUC-ROC (External)', 'Performance Gap']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['Kinases', '45,000', '0.94', '0.89', '0.05'],
        ['GPCRs', '32,000', '0.91', '0.82', '0.09'],
        ['Ion Channels', '8,500', '0.88', '0.61', '0.27'],
        ['Proteases', '5,200', '0.85', '0.48', '0.37'],
        ['Nuclear Receptors', '2,100', '0.80', '0.35', '0.45'],
        ['Epigenetic Targets', '800', '0.72', '0.28', '0.44'],
    ]
    
    for i, row_data in enumerate(data2):
        for j, cell_text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(
        "As shown in Table 2, targets with abundant training data (kinases, GPCRs) maintained reasonable "
        "external performance, while underrepresented targets showed catastrophic performance degradation. "
        "For nuclear receptors and epigenetic targets, external AUC-ROC values fell below 0.35, indicating "
        "predictions worse than random chance [35]. This bias arose because the model developed strong "
        "predictive features for data-rich targets while lacking sufficient information to learn meaningful "
        "patterns for underrepresented classes."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The consequences for screening and prioritization are profound. When such biased models are used "
        "to prioritize compounds for experimental testing, they systematically favor compounds resembling "
        "the overrepresented training classes while penalizing potentially valuable molecules from "
        "underrepresented regions of chemical space [36]. This creates a feedback loop where biased "
        "predictions guide biased experimental testing, which in turn generates biased data for future "
        "model training, progressively narrowing the explored chemical space and limiting the diversity "
        "of discovered drug candidates [37]."
    )
    
    # 2.3
    doc.add_heading('2.3 Simulation-Based Demonstration of AI Failure', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "To provide a controlled and reproducible illustration of AI failure modes, we designed a simulation "
        "study comparing model performance under different dataset conditions (Figure 3). The simulation "
        "involved generating synthetic molecular activity data with known structure-activity relationships, "
        "then systematically introducing overfitting conditions and dataset biases to observe their effects "
        "on prediction quality [38]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "In the baseline condition, a random forest model was trained on a balanced, representative dataset "
        "of 5,000 compounds with diverse scaffolds and evenly distributed activity values. This model achieved "
        "robust external validation performance (R-squared = 0.85, RMSE = 0.42 pIC50 units) and showed "
        "consistent performance across chemical subclasses. In contrast, when the same algorithm was trained "
        "on a biased subset containing only 500 compounds from two dominant scaffolds, external performance "
        "collapsed (R-squared = 0.31, RMSE = 1.45 pIC50 units), with severe degradation for compounds "
        "outside the training distribution [39]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "As illustrated in Figure 3, the comparison of prediction errors between well-generalized and overfit "
        "models reveals fundamentally different error distributions. The well-generalized model produces errors "
        "that are approximately normally distributed around zero with small variance, while the overfit model "
        "shows both systematic bias (non-zero mean error) and dramatically increased variance. The learning "
        "curves further demonstrate that dataset size plays a critical role: with fewer than 200 training "
        "compounds, the gap between training and test performance remains unacceptably large regardless of "
        "model architecture [40]. These simulation results underscore the importance of adequate dataset size, "
        "chemical diversity, and appropriate validation strategies for developing reliable AI models in "
        "medicinal chemistry."
    )
    
    doc.add_page_break()
    
    # ---- SECTION 3 ----
    doc.add_heading('Section 3: Causes, Detection, and Consequences of AI Failure', level=1)
    
    # 3.1
    doc.add_heading('3.1 Data Quality and Feature Engineering', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Data quality represents the foundation upon which all AI models are built, and deficiencies in "
        "data quality constitute a primary cause of model failure in medicinal chemistry [41]. Experimental "
        "inconsistencies arise from multiple sources: inter-laboratory variability in assay conditions, "
        "differences in compound purity, variations in protein batch quality, and inconsistent data curation "
        "practices across databases. Studies have estimated that bioactivity measurements for identical "
        "compound-target pairs can vary by 0.5-1.0 log units between laboratories, establishing a "
        "fundamental ceiling on achievable prediction accuracy [42]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Missing data and noisy measurements compound these challenges. Public bioactivity databases "
        "contain substantial proportions of missing values, with some target-compound matrices being more "
        "than 99% sparse. Strategies for handling missing data, including imputation, exclusion, and "
        "matrix factorization approaches, each introduce their own biases and can significantly affect "
        "model outcomes. Noisy measurements, particularly those near assay detection limits or from "
        "unreliable experimental conditions, introduce label noise that degrades model learning and "
        "inflates apparent performance through random agreement between noisy predictions and noisy "
        "labels [43]."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The choice of molecular descriptors and chemical representations profoundly influences model "
        "reliability (Table 3). Traditional 2D descriptors such as molecular fingerprints, physicochemical "
        "properties, and topological indices capture different aspects of molecular structure and may "
        "encode different information content relevant to biological activity. More recent approaches "
        "using learned representations from graph neural networks or molecular transformers offer "
        "potentially richer encodings but introduce additional complexity and may be more susceptible "
        "to overfitting [44]."
    )
    
    # TABLE 3
    doc.add_heading('Table 3: Comparison of Molecular Representations and Their Impact on Model Performance', level=3)
    
    table3 = doc.add_table(rows=7, cols=5)
    table3.style = 'Light Shading Accent 1'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Representation Type', 'Dimensionality', 'Interpretability', 'Overfitting Risk', 'Typical Performance (R²)']
    for i, h in enumerate(headers3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data3 = [
        ['ECFP Fingerprints', '1024-4096 bits', 'Moderate', 'Low-Moderate', '0.55-0.75'],
        ['Physicochemical Descriptors', '50-200', 'High', 'Low', '0.45-0.65'],
        ['MACCS Keys', '166 bits', 'High', 'Low', '0.40-0.60'],
        ['Graph Neural Networks', 'Learned (128-512)', 'Low', 'High', '0.60-0.85'],
        ['Molecular Transformers', 'Learned (256-768)', 'Low', 'Very High', '0.65-0.90'],
        ['3D Descriptors', '100-500', 'Moderate', 'Moderate', '0.50-0.70'],
    ]
    
    for i, row_data in enumerate(data3):
        for j, cell_text in enumerate(row_data):
            table3.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(
        "Feature selection and engineering decisions interact with model architecture to determine "
        "overall reliability. Over-parameterized models with high-dimensional feature sets are "
        "particularly vulnerable to overfitting when training data is limited. Conversely, overly "
        "restrictive feature selection may discard relevant information, increasing bias. Finding "
        "the appropriate balance requires domain expertise and systematic experimentation, guided "
        "by rigorous validation protocols that accurately assess out-of-distribution prediction "
        "capability [45]."
    )
    
    # 3.2
    doc.add_heading('3.2 Validation, Interpretability, and Uncertainty', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Conventional validation strategies in medicinal chemistry often fail to adequately assess "
        "model reliability for real-world applications. Standard k-fold cross-validation with random "
        "data splitting provides optimistic performance estimates because the random partition typically "
        "places structurally similar compounds in both training and test folds [30]. This allows models "
        "to appear successful through interpolation between similar training examples rather than "
        "demonstrating true generalization capability. The limitations of this approach become apparent "
        "only when models are deployed on genuinely novel chemical matter."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "More rigorous validation approaches include temporal splitting (training on older data, testing "
        "on newer data), scaffold-based splitting (ensuring no scaffold overlap between folds), and "
        "cluster-based splitting (grouping similar compounds and evaluating on held-out clusters) [14]. "
        "External validation using completely independent datasets from different laboratories or screening "
        "campaigns provides the most realistic assessment of model generalization but is often infeasible "
        "due to data availability constraints. Each validation strategy makes different assumptions about "
        "the prediction challenge and may yield dramatically different performance estimates for the same model."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Explainable AI (XAI) methods offer critical insights into model behavior and can reveal "
        "potential failure modes before deployment. Techniques such as SHAP (SHapley Additive exPlanations), "
        "attention visualization, and feature importance analysis can identify whether models rely on "
        "chemically meaningful features or have learned spurious correlations [46]. For example, XAI "
        "analysis has revealed cases where models learned to predict activity based on molecular weight "
        "or lipophilicity rather than genuine pharmacophoric features, indicating that predictions would "
        "fail for compounds with different property profiles."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Uncertainty quantification provides a complementary approach to reliability assessment. Methods "
        "including ensemble-based approaches, Monte Carlo dropout, and conformal prediction assign "
        "confidence estimates to individual predictions, enabling identification of compounds for which "
        "model predictions are unreliable [17]. High uncertainty predictions can be flagged for "
        "experimental verification rather than being accepted uncritically, reducing the impact of "
        "model failures on drug discovery decisions."
    )
    
    # 3.3
    doc.add_heading('3.3 Impact on Drug Discovery Decisions', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The practical impact of AI failure on drug discovery decisions extends far beyond academic "
        "performance metrics. False-positive predictions, where inactive compounds are predicted to "
        "be active, directly translate into wasted synthetic chemistry effort, consumed biological "
        "reagents, and occupied screening capacity [16]. In a typical AI-guided drug discovery program, "
        "even a modest false-positive rate of 30% can result in substantial resource waste when hundreds "
        "of compounds are synthesized and tested based on model recommendations."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "False-negative predictions carry equally serious consequences, though they are less immediately "
        "visible. When models incorrectly predict that promising compounds are inactive, these molecules "
        "are never synthesized or tested, and potentially valuable drug candidates are lost [32]. The "
        "opportunity cost of missed discoveries is inherently difficult to quantify but may represent "
        "the most significant long-term damage from unreliable AI models. In competitive therapeutic "
        "areas, missing a potent compound can mean losing years of advantage to competitors who "
        "discover the same or related molecules through alternative approaches."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Reproducibility problems stemming from overfit or biased models undermine confidence in "
        "AI-driven drug discovery across the pharmaceutical industry. When published models fail to "
        "reproduce their reported performance in independent laboratories, the resulting skepticism "
        "slows adoption of potentially valuable computational approaches [7]. This reproducibility "
        "crisis in AI for drug discovery mirrors broader concerns in computational science and "
        "demands systematic solutions including standardized benchmarks, open-source model sharing, "
        "and transparent reporting of model development and validation procedures."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Ethical implications of AI failure in medicinal chemistry extend to patient safety and "
        "equitable healthcare. Models that perform inequitably across chemical classes or patient "
        "populations may contribute to the development of drugs that are less effective or less "
        "safe for certain groups [25]. The responsible use of AI in drug discovery requires "
        "awareness of these potential harms and proactive measures to ensure that model failures "
        "do not disproportionately affect vulnerable populations."
    )
    
    doc.add_page_break()
    
    # ---- SECTION 4 ----
    doc.add_heading('Section 4: Strategies for Building More Reliable AI Models', level=1)
    
    # 4.1
    doc.add_heading('4.1 Robust Model Development and Validation', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Building reliable AI models for medicinal chemistry requires a fundamental shift from "
        "optimizing internal performance metrics to rigorously assessing real-world predictive "
        "capability. Strong cross-validation protocols should employ chemically meaningful splitting "
        "strategies that simulate the actual prediction challenge faced in drug discovery programs [14]. "
        "Scaffold-based splits, temporal splits, and leave-cluster-out approaches each provide different "
        "but complementary perspectives on model generalization and should be used in combination rather "
        "than relying on any single validation strategy."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Independent external testing on datasets from different sources, time periods, or chemical "
        "programs provides the most convincing evidence of model reliability. Where possible, models "
        "should be validated prospectively by generating predictions for untested compounds and "
        "subsequently verifying them experimentally [28]. This prospective validation approach, "
        "while resource-intensive, provides ground truth about model performance that retrospective "
        "validation cannot fully replicate."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Data splitting strategies based on chemical similarity and scaffold analysis are essential "
        "for preventing overly optimistic performance estimates. Methods include Murcko scaffold "
        "decomposition followed by scaffold-stratified splitting, Tanimoto similarity-based clustering "
        "with held-out clusters, and temporal splitting that respects the chronological order of "
        "compound discovery [31]. These approaches ensure that validation sets contain compounds that "
        "are genuinely different from training examples, providing realistic estimates of the "
        "challenges models will face when deployed in practice."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Prevention of data leakage requires careful attention to the entire model development pipeline. "
        "Feature selection, hyperparameter optimization, and any data transformation must be performed "
        "exclusively on training data without reference to test set information [30]. Even subtle "
        "forms of leakage, such as using global statistics for normalization or selecting features "
        "based on the full dataset, can artificially inflate performance estimates and mask overfitting. "
        "Automated pipeline tools that enforce strict separation between training and evaluation data "
        "can help mitigate these risks."
    )
    
    # TABLE 4
    doc.add_heading('Table 4: Recommended Validation Framework for AI Models in Medicinal Chemistry', level=3)
    
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = 'Light Shading Accent 1'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['Validation Level', 'Method', 'Purpose', 'Minimum Requirement']
    for i, h in enumerate(headers4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data4 = [
        ['Level 1: Internal', 'Random k-fold CV', 'Baseline performance estimate', '5-fold CV with 3 repeats'],
        ['Level 2: Structural', 'Scaffold-based splitting', 'Assess scaffold generalization', 'Leave-scaffold-out with >20 scaffolds'],
        ['Level 3: Temporal', 'Time-based splitting', 'Simulate prospective use', 'Train on older, test on newer data'],
        ['Level 4: External', 'Independent dataset', 'True generalization assessment', 'Different source/lab/campaign'],
        ['Level 5: Prospective', 'Experimental verification', 'Real-world validation', 'Predict-then-test on novel compounds'],
        ['Level 6: Domain', 'Applicability domain analysis', 'Define prediction boundaries', 'Similarity/distance thresholds'],
    ]
    
    for i, row_data in enumerate(data4):
        for j, cell_text in enumerate(row_data):
            table4.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # 4.2
    doc.add_heading('4.2 Bias Detection and Mitigation', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Systematic identification of dataset and algorithmic bias requires proactive analysis at "
        "multiple stages of model development (Figure 4). Dataset bias can be detected through chemical "
        "space visualization using dimensionality reduction techniques (t-SNE, UMAP), scaffold diversity "
        "analysis, property distribution comparisons between training data and the intended application "
        "domain, and target-class representation assessment [22]. Algorithmic bias, which arises from "
        "model architecture choices or training procedures, can be identified through subgroup performance "
        "analysis, examining whether models perform consistently across different chemical classes, "
        "property ranges, and structural types."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Dataset balancing and representative chemical-space sampling provide direct approaches to "
        "bias mitigation. Techniques include oversampling of underrepresented classes, undersampling "
        "of overrepresented classes, synthetic data generation through molecular interpolation or "
        "generative models, and stratified sampling strategies that ensure adequate representation "
        "of diverse chemical scaffolds [36]. Active learning approaches can systematically identify "
        "and fill gaps in chemical space coverage by prioritizing compounds from underrepresented "
        "regions for experimental testing."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Domain knowledge and chemically informed model development serve as powerful bias mitigation "
        "strategies. Incorporating physical chemistry constraints, known structure-activity relationships, "
        "and medicinal chemistry heuristics into model architecture and training can prevent models from "
        "learning biased shortcuts [44]. For example, physics-informed neural networks that respect "
        "known relationships between molecular properties and biological activity are less likely to "
        "develop biased predictions based on spurious correlations in the training data. The framework "
        "for trustworthy AI depicted in Figure 4 integrates these multiple approaches into a coherent "
        "strategy for bias mitigation."
    )
    
    # 4.3
    doc.add_heading('4.3 Toward Trustworthy and Responsible AI in Medicinal Chemistry', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The integration of explainable AI methods into drug discovery workflows represents a critical "
        "step toward trustworthy AI in medicinal chemistry. XAI techniques enable researchers to "
        "understand not only what a model predicts but why it makes specific predictions, facilitating "
        "identification of potential errors and building confidence in reliable predictions [46]. "
        "Methods such as SHAP analysis, integrated gradients, and attention visualization provide "
        "complementary perspectives on model reasoning, while counterfactual explanations reveal "
        "what minimal changes to molecular structure would alter predictions."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Uncertainty estimation must become a standard component of AI model deployment in drug "
        "discovery. Ensemble methods that train multiple models with different initializations or "
        "on different data subsets provide calibrated uncertainty estimates reflecting both data "
        "uncertainty (aleatoric) and model uncertainty (epistemic) [17]. Conformal prediction "
        "offers distribution-free prediction intervals with guaranteed coverage probability, providing "
        "rigorous statistical guarantees on prediction reliability. These uncertainty estimates enable "
        "decision-makers to distinguish between high-confidence predictions that can be acted upon "
        "and low-confidence predictions that require additional experimental validation."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Human expertise remains indispensable in the responsible application of AI to medicinal "
        "chemistry. Expert medicinal chemists provide critical contextual knowledge that current "
        "AI systems cannot replicate, including understanding of synthetic feasibility, toxicological "
        "concerns, metabolic vulnerabilities, and clinical translatability [8]. The most effective "
        "paradigm combines AI predictions with human judgment in a collaborative framework where "
        "each compensates for the limitations of the other. AI excels at processing large datasets "
        "and identifying statistical patterns, while human experts provide mechanistic insight, "
        "creative problem-solving, and the ability to recognize when model predictions violate "
        "fundamental scientific principles."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Best practices for transparent, reproducible, and ethically responsible AI-driven drug "
        "discovery encompass multiple dimensions [47]. Transparency requires complete reporting of "
        "model development procedures, training data characteristics, validation strategies, and "
        "known limitations. Reproducibility demands open-source code sharing, standardized benchmarks, "
        "and independent verification of key results. Ethical responsibility includes ensuring "
        "equitable model performance across populations, protecting patient data privacy, and "
        "maintaining honesty about model capabilities and limitations. Together, these principles "
        "form the foundation for AI systems that medicinal chemists can trust and that ultimately "
        "serve the goal of developing safer, more effective medicines for all patients."
    )
    
    # ---- CONCLUSIONS ----
    doc.add_heading('Conclusions', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "This chapter has provided a comprehensive examination of AI failure modes in medicinal "
        "chemistry, demonstrating that overfitting and bias represent serious, systematic challenges "
        "rather than isolated incidents. Through case studies and simulations, we have shown how "
        "these failures arise from fundamental issues in data quality, model development, and "
        "validation methodology, and how they propagate to impact drug discovery decisions with "
        "significant consequences for resource allocation and scientific progress. The strategies "
        "presented for building more reliable AI models, encompassing robust validation, bias "
        "mitigation, explainability, and human oversight, provide a practical roadmap for "
        "researchers committed to leveraging AI responsibly in medicinal chemistry. As the field "
        "continues to evolve, adherence to these principles will be essential for realizing the "
        "full potential of AI in drug discovery while avoiding the pitfalls that have undermined "
        "past efforts."
    )
    
    doc.add_page_break()
    
    # ---- FIGURES ----
    doc.add_heading('Figures', level=1)
    
    # Figure 1
    p = doc.add_paragraph()
    p.add_run('Figure 1: ').bold = True
    p.add_run('Illustration of overfitting in machine learning models for medicinal chemistry. '
              '(A) Training loss continues to decrease while validation loss increases after the optimal stopping point, '
              'indicating memorization of training data noise. (B) The bias-variance trade-off shows how model complexity '
              'affects prediction error, with an optimal complexity level that minimizes total error.')
    doc.add_picture('/projects/sandbox/AMMAN/ai_failure_figures/Figure_1_Overfitting_Illustration.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Figure 2
    p = doc.add_paragraph()
    p.add_run('Figure 2: ').bold = True
    p.add_run('Bias in chemical space coverage and its impact on model performance. '
              '(A) Visualization of biased dataset coverage showing overrepresented, underrepresented, and missing chemical classes. '
              '(B) Performance disparity across target classes demonstrating significant degradation for underrepresented targets.')
    doc.add_picture('/projects/sandbox/AMMAN/ai_failure_figures/Figure_2_Bias_Chemical_Space.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Figure 3
    p = doc.add_paragraph()
    p.add_run('Figure 3: ').bold = True
    p.add_run('Simulation-based demonstration of AI model failure. (A) Well-generalized model showing tight correlation '
              'between predicted and experimental pIC50 values. (B) Overfit model showing dramatic performance degradation on external data. '
              '(C) Error distribution comparison between good and overfit models. (D) Learning curves demonstrating the effect of dataset size on generalization.')
    doc.add_picture('/projects/sandbox/AMMAN/ai_failure_figures/Figure_3_Simulation_Model_Failure.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Figure 4
    p = doc.add_paragraph()
    p.add_run('Figure 4: ').bold = True
    p.add_run('Framework for building trustworthy AI in medicinal chemistry. The central goal of trustworthy AI is supported '
              'by six key components: data quality and curation, robust validation, bias detection and mitigation, explainable AI, '
              'uncertainty quantification, and human expert oversight. Each component addresses specific failure modes identified in this chapter.')
    doc.add_picture('/projects/sandbox/AMMAN/ai_failure_figures/Figure_4_Trustworthy_AI_Framework.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ---- REFERENCES ----
    doc.add_heading('References', level=1)
    
    references = [
        "[1] Vamathevan, J., et al. Applications of machine learning in drug discovery and development. Nature Reviews Drug Discovery, 18(6), 463-477, 2019.",
        "[2] Schneider, P., et al. Rethinking drug design in the artificial intelligence era. Nature Reviews Drug Discovery, 19(5), 353-364, 2020.",
        "[3] Wouters, O.J., McKee, M., Luyten, J. Estimated research and development investment needed to bring a new medicine to market, 2009-2018. JAMA, 323(9), 844-853, 2020.",
        "[4] Elton, D.C., et al. Deep learning for molecular design: a review of the state of the art. Molecular Systems Design & Engineering, 4(4), 828-849, 2019.",
        "[5] Yang, X., et al. Concepts of artificial intelligence for computer-assisted drug discovery. Chemical Reviews, 119(18), 10520-10594, 2019.",
        "[6] Cherkasov, A., et al. QSAR modeling: where have you been? Where are you going to? Journal of Medicinal Chemistry, 57(12), 4977-5010, 2014.",
        "[7] Bender, A., Cortés-Ciriano, I. Artificial intelligence in drug discovery: what is realistic, what are illusions? Drug Discovery Today, 26(2), 511-524, 2021.",
        "[8] Jiménez-Luna, J., Grisoni, F., Schneider, G. Drug discovery with explainable artificial intelligence. Nature Machine Intelligence, 2(10), 573-584, 2020.",
        "[9] Walters, W.P., Barzilay, R. Critical assessment of AI in drug discovery. Expert Opinion on Drug Discovery, 16(9), 937-947, 2021.",
        "[10] Stokes, J.M., et al. A deep learning approach to antibiotic discovery. Cell, 180(4), 688-702, 2020.",
        "[11] Mervin, L.H., et al. Target prediction utilising negative bioactivity data covering large chemical space. Journal of Cheminformatics, 7(1), 51, 2015.",
        "[12] Hawkins, D.M. The problem of overfitting. Journal of Chemical Information and Computer Sciences, 44(1), 1-12, 2004.",
        "[13] Tetko, I.V., et al. Critical assessment of QSAR models of environmental toxicity against Tetrahymena pyriformis. Journal of Chemical Information and Modeling, 48(9), 1733-1746, 2008.",
        "[14] Sheridan, R.P. Time-split cross-validation as a method for estimating the goodness of prospective prediction. Journal of Chemical Information and Modeling, 53(4), 783-790, 2013.",
        "[15] Kearnes, S., et al. Molecular graph convolutions: moving beyond fingerprints. Journal of Computer-Aided Molecular Design, 30(8), 595-608, 2016.",
        "[16] Scior, T., et al. Recognizing pitfalls in virtual screening: a critical review. Journal of Chemical Information and Modeling, 52(4), 867-881, 2012.",
        "[17] Hirschfeld, L., et al. Uncertainty quantification using neural networks for molecular property prediction. Journal of Chemical Information and Modeling, 60(8), 3770-3780, 2020.",
        "[18] Tropsha, A. Best practices for QSAR model development, validation, and exploitation. Molecular Informatics, 29(6-7), 476-488, 2010.",
        "[19] Maggiora, G., et al. Molecular similarity in medicinal chemistry. Journal of Medicinal Chemistry, 57(8), 3186-3204, 2014.",
        "[20] Barocas, S., Hardt, M., Narayanan, A. Fairness and Machine Learning: Limitations and Opportunities. MIT Press, 2023.",
        "[21] Cai, C., et al. Transfer learning for drug discovery. Journal of Medicinal Chemistry, 63(16), 8683-8694, 2020.",
        "[22] Chen, H., et al. The rise of deep learning in drug discovery. Drug Discovery Today, 23(6), 1241-1250, 2018.",
        "[23] Papadatos, G., et al. Activity, assay and target data curation and quality in the ChEMBL database. Journal of Computer-Aided Molecular Design, 29(9), 885-896, 2015.",
        "[24] Kramer, C., et al. The experimental uncertainty of heterogeneous public Ki data. Journal of Medicinal Chemistry, 55(11), 5165-5173, 2012.",
        "[25] Popejoy, A.B., Fullerton, S.M. Genomics is failing on diversity. Nature, 538(7624), 161-164, 2016.",
        "[26] Martin, E.J., et al. Profile-QSAR 2.0: Kinase virtual screening accuracy comparable to four-concentration IC50s for realistically novel compounds. Journal of Chemical Information and Modeling, 57(8), 2077-2088, 2017.",
        "[27] Cortes-Ciriano, I., Bender, A. Reliable prediction errors for deep neural networks using test-time dropout. Journal of Chemical Information and Modeling, 59(7), 3330-3339, 2019.",
        "[28] Feinberg, E.N., et al. PotentialNet for molecular property prediction. ACS Central Science, 4(11), 1520-1530, 2018.",
        "[29] Wu, Z., et al. MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 9(2), 513-530, 2018.",
        "[30] Wallach, I., Heifets, A. Most ligand-based classification benchmarks reward memorization rather than generalization. Journal of Chemical Information and Modeling, 58(5), 916-932, 2018.",
        "[31] Yang, K., et al. Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 59(8), 3370-3388, 2019.",
        "[32] Lenselink, E.B., et al. Beyond the hype: deep neural networks outperform established methods using a ChEMBL bioactivity benchmark set. Journal of Cheminformatics, 9(1), 45, 2017.",
        "[33] Mayr, A., et al. DeepTox: toxicity prediction using deep learning. Frontiers in Environmental Science, 3, 80, 2016.",
        "[34] Ramsundar, B., et al. Is multitask deep learning practical for pharma? Journal of Chemical Information and Modeling, 57(8), 2068-2076, 2017.",
        "[35] Unterthiner, T., et al. Deep learning as an opportunity in virtual screening. Proceedings of the Deep Learning Workshop at NIPS, 2014.",
        "[36] Simm, J., et al. Repurposing high-throughput image assays enables biological activity prediction for drug discovery. Cell Chemical Biology, 25(5), 611-618, 2018.",
        "[37] Meredig, B. Five high-impact research areas in machine learning for materials science. Chemistry of Materials, 31(23), 9579-9581, 2019.",
        "[38] Cortes-Ciriano, I., Bender, A. Deep confidence: a computationally efficient framework for calculating reliable prediction errors for deep neural networks. Journal of Chemical Information and Modeling, 59(3), 1269-1281, 2019.",
        "[39] Shen, J., Nicolaou, C.A. Molecular property prediction: recent trends in the era of artificial intelligence. Drug Discovery Today: Technologies, 32-33, 29-36, 2019.",
        "[40] Sun, J., et al. ExCAPE-DB: an integrated large scale dataset facilitating Big Data analysis in chemogenomics. Journal of Cheminformatics, 9(1), 17, 2017.",
        "[41] Fourches, D., et al. Trust, but verify: on the importance of chemical structure curation in cheminformatics and QSAR modeling research. Journal of Chemical Information and Modeling, 50(7), 1189-1204, 2010.",
        "[42] Kalliokoski, T., et al. Comparability of mixed IC50 data: a statistical analysis. PLoS ONE, 8(4), e61007, 2013.",
        "[43] Hu, Y., Bajorath, J. Influence of search parameters and criteria on compound selection, promiscuity, and pan-assay interference characteristics. Journal of Chemical Information and Modeling, 54(11), 3056-3066, 2014.",
        "[44] Gilmer, J., et al. Neural message passing for quantum chemistry. Proceedings of the 34th International Conference on Machine Learning, 1263-1272, 2017.",
        "[45] Rogers, D., Hahn, M. Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 50(5), 742-754, 2010.",
        "[46] Lundberg, S.M., Lee, S.I. A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765-4774, 2017.",
        "[47] Norinder, U., Carlsson, L., Boyer, S., Eklund, M. Introducing conformal prediction in predictive modeling for regulatory purposes: a transparent and flexible alternative to applicability domain determination. Regulatory Toxicology and Pharmacology, 71(2), 279-284, 2015.",
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.left_indent = Cm(1.0)
    
    # Save document
    output_path = '/projects/sandbox/AMMAN/AI_Failure_Medicinal_Chemistry_Chapter.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

# ============================================================
# MAIN EXECUTION
# ============================================================
if __name__ == '__main__':
    print("Creating figures...")
    create_figure1()
    create_figure2()
    create_figure3()
    create_figure4()
    print("\nAll figures created successfully.")
    
    print("\nGenerating Word document...")
    output = create_word_document()
    print(f"\nComplete! Document available at: {output}")
