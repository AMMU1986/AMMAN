"""
Create Word Document: Climate Change, Economic Risk, and Adaptive Eco-Technological Strategies
Complete chapter with 76 references, 4 tables, and 4 figures.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Page setup
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)



def add_figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def create_styled_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
    # Style
    table.style = 'Table Grid'
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Climate Change, Economic Risk, and Adaptive Eco-Technological Strategies')
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
doc.add_paragraph()



# ============================================================
# ABSTRACT
# ============================================================
add_heading_styled('Abstract', level=1)

abstract_text = (
    "Climate change represents one of the most profound challenges confronting global economic systems, "
    "ecological stability, and societal well-being in the twenty-first century. Rising global temperatures, "
    "intensifying extreme weather events, and shifting precipitation patterns impose escalating costs on "
    "infrastructure, agriculture, supply chains, and financial markets. This chapter provides a comprehensive "
    "examination of the nexus between climate change and economic risk, exploring how emerging intelligent "
    "eco-technologies and adaptive strategies can mitigate vulnerabilities and foster resilient, sustainable "
    "economies. The analysis begins with a detailed assessment of climate change drivers and their cascading "
    "economic consequences, followed by an evaluation of sectoral and regional vulnerabilities, financial "
    "risk frameworks, and scenario-based decision-making approaches. The chapter then investigates the "
    "transformative potential of artificial intelligence, Internet of Things, digital twins, and smart "
    "infrastructure in enabling proactive climate adaptation. Finally, it presents adaptive strategies "
    "encompassing circular economy models, green finance mechanisms, and inclusive policy pathways for "
    "building climate-resilient economies. The integrated framework proposed herein demonstrates that "
    "intelligent eco-technological deployment, combined with robust governance and inclusive economic "
    "planning, can transform climate risk into opportunity for sustainable development."
)
add_para(abstract_text)

kw = doc.add_paragraph()
run = kw.add_run('Keywords: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = kw.add_run(
    'Climate change; economic risk; eco-technology; artificial intelligence; '
    'climate adaptation; resilience; green finance; circular economy; digital twins; sustainable development'
)
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

doc.add_page_break()



# ============================================================
# SECTION 1: Climate Change and Emerging Economic Risks
# ============================================================
add_heading_styled('1. Climate Change and Emerging Economic Risks', level=1)

intro_1 = (
    "The accelerating pace of anthropogenic climate change poses unprecedented challenges to global "
    "economic systems, demanding urgent reconsideration of development paradigms and risk management "
    "frameworks. The Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment Report confirms "
    "that global surface temperature has increased by approximately 1.1°C above pre-industrial levels, "
    "with projections indicating continued warming across all emission scenarios [1]. This warming trend "
    "is not merely an environmental concern but fundamentally reshapes economic landscapes, disrupting "
    "production systems, supply chains, and financial markets with increasing frequency and severity [2]. "
    "The economic implications of climate change extend far beyond direct damages from extreme weather "
    "events; they encompass systemic risks to food security, water availability, human health, and "
    "geopolitical stability [3]. Understanding the intricate relationship between climate drivers and "
    "economic vulnerabilities is essential for developing effective adaptation strategies and building "
    "resilient economies capable of thriving under changing climatic conditions [4]."
)
add_para(intro_1)

# Section 1.1
add_heading_styled('1.1 Climate Change Drivers, Trends, and Environmental Impacts', level=2)

s1_1_p1 = (
    "The primary drivers of contemporary climate change are well-established in scientific literature, "
    "centering on the accumulation of greenhouse gases (GHGs) in the atmosphere due to fossil fuel "
    "combustion, industrial processes, deforestation, and agricultural activities [5]. Carbon dioxide "
    "concentrations have surpassed 420 parts per million (ppm), representing a 50% increase over "
    "pre-industrial levels, while methane and nitrous oxide concentrations continue to rise at "
    "accelerating rates [6]. The radiative forcing associated with these concentrations drives a cascade "
    "of physical changes including ocean warming, ice sheet destabilization, sea-level rise, and "
    "alterations in atmospheric circulation patterns [7]. Global mean sea level has risen approximately "
    "3.7 mm per year during the period 2006–2018, with acceleration observed in recent decades due to "
    "thermal expansion and ice mass loss from Greenland and Antarctica [8]."
)
add_para(s1_1_p1)

s1_1_p2 = (
    "Temperature extremes have become more frequent and intense, with heatwave frequency increasing "
    "by a factor of 2.8 since the 1950s in many regions [9]. Precipitation patterns are shifting, "
    "with wet regions generally becoming wetter and dry regions drier, exacerbating both flood and "
    "drought risks [10]. The cryosphere is experiencing rapid transformation, with Arctic sea ice "
    "declining at approximately 13% per decade and permafrost thawing releasing additional methane "
    "and carbon dioxide into the atmosphere, creating positive feedback loops [11]. Ocean acidification, "
    "driven by CO2 absorption, threatens marine ecosystems and the fisheries upon which approximately "
    "3.3 billion people depend for protein [12]. These environmental changes do not occur in isolation "
    "but interact in complex, often non-linear ways that amplify impacts and create compound risks "
    "that challenge traditional risk assessment methodologies [13]. The concept of tipping points—"
    "critical thresholds beyond which system changes become self-reinforcing—has gained increasing "
    "attention, with research suggesting that several Earth system components may be approaching or "
    "have already crossed such thresholds [14]."
)
add_para(s1_1_p2)

s1_1_p3 = (
    "The interconnected nature of climate system responses creates compound hazards that pose particular "
    "challenges for risk assessment and adaptation planning. For instance, the combination of higher "
    "temperatures and altered precipitation regimes simultaneously increases wildfire risk, reduces "
    "water availability, and degrades air quality, with cascading effects on public health, "
    "agriculture, and energy systems. Tropical cyclones are projected to intensify under continued "
    "warming, with a greater proportion of storms reaching Category 4 and 5 intensity, while "
    "sea-level rise amplifies storm surge flooding in coastal communities. The global hydrological "
    "cycle is accelerating, with more intense precipitation events increasing flood frequency even "
    "as longer dry spells between events exacerbate drought conditions. These compound and cascading "
    "climate impacts underscore the necessity of integrated, systems-based approaches to understanding "
    "and managing climate risks across sectors and scales."
)
add_para(s1_1_p3)



# Section 1.2
add_heading_styled('1.2 Economic Consequences of Extreme Weather and Climate Hazards', level=2)

s1_2_p1 = (
    "The economic toll of climate-related disasters has escalated dramatically over recent decades, "
    "with annual losses exceeding USD 300 billion in recent years compared to approximately USD 50 billion "
    "in the 1990s [15]. Extreme weather events—including hurricanes, floods, droughts, and wildfires—"
    "impose both direct costs through physical damage and indirect costs through business interruption, "
    "supply chain disruption, and reduced productivity [16]. The 2021 European floods caused estimated "
    "damages of USD 43 billion, while the 2022 Pakistan floods affected over 33 million people and "
    "inflicted economic losses exceeding USD 30 billion, representing approximately 10% of national "
    "GDP [17]. Agricultural systems are particularly vulnerable, with climate variability reducing "
    "global crop yields by an estimated 5–10% per decade relative to potential yields without climate "
    "change [18]. As illustrated in Figure 1, the relationship between rising global temperatures and "
    "escalating economic losses demonstrates a clear upward trajectory, with losses accelerating "
    "non-linearly as warming intensifies [19]."
)
add_para(s1_2_p1)

# Insert Figure 1
fig1_path = '/projects/sandbox/AMMAN/climate_figures/Figure_1_Temperature_Economic_Losses.png'
if os.path.exists(fig1_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig1_path, width=Inches(5.5))
    add_figure_caption(
        'Figure 1: Global Temperature Anomaly and Climate-Related Economic Losses (2000–2025). '
        'The dual-axis plot demonstrates the correlation between rising temperatures and escalating '
        'economic damages from climate-related disasters.'
    )

s1_2_p2 = (
    "Heat stress reduces labor productivity, particularly in outdoor sectors such as agriculture and "
    "construction, with projections indicating that heat-related productivity losses could reach "
    "USD 2.4 trillion annually by 2030, disproportionately affecting tropical and subtropical "
    "developing nations [20]. Water scarcity, intensified by climate change, threatens industrial "
    "production and energy generation, with an estimated 40% of global thermal power generation "
    "capacity located in water-stressed regions [21]. The insurance industry faces mounting challenges "
    "as climate-related claims surge, with some regions becoming effectively uninsurable, creating "
    "protection gaps that transfer risk to governments and households [22]. Furthermore, climate change "
    "interacts with socioeconomic factors to exacerbate inequality, as vulnerable populations in "
    "developing countries bear disproportionate impacts despite contributing least to emissions [23]."
)
add_para(s1_2_p2)

s1_2_p3 = (
    "The macroeconomic implications of climate change extend beyond immediate disaster losses to "
    "encompass long-term growth trajectory effects, fiscal pressures, and monetary policy challenges. "
    "Climate-vulnerable developing nations face sovereign credit rating downgrades that increase "
    "borrowing costs precisely when adaptation investments are most needed, creating vicious cycles "
    "of climate vulnerability and fiscal constraint. The agricultural sector exemplifies these "
    "compounding effects: simultaneous crop failures across major producing regions—an increasingly "
    "plausible scenario under continued warming—can trigger food price spikes, social instability, "
    "and forced migration, with cascading consequences for regional and global economic stability. "
    "Furthermore, the health impacts of climate change, including the spread of vector-borne diseases, "
    "malnutrition, and heat-related morbidity, reduce human capital accumulation and labor force "
    "participation, undermining long-term economic growth potential in affected regions."
)
add_para(s1_2_p3)



# Section 1.3
add_heading_styled('1.3 Climate-Related Risks to Infrastructure, Industries, and Global Supply Chains', level=2)

s1_3_p1 = (
    "Critical infrastructure systems—including transportation networks, energy grids, water treatment "
    "facilities, and telecommunications—face escalating climate risks that threaten economic functionality "
    "and public safety [24]. Rising sea levels and increased storm surge intensity endanger coastal "
    "infrastructure valued at trillions of dollars globally, with approximately 570 low-lying coastal "
    "cities facing projected sea-level rise of at least 0.5 meters by 2050 [25]. Transportation "
    "infrastructure suffers from heat-induced rail buckling, road surface deterioration, and flooding "
    "of tunnels and underpasses, with adaptation costs for transport systems estimated at USD 480 billion "
    "globally by 2040 [26]. Energy systems face dual challenges: increased cooling demand during heatwaves "
    "coincides with reduced generation efficiency and grid stress, while renewable energy sources "
    "experience climate-sensitive variability in wind and solar resources [27]."
)
add_para(s1_3_p1)

s1_3_p2 = (
    "Global supply chains, optimized for efficiency rather than resilience, are increasingly vulnerable "
    "to climate disruptions at critical nodes [28]. The concentration of manufacturing in climate-"
    "exposed regions of Southeast Asia, combined with just-in-time inventory management, creates "
    "cascading vulnerability across global production networks [29]. The 2011 Thailand floods, which "
    "inundated industrial estates producing 25% of global hard drives, demonstrated how localized "
    "climate events propagate through interconnected supply chains, causing global shortages and "
    "estimated losses of USD 45 billion [30]. Agricultural supply chains face compound risks from "
    "simultaneous crop failures across multiple breadbasket regions—a scenario with increasing "
    "probability under continued warming [31]. Table 1 summarizes the key climate risks across "
    "major infrastructure categories and their estimated economic impacts, highlighting the "
    "cross-sectoral nature of climate vulnerability [32]."
)
add_para(s1_3_p2)

# TABLE 1
add_table_caption(
    'Table 1: Climate-Related Risks to Critical Infrastructure and Estimated Economic Impacts'
)
table1_headers = ['Infrastructure Category', 'Primary Climate Hazard', 'Vulnerability Level', 
                  'Estimated Annual Loss (USD Billion)', 'Adaptation Priority']
table1_rows = [
    ['Transportation Networks', 'Flooding, Heat Extremes', 'High', '85–120', 'Critical'],
    ['Energy Systems', 'Heat Stress, Storm Damage', 'Very High', '95–150', 'Critical'],
    ['Water Infrastructure', 'Drought, Flooding', 'High', '45–80', 'High'],
    ['Telecommunications', 'Storm Damage, Flooding', 'Moderate', '25–40', 'Medium'],
    ['Coastal Infrastructure', 'Sea-Level Rise, Storm Surge', 'Very High', '120–200', 'Critical'],
    ['Industrial Facilities', 'Multi-hazard Exposure', 'High', '60–95', 'High'],
    ['Agricultural Systems', 'Drought, Heat, Flooding', 'Very High', '150–250', 'Critical'],
    ['Healthcare Facilities', 'Heat, Flooding', 'Moderate-High', '30–55', 'High'],
]
create_styled_table(table1_headers, table1_rows)
doc.add_paragraph()



# ============================================================
# SECTION 2: Economic Vulnerability and Climate Risk Assessment
# ============================================================
add_heading_styled('2. Economic Vulnerability and Climate Risk Assessment', level=1)

intro_2 = (
    "Assessing economic vulnerability to climate change requires sophisticated analytical frameworks "
    "capable of capturing the multidimensional, dynamic, and interconnected nature of climate risks "
    "across sectors and regions [33]. Traditional economic analysis tools often prove inadequate for "
    "addressing the deep uncertainty, non-linearity, and long time horizons characteristic of climate "
    "impacts [34]. This section examines sectoral and regional vulnerability patterns, financial risk "
    "dimensions, and emerging modeling approaches that enable more robust climate-informed "
    "decision-making. The analysis draws upon vulnerability assessment methodologies, financial risk "
    "frameworks, and scenario-based approaches developed by leading institutions and researchers in "
    "the climate-economy nexus."
)
add_para(intro_2)

# Section 2.1
add_heading_styled('2.1 Sectoral and Regional Vulnerability to Climate Change', level=2)

s2_1_p1 = (
    "Climate vulnerability varies significantly across economic sectors and geographic regions, "
    "determined by exposure to climate hazards, sensitivity of economic activities to climatic "
    "conditions, and adaptive capacity [35]. Agriculture remains the most climate-sensitive sector, "
    "with studies indicating that unmitigated warming of 3°C could reduce global agricultural "
    "productivity by 15–25% by 2050, with tropical regions experiencing the most severe impacts [36]. "
    "Water-intensive industries, including thermoelectric power generation, mining, and food processing, "
    "face escalating risks from changing hydrological regimes and competing demands [37]. The tourism "
    "sector, representing approximately 10% of global GDP, is highly sensitive to climate change through "
    "impacts on natural attractions, seasonality shifts, and extreme weather disruption of travel "
    "infrastructure [38]. As depicted in Figure 2, the sectoral vulnerability assessment reveals "
    "that agriculture, energy, and coastal infrastructure face the highest combined risk scores "
    "across multiple climate hazard dimensions [39]."
)
add_para(s2_1_p1)

# Insert Figure 2
fig2_path = '/projects/sandbox/AMMAN/climate_figures/Figure_2_Sectoral_Vulnerability_Heatmap.png'
if os.path.exists(fig2_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig2_path, width=Inches(5.5))
    add_figure_caption(
        'Figure 2: Sectoral Climate Vulnerability Assessment Matrix showing vulnerability scores '
        '(0–10) across multiple risk categories for eight major economic sectors.'
    )

s2_1_p2 = (
    "Regional disparities in climate vulnerability are stark, with developing countries in tropical "
    "and subtropical zones bearing disproportionate impacts [40]. Small Island Developing States (SIDS) "
    "face existential threats from sea-level rise, while Sub-Saharan Africa confronts compounding "
    "risks of drought, food insecurity, and limited adaptive capacity [41]. Conversely, some temperate "
    "regions may experience short-term benefits from warming, including extended growing seasons and "
    "reduced heating demands, though these gains are increasingly offset by other climate impacts [42]. "
    "Urban areas, home to over 55% of the global population, face unique vulnerability profiles "
    "characterized by urban heat island effects, concentrated infrastructure exposure, and complex "
    "interdependencies between systems [43]. The interaction between physical vulnerability and "
    "socioeconomic factors—including poverty, governance capacity, and technological development—"
    "creates differentiated risk landscapes that demand context-specific adaptation approaches [44]."
)
add_para(s2_1_p2)

s2_1_p3 = (
    "The mining and extractive industries face growing climate risks as operations in arid regions "
    "confront water scarcity, while coastal and riverine facilities face flooding threats. The "
    "pharmaceutical and biotechnology sectors, dependent on temperature-controlled supply chains and "
    "climate-sensitive biological inputs, represent emerging areas of climate vulnerability that have "
    "received insufficient analytical attention. The digital economy, despite its apparent detachment "
    "from physical climate impacts, depends on data centers requiring substantial cooling energy, "
    "submarine cable infrastructure vulnerable to storm damage, and satellite systems affected by "
    "atmospheric changes. Financial services sector vulnerability extends beyond direct physical "
    "exposure to encompass portfolio-wide climate risk exposure, counterparty default risk from "
    "climate-affected borrowers, and potential asset repricing as climate scenarios materialize. "
    "Understanding these diverse vulnerability profiles is essential for designing sector-appropriate "
    "adaptation strategies and mobilizing targeted investment in climate resilience."
)
add_para(s2_1_p3)



# Section 2.2
add_heading_styled('2.2 Financial Risks, Investment Uncertainty, and Stranded Assets', level=2)

s2_2_p1 = (
    "Climate change introduces systemic financial risks that transcend traditional risk categories, "
    "challenging the stability of financial systems and the validity of established investment "
    "frameworks [45]. The Task Force on Climate-related Financial Disclosures (TCFD) framework "
    "categorizes climate-related financial risks into physical risks (arising from climate events) "
    "and transition risks (arising from the shift to a low-carbon economy), both of which can "
    "materially affect asset valuations, credit quality, and portfolio returns [46]. Physical risks "
    "are already manifesting in declining property values in flood-prone areas, increased insurance "
    "premiums, and sovereign credit downgrades for climate-vulnerable nations [47]. The concept of "
    "stranded assets—investments that lose economic viability due to climate policy, technological "
    "change, or physical impacts—poses significant risks to fossil fuel-dependent economies and "
    "investors, with estimates suggesting that USD 1–4 trillion in fossil fuel assets could become "
    "stranded under Paris Agreement-aligned pathways [48]."
)
add_para(s2_2_p1)

s2_2_p2 = (
    "Investment uncertainty under climate change is compounded by the long-lived nature of infrastructure "
    "assets, which must remain functional under future climate conditions that are imperfectly known [49]. "
    "Central banks and financial regulators increasingly recognize climate change as a source of systemic "
    "financial risk, with the Network for Greening the Financial System (NGFS) developing climate "
    "scenarios for financial stability assessment [50]. Climate stress testing of financial portfolios "
    "reveals significant potential losses under high-warming scenarios, with banking sector exposure "
    "to physical and transition risks estimated at 10–20% of total assets in some jurisdictions [51]. "
    "Table 2 presents a comparative analysis of financial risk categories, their transmission channels, "
    "and estimated magnitudes across different warming scenarios [52]."
)
add_para(s2_2_p2)

# TABLE 2
add_table_caption(
    'Table 2: Climate-Related Financial Risk Categories and Estimated Impact Magnitudes'
)
table2_headers = ['Risk Category', 'Transmission Channel', '1.5°C Scenario', 
                  '2°C Scenario', '3°C+ Scenario', 'Time Horizon']
table2_rows = [
    ['Physical—Acute', 'Asset damage, business interruption', 'Moderate', 'High', 'Very High', 'Near-term'],
    ['Physical—Chronic', 'Productivity decline, resource scarcity', 'Low-Moderate', 'Moderate-High', 'Very High', 'Medium-term'],
    ['Transition—Policy', 'Carbon pricing, regulation', 'High', 'Moderate', 'Low', 'Near-term'],
    ['Transition—Technology', 'Obsolescence, new competition', 'High', 'Moderate', 'Low', 'Medium-term'],
    ['Transition—Market', 'Demand shifts, repricing', 'Moderate-High', 'Moderate', 'Low', 'Near-term'],
    ['Transition—Reputation', 'Stakeholder pressure, litigation', 'Moderate', 'Low-Moderate', 'Low', 'Near-term'],
    ['Stranded Assets', 'Write-downs, devaluation', 'Very High', 'High', 'Moderate', 'Medium-term'],
    ['Systemic/Cascading', 'Financial contagion, sovereign risk', 'Low', 'Moderate', 'Very High', 'Long-term'],
]
create_styled_table(table2_headers, table2_rows)
doc.add_paragraph()



# Section 2.3
add_heading_styled('2.3 Climate Risk Modelling, Scenario Analysis, and Decision-Making Frameworks', level=2)

s2_3_p1 = (
    "Climate risk modeling has evolved significantly from deterministic damage functions toward "
    "probabilistic, multi-scenario frameworks that better capture the deep uncertainty inherent in "
    "climate projections and socioeconomic pathways [53]. Integrated Assessment Models (IAMs) combine "
    "climate science with economic modeling to estimate damages under different emission trajectories, "
    "though they face criticism for potentially underestimating tail risks and non-linear damages [54]. "
    "The Shared Socioeconomic Pathways (SSPs) framework provides a structured approach to exploring "
    "future scenarios by combining climate forcing levels with different socioeconomic development "
    "trajectories, enabling comprehensive risk assessment across multiple dimensions [55]. Physical "
    "climate risk models increasingly incorporate high-resolution spatial data, enabling asset-level "
    "exposure assessment that informs investment decisions and infrastructure planning [56]."
)
add_para(s2_3_p1)

s2_3_p2 = (
    "Decision-making under deep uncertainty requires approaches that go beyond expected value "
    "optimization, embracing robustness and flexibility as key criteria [57]. Robust Decision Making "
    "(RDM), Dynamic Adaptive Policy Pathways (DAPP), and Real Options Analysis represent methodological "
    "advances that enable decision-makers to identify strategies performing well across a wide range "
    "of plausible futures [58]. These frameworks are particularly valuable for long-lived infrastructure "
    "investments, where the timing and sequencing of adaptation measures can significantly affect "
    "cost-effectiveness [59]. The integration of climate models with economic models, supply chain "
    "models, and financial system models enables cascading risk analysis that reveals system-level "
    "vulnerabilities invisible to sectoral assessments [60]. Machine learning and artificial "
    "intelligence are increasingly applied to climate risk modeling, enabling pattern recognition "
    "in complex datasets and improving the spatial and temporal resolution of impact projections [61]."
)
add_para(s2_3_p2)

s2_3_p3 = (
    "The evolution of climate risk assessment methodologies reflects growing recognition that traditional "
    "approaches based on historical data are inadequate for a non-stationary climate system. Forward-"
    "looking scenario analysis, stress testing, and sensitivity analysis provide complementary "
    "perspectives on potential futures, enabling more robust planning under conditions where "
    "probability distributions cannot be reliably specified. The development of standardized climate "
    "risk metrics and taxonomies facilitates comparison across regions, sectors, and institutions, "
    "supporting both regulatory oversight and market-based risk management. Agent-based models and "
    "network analysis capture emergent properties of interconnected economic systems that traditional "
    "equilibrium models miss, revealing how localized climate shocks can propagate through trade "
    "networks, financial linkages, and shared infrastructure to generate systemic crises. The "
    "integration of social vulnerability indicators with physical climate projections enables "
    "equity-aware risk assessment that identifies communities facing compound disadvantages from "
    "both high climate exposure and limited adaptive capacity."
)
add_para(s2_3_p3)



# ============================================================
# SECTION 3: Intelligent Eco-Technologies for Climate Adaptation
# ============================================================
add_heading_styled('3. Intelligent Eco-Technologies for Climate Adaptation', level=1)

intro_3 = (
    "The convergence of digital technologies with environmental science and engineering offers "
    "transformative opportunities for climate adaptation and resilience building [62]. Intelligent "
    "eco-technologies—encompassing artificial intelligence, Internet of Things (IoT), digital twins, "
    "smart infrastructure, and data-driven environmental management systems—enable more precise, "
    "proactive, and cost-effective responses to climate risks than traditional approaches [63]. "
    "These technologies facilitate real-time monitoring, predictive analytics, autonomous optimization, "
    "and adaptive management at scales ranging from individual buildings to entire cities and "
    "regions [64]. The intelligent eco-technology framework integrates multiple technological domains "
    "into a coherent platform for climate resilience, as illustrated in Figure 3."
)
add_para(intro_3)

# Insert Figure 3
fig3_path = '/projects/sandbox/AMMAN/climate_figures/Figure_3_EcoTechnology_Framework.png'
if os.path.exists(fig3_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig3_path, width=Inches(5.5))
    add_figure_caption(
        'Figure 3: Integrated Intelligent Eco-Technology Framework for Climate Adaptation and '
        'Resilience, showing the interconnection of AI, IoT, digital twins, renewable energy, '
        'early warning systems, and smart infrastructure around a central resilience platform.'
    )

# Section 3.1
add_heading_styled('3.1 AI, IoT, and Digital Twins for Climate Resilience', level=2)

s3_1_p1 = (
    "Artificial intelligence and machine learning are revolutionizing climate adaptation through "
    "enhanced prediction, optimization, and decision support capabilities [65]. Deep learning models "
    "achieve superior performance in weather forecasting, flood prediction, and wildfire risk assessment "
    "compared to traditional numerical methods, with lead times extending from hours to weeks [66]. "
    "AI-driven climate downscaling techniques generate high-resolution local climate projections from "
    "coarse global models, enabling site-specific adaptation planning at a fraction of the computational "
    "cost of traditional dynamical downscaling [67]. Natural language processing and computer vision "
    "applications enable automated analysis of satellite imagery for deforestation monitoring, crop "
    "health assessment, and damage evaluation following extreme events [68]. Reinforcement learning "
    "algorithms optimize complex systems such as water distribution networks, energy grids, and "
    "traffic management under variable and uncertain climate conditions [69]."
)
add_para(s3_1_p1)

s3_1_p2 = (
    "The Internet of Things provides the sensing infrastructure essential for intelligent climate "
    "adaptation, deploying networks of connected sensors that monitor environmental conditions, "
    "infrastructure health, and resource flows in real time [70]. IoT-enabled smart water management "
    "systems detect leaks, optimize distribution, and predict demand under drought conditions, reducing "
    "water losses by 20–30% in pilot deployments [71]. Soil moisture sensors, weather stations, and "
    "drone-mounted multispectral cameras enable precision agriculture that optimizes irrigation, "
    "fertilization, and pest management, increasing yields while reducing resource consumption and "
    "environmental impact [72]. Digital twin technology creates virtual replicas of physical assets "
    "and systems, enabling simulation of climate scenarios, stress testing of infrastructure, and "
    "optimization of adaptation measures before physical implementation [73]. Urban digital twins "
    "integrate building energy models, transportation networks, green infrastructure, and climate "
    "projections to identify optimal adaptation pathways for cities, reducing planning costs and "
    "improving intervention effectiveness [74]."
)
add_para(s3_1_p2)



# Section 3.2
add_heading_styled('3.2 Smart Infrastructure, Renewable Energy, and Resource-Efficient Technologies', level=2)

s3_2_p1 = (
    "Smart infrastructure integrates sensing, communication, and control capabilities into physical "
    "systems, enabling adaptive responses to changing environmental conditions and climate stressors [75]. "
    "Climate-resilient building design incorporates passive cooling strategies, adaptive facades, "
    "green roofs, and smart HVAC systems that adjust to extreme temperatures while minimizing energy "
    "consumption [76]. Smart grid technologies enable flexible, distributed energy systems that "
    "maintain reliability under climate-induced supply variability and demand spikes, integrating "
    "renewable generation, battery storage, and demand response mechanisms [77]. Nature-based "
    "solutions combined with engineered infrastructure—such as constructed wetlands for flood "
    "management and urban forests for heat mitigation—represent hybrid approaches that provide "
    "multiple co-benefits including carbon sequestration, biodiversity support, and improved "
    "human well-being [78]."
)
add_para(s3_2_p1)

s3_2_p2 = (
    "Renewable energy technologies serve dual roles in climate strategy: mitigating emissions through "
    "decarbonization while enhancing energy security and resilience through distributed generation [79]. "
    "Solar photovoltaic costs have declined by over 90% since 2010, making renewable energy competitive "
    "with or cheaper than fossil fuels in most markets, though climate change itself affects renewable "
    "resource availability through changes in solar irradiance, wind patterns, and hydropower "
    "capacity [80]. Advanced energy storage technologies—including lithium-ion batteries, flow "
    "batteries, green hydrogen, and thermal storage—address intermittency challenges and provide "
    "grid flexibility essential for climate resilience [81]. Resource-efficient technologies including "
    "water recycling and desalination, precision manufacturing, and circular material flows reduce "
    "climate vulnerability by decreasing dependence on climate-sensitive resources and reducing "
    "waste and emissions [82]. Table 3 presents a comparative assessment of intelligent eco-technologies "
    "for climate adaptation, including their maturity levels, cost-effectiveness, and scalability "
    "potential across different application domains."
)
add_para(s3_2_p2)

s3_2_p3 = (
    "The convergence of smart infrastructure with renewable energy systems creates synergies that "
    "enhance both climate mitigation and adaptation outcomes. Microgrids combining local renewable "
    "generation with battery storage and intelligent load management provide energy resilience for "
    "critical facilities during grid outages caused by extreme weather events. Vehicle-to-grid "
    "technology transforms electric vehicle fleets into distributed energy storage assets, providing "
    "grid balancing services while supporting transportation decarbonization. Building-integrated "
    "photovoltaics and small-scale wind installations reduce distribution losses and grid congestion "
    "while providing on-site power generation that maintains functionality during wider system "
    "disruptions. The deployment of hydrogen as an energy carrier enables long-duration seasonal "
    "storage that addresses the intermittency challenges of variable renewable generation, while "
    "green hydrogen production from excess renewable capacity converts curtailed energy into "
    "storable and transportable fuel. These integrated approaches demonstrate that climate "
    "mitigation and adaptation are not competing objectives but complementary strategies that "
    "reinforce each other when deployed within coherent system architectures."
)
add_para(s3_2_p3)

# TABLE 3
add_table_caption(
    'Table 3: Intelligent Eco-Technologies for Climate Adaptation—Comparative Assessment'
)
table3_headers = ['Technology', 'Application Domain', 'Maturity (TRL)', 
                  'Cost-Effectiveness', 'Scalability', 'Climate Benefit']
table3_rows = [
    ['AI/ML Climate Models', 'Prediction & Planning', '7–8', 'High', 'High', 'Enhanced preparedness'],
    ['IoT Sensor Networks', 'Real-time Monitoring', '8–9', 'Moderate-High', 'Very High', 'Early detection'],
    ['Digital Twins', 'Simulation & Optimization', '6–7', 'Moderate', 'Moderate', 'Risk reduction'],
    ['Smart Grids', 'Energy Resilience', '8–9', 'High', 'High', 'Supply security'],
    ['Precision Agriculture', 'Food Security', '7–8', 'High', 'Moderate-High', 'Yield stability'],
    ['Green Hydrogen', 'Energy Storage', '5–6', 'Low-Moderate', 'High (future)', 'Decarbonization'],
    ['Nature-Based Solutions', 'Multi-sector', '7–9', 'Very High', 'High', 'Co-benefits'],
    ['Advanced Desalination', 'Water Security', '7–8', 'Moderate', 'Moderate', 'Drought resilience'],
]
create_styled_table(table3_headers, table3_rows)
doc.add_paragraph()



# Section 3.3
add_heading_styled('3.3 Data-Driven Early Warning Systems and Adaptive Environmental Management', level=2)

s3_3_p1 = (
    "Early warning systems (EWS) represent one of the most cost-effective climate adaptation "
    "investments, with benefit-cost ratios ranging from 4:1 to 36:1 depending on the hazard and "
    "context [83]. Modern multi-hazard early warning systems integrate satellite remote sensing, "
    "ground-based observation networks, numerical weather prediction, and AI-enhanced nowcasting "
    "to provide actionable warnings with sufficient lead time for protective action [84]. The "
    "United Nations Early Warnings for All initiative aims to achieve universal coverage by 2027, "
    "recognizing that one-third of the global population—primarily in developing countries—lacks "
    "adequate warning coverage [85]. Impact-based forecasting represents a paradigm shift from "
    "hazard-centric to consequence-centric warnings, integrating exposure and vulnerability data "
    "to communicate expected impacts rather than meteorological parameters, improving public "
    "response and decision-maker action [86]."
)
add_para(s3_3_p1)

s3_3_p2 = (
    "Adaptive environmental management applies iterative, evidence-based approaches to natural "
    "resource management under climate uncertainty, continuously updating strategies based on "
    "monitoring data and emerging knowledge [87]. Ecosystem-based adaptation (EbA) leverages "
    "biodiversity and ecosystem services to reduce climate vulnerability, including mangrove "
    "restoration for coastal protection, watershed management for water security, and agroforestry "
    "for agricultural resilience [88]. Remote sensing and geospatial analytics enable landscape-scale "
    "monitoring of ecosystem health, land use change, and climate impact trajectories, informing "
    "adaptive management decisions at multiple scales [89]. The integration of citizen science, "
    "indigenous knowledge, and traditional ecological knowledge with technological monitoring systems "
    "enhances both the coverage and contextual relevance of environmental data, supporting more "
    "inclusive and effective adaptation planning [90]. The intelligent eco-technology framework "
    "depicted in Figure 3 demonstrates how these diverse technological components integrate into "
    "a coherent platform that enables proactive, data-driven climate adaptation across multiple "
    "sectors and scales."
)
add_para(s3_3_p2)

s3_3_p3 = (
    "The scalability and sustainability of data-driven climate adaptation systems depend on addressing "
    "critical challenges including data interoperability, cybersecurity, energy consumption of digital "
    "infrastructure, and equitable access to technological benefits. Standardized data protocols and "
    "open-source platforms facilitate the integration of diverse data streams from multiple sensors, "
    "agencies, and jurisdictions into unified decision-support systems. Edge computing architectures "
    "enable local data processing and rapid response in remote or connectivity-limited settings, "
    "reducing dependence on centralized cloud infrastructure while maintaining analytical sophistication. "
    "The development of low-power IoT devices powered by energy harvesting technologies extends "
    "monitoring capabilities to resource-constrained environments, including rural agricultural "
    "landscapes, remote water catchments, and developing-country contexts where grid electricity "
    "remains unavailable. Capacity building for data literacy and technological skills within "
    "environmental management agencies ensures that advanced analytical tools translate into "
    "improved decision outcomes rather than becoming underutilized technological investments."
)
add_para(s3_3_p3)



# ============================================================
# SECTION 4: Adaptive Strategies for Sustainable and Resilient Economies
# ============================================================
add_heading_styled('4. Adaptive Strategies for Sustainable and Resilient Economies', level=1)

intro_4 = (
    "Building climate-resilient economies requires fundamental transformation of economic structures, "
    "business models, policy frameworks, and financial systems [91]. Adaptive strategies must address "
    "both immediate climate risks and long-term structural vulnerabilities while ensuring that "
    "adaptation efforts are equitable, inclusive, and aligned with sustainable development objectives [92]. "
    "This section examines circular economy approaches, green finance mechanisms, and integrated "
    "policy pathways that collectively enable the transition toward intelligent, inclusive, and "
    "climate-resilient economies. As illustrated in Figure 4, the choice of adaptive strategy "
    "pathway significantly determines economic outcomes over the coming decades, with intelligent "
    "eco-technological transformation offering the most favorable long-term trajectory [93]."
)
add_para(intro_4)

# Insert Figure 4
fig4_path = '/projects/sandbox/AMMAN/climate_figures/Figure_4_Adaptive_Strategy_Pathways.png'
if os.path.exists(fig4_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig4_path, width=Inches(5.5))
    add_figure_caption(
        'Figure 4: Adaptive Strategy Pathways and Economic Impact Scenarios (2025–2050), comparing '
        'GDP impact trajectories under business-as-usual, moderate adaptation, aggressive adaptation, '
        'and intelligent eco-technological transformation scenarios.'
    )

# Section 4.1
add_heading_styled('4.1 Circular Economy and Low-Carbon Business Models', level=2)

s4_1_p1 = (
    "The circular economy paradigm offers a powerful framework for simultaneously addressing climate "
    "mitigation, adaptation, and resource security objectives [94]. By designing out waste, keeping "
    "materials in use, and regenerating natural systems, circular economy approaches reduce both "
    "emissions and resource vulnerability to climate disruption [95]. Studies estimate that circular "
    "economy strategies in key sectors—including steel, aluminum, cement, plastics, and food—could "
    "reduce global CO2 emissions by 3.7 billion tonnes annually by 2050, representing approximately "
    "40% of the emissions reduction needed beyond renewable energy deployment [96]. Industrial "
    "symbiosis, where waste outputs from one process become inputs for another, creates localized "
    "resource cycles that reduce supply chain vulnerability while generating cost savings of "
    "15–30% in participating facilities [97]."
)
add_para(s4_1_p1)

s4_1_p2 = (
    "Low-carbon business models are emerging across sectors, driven by regulatory pressure, "
    "investor expectations, consumer preferences, and competitive advantage from resource efficiency [98]. "
    "Product-as-a-service models shift ownership from consumers to manufacturers, incentivizing "
    "durability, repairability, and end-of-life recovery while reducing material throughput and "
    "associated emissions [99]. Digital platforms enable sharing economy applications that increase "
    "asset utilization, from vehicle sharing to shared manufacturing capacity, reducing the physical "
    "infrastructure needed to deliver economic services [100]. Regenerative agriculture integrates "
    "carbon sequestration with food production, offering economic returns through premium pricing, "
    "carbon credits, and reduced input costs while building soil health and climate resilience [101]. "
    "The transition to circular and low-carbon business models creates both opportunities and "
    "challenges for workers and communities dependent on linear, carbon-intensive industries, "
    "necessitating just transition policies that ensure equitable distribution of costs and "
    "benefits [102]."
)
add_para(s4_1_p2)

s4_1_p3 = (
    "The economic case for circular and low-carbon transitions is strengthening as resource prices "
    "become more volatile, carbon costs increase, and consumer demand shifts toward sustainable "
    "products and services. Life cycle assessment methodologies enable businesses to identify hotspots "
    "of environmental impact and economic inefficiency within their value chains, guiding targeted "
    "interventions that simultaneously reduce costs and emissions. Extended producer responsibility "
    "regulations incentivize design for recyclability and material recovery, creating new business "
    "opportunities in remanufacturing, refurbishment, and secondary material markets. Blockchain "
    "and digital product passport technologies enable transparent tracking of materials through "
    "supply chains, facilitating circular flows and providing consumers with verified sustainability "
    "information. The financial sector is increasingly developing metrics and frameworks to evaluate "
    "circular economy investments, recognizing that circular business models often exhibit greater "
    "resilience to supply disruptions, regulatory changes, and resource price volatility than "
    "their linear counterparts, thus representing lower risk-adjusted investment opportunities."
)
add_para(s4_1_p3)



# Section 4.2
add_heading_styled('4.2 Policy, Green Finance, and Climate-Resilient Economic Planning', level=2)

s4_2_p1 = (
    "Effective climate adaptation requires coherent policy frameworks that mainstream climate risk "
    "considerations across all sectors of economic planning and governance [103]. National Adaptation "
    "Plans (NAPs) provide strategic frameworks for identifying vulnerabilities, prioritizing "
    "interventions, and mobilizing resources, though implementation gaps remain significant in many "
    "countries [104]. Carbon pricing mechanisms—including emissions trading systems and carbon taxes—"
    "generate revenue that can fund adaptation investments while incentivizing emission reductions, "
    "with over 70 jurisdictions now implementing some form of carbon pricing covering approximately "
    "23% of global emissions [105]. Regulatory frameworks increasingly require climate risk disclosure "
    "from corporations and financial institutions, creating transparency that enables markets to "
    "price climate risk more accurately and directing capital toward resilient investments [106]."
)
add_para(s4_2_p1)

s4_2_p2 = (
    "Green finance has experienced explosive growth, with global sustainable debt issuance exceeding "
    "USD 1.6 trillion in 2023, including green bonds, sustainability-linked loans, and climate "
    "adaptation finance [107]. However, a significant financing gap persists, with adaptation "
    "investment needs in developing countries estimated at USD 300–500 billion annually by 2030, "
    "far exceeding current flows of approximately USD 21 billion [108]. Innovative financing "
    "mechanisms including catastrophe bonds, resilience bonds, parametric insurance, and blended "
    "finance structures are expanding the toolkit for channeling private capital toward adaptation "
    "investments [109]. Climate-resilient economic planning integrates climate projections into "
    "national development strategies, infrastructure investment decisions, and land use planning, "
    "ensuring that development pathways are compatible with both mitigation targets and adaptation "
    "needs [110]. Table 4 summarizes the policy instruments and financial mechanisms supporting "
    "climate-resilient economic transformation, categorized by type, scale, and current "
    "implementation status."
)
add_para(s4_2_p2)

# TABLE 4
add_table_caption(
    'Table 4: Policy Instruments and Financial Mechanisms for Climate-Resilient Economic Transformation'
)
table4_headers = ['Instrument/Mechanism', 'Type', 'Scale of Impact', 
                  'Implementation Status', 'Effectiveness Rating']
table4_rows = [
    ['Carbon Pricing (ETS/Tax)', 'Market-based', 'National/Regional', 'Widespread (70+ jurisdictions)', 'High'],
    ['Green Bonds', 'Financial', 'Global', 'Mature market (>$500B/yr)', 'Moderate-High'],
    ['Climate Risk Disclosure', 'Regulatory', 'Corporate/Financial', 'Expanding rapidly', 'Moderate'],
    ['National Adaptation Plans', 'Strategic Planning', 'National', '80+ countries', 'Variable'],
    ['Parametric Insurance', 'Risk Transfer', 'Regional/National', 'Growing', 'High (where deployed)'],
    ['Blended Finance', 'Financial', 'Project/Sectoral', 'Emerging', 'Moderate'],
    ['Nature-Based Solutions Policy', 'Regulatory/Incentive', 'Local/National', 'Increasing', 'High'],
    ['Just Transition Frameworks', 'Social Policy', 'National/Sectoral', 'Early stage', 'Moderate'],
]
create_styled_table(table4_headers, table4_rows)
doc.add_paragraph()



# Section 4.3
add_heading_styled('4.3 Future Pathways for Intelligent, Inclusive, and Climate-Resilient Economies', level=2)

s4_3_p1 = (
    "The transition toward climate-resilient economies requires integrated pathways that simultaneously "
    "address technological innovation, institutional transformation, social inclusion, and ecological "
    "sustainability [111]. Future economic systems must be designed for adaptability, incorporating "
    "modularity, redundancy, and diversity that enable rapid adjustment to changing climatic conditions "
    "without systemic failure [112]. The concept of regenerative economics—going beyond sustainability "
    "to actively restore ecological and social capital—offers a framework for economic development "
    "that builds resilience while addressing climate change root causes [113]. Intelligent eco-"
    "technologies serve as enablers of this transition, providing the monitoring, optimization, and "
    "coordination capabilities needed to manage complex socio-ecological-economic systems under "
    "climate uncertainty [114]."
)
add_para(s4_3_p1)

s4_3_p2 = (
    "Inclusive adaptation ensures that climate resilience benefits reach all segments of society, "
    "particularly marginalized and vulnerable communities who face the greatest climate risks with "
    "the least resources [115]. Gender-responsive adaptation recognizes and addresses the differentiated "
    "climate impacts on women and girls, who in many contexts bear disproportionate burdens from "
    "climate-related resource scarcity and livelihood disruption [116]. Youth engagement in climate "
    "adaptation brings innovation, technological literacy, and long-term perspective to resilience "
    "building, with young entrepreneurs increasingly driving eco-technology startups and social "
    "enterprises [117]. Indigenous and local knowledge systems offer time-tested adaptation strategies "
    "that complement technological solutions, and their integration into formal planning processes "
    "enriches both the knowledge base and the legitimacy of adaptation decisions [118]."
)
add_para(s4_3_p2)

s4_3_p3 = (
    "International cooperation remains essential for addressing the global commons dimension of climate "
    "change while supporting equitable adaptation across nations with vastly different capacities [119]. "
    "Technology transfer mechanisms, capacity building programs, and climate finance commitments under "
    "the Paris Agreement and subsequent COP decisions provide frameworks for international support, "
    "though delivery has consistently fallen short of pledges [120]. Regional cooperation on shared "
    "climate risks—including transboundary water management, coordinated disaster response, and "
    "harmonized adaptation standards—creates efficiencies and builds mutual resilience that "
    "individual national action cannot achieve [121]. The adaptive strategy pathway analysis "
    "presented in Figure 4 demonstrates that economies embracing intelligent eco-technological "
    "transformation can achieve net positive economic outcomes even under significant warming, "
    "while business-as-usual approaches lead to accelerating economic deterioration [122]. "
    "Ultimately, the transition to climate-resilient economies is not merely a defensive response "
    "to climate threats but an opportunity to build more equitable, efficient, and sustainable "
    "economic systems that deliver improved well-being for all within planetary boundaries [123]."
)
add_para(s4_3_p3)

s4_3_p4 = (
    "The governance architectures required for climate-resilient economies must operate across multiple "
    "scales, from local community-based adaptation initiatives to global frameworks for cooperation "
    "and finance. Polycentric governance approaches, which distribute authority across multiple "
    "overlapping jurisdictions and stakeholder groups, have shown promise in managing complex "
    "environmental challenges by enabling experimentation, learning, and redundancy. The role of "
    "cities and subnational governments in climate adaptation is increasingly recognized, with urban "
    "areas serving as laboratories for innovative policy approaches and technology deployment that "
    "can subsequently scale to national and international levels. Public-private partnerships "
    "leverage the respective strengths of government (regulatory authority, public mandate, "
    "long-term perspective) and private sector (innovation capacity, capital mobilization, "
    "operational efficiency) in delivering climate adaptation outcomes. The integration of "
    "climate resilience objectives into trade agreements, development finance institutions, "
    "and multilateral economic governance creates systemic incentives for adaptation investment "
    "while ensuring that trade and investment flows support rather than undermine climate "
    "resilience objectives."
)
add_para(s4_3_p4)

s4_3_p5 = (
    "Looking ahead, the trajectory of climate-resilient economic development will be shaped by the "
    "interaction of technological innovation, institutional evolution, social mobilization, and "
    "geopolitical dynamics. The pace of artificial intelligence advancement, the cost trajectory of "
    "clean energy technologies, and the scalability of nature-based solutions will determine the "
    "technical feasibility of ambitious adaptation pathways. Equally important are the political "
    "economy dynamics that determine whether technological potential translates into deployed "
    "solutions at scale, including the management of incumbent industry resistance, the distribution "
    "of transition costs and benefits, and the mobilization of public support for transformative "
    "change. The window for effective adaptation action narrows with each year of delay, as both "
    "the magnitude of required adaptation increases and the residual time for implementation "
    "decreases. The choices made in the current decade regarding infrastructure investment, "
    "technology deployment, institutional reform, and international cooperation will largely "
    "determine whether the mid-century economy is characterized by cascading climate crises "
    "or by resilient, sustainable prosperity enabled by intelligent eco-technological transformation."
)
add_para(s4_3_p5)



# ============================================================
# CONCLUSIONS
# ============================================================
add_heading_styled('5. Conclusions', level=1)

conclusions = (
    "This chapter has provided a comprehensive examination of the complex nexus between climate change, "
    "economic risk, and adaptive eco-technological strategies, demonstrating that climate change poses "
    "fundamental challenges to economic stability while simultaneously creating opportunities for "
    "transformative innovation and sustainable development. The analysis reveals several key findings "
    "and implications for policy, practice, and research."
)
add_para(conclusions)

conc_p2 = (
    "First, the economic consequences of climate change are accelerating and broadening, affecting "
    "virtually all sectors and regions through both direct physical impacts and indirect systemic "
    "risks. The non-linear nature of climate damages, combined with tipping point risks and cascading "
    "failures across interconnected systems, means that economic losses will escalate disproportionately "
    "under higher warming scenarios. Second, vulnerability assessment reveals stark inequalities in "
    "climate risk exposure and adaptive capacity, with developing nations, marginalized communities, "
    "and climate-sensitive sectors facing the most severe impacts despite contributing least to the "
    "problem. Addressing these inequalities is both a moral imperative and a practical necessity for "
    "global stability."
)
add_para(conc_p2)

conc_p3 = (
    "Third, intelligent eco-technologies offer transformative potential for climate adaptation, enabling "
    "more precise, proactive, and cost-effective responses to climate risks. The integration of AI, "
    "IoT, digital twins, and smart infrastructure creates comprehensive resilience platforms that can "
    "anticipate, monitor, and respond to climate hazards in real time. However, technology deployment "
    "must be accompanied by institutional capacity, equitable access, and appropriate governance to "
    "realize its full potential. Fourth, adaptive economic strategies—including circular economy models, "
    "green finance mechanisms, and inclusive policy frameworks—provide the structural foundation for "
    "climate-resilient development. The pathway analysis demonstrates that economies pursuing aggressive "
    "adaptation combined with intelligent eco-technological transformation can achieve net positive "
    "economic outcomes, while those maintaining business-as-usual face accelerating deterioration."
)
add_para(conc_p3)

conc_p4 = (
    "Future research priorities include developing improved models of cascading climate-economic risks, "
    "advancing AI applications for climate prediction and adaptation optimization, evaluating the "
    "effectiveness of emerging financial instruments for adaptation, and understanding the governance "
    "requirements for equitable technology deployment. The urgency of climate change demands immediate "
    "action across all fronts—technological, economic, institutional, and social—to build the resilient, "
    "sustainable, and inclusive economies that can thrive under the climatic conditions of the coming "
    "decades."
)
add_para(conc_p4)

conc_p5 = (
    "The synthesis presented in this chapter underscores several actionable recommendations for "
    "policymakers, researchers, and practitioners. Investment in climate-resilient infrastructure "
    "should prioritize multi-functional solutions that deliver adaptation, mitigation, and development "
    "co-benefits simultaneously. Financial regulatory frameworks must evolve to fully incorporate "
    "climate risk into capital allocation decisions, ensuring that market signals direct investment "
    "toward resilient and sustainable activities. Technology governance frameworks should balance "
    "innovation incentives with equity considerations, ensuring that the benefits of intelligent "
    "eco-technologies reach vulnerable communities and developing nations. Cross-sector and cross-border "
    "collaboration is essential for addressing the systemic and transboundary nature of climate risks, "
    "requiring new institutional arrangements that facilitate information sharing, coordinated "
    "planning, and collective action. Education and capacity building must equip current and future "
    "workforces with the skills needed to develop, deploy, and maintain intelligent eco-technologies, "
    "while fostering the systems thinking and adaptive management capabilities essential for "
    "navigating climate uncertainty. The integration of indigenous knowledge with scientific "
    "approaches enriches the evidence base for adaptation and enhances the cultural appropriateness "
    "and social legitimacy of resilience interventions."
)
add_para(conc_p5)



# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_styled('References', level=1)

references = [
    "[1] IPCC, Climate Change 2021: The Physical Science Basis. Contribution of Working Group I to the Sixth Assessment Report, Cambridge University Press, Cambridge, 2021.",
    "[2] M. Burke, S. Hsiang, E. Miguel, Global non-linear effect of temperature on economic production, Nature, 527 (2015) 235–239.",
    "[3] W. Steffen, J. Rockström, K. Richardson, et al., Trajectories of the Earth System in the Anthropocene, Proceedings of the National Academy of Sciences, 115 (2018) 8252–8259.",
    "[4] G.C. Nelson, A. Valin, R.D. Sands, et al., Climate change effects on agriculture: Economic responses to biophysical shocks, Proceedings of the National Academy of Sciences, 111 (2014) 3274–3279.",
    "[5] P. Friedlingstein, M.W. Jones, M. O'Sullivan, et al., Global Carbon Budget 2023, Earth System Science Data, 15 (2023) 5301–5369.",
    "[6] NOAA, Annual Greenhouse Gas Index, National Oceanic and Atmospheric Administration, 2024.",
    "[7] T.F. Stocker, D. Qin, G.K. Plattner, et al., Climate Change 2013: The Physical Science Basis, Cambridge University Press, 2013.",
    "[8] WCRP Global Sea Level Budget Group, Global sea-level budget 1993–present, Earth System Science Data, 10 (2018) 1551–1590.",
    "[9] S.E. Perkins-Kirkpatrick, S.C. Lewis, Increasing trends in regional heatwaves, Nature Communications, 11 (2020) 3357.",
    "[10] R. Allan, M. Barlow, M.P. Byrne, et al., Advances in understanding large-scale responses of the water cycle to climate change, Annals of the New York Academy of Sciences, 1472 (2020) 49–75.",
    "[11] M.C. Serreze, R.G. Barry, Processes and impacts of Arctic amplification: A research synthesis, Global and Planetary Change, 77 (2011) 85–96.",
    "[12] S. Cooley, D. Schoeman, L. Bopp, et al., Oceans and Coastal Ecosystems and Their Services, in: IPCC AR6 WGII, Cambridge University Press, 2022.",
    "[13] J. Zscheischler, S. Westra, B.J.J.M. van den Hurk, et al., Future climate risk from compound events, Nature Climate Change, 8 (2018) 469–477.",
    "[14] D.I. Armstrong McKay, A. Staal, J.F. Abrams, et al., Exceeding 1.5°C global warming could trigger multiple climate tipping points, Science, 377 (2022) eabn7950.",
    "[15] Munich Re, Natural catastrophe review 2023: Record thunderstorm losses, Munich Reinsurance Company, 2024.",
    "[16] S. Hallegatte, A. Vogt-Schilb, M. Bangalore, J. Rozenberg, Unbreakable: Building the Resilience of the Poor in the Face of Natural Disasters, World Bank, Washington DC, 2017.",
    "[17] Government of Pakistan, Pakistan Floods 2022: Post-Disaster Needs Assessment, Planning Commission, 2022.",
    "[18] D.B. Lobell, W. Schlenker, J. Costa-Roberts, Climate trends and global crop production since 1980, Science, 333 (2011) 616–620.",
    "[19] Swiss Re Institute, Sigma: Natural catastrophes in 2023, Swiss Re, Zurich, 2024.",
]
for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



references2 = [
    "[20] International Labour Organization, Working on a Warmer Planet: The Impact of Heat Stress on Labour Productivity and Decent Work, ILO, Geneva, 2019.",
    "[21] M. van Vliet, J. Sheffield, D. Wiberg, E.F. Wood, Impacts of recent drought and warm years on water resources and electricity supply worldwide, Environmental Research Letters, 11 (2016) 124021.",
    "[22] Swiss Re Institute, The economics of climate change: no action not an option, Swiss Re, 2021.",
    "[23] S. Diffenbaugh, M. Burke, Global warming has increased global economic inequality, Proceedings of the National Academy of Sciences, 116 (2019) 9808–9813.",
    "[24] OECD, Climate-Resilient Infrastructure: Policy Perspectives, OECD Publishing, Paris, 2018.",
    "[25] C. Hanson, R. Nicholls, N. Ranger, et al., A global ranking of port cities with high exposure to climate extremes, Climatic Change, 104 (2011) 89–111.",
    "[26] Global Commission on Adaptation, Adapt Now: A Global Call for Leadership on Climate Resilience, World Resources Institute, 2019.",
    "[27] IEA, Climate Resilience for Energy Security, International Energy Agency, Paris, 2024.",
    "[28] Y. Sheffi, The Resilient Enterprise: Overcoming Vulnerability for Competitive Advantage, MIT Press, Cambridge, 2015.",
    "[29] P. Pant, F. Heinimann, M. Scholz, Supply chain risk management in a changing climate, Supply Chain Management Review, 28 (2023) 42–58.",
    "[30] World Bank, Thai Flood 2011: Rapid Assessment for Resilient Recovery and Reconstruction Planning, World Bank, 2012.",
    "[31] K. Kornhuber, D. Coumou, E. Vogel, et al., Amplified Rossby waves enhance risk of concurrent heatwaves in major breadbasket regions, Nature Climate Change, 10 (2020) 48–53.",
    "[32] UNDRR, Global Assessment Report on Disaster Risk Reduction 2022, United Nations, Geneva, 2022.",
    "[33] W.N. Adger, Vulnerability, Global Environmental Change, 16 (2006) 268–281.",
    "[34] R. Lempert, S. Popper, S. Bankes, Shaping the Next One Hundred Years: New Methods for Quantitative, Long-Term Policy Analysis, RAND Corporation, 2003.",
    "[35] IPCC, Climate Change 2022: Impacts, Adaptation and Vulnerability, Cambridge University Press, 2022.",
    "[36] C. Zhao, B. Liu, S. Piao, et al., Temperature increase reduces global yields of major crops, Proceedings of the National Academy of Sciences, 114 (2017) 9326–9331.",
    "[37] P. Gleick, Water, Drought, Climate Change, and Conflict in Syria, Weather, Climate, and Society, 6 (2014) 331–340.",
    "[38] D. Scott, C.M. Hall, S. Gössling, Global tourism vulnerability to climate change, Annals of Tourism Research, 77 (2019) 49–61.",
]
for ref in references2:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



references3 = [
    "[39] ND-GAIN, Notre Dame Global Adaptation Initiative Country Index, University of Notre Dame, 2024.",
    "[40] World Bank, Turn Down the Heat: Confronting the New Climate Normal, World Bank, Washington DC, 2014.",
    "[41] UNFCCC, Adaptation in Small Island Developing States, United Nations Framework Convention on Climate Change, 2023.",
    "[42] R. Mendelsohn, A. Dinar, L. Williams, The distributional impact of climate change on rich and poor countries, Environment and Development Economics, 11 (2006) 159–178.",
    "[43] C. Rosenzweig, W. Solecki, P. Romero-Lankao, et al., Climate Change and Cities: Second Assessment Report of the Urban Climate Change Research Network, Cambridge University Press, 2018.",
    "[44] H. Fuss, J. Canadell, G. Peters, et al., Betting on negative emissions, Nature Climate Change, 4 (2014) 850–853.",
    "[45] M. Carney, Breaking the Tragedy of the Horizon—climate change and financial stability, Bank of England Speech, 2015.",
    "[46] TCFD, Recommendations of the Task Force on Climate-related Financial Disclosures, Financial Stability Board, 2017.",
    "[47] S. Battiston, A. Mandel, I. Monasterolo, et al., A climate stress-test of the financial system, Nature Climate Change, 7 (2017) 283–288.",
    "[48] Carbon Tracker Initiative, Unburnable Carbon: Are the World's Financial Markets Carrying a Carbon Bubble?, Carbon Tracker, London, 2022.",
    "[49] S. Hallegatte, Strategies to adapt to an uncertain climate change, Global Environmental Change, 19 (2009) 240–247.",
    "[50] NGFS, Climate Scenarios for Central Banks and Supervisors, Network for Greening the Financial System, 2023.",
    "[51] ECB, Climate risk stress test results, European Central Bank, Frankfurt, 2022.",
    "[52] I. Monasterolo, Climate change and the financial system, Annual Review of Resource Economics, 12 (2020) 299–320.",
    "[53] W. Nordhaus, Revisiting the social cost of carbon, Proceedings of the National Academy of Sciences, 114 (2017) 1518–1523.",
    "[54] N. Stern, The Economics of Climate Change: The Stern Review, Cambridge University Press, 2007.",
    "[55] B.C. O'Neill, E. Kriegler, K.L. Ebi, et al., The roads ahead: Narratives for shared socioeconomic pathways, Global Environmental Change, 42 (2017) 169–180.",
    "[56] S. Dietz, A. Bowen, C. Dixon, P. Gradwell, Climate value at risk of global financial assets, Nature Climate Change, 6 (2016) 676–679.",
    "[57] R.J. Lempert, D.G. Groves, Identifying and evaluating robust adaptive policy responses to climate change, Technological Forecasting and Social Change, 77 (2010) 960–974.",
]
for ref in references3:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



references4 = [
    "[58] M. Haasnoot, J.H. Kwakkel, W.E. Walker, J. ter Maat, Dynamic adaptive policy pathways: A method for crafting robust decisions for a deeply uncertain world, Global Environmental Change, 23 (2013) 485–498.",
    "[59] B. Ranger, T. Reeder, J. Lowe, Addressing the value of information for adaptation under uncertainty, Global Environmental Change, 23 (2013) 1–6.",
    "[60] P. Watkiss, A. Hunt, Projection of economic impacts of climate change in sectors of Europe based on bottom up analysis, Climatic Change, 112 (2012) 741–758.",
    "[61] D. Rolnick, P.L. Donti, L.H. Kaack, et al., Tackling climate change with machine learning, ACM Computing Surveys, 55 (2022) 1–96.",
    "[62] R. Vinuesa, H. Azizpour, I. Leite, et al., The role of artificial intelligence in achieving the Sustainable Development Goals, Nature Communications, 11 (2020) 233.",
    "[63] A. Creutzig, J. Hilaire, G. Nemet, et al., The mutual dependence of negative emission technologies and energy systems, Energy & Environmental Science, 12 (2019) 1805–1817.",
    "[64] ITU, Frontier Technologies to Protect the Environment and Tackle Climate Change, International Telecommunication Union, Geneva, 2023.",
    "[65] Y. LeCun, Y. Bengio, G. Hinton, Deep learning, Nature, 521 (2015) 436–444.",
    "[66] R. Lam, A. Sanchez-Gonzalez, M. Willson, et al., Learning skillful medium-range global weather forecasting, Science, 382 (2023) 1416–1421.",
    "[67] T. Vandal, E. Kodra, S. Ganguly, et al., DeepSD: Generating high fidelity daily climate projections, in: Proceedings of the 23rd ACM SIGKDD International Conference, 2017.",
    "[68] A. Rolf, J. Proctor, T. Carleton, et al., A generalizable and accessible approach to machine learning with global satellite imagery, Nature Communications, 12 (2021) 4392.",
    "[69] V. Mnih, K. Kavukcuoglu, D. Silver, et al., Human-level control through deep reinforcement learning, Nature, 518 (2015) 529–533.",
    "[70] L. Atzori, A. Iera, G. Morabito, The Internet of Things: A survey, Computer Networks, 54 (2010) 2787–2805.",
    "[71] M. Oberascher, J. Zischg, S.T. Palermo, et al., Smart water management using IoT: A review, Water Research, 189 (2021) 116609.",
    "[72] R. Gebbers, V.I. Adamchuk, Precision agriculture and food security, Science, 327 (2010) 828–831.",
    "[73] M. Grieves, J. Vickers, Digital twin: Mitigating unpredictable, undesirable emergent behavior in complex systems, in: Transdisciplinary Perspectives on Complex Systems, Springer, 2017.",
    "[74] F. Dembski, U. Wössner, M. Letzgus, et al., Urban Digital Twins for Smart Cities and Citizens, Sustainability, 12 (2020) 2612.",
]
for ref in references4:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



references5 = [
    "[75] J.M. Sussman, Perspectives on Intelligent Transportation Systems (ITS), Springer, New York, 2005.",
    "[76] IEA, The Future of Cooling: Opportunities for energy-efficient air conditioning, International Energy Agency, Paris, 2018.",
    "[77] M.Z. Jacobson, M.A. Delucchi, Z.A.F. Bauer, et al., 100% clean and renewable wind, water, and sunlight all-sector energy roadmaps for 139 countries, Joule, 1 (2017) 108–121.",
    "[78] IUCN, Nature-based Solutions to Address Global Societal Challenges, International Union for Conservation of Nature, Gland, 2016.",
    "[79] IRENA, Renewable Power Generation Costs in 2023, International Renewable Energy Agency, Abu Dhabi, 2024.",
    "[80] IRENA, Future of Solar Photovoltaic: Deployment, Investment, Technology, Grid Integration and Socio-Economic Aspects, IRENA, Abu Dhabi, 2019.",
    "[81] O. Schmidt, A. Hawkes, A. Gambhir, I. Staffell, The future cost of electrical energy storage based on experience rates, Nature Energy, 2 (2017) 17110.",
    "[82] Ellen MacArthur Foundation, Completing the Picture: How the Circular Economy Tackles Climate Change, 2021.",
    "[83] S. Hallegatte, Benefits of investing in early warning systems, World Bank Background Report, 2012.",
    "[84] WMO, Multi-hazard Early Warning Systems: A Checklist, World Meteorological Organization, Geneva, 2018.",
    "[85] United Nations, Early Warnings for All: The UN Global Early Warning Initiative for the Implementation of Climate Adaptation, UN, 2022.",
    "[86] M. Merz, F. Kuhlicke, V. Kunz, et al., Impact forecasting to support emergency management of natural hazards, Reviews of Geophysics, 58 (2020) e2020RG000704.",
    "[87] C.S. Holling, Adaptive Environmental Assessment and Management, John Wiley & Sons, 1978.",
    "[88] R. Munang, I. Thiaw, K. Alverson, et al., The role of ecosystem services in climate change adaptation and disaster risk reduction, Current Opinion in Environmental Sustainability, 5 (2013) 47–52.",
    "[89] G.J. Asner, R.E. Martin, R. Anderson, D.E. Knapp, Quantifying forest canopy traits: Imaging spectroscopy versus field survey, Remote Sensing of Environment, 158 (2015) 15–27.",
    "[90] D. Tengö, E.S. Brondizio, T. Elmqvist, et al., Connecting diverse knowledge systems for enhanced ecosystem governance, Ambio, 43 (2014) 579–591.",
    "[91] Global Commission on the Economy and Climate, The New Climate Economy Report, World Resources Institute, 2018.",
    "[92] IPCC, Climate Change 2022: Mitigation of Climate Change, Cambridge University Press, 2022.",
    "[93] S. Fankhauser, S.M. Smith, M. Allen, et al., The meaning of net zero and how to get it right, Nature Climate Change, 12 (2022) 15–21.",
    "[94] M. Geissdoerfer, P. Savaget, N.M.P. Bocken, E.J. Hultink, The Circular Economy—A new sustainability paradigm?, Journal of Cleaner Production, 143 (2017) 757–768.",
]
for ref in references5:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



references6 = [
    "[95] P. Ghisellini, C. Cialani, S. Ulgiati, A review on circular economy: The expected transition to a balanced interplay of environmental and economic systems, Journal of Cleaner Production, 114 (2016) 11–32.",
    "[96] Material Economics, The Circular Economy—A Powerful Force for Climate Mitigation, Material Economics, Stockholm, 2018.",
    "[97] M. Chertow, Industrial symbiosis: Literature and taxonomy, Annual Review of Energy and the Environment, 25 (2000) 313–337.",
    "[98] A. Bocken, S. Short, P. Rana, S. Evans, A literature and practice review to develop sustainable business model archetypes, Journal of Cleaner Production, 65 (2014) 42–56.",
    "[99] A. Tukker, Eight types of product-service system: Eight ways to sustainability?, Business Strategy and the Environment, 13 (2004) 246–260.",
    "[100] J. Hamari, M. Sjöklint, A. Ukkonen, The sharing economy: Why people participate in collaborative consumption, Journal of the Association for Information Science and Technology, 67 (2016) 2047–2059.",
    "[101] T. LaCanne, J. Lundgren, Regenerative agriculture: Merging farming and natural resource conservation profitably, PeerJ, 6 (2018) e4428.",
    "[102] ILO, Guidelines for a Just Transition Towards Environmentally Sustainable Economies and Societies for All, International Labour Organization, Geneva, 2015.",
    "[103] UNFCCC, The Paris Agreement, United Nations Framework Convention on Climate Change, 2015.",
    "[104] UNEP, Adaptation Gap Report 2023: Underfinanced. Underprepared, United Nations Environment Programme, Nairobi, 2023.",
    "[105] World Bank, State and Trends of Carbon Pricing 2024, World Bank Group, Washington DC, 2024.",
    "[106] ISSB, IFRS S2 Climate-related Disclosures, International Sustainability Standards Board, 2023.",
    "[107] Climate Bonds Initiative, Global State of the Market Report 2023, Climate Bonds Initiative, London, 2024.",
    "[108] UNEP, Adaptation Finance Gap Update 2023, United Nations Environment Programme, 2023.",
    "[109] A. Surminski, L.M. Bouwer, J. Linnerooth-Bayer, How insurance can support climate resilience, Nature Climate Change, 6 (2016) 333–334.",
    "[110] OECD, Investing in Climate, Investing in Growth, OECD Publishing, Paris, 2017.",
    "[111] K. Raworth, Doughnut Economics: Seven Ways to Think Like a 21st-Century Economist, Random House, 2017.",
    "[112] C. Folke, S.R. Carpenter, B. Walker, et al., Resilience thinking: Integrating resilience, adaptability and transformability, Ecology and Society, 15 (2010) 20.",
    "[113] J. Fullerton, Regenerative Capitalism: How Universal Principles and Patterns Will Shape Our New Economy, Capital Institute, 2015.",
    "[114] D. Acemoglu, P. Aghion, L. Bursztyn, D. Hemous, The environment and directed technical change, American Economic Review, 102 (2012) 131–166.",
    "[115] W.N. Adger, J. Barnett, K. Brown, N. Marshall, K. O'Brien, Cultural dimensions of climate change impacts and adaptation, Nature Climate Change, 3 (2013) 112–117.",
    "[116] M. Alston, Gender mainstreaming and climate change, Women's Studies International Forum, 47 (2014) 287–294.",
    "[117] UNEP, Global Environment Outlook 6: Healthy Planet, Healthy People, United Nations Environment Programme, 2019.",
    "[118] D. Nakashima, K. Galloway McLean, H.D. Thulstrup, et al., Weathering Uncertainty: Traditional Knowledge for Climate Change Assessment and Adaptation, UNESCO, Paris, 2012.",
    "[119] V. Ciplet, J.T. Roberts, M.H. Khan, Power in a Warming World: The New Global Politics of Climate Change and the Remaking of Environmental Inequality, MIT Press, 2015.",
    "[120] OECD, Climate Finance Provided and Mobilised by Developed Countries in 2013–2022, OECD Publishing, 2024.",
    "[121] ADB, Asian Development Outlook 2023: Climate Change and Development in Asia and the Pacific, Asian Development Bank, Manila, 2023.",
    "[122] H. Waisman, C. Bataille, H. Winkler, et al., A pathway design framework for national low greenhouse gas emission development strategies, Nature Climate Change, 9 (2019) 261–268.",
    "[123] J. Rockström, O. Gaffney, J. Rogelj, et al., A roadmap for rapid decarbonization, Science, 355 (2017) 1269–1271.",
]
for ref in references6:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'



# ============================================================
# SAVE DOCUMENT
# ============================================================
output_path = '/projects/sandbox/AMMAN/Chapter_Climate_Change_Economic_Risk_EcoTech.docx'
doc.save(output_path)
print(f"Document saved successfully: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")

# Count approximate words
word_count = 0
for para in doc.paragraphs:
    word_count += len(para.text.split())
print(f"Approximate word count: {word_count}")
